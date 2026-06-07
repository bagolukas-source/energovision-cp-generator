#!/usr/bin/env python3
"""
Validátor Energovision posudku — kontrola subjektívnych formulácií,
interných údajov a sanity check čísel.

Príklad:
    python validate.py Posudok_FVE_BESS_KLIENT.docx
"""
import argparse
import re
import sys
from pathlib import Path
from docx import Document


SUBJEKTIVNE = [
    'najlepší', 'najsilnejší', 'najslabší', 'najvýhodnejší',
    'vynikajúc', 'skvelý', 'vynikajúce',
    'odporúčame klientovi', 'odporúčame realizovať',
    'strategicky výhodné', 'kľúčový prínos',
]

INTERNE_UDAJE = [
    'Hagard', 'Wattstor', 'feed.xml', 'marže',
    'INTERNÁ', 'energovision_zľava', '35 % zľava',
    'cena modulu', 'efektívna cena modulu',
    'nákupná cena', 'veľkoobchodný',
]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx', help='Cesta k DOCX súboru')
    ap.add_argument('--strict', action='store_true',
                    help='Failne pri akejkoľvek issue (exit 1)')
    args = ap.parse_args()

    fn = Path(args.docx)
    if not fn.exists():
        print(f"⚠ Súbor neexistuje: {fn}", file=sys.stderr)
        sys.exit(1)

    d = Document(str(fn))
    text = ' '.join(p.text for p in d.paragraphs)
    for tab in d.tables:
        for row in tab.rows:
            for c in row.cells:
                text += ' ' + c.text

    text_lower = text.lower()

    print(f"Súbor: {fn.name}")
    print(f"  Odsekov: {len(d.paragraphs)}, Tabuliek: {len(d.tables)}")
    print(f"  Dĺžka textu: {len(text):,} znakov")
    print()

    issues_count = 0

    # 1. Subjektívne formulácie
    found_subj = [w for w in SUBJEKTIVNE if w.lower() in text_lower]
    if found_subj:
        print(f"⚠ SUBJEKTÍVNE FORMULÁCIE ({len(found_subj)}):")
        for w in found_subj:
            print(f'    – „{w}"')
        issues_count += len(found_subj)
    else:
        print("✓ Žiadne subjektívne formulácie")

    # 2. Interné údaje
    found_int = [w for w in INTERNE_UDAJE if w.lower() in text_lower]
    if found_int:
        print(f"\n⚠ INTERNÉ ÚDAJE ({len(found_int)}):")
        for w in found_int:
            print(f'    – „{w}"')
        issues_count += len(found_int)
    else:
        print("✓ Žiadne interné údaje (ceny komponentov, zľavy, marže)")

    # 3. Sanity check čísel — návratnosť
    payback_matches = re.findall(r'(?:návratnosť|návrat\.?\s*)(?:[^\d]*?)(\d+[,.]?\d*)\s*r(?:oka|oky|okov)?', text)
    if payback_matches:
        paybacks = [float(p.replace(',', '.')) for p in payback_matches]
        out_of_range = [p for p in paybacks if p < 1 or p > 15]
        if out_of_range:
            print(f"\n⚠ NÁVRATNOSŤ MIMO ROZSAHU 1–15 r: {out_of_range}")
            issues_count += 1
        else:
            print(f"\n✓ Návratnosti v rozsahu (min {min(paybacks):.1f}, max {max(paybacks):.1f} r)")

    # 4. NPV check — má byť uvedené
    npv_present = bool(re.search(r'NPV', text))
    if npv_present:
        print("✓ NPV uvedené v posudku")
    else:
        print("⚠ NPV nie je uvedené")
        issues_count += 1

    # 5. CAPEX check
    capex_present = bool(re.search(r'CAPEX|cena\s+diela', text, re.IGNORECASE))
    if capex_present:
        print("✓ CAPEX uvedené")
    else:
        print("⚠ CAPEX nie je uvedené")
        issues_count += 1

    # 6. Tarif transparentnosť
    if 'odhad' in text_lower and 'tarif' in text_lower:
        print("✓ Tarif uvedený s upozornením na odhad")
    elif re.search(r'\d[,.]?\d*\s*€/kWh', text):
        print("ℹ Tarif uvedený (over či je to skutočná hodnota alebo odhad)")

    print(f"\n{'=' * 50}")
    if issues_count == 0:
        print("✓ VALIDÁCIA OK — bez issue")
        sys.exit(0)
    else:
        print(f"⚠ ISSUES: {issues_count}")
        if args.strict:
            sys.exit(1)


if __name__ == '__main__':
    main()
