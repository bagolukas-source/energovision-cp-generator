"""
B2B Document Generation
========================

Generuje docx dokumenty pre B2B projekty:
- Nová klasická ZoD (Zmluva o dielo) — z templates_b2b/Nova_klasicka_ZOD.docx
- Energovision Faktúra — z templates_b2b/Energovision_Faktura_B2B.docx

Šablóny používajú Flowii-style placeholdery: {companyName}, {totalAmount}, atď.
Faktúra používa {{double brace}} placeholdery.

Vstupy:
- project_id (Supabase UUID) — načíta projects + customers + project_milestones
- Voliteľne fa_no (1/2/3) pre faktúru — určuje milestone

Výstup:
- Cesta k vyplnenému .docx súboru
"""

from __future__ import annotations
import os
import re
import logging
from pathlib import Path
from datetime import datetime, date
from docx import Document

log = logging.getLogger("generuj_b2b")

TEMPLATES_DIR = Path(__file__).parent / "templates_b2b"


def _fmt_eur(amount: float | int | None) -> str:
    if amount is None:
        return "0,00"
    try:
        n = float(amount)
    except (TypeError, ValueError):
        return "0,00"
    return f"{n:,.2f}".replace(",", " ").replace(".", ",")


def _fmt_date_sk(d: str | date | datetime | None) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        try:
            d = datetime.fromisoformat(d.replace("Z", "+00:00")).date()
        except Exception:
            return d
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%d.%m.%Y")


def _replace_in_paragraph(paragraph, mapping: dict):
    """Replace placeholders in a paragraph, preserving formatting where possible."""
    # Najprv full text per paragraph, hľadáme placeholders
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return
    new_full = full
    for key, val in mapping.items():
        if key in new_full:
            new_full = new_full.replace(key, str(val) if val is not None else "")
    if new_full == full:
        return
    # Najdeme prvý nesplnený run a do neho dáme všetko, ostatné runs vymažeme
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""


def _replace_in_doc(doc: Document, mapping: dict):
    """Replace placeholders v paragraphoch + tabuľkách + headeroch/footeroch."""
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)
                # Vnorené tabuľky
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                _replace_in_paragraph(np, mapping)
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            for p in hdr.paragraphs:
                _replace_in_paragraph(p, mapping)
            for tbl in hdr.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            _replace_in_paragraph(p, mapping)


# =========================================================
# Nová klasická ZoD
# =========================================================
def generuj_zod(*, project: dict, customer: dict, milestones: list[dict], out_dir: str | Path) -> str:
    """
    Vyplní ZoD šablónu dátami z B2B projektu.

    project (zo Supabase `projects`): id, name, project_code, scale_kwp, contract_value_no_vat,
                                       contract_signed_at, custom_fields (jsonb)
    customer: company_name, ico, dic, ic_dph, billing_street/city/postal_code,
              first_name, last_name, email, phone, title_before
    milestones: zoznam project_milestones (pre dátumy FA1/FA2/FA3)
    """
    template_path = TEMPLATES_DIR / "Nova_klasicka_ZOD.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Šablóna nenájdená: {template_path}")

    cf = (project.get("custom_fields") or {}) if isinstance(project, dict) else {}

    # Flowii-style placeholders → CRM dáta
    mapping = {
        # Spoločnosť (Objednávateľ)
        "{companyName}": customer.get("company_name") or "",
        "{companyRegNumber}": customer.get("ico") or "",
        "{companyTaxNumber}": customer.get("dic") or "",
        "{companyStreet}": customer.get("billing_street") or customer.get("installation_street") or "",
        "{companyCity}": customer.get("billing_city") or customer.get("installation_city") or "",
        "{companyZipCode}": customer.get("billing_postal_code") or customer.get("installation_postal_code") or "",
        # Konajúca osoba
        "{businessCaseTitul_pred_0cfa0}": customer.get("title_before") or "",
        "{businessCaseMeno_03a27}": customer.get("first_name") or "",
        "{businessCasePriezvisko_8375b}": customer.get("last_name") or "",
        "{businessCaseEmail_2c918}": customer.get("email") or project.get("client_pm_email") or "",
        "{businessCaseTelefonne__fa5fd}": customer.get("phone") or project.get("client_pm_phone") or "",
        # Projekt — energetické dáta
        "{businessCaseEIC_0664b}": cf.get("eic_spotreba") or cf.get("eic") or "",
        "{businessCaseCislo_obch_d6db3}": cf.get("cislo_obchodneho_partnera") or "",
        "{businessCaseHodnota_hl_a04dc}": cf.get("hodnota_hl_istica") or "",
        "{Vykon_AC_c49e0}": str(cf.get("ac_kw") or project.get("scale_kwp") or ""),
        "{Vykon_DC_1f5f9}": str(project.get("scale_kwp") or cf.get("dc_kwp") or ""),
        # Cena + dátumy
        "{totalAmount}": _fmt_eur(project.get("contract_value_no_vat")),
        "{createdAtDate}": _fmt_date_sk(project.get("contract_signed_at") or date.today()),
        "{validFrom}": _fmt_date_sk(project.get("contract_signed_at") or date.today()),
        # Kód
        "{code}": project.get("project_code") or project.get("name") or "",
        # Splatnosť B2B = 14 dní (zladenie so spôsobom fakturácie; šablóna má default 7)
        "Lehota splatnosti faktúr je 7 dní.": "Lehota splatnosti faktúr je 14 dní.",
    }

    log.info("[generuj_zod] project=%s, mapping keys=%d", project.get("id"), len(mapping))

    doc = Document(str(template_path))
    _replace_in_doc(doc, mapping)

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]+", "_", (customer.get("company_name") or "Klient")).strip("_")[:50]
    out_path = out_dir / f"ZoD_{project.get('project_code') or 'B2B'}_{safe_name}.docx"
    doc.save(str(out_path))
    log.info("[generuj_zod] saved: %s", out_path)
    return str(out_path)


# =========================================================
# Energovision Faktúra B2B
# =========================================================
def generuj_faktura(
    *,
    project: dict,
    customer: dict,
    milestone: dict,
    faktura_cislo: str,
    variabilny_symbol: str,
    out_dir: str | Path,
) -> str:
    """
    Vyplní faktúru pre konkrétny milestone (FA1/FA2/FA3).
    Faktúra je iba "podklad" — Lukáš/admin si ju nahodí do Flowii.
    """
    template_path = TEMPLATES_DIR / "Energovision_Faktura_B2B.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Šablóna nenájdená: {template_path}")

    cf = (project.get("custom_fields") or {}) if isinstance(project, dict) else {}
    today = date.today()

    payment_amount = milestone.get("payment_amount")
    payment_pct = milestone.get("payment_pct")
    fa_no = milestone.get("fa_no", 1)
    dph_pct = 0.23
    cena_bez_dph = float(payment_amount or 0)
    dph_suma = round(cena_bez_dph * dph_pct, 2)
    cena_s_dph = round(cena_bez_dph + dph_suma, 2)

    predmet = (
        f"FA{fa_no} — {milestone.get('title') or 'Inštalácia FVE'}"
        f" — projekt {project.get('project_code') or project.get('name', '')}"
        f" (zmluva o dielo, {payment_pct or 0} % z celkovej ceny)"
    )

    mapping = {
        "{{faktura_cislo}}": faktura_cislo,
        # Banka (placeholder — Lukáš si vyplní)
        "{{banka_nazov}}": cf.get("banka_nazov") or "{vyplniť — bankové spojenie}",
        "{{iban}}": cf.get("iban") or "{vyplniť — IBAN}",
        "{{swift}}": cf.get("swift") or "{vyplniť — SWIFT}",
        # Odberateľ
        "{{odberatel_nazov}}": customer.get("company_name") or "",
        "{{odberatel_ulica}}": customer.get("billing_street") or customer.get("installation_street") or "",
        "{{odberatel_psc}}": customer.get("billing_postal_code") or customer.get("installation_postal_code") or "",
        "{{odberatel_mesto}}": customer.get("billing_city") or customer.get("installation_city") or "",
        "{{odberatel_ico}}": customer.get("ico") or "",
        "{{odberatel_dic}}": customer.get("dic") or "",
        "{{odberatel_ic_dph}}": customer.get("ic_dph") or "",
        "{{odberatel_kontakt_meno}}": f"{customer.get('first_name') or ''} {customer.get('last_name') or ''}".strip(),
        "{{odberatel_email}}": customer.get("email") or project.get("client_pm_email") or "",
        "{{odberatel_telefon}}": customer.get("phone") or project.get("client_pm_phone") or "",
        # Dátumy
        "{{variabilny_symbol}}": variabilny_symbol,
        "{{specificky_symbol}}": project.get("project_code") or "",
        "{{datum_vystavenia}}": _fmt_date_sk(today),
        "{{datum_dodania}}": _fmt_date_sk(milestone.get("completed_at") or today),
        "{{datum_splatnosti}}": _fmt_date_sk(milestone.get("due_date") or today),
        # Predmet
        "{{predmet_faktury}}": predmet,
        # Položky (zatial 1 riadok = celý milestone)
        "{{polozka_popis}}": predmet,
        "{{polozka_mnozstvo}}": "1",
        "{{polozka_mj}}": "ks",
        "{{polozka_cena_jednotka}}": _fmt_eur(cena_bez_dph),
        "{{polozka_cena_spolu}}": _fmt_eur(cena_bez_dph),
        # Sumy
        "{{cena_bez_dph}}": _fmt_eur(cena_bez_dph),
        "{{dph_suma}}": _fmt_eur(dph_suma),
        "{{cena_s_dph}}": _fmt_eur(cena_s_dph),
        # Poznámka
        "{{poznamka}}": (
            f"Faktúra k zmluve o dielo na realizáciu fotovoltického zariadenia. "
            f"V prípade omeškania s úhradou je zhotoviteľ oprávnený účtovať úrok z omeškania v zmysle zmluvy."
        ),
    }

    log.info("[generuj_faktura] project=%s, fa_no=%s, cena=%s", project.get("id"), fa_no, cena_bez_dph)

    doc = Document(str(template_path))
    _replace_in_doc(doc, mapping)

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]+", "_", (customer.get("company_name") or "Klient")).strip("_")[:50]
    out_path = out_dir / f"Faktura_{faktura_cislo.replace('/', '_')}_{safe_name}.docx"
    doc.save(str(out_path))
    log.info("[generuj_faktura] saved: %s", out_path)
    return str(out_path)
