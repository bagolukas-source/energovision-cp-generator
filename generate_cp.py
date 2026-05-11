"""
generate_cp.py — Energovision B2C CP generator
Vstup: lead.json (alebo dict)
Výstupy: CP_*.docx, CP_*.pdf, CP_*.eml, kalkulacia_*.xlsx, 3 grafy

Použitie:
  python3 generate_cp.py lead_sedlar.json
"""

import json, sys, os, base64, datetime, subprocess
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.utils import formatdate

# Brand
GREEN = "92D050"
GREEN_RGB = (0x92, 0xD0, 0x50)
BLACK = "000000"
GRAY = "666666"
LIGHT_GREEN = "E8F5D8"
FONT = "Arial"
_SD = os.path.dirname(os.path.abspath(__file__))
BRAND_HEADER = os.path.join(_SD, "energovision_header.png")
BRAND_FOOTER = os.path.join(_SD, "energovision_footer.png")
CENNIK = os.path.join(_SD, "Cennik_v2.xlsx")

# Energovision firma (z brand assetov)
ZHOTOVITEL = {
    "nazov": "Energovision, s.r.o.",
    "adresa": "Lamačská cesta 1738/111",
    "psc_mesto": "841 03 Bratislava",
    "ico": "53 036 280",
    "tel": "+421 918 950 776",
    "email": "obchod@energovision.sk",
    "web": "www.energovision.sk",
}

# Defaultné konštanty (môžu byť prepísané v lead-e)
DEFAULTS = {
    "marza_pct": 30,                # marža v %
    "rezerva_pct": 5,               # rezerva v %
    "dph_pct": 23,                  # DPH B2C SK
    "platnost_dni": 30,
    "cena_el_eur_kwh": 0.16,        # priemer SK 2026
    "vykupna_cena_eur_kwh": 0.05,   # prebytky do siete
    "samospotreba_pct": 70,         # typicky 60-80% pre RD s batériou
    "degradacia_pct_rok": 0.5,      # ročná degradácia panelov %
    "vyroba_kwh_per_kwp": 1075,     # SR priemer pre J orientáciu
    "narast_cien_el_pct_rok": 3.0,  # ročný nárast ceny elektriny
    "dotacia_flat_eur": 1500,       # Dotácia Zelená domácnostiam — flat MAX 1 500 € (od 2026-05-11)
    "obchodnik": {
        "meno": "Dominik Galaba",
        "funkcia": "Office & Administration Manager",
        "tel": "+421 917 424 564",
        "email": "dominik.galaba@energovision.sk",
    },
}


def load_cennik():
    """Načíta Cennik_v2.xlsx do dictu kód → {nazov, mj, cena}."""
    from openpyxl import load_workbook
    wb = load_workbook(CENNIK, data_only=True)
    ws = wb.active
    out = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        kod, kat, naz, mj, cena, zdr, dat, pozn = row
        if kod and cena is not None and isinstance(cena, (int, float)):
            out[kod] = {"nazov": naz, "kategoria": kat, "mj": mj, "cena": float(cena)}
    return out


def vyrataj_konfig(lead, cennik):
    """Z parametrov leadu zostaví BOM a vyráta ceny."""
    vykon = lead["vykon_kwp"]
    panel_kod = lead.get("panel_kod", "PAN-002")  # default LONGi 540 Wp
    panel = cennik[panel_kod]
    # počet panelov: výkon kWp / Wp panela (z názvu vyextrahuj 535-545 → 540)
    if "470" in panel["nazov"]: wp = 470
    elif "535" in panel["nazov"] or "540" in panel["nazov"]: wp = 540
    else: wp = 500
    pocet_panelov = round(vykon * 1000 / wp)

    # menič (default Solinteg 10K alebo Huawei podľa kWp)
    inv_kod = lead.get("invertor_kod")
    if not inv_kod:
        if vykon <= 5: inv_kod = "INV-002"     # Huawei 5K
        elif vykon <= 6: inv_kod = "INV-003"
        elif vykon <= 8: inv_kod = "INV-004"
        elif vykon <= 10: inv_kod = "INV-005"  # Huawei 10K (preferovaný)
        else: inv_kod = "INV-001"              # Solinteg 10K
    menic = cennik[inv_kod]

    # konštrukcia
    kon_kod = lead.get("konstrukcia_kod", "KON-001")  # default škridla
    konstrukcia = cennik[kon_kod]

    # batéria (ak požadovaná)
    bat_items = []
    if lead.get("bateria_kwh", 0) > 0:
        bat_kwh = lead["bateria_kwh"]
        bat_kod = lead.get("bateria_kod", "BAT-001")  # default Pylontech 5.12 modul
        bat = cennik[bat_kod]
        # počet modulov
        if "5,12" in bat["nazov"]: pocet_bat = max(2, round(bat_kwh / 5.12))
        elif "10,24" in bat["nazov"]: pocet_bat = max(1, round(bat_kwh / 10.24))
        elif "5 kWh" in bat["nazov"]: pocet_bat = max(1, round(bat_kwh / 5))
        elif "7 kWh" in bat["nazov"]: pocet_bat = max(1, round(bat_kwh / 7))
        else: pocet_bat = 2
        bat_items.append((bat_kod, pocet_bat, bat["nazov"], bat["mj"], bat["cena"]))
        # BMS riadiaca jednotka pre Pylontech a Huawei
        if "Pylontech" in bat["nazov"]:
            bms = cennik["BAT-002"]
            bat_items.append(("BAT-002", 1, bms["nazov"], bms["mj"], bms["cena"]))
        elif "LUNA" in bat["nazov"]:
            bms = cennik["BAT-007"]
            bat_items.append(("BAT-007", 1, bms["nazov"], bms["mj"], bms["cena"]))

    # wallbox (ak požadovaný)
    wb_items = []
    if lead.get("wallbox", False):
        wb_kod = lead.get("wallbox_kod", "WBX-002")  # default Solinteg 11K
        w = cennik[wb_kod]
        wb_items.append((wb_kod, 1, w["nazov"], w["mj"], w["cena"]))
        # kabeláž do 10 m
        wb_items.append(("KAB-005", 10, "Kábel pre EV nabíjačku (10 m)", "m", cennik["KAB-005"]["cena"]))
        # spotrebný materiál
        wb_items.append(("PRC-004", 1, "Spotrebný materiál a práca pre wallbox", "ks", cennik["PRC-004"]["cena"]))

    # MC4 konektory: pri 24 paneloch potreba 24 párov (samec+samica) + 4-6 ks rezerva
    # Cena MC4 je za pár (~2 € za samec+samica priemer)
    mc4_pocet = max(8, pocet_panelov + 4)  # min 8, alebo počet panelov + rezerva

    # BOM (interný — nákupné ceny)
    bom = [
        # (kod, počet, názov, mj, nákupná cena)
        (panel_kod, pocet_panelov, panel["nazov"], panel["mj"], panel["cena"]),
        (inv_kod, 1, menic["nazov"], menic["mj"], menic["cena"]),
        (kon_kod, pocet_panelov, konstrukcia["nazov"], "modul", konstrukcia["cena"]),
        ("KAB-001", 50, "FVE solárny kábel červený 6 mm² (50 m)", "m", cennik["KAB-001"]["cena"]),
        ("KAB-002", 50, "FVE solárny kábel čierny 6 mm² (50 m)", "m", cennik["KAB-002"]["cena"]),
        ("KAB-003", 30, "Kábel CYKY-J 5x4 mm² (30 m)", "m", cennik["KAB-003"]["cena"]),
        ("KAB-004", mc4_pocet, f"Konektor MC4 ({mc4_pocet} ks)", "ks", cennik["KAB-004"]["cena"]),
        ("RVZ-003", 1, cennik["RVZ-003"]["nazov"], "ks", cennik["RVZ-003"]["cena"]),
    ]
    # AC rozvádzač podľa distribučky
    dist = lead.get("distribucka", "ZSD")
    rvz_kod = {"ZSD": "RVZ-004", "SSD": "RVZ-005", "VSD": "RVZ-006"}[dist]
    bom.append((rvz_kod, 1, cennik[rvz_kod]["nazov"], "ks", cennik[rvz_kod]["cena"]))
    # smart meter
    bom.append(("SMT-001", 1, cennik["SMT-001"]["nazov"], "ks", cennik["SMT-001"]["cena"]))
    bom.append(("SMT-002", 1, cennik["SMT-002"]["nazov"], "ks", cennik["SMT-002"]["cena"]))
    # spotrebný materiál (drobný — svorky, lišty, izolačka, šrubky, atď.)
    bom.append(("PRC-007", 1, "Spotrebný materiál (drobný)", "ks", cennik["PRC-007"]["cena"]))
    # batéria, wallbox
    bom.extend(bat_items)
    bom.extend(wb_items)

    # PRÁCA
    praca = [
        ("PRC-001", vykon, "Inštalácia fotovoltiky", "kWp", cennik["PRC-001"]["cena"]),
        ("PRC-005", 1, "Inžiniering: revízia, žiadosti DIS+SIEA+dodávateľ el.", "projekt", cennik["PRC-005"]["cena"]),
    ]
    if lead.get("bateria_kwh", 0) > 0:
        praca.append(("PRC-002", 1, "Doplnenie batérie (paušál)", "kpl", cennik["PRC-002"]["cena"]))
    # doprava — default 100 km (typická SK trasa)
    km = lead.get("doprava_km", 100)
    praca.append(("PRC-006", km, f"Doprava materiálu ({km} km)", "km", cennik["PRC-006"]["cena"]))

    return {
        "vykon_kwp": vykon,
        "pocet_panelov": pocet_panelov,
        "panel": panel["nazov"],
        "menic": menic["nazov"],
        "konstrukcia": konstrukcia["nazov"],
        "bom": bom,
        "praca": praca,
        "ma_bateriu": lead.get("bateria_kwh", 0) > 0,
        "bateria_kwh": lead.get("bateria_kwh", 0),
        "ma_wallbox": lead.get("wallbox", False),
    }


def vyrataj_ceny(konfig, lead):
    marza = lead.get("marza_pct", DEFAULTS["marza_pct"]) / 100
    rezerva = lead.get("rezerva_pct", DEFAULTS["rezerva_pct"]) / 100
    dph = DEFAULTS["dph_pct"] / 100

    nakupna_material = sum(p[1] * p[4] for p in konfig["bom"])
    nakupna_praca = sum(p[1] * p[4] for p in konfig["praca"])
    nakupna_spolu = nakupna_material + nakupna_praca
    rezerva_eur = nakupna_spolu * rezerva
    marza_eur = (nakupna_spolu + rezerva_eur) * marza
    cena_bez_dph = nakupna_spolu + rezerva_eur + marza_eur
    cena_s_dph = cena_bez_dph * (1 + dph)

    # Dotácia Zelená domácnostiam — flat 1 500 € (3 kW × 500 €/kW)
    # Okresové zvýhodnenia 575/900 €/kW už NEPLATIA (update 2026-05-11)
    if lead.get("dotacia", True):
        dotacia = DEFAULTS["dotacia_flat_eur"]
    else:
        dotacia = 0

    cena_po_dotacii = cena_s_dph - dotacia
    zlava_eur = lead.get("zlava_eur", 0)
    cena_finalna = cena_po_dotacii - zlava_eur

    return {
        "nakupna_material": nakupna_material,
        "nakupna_praca": nakupna_praca,
        "nakupna_spolu": nakupna_spolu,
        "rezerva_eur": rezerva_eur,
        "marza_eur": marza_eur,
        "cena_bez_dph": cena_bez_dph,
        "cena_s_dph": cena_s_dph,
        "dotacia": dotacia,
        "cena_po_dotacii": cena_po_dotacii,
        "zlava_eur": zlava_eur,
        "cena_finalna": cena_finalna,
        "marza_pct": lead.get("marza_pct", DEFAULTS["marza_pct"]),
        "zisk": cena_bez_dph - nakupna_spolu - rezerva_eur,
    }


def vyrataj_navratnost(konfig, ceny, lead):
    vykon = konfig["vykon_kwp"]
    rocna_vyroba = vykon * DEFAULTS["vyroba_kwh_per_kwp"]
    cena_el = lead.get("cena_el_eur_kwh", DEFAULTS["cena_el_eur_kwh"])
    samosp = (lead.get("samospotreba_pct") or
              (DEFAULTS["samospotreba_pct"] + (10 if konfig["ma_bateriu"] else 0))) / 100
    if samosp > 1: samosp = 0.85
    vykupna = DEFAULTS["vykupna_cena_eur_kwh"]
    naras_pct = DEFAULTS["narast_cien_el_pct_rok"] / 100
    deg_pct = DEFAULTS["degradacia_pct_rok"] / 100

    rocne_uspora = rocna_vyroba * (samosp * cena_el + (1 - samosp) * vykupna)
    rocne_naklady_bez_fve = lead["rocna_spotreba_kwh"] * cena_el

    # 25-rocna kumulativna
    kumul = []
    suma = 0
    for r in range(1, 26):
        deg = (1 - deg_pct) ** (r - 1)
        cena = cena_el * (1 + naras_pct) ** (r - 1)
        usp = rocna_vyroba * deg * (samosp * cena + (1 - samosp) * vykupna)
        suma += usp
        kumul.append((r, usp, suma))

    navratnost_rokov = ceny["cena_po_dotacii"] / rocne_uspora if rocne_uspora > 0 else 0
    return {
        "rocna_vyroba_kwh": rocna_vyroba,
        "rocne_uspora_eur": rocne_uspora,
        "rocne_naklady_bez_fve_eur": rocne_naklady_bez_fve,
        "navratnost_rokov": navratnost_rokov,
        "uspora_25_rokov": suma,
        "kumul_25": kumul,
        "samospotreba_pct": samosp * 100,
        "kg_co2_rok": rocna_vyroba * 0.4,  # SK mix ~0.4 kg CO2/kWh
    }


def vyrob_grafy(navratnost, lead, out_dir, prefix):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["axes.spines.top"] = False
    plt.rcParams["axes.spines.right"] = False
    paths = {}

    # Graf 1: 25-ročná kumulatívna úspora
    fig, ax = plt.subplots(figsize=(8, 5), dpi=120)
    roky = [k[0] for k in navratnost["kumul_25"]]
    suma = [k[2] for k in navratnost["kumul_25"]]
    ax.plot(roky, suma, color="#92D050", linewidth=2.5)
    ax.fill_between(roky, suma, alpha=0.2, color="#92D050")
    nv = navratnost["navratnost_rokov"]
    if 0 < nv < 25:
        ax.axvline(nv, color="#666666", linestyle="--", linewidth=1)
        ax.text(nv + 0.3, max(suma) * 0.05, f"Návratnosť\n{nv:.1f} rokov",
                fontsize=9, color="#666666")
    ax.set_xlabel("Rok prevádzky")
    ax.set_ylabel("Kumulatívna úspora (€)")
    ax.set_title("25-ročná kumulatívna úspora na elektrine", fontsize=11, weight="bold")
    ax.grid(True, alpha=0.3)
    ax.set_xlim(1, 25)
    fig.tight_layout()
    p1 = f"{out_dir}/{prefix}_graf_uspora.png"
    fig.savefig(p1, dpi=120, bbox_inches="tight")
    plt.close(fig)
    paths["uspora"] = p1

    # Graf 2: Mesačná výroba kWh (typický profil SR)
    profil = [0.025, 0.045, 0.085, 0.115, 0.135, 0.140, 0.140, 0.125, 0.090, 0.060, 0.025, 0.015]
    mesiace = ["Jan", "Feb", "Mar", "Apr", "Máj", "Jún", "Júl", "Aug", "Sep", "Okt", "Nov", "Dec"]
    rocna = navratnost["rocna_vyroba_kwh"]
    vyroba = [rocna * p for p in profil]
    fig, ax = plt.subplots(figsize=(8, 5), dpi=120)
    bars = ax.bar(mesiace, vyroba, color="#92D050", edgecolor="white")
    for b, v in zip(bars, vyroba):
        ax.text(b.get_x() + b.get_width() / 2, b.get_height() + max(vyroba) * 0.02,
                f"{v:.0f}", ha="center", fontsize=8, color="#444")
    ax.set_ylabel("Výroba (kWh)")
    ax.set_title(f"Predpokladaná mesačná výroba — celkom {rocna:.0f} kWh/rok",
                 fontsize=11, weight="bold")
    ax.set_ylim(0, max(vyroba) * 1.15)
    fig.tight_layout()
    p2 = f"{out_dir}/{prefix}_graf_vyroba.png"
    fig.savefig(p2, dpi=120, bbox_inches="tight")
    plt.close(fig)
    paths["vyroba"] = p2

    # Graf 3: Porovnanie nákladov bez/s FVE (10 rokov)
    cena_el = lead.get("cena_el_eur_kwh", DEFAULTS["cena_el_eur_kwh"])
    naras_pct = DEFAULTS["narast_cien_el_pct_rok"] / 100
    samosp = navratnost["samospotreba_pct"] / 100
    deg_pct = DEFAULTS["degradacia_pct_rok"] / 100
    spotreba = lead["rocna_spotreba_kwh"]
    rocna = navratnost["rocna_vyroba_kwh"]

    bez_fve_per_year = []
    s_fve_per_year = []
    # Fixné poplatky distribučke (jistič, distribúcia, OZE poplatok) ostávajú vždy ~25 % pôvodných nákladov
    fixne_pct = 0.25
    for r in range(1, 11):
        cena = cena_el * (1 + naras_pct) ** (r - 1)
        deg = (1 - deg_pct) ** (r - 1)
        bez = spotreba * cena
        usp = rocna * deg * (samosp * cena + (1 - samosp) * DEFAULTS["vykupna_cena_eur_kwh"])
        # S FVE = ušetríme spotrebu, ale fixné poplatky platíme stále
        variable_part = max(0, bez * (1 - fixne_pct) - usp)
        s_fve = bez * fixne_pct + variable_part
        bez_fve_per_year.append(bez)
        s_fve_per_year.append(s_fve)

    fig, ax = plt.subplots(figsize=(7, 3.5), dpi=120)
    x = list(range(1, 11))
    width = 0.4
    ax.bar([i - width/2 for i in x], bez_fve_per_year, width=width, label="Bez FVE", color="#888")
    ax.bar([i + width/2 for i in x], s_fve_per_year, width=width, label="S FVE", color="#92D050")
    ax.set_xlabel("Rok")
    ax.set_ylabel("Ročné náklady na elektrinu (€)")
    ax.set_title("Porovnanie ročných nákladov bez FVE vs. s FVE (10 rokov)",
                 fontsize=11, weight="bold")
    ax.legend(loc="upper left")
    ax.set_xticks(x)
    fig.tight_layout()
    p3 = f"{out_dir}/{prefix}_graf_porovnanie.png"
    fig.savefig(p3, dpi=120, bbox_inches="tight")
    plt.close(fig)
    paths["porovnanie"] = p3

    return paths


def vyrob_docx(lead, konfig, ceny, navratnost, grafy, out_path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    # margins
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

    def set_run(run, text, size=10, bold=False, color=None, italic=False):
        run.text = text
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color: run.font.color.rgb = RGBColor(*color)

    def add_p(text="", size=10, bold=False, color=None, align=None, space_after=4, italic=False):
        p = doc.add_paragraph()
        if align: p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run()
            set_run(r, text, size, bold, color, italic)
        return p

    def add_h(text, size=14, color=GREEN_RGB):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
        r = p.add_run()
        set_run(r, text, size, bold=True, color=color)
        return p

    def cell_set(cell, text, size=10, bold=False, color=None, fill=None, align="left"):
        cell.text = ""
        p = cell.paragraphs[0]
        if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run()
        set_run(r, text, size, bold, color)
        if fill:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
            tc_pr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ==== STRANA 1 — TITULKA ====
    if Path(BRAND_HEADER).exists():
        doc.add_picture(BRAND_HEADER, width=Cm(17))
    add_p("", space_after=20)
    add_p("CENOVÁ PONUKA", size=24, bold=True, color=GREEN_RGB,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p("Fotovoltická elektráreň pre Váš dom", size=14, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    t = doc.add_table(rows=4, cols=2)
    t.columns[0].width = Cm(5); t.columns[1].width = Cm(11)
    cell_set(t.cell(0,0), "Zákazník:", bold=True)
    cell_set(t.cell(0,1), f"{lead['meno']}, {lead['mesto']}")
    cell_set(t.cell(1,0), "Číslo ponuky:", bold=True)
    cell_set(t.cell(1,1), lead.get("cislo_ponuky", f"PON-{datetime.date.today():%Y-%m%d}"))
    cell_set(t.cell(2,0), "Dátum vystavenia:", bold=True)
    cell_set(t.cell(2,1), datetime.date.today().strftime("%d. %m. %Y"))
    cell_set(t.cell(3,0), "Platnosť ponuky:", bold=True)
    platnost = datetime.date.today() + datetime.timedelta(days=lead.get("platnost_dni", DEFAULTS["platnost_dni"]))
    cell_set(t.cell(3,1), platnost.strftime("%d. %m. %Y"))
    add_p("", space_after=30)

    add_p("Vypracoval:", bold=True, size=11)
    obch = lead.get("obchodnik", DEFAULTS["obchodnik"])
    add_p(f"{obch['meno']} — {obch['funkcia']}, Energovision", size=11)
    add_p(f"{obch['tel']}  |  {obch['email']}", size=10, color=(0x66,0x66,0x66))

    doc.add_page_break()

    # ==== STRANA 2 — ZHRNUTIE RIEŠENIA ====
    add_h("1. Zhrnutie navrhovaného riešenia")

    add_p("Čo navrhujeme:", bold=True, size=12, space_after=6)
    sumar = doc.add_table(rows=0, cols=2)
    sumar.columns[0].width = Cm(7); sumar.columns[1].width = Cm(9)
    def srow(label, value, bold=False):
        row = sumar.add_row()
        cell_set(row.cells[0], label, size=10, fill="F5F5F5")
        cell_set(row.cells[1], value, size=10, bold=bold)
    srow("Výkon FVE", f"{konfig['vykon_kwp']:.2f} kWp".replace(".", ","), bold=True)
    srow("Počet panelov", f"{konfig['pocet_panelov']} ks")
    srow("Typ panelov", konfig["panel"])
    srow("Menič", konfig["menic"])
    srow("Konštrukcia", konfig["konstrukcia"])
    if konfig["ma_bateriu"]:
        srow("Batériové úložisko", f"{konfig['bateria_kwh']:.2f} kWh")
    if konfig["ma_wallbox"]:
        srow("Wallbox pre EV", "áno (v cene)")
    srow("Predpokladaná ročná výroba", f"{navratnost['rocna_vyroba_kwh']:,.0f} kWh".replace(",", " "), bold=True)
    add_p("", space_after=12)

    add_p("Čo Vám to prinesie:", bold=True, size=12, space_after=6)
    prinos = doc.add_table(rows=0, cols=2)
    prinos.columns[0].width = Cm(7); prinos.columns[1].width = Cm(9)
    def prow(label, value):
        row = prinos.add_row()
        cell_set(row.cells[0], label, size=10, fill="F5F5F5")
        cell_set(row.cells[1], value, size=11, bold=True, color=GREEN_RGB)
    prow("Predpokladaná ročná úspora", f"{navratnost['rocne_uspora_eur']:,.0f} €/rok".replace(",", " "))
    prow("Pokrytie spotreby", f"{min(100, navratnost['rocna_vyroba_kwh']/lead['rocna_spotreba_kwh']*100):.0f} %")
    prow("Návratnosť investície", f"{navratnost['navratnost_rokov']:.1f} rokov")
    prow("25-ročná úspora", f"{navratnost['uspora_25_rokov']:,.0f} €".replace(",", " "))
    prow("Príspevok k ochrane klímy", f"{navratnost['kg_co2_rok']:,.0f} kg CO₂ / rok".replace(",", " "))

    doc.add_page_break()

    # ==== STRANA 3 — CENA + GRAFY ====
    add_h("2. Cena")

    cena_t = doc.add_table(rows=0, cols=2)
    cena_t.columns[0].width = Cm(10); cena_t.columns[1].width = Cm(6)
    def cprow(label, value, bold=False, hi=False, size=11):
        row = cena_t.add_row()
        cell_set(row.cells[0], label, size=size, bold=bold)
        cell_set(row.cells[1], value, size=size, bold=bold, align="right",
                 fill=(GREEN if hi else None), color=((0,0,0) if hi else None))
    cprow("Cena fotovoltaického zariadenia s DPH 23 %", f"{ceny['cena_s_dph']:,.2f} €".replace(",", "X").replace(".", ",").replace("X", " "), bold=True)
    if ceny["dotacia"] > 0:
        cprow("Dotácia Zelená domácnostiam", f"− {ceny['dotacia']:,.0f} €".replace(",", " "))
    if ceny["zlava_eur"] > 0:
        cprow("Zľava", f"− {ceny['zlava_eur']:,.2f} €".replace(",", "X").replace(".", ",").replace("X", " "))
    cprow("Cena po dotácii", f"{ceny['cena_finalna']:,.2f} €".replace(",", "X").replace(".", ",").replace("X", " "), bold=True, hi=True, size=14)
    add_p("", space_after=10)
    add_p(f"Sadzba DPH 23 % podľa platnej legislatívy SR.", size=9, color=(0x88,0x88,0x88), italic=True)

    add_h("3. Návratnosť a úspora", size=14)

    if Path(grafy["uspora"]).exists():
        doc.add_picture(grafy["uspora"], width=Cm(16.5))
    add_p("", space_after=8)
    if Path(grafy["porovnanie"]).exists():
        doc.add_picture(grafy["porovnanie"], width=Cm(16.5))

    doc.add_page_break()

    if Path(grafy["vyroba"]).exists():
        add_h("4. Predpokladaná výroba elektriny", size=14)
        doc.add_picture(grafy["vyroba"], width=Cm(16.5))
        add_p("", space_after=8)
        add_p("Profil výroby je orientačný, pre J orientáciu na strednom Slovensku. "
              "Skutočná výroba závisí od orientácie strechy, sklonu, tienenia a počasia v danom roku.",
              size=9, color=(0x88,0x88,0x88), italic=True)

    add_h("5. Čo je v cene", size=14)
    items = [
        "Komponenty (panely, menič, konštrukcia, kabeláž, ochrany, smartmeter)",
        "Montáž a uvedenie do prevádzky vlastným tímom",
        "Revízna správa elektroinštalácie",
        "Vybavenie pripojenia do distribučnej siete",
    ]
    if lead.get("dotacia", True):
        items.append("Vybavenie žiadosti o dotáciu Zelená domácnostiam")
    items.append("Záruka 12 rokov na panely (produkt) a 25 rokov na lineárny pokles výkonu")
    items.append("Záruka 10 rokov na fotovoltický menič a 12 rokov na montáž")
    if konfig["ma_bateriu"]:
        items.append("Záruka 10 rokov na batériové úložisko")
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run()
        set_run(r, f"✓  {it}", size=10)

    doc.add_page_break()

    # ==== STRANA 4 — HARMONOGRAM, PODMIENKY, KONTAKT ====
    add_h("6. Harmonogram realizácie", size=14)

    harm = [
        ("D", "Podpis zmluvy"),
        ("D + 1–2 týždne", "Obhliadka a projekčné práce"),
        ("D + 4–8 týždňov", "Vybavenie pripojenia do distribučnej siete"),
        ("D + 8–10 týždňov", "Montáž (1–3 dni)"),
        ("D + 10–11 týždňov", "Revízia a uvedenie do prevádzky"),
        ("D + 11+ týždňov", "Spustenie produkcie a sledovanie cez aplikáciu meniča"),
    ]
    ht = doc.add_table(rows=0, cols=2)
    ht.columns[0].width = Cm(5.5); ht.columns[1].width = Cm(11)
    for k, v in harm:
        row = ht.add_row()
        cell_set(row.cells[0], k, size=10, bold=True, color=GREEN_RGB)
        cell_set(row.cells[1], v, size=10)

    add_h("7. Platobné podmienky", size=14)
    pp = lead.get("platby", "60 % zálohová faktúra vopred  /  30 % po nainštalovaní elektrárne  /  10 % po protokolárnom odovzdaní")
    add_p(pp, size=10)

    add_h("8. Doplnkové služby", size=14)
    dop = []
    if not konfig["ma_bateriu"]:
        dop.append("Doplnenie batériového úložiska (zvýšenie samospotreby z ~70 % na ~90 %)")
    if not konfig["ma_wallbox"]:
        dop.append("Wallbox pre elektromobil (priama integrácia s FVE — nabíjanie zo slnka)")
    dop += [
        "Pravidelná revízia elektroinštalácie a FVE (každé 4 roky)",
        "Údržbová zmluva na FVE — kontrola výkonu, čistenie panelov pri potrebe",
        "Bleskozvod pre rodinný dom (ak chýba)",
    ]
    for d in dop:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(); set_run(r, f"•  {d}", size=10)
    add_p("Radi pripravíme samostatnú ponuku — ozvite sa.", size=10, italic=True,
          color=(0x66,0x66,0x66))

    add_h("9. Kontakt", size=14)
    add_p(f"{obch['meno']}", size=11, bold=True)
    add_p(f"{obch['funkcia']}, {ZHOTOVITEL['nazov']}", size=10)
    add_p(f"Tel: {obch['tel']}", size=10)
    add_p(f"E-mail: {obch['email']}", size=10)
    add_p(f"Web: {ZHOTOVITEL['web']}", size=10)
    add_p("", space_after=20)

    add_p(f"Vystavené v Bratislave, {datetime.date.today():%d. %m. %Y}.", size=9, italic=True, color=(0x88,0x88,0x88))
    add_p(f"{ZHOTOVITEL['nazov']}, {ZHOTOVITEL['adresa']}, {ZHOTOVITEL['psc_mesto']}, IČO: {ZHOTOVITEL['ico']}",
          size=9, italic=True, color=(0x88,0x88,0x88))

    doc.save(out_path)


def vyrob_internu_kalkulaciu(lead, konfig, ceny, navratnost, out_path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    THIN = Side(border_style="thin", color="CCCCCC")
    B = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    HDR = PatternFill("solid", start_color=GREEN)
    SUB = PatternFill("solid", start_color=LIGHT_GREEN)

    # Sheet 1: Súhrn
    ws = wb.active; ws.title = "Súhrn"
    ws.column_dimensions["A"].width = 40
    ws.column_dimensions["B"].width = 18

    def s(r, l, v, b=False, fill=None, fmt=None):
        ws.cell(r, 1, l).font = Font(name=FONT, size=11, bold=b)
        c = ws.cell(r, 2, v); c.font = Font(name=FONT, size=11, bold=b)
        if fmt: c.number_format = fmt
        if fill:
            ws.cell(r, 1).fill = fill; c.fill = fill
        ws.cell(r, 1).border = B; ws.cell(r, 2).border = B

    ws.cell(1, 1, "INTERNÁ KALKULÁCIA").font = Font(name=FONT, size=14, bold=True)
    ws.cell(2, 1, f"Zákazník: {lead['meno']}, {lead['mesto']}").font = Font(name=FONT, size=10)
    ws.cell(3, 1, f"Vytvorené: {datetime.date.today()}").font = Font(name=FONT, size=10)

    r = 5
    s(r, "Konfigurácia", "", b=True, fill=SUB); r += 1
    s(r, "Výkon FVE (kWp)", konfig["vykon_kwp"], fmt="0.00"); r += 1
    s(r, "Počet panelov", konfig["pocet_panelov"]); r += 1
    s(r, "Batéria (kWh)", konfig["bateria_kwh"]); r += 1

    r += 1
    s(r, "Náklady", "", b=True, fill=SUB); r += 1
    s(r, "Materiál (nákupná cena)", ceny["nakupna_material"], fmt='#,##0.00 €'); r += 1
    s(r, "Práca (nákupná cena)", ceny["nakupna_praca"], fmt='#,##0.00 €'); r += 1
    s(r, "Spolu nákupná cena", ceny["nakupna_spolu"], b=True, fmt='#,##0.00 €'); r += 1

    r += 1
    s(r, "Marža a rezerva", "", b=True, fill=SUB); r += 1
    s(r, f"Rezerva ({DEFAULTS['rezerva_pct']} %)", ceny["rezerva_eur"], fmt='#,##0.00 €'); r += 1
    s(r, f"Marža ({ceny['marza_pct']} %)", ceny["marza_eur"], fmt='#,##0.00 €'); r += 1
    s(r, "Zisk (marža)", ceny["zisk"], b=True, fmt='#,##0.00 €'); r += 1

    r += 1
    s(r, "Cena pre zákazníka", "", b=True, fill=SUB); r += 1
    s(r, "Cena bez DPH", ceny["cena_bez_dph"], fmt='#,##0.00 €'); r += 1
    s(r, "DPH 23 %", ceny["cena_s_dph"] - ceny["cena_bez_dph"], fmt='#,##0.00 €'); r += 1
    s(r, "Cena s DPH", ceny["cena_s_dph"], b=True, fmt='#,##0.00 €'); r += 1
    s(r, "Dotácia Zelená domácnostiam", -ceny["dotacia"], fmt='#,##0.00 €'); r += 1
    s(r, "Cena po dotácii", ceny["cena_po_dotacii"], b=True, fmt='#,##0.00 €'); r += 1
    if ceny["zlava_eur"] > 0:
        s(r, "Zľava", -ceny["zlava_eur"], fmt='#,##0.00 €'); r += 1
        s(r, "FINÁLNA cena pre zákazníka", ceny["cena_finalna"], b=True, fmt='#,##0.00 €')

    # Sheet 2: BOM
    ws2 = wb.create_sheet("BOM (materiál)")
    headers = ["Kód", "Názov", "M.J.", "Počet", "Nákupná cena/MJ", "Spolu nákupne"]
    for i, h in enumerate(headers, 1):
        c = ws2.cell(1, i, h); c.font = Font(name=FONT, size=11, bold=True); c.fill = HDR; c.border = B
    widths = [12, 50, 8, 10, 18, 18]
    for i, w in enumerate(widths, 1):
        ws2.column_dimensions[chr(64+i)].width = w
    for ri, item in enumerate(konfig["bom"], 2):
        kod, pocet, naz, mj, cena = item
        spolu = pocet * cena
        ws2.cell(ri, 1, kod).font = Font(name=FONT, size=10)
        ws2.cell(ri, 2, naz).font = Font(name=FONT, size=10)
        ws2.cell(ri, 3, mj).font = Font(name=FONT, size=10)
        ws2.cell(ri, 4, pocet).font = Font(name=FONT, size=10)
        c = ws2.cell(ri, 5, cena); c.number_format = '#,##0.00 €'
        c = ws2.cell(ri, 6, spolu); c.number_format = '#,##0.00 €'

    # Sheet 3: Práca
    ws3 = wb.create_sheet("Práca")
    for i, h in enumerate(headers, 1):
        c = ws3.cell(1, i, h); c.font = Font(name=FONT, size=11, bold=True); c.fill = HDR; c.border = B
    for i, w in enumerate(widths, 1):
        ws3.column_dimensions[chr(64+i)].width = w
    for ri, item in enumerate(konfig["praca"], 2):
        kod, pocet, naz, mj, cena = item
        spolu = pocet * cena
        ws3.cell(ri, 1, kod); ws3.cell(ri, 2, naz); ws3.cell(ri, 3, mj); ws3.cell(ri, 4, pocet)
        c = ws3.cell(ri, 5, cena); c.number_format = '#,##0.00 €'
        c = ws3.cell(ri, 6, spolu); c.number_format = '#,##0.00 €'

    wb.save(out_path)


def vyrob_eml(lead, konfig, ceny, navratnost, pdf_path, out_path):
    obch = lead.get("obchodnik", DEFAULTS["obchodnik"])
    body = f"""Dobrý deň, pán {lead['meno'].split()[-1]},

ďakujem za Váš záujem o fotovoltickú elektráreň pre Váš dom v {lead['mesto']}.
V prílohe Vám posielam cenovú ponuku spracovanú na základe údajov, ktoré ste mi poskytli.

Krátko zhrnuté:
• Výkon FVE: {konfig['vykon_kwp']:.2f} kWp ({konfig['pocet_panelov']} ks panelov)
• Predpokladaná ročná výroba: {navratnost['rocna_vyroba_kwh']:,.0f} kWh
• Predpokladaná ročná úspora: {navratnost['rocne_uspora_eur']:,.0f} €
• Cena s DPH: {ceny['cena_s_dph']:,.2f} €
• Cena po dotácii Zelená domácnostiam: {ceny['cena_finalna']:,.2f} €
• Návratnosť: cca {navratnost['navratnost_rokov']:.1f} rokov

Ako ďalší krok navrhujem bezplatnú obhliadku, kde upresníme technické detaily
a finalizujeme cenu. Stačí zavolať alebo odpovedať na tento email.

Ponuka je platná {lead.get('platnost_dni', 30)} dní.

Ak máte akékoľvek otázky, som Vám k dispozícii.

S pozdravom,
{obch['meno']}
{obch['funkcia']}, {ZHOTOVITEL['nazov']}
{obch['tel']}  |  {obch['email']}
""".replace(",", " ")

    msg = MIMEMultipart()
    msg['From'] = ''
    msg['To'] = lead.get("email", "")
    msg['Subject'] = f"Cenová ponuka FVE pre Váš dom — {ZHOTOVITEL['nazov'].split(',')[0]}"
    msg['Date'] = formatdate(localtime=True)
    msg['X-Unsent'] = '1'
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    if pdf_path and Path(pdf_path).exists():
        with open(pdf_path, 'rb') as f:
            part = MIMEBase('application', 'pdf')
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename="{Path(pdf_path).name}"')
            msg.attach(part)
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(msg.as_string())


def docx_to_pdf(docx_path, out_dir):
    """Konvertuje .docx → .pdf cez LibreOffice (musí byť nainštalovaný)."""
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            check=True, capture_output=True, timeout=60
        )
        return Path(out_dir) / (Path(docx_path).stem + ".pdf")
    except Exception as e:
        print(f"  ⚠️  PDF konverzia zlyhala: {e}")
        return None


def main(lead_path):
    with open(lead_path, encoding="utf-8") as f:
        lead = json.load(f)

    print(f"📋 Načítaný lead: {lead['meno']} ({lead['mesto']})")
    cennik = load_cennik()
    print(f"📚 Cenník: {len(cennik)} položiek")

    konfig = vyrataj_konfig(lead, cennik)
    ceny = vyrataj_ceny(konfig, lead)
    navratnost = vyrataj_navratnost(konfig, ceny, lead)

    priezv = lead["meno"].split()[-1].replace(" ", "_")
    mesto = lead["mesto"].split(",")[0].replace(" ", "_")
    base = f"CP_{priezv}_{mesto}_v1"
    interna = f"kalkulacia_{priezv}_{mesto}_v1"
    out_dir = lead.get("out_dir", "/sessions/magical-eager-gates/mnt/outputs")

    print(f"📊 Generujem grafy...")
    grafy = vyrob_grafy(navratnost, lead, out_dir, base)

    print(f"📄 Generujem CP {base}.docx ...")
    docx_path = f"{out_dir}/{base}.docx"
    vyrob_docx(lead, konfig, ceny, navratnost, grafy, docx_path)

    print(f"💼 Generujem internú kalkuláciu {interna}.xlsx ...")
    vyrob_internu_kalkulaciu(lead, konfig, ceny, navratnost, f"{out_dir}/{interna}.xlsx")

    print(f"📑 Konvertujem .docx → .pdf ...")
    pdf_path = docx_to_pdf(docx_path, out_dir)

    print(f"✉️  Generujem .eml draft ...")
    vyrob_eml(lead, konfig, ceny, navratnost,
              str(pdf_path) if pdf_path else None,
              f"{out_dir}/{base}.eml")

    print(f"\n✅ Hotovo:")
    print(f"   {out_dir}/{base}.docx")
    if pdf_path: print(f"   {pdf_path}")
    print(f"   {out_dir}/{base}.eml")
    print(f"   {out_dir}/{interna}.xlsx")
    print(f"\n💰 Súhrn cien:")
    print(f"   Cena s DPH:        {ceny['cena_s_dph']:>12,.2f} €".replace(",", "X").replace(".", ",").replace("X", " "))
    print(f"   Dotácia:           {-ceny['dotacia']:>12,.0f} €".replace(",", " "))
    print(f"   Cena po dotácii:   {ceny['cena_po_dotacii']:>12,.2f} €".replace(",", "X").replace(".", ",").replace("X", " "))
    print(f"   Zisk (interný):    {ceny['zisk']:>12,.2f} €".replace(",", "X").replace(".", ",").replace("X", " "))
    print(f"   Návratnosť:        {navratnost['navratnost_rokov']:>12.1f} rokov")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "lead_sedlar.json")
