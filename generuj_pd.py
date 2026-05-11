"""
Generuj projektovú dokumentáciu (PD) pre malé zdroje FVE do 10 kW.

Generuje programaticky 5 dokumentov:
1. Krycí list
2. Zoznam dokumentácie
3. Technická správa (zjednodušená pre B2C rodinný dom)
4. Protokol o určení vonkajších vplyvov (PoUVV)
5. Súhrnná technická správa

Žiadne templaty — všetko programaticky cez python-docx.
Technické parametre panelu a striedača: zabudovaný cennik (synced z Make Data Store).

Komisia:
- Vypracoval: Lukáš Bago
- Kontroloval: Matej Horváth
- Zodpovedný projektant: Ing. Pavol Kaprál
"""
import os
import re
import logging
from io import BytesIO
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

log = logging.getLogger("generuj_pd")


# ============================================================
# KOMISIA
# ============================================================
KOMISIA = {
    "vypracoval": "Lukáš Bago",
    "vypracoval_funkcia": "Projektant",
    "kontroloval": "Matej Horváth",
    "kontroloval_funkcia": "Kontrolor",
    "zodpovedny_projektant": "Ing. Pavol Kaprál",
    "zodpovedny_projektant_funkcia": "Zodpovedný projektant",
}


# ============================================================
# TECHNICKÉ CENNÍKY — synced z Make Data Store
# ============================================================
PANELY = {
    "JAM72S20 460MR": {
        "Manufacturer": "JA Solar", "Type": "JAM72S20 460MR", "Dimensions_WxHxD": "2120x1052x40mm",
        "Weight": "25", "IP_Ingress_Protection": "IP68", "Ambient_Temperature": "-40÷85°C",
        "Classification": "Trieda II", "Cell_Type": "6x24 monokryštál", "Design_Load": "3600Pa",
        "Design_Pull": "1600Pa", "Cable_Connection": "MC4 4 mm² 2x0,3m", "UN_MAX": "1500",
        "IREV_MAX": "20", "PMPP": "460", "ISC": "11,45", "UOC": "50,01",
        "IMPP": "10,92", "UMPP": "42,13", "Efficiency": "20,6"
    },
    "LR7-60HVH-535M": {
        "Manufacturer": "LONGi", "Type": "LR7-60HVH-535M", "Dimensions_WxHxD": "1990x1134x30mm",
        "Weight": "25", "IP_Ingress_Protection": "IP68", "Ambient_Temperature": "-40÷85°C",
        "Classification": "Trieda II", "Cell_Type": "6x20 monokryštál", "Design_Load": "5400Pa",
        "Design_Pull": "2400Pa", "Cable_Connection": "MC4 4 mm² 2x1,4m", "UN_MAX": "1500",
        "IREV_MAX": "25", "PMPP": "535", "ISC": "15,15", "UOC": "44,78",
        "IMPP": "14,46", "UMPP": "37,01", "Efficiency": "23,7"
    },
    "LR7-60HVH-540M": {
        "Manufacturer": "LONGi", "Type": "LR7-60HVH-540M", "Dimensions_WxHxD": "1990x1134x30mm",
        "Weight": "25", "IP_Ingress_Protection": "IP68", "Ambient_Temperature": "-40÷85°C",
        "Classification": "Trieda II", "Cell_Type": "6x20 monokryštál", "Design_Load": "5400Pa",
        "Design_Pull": "2400Pa", "Cable_Connection": "MC4 4 mm² 2x1,4m", "UN_MAX": "1500",
        "IREV_MAX": "25", "PMPP": "540", "ISC": "15,25", "UOC": "44,88",
        "IMPP": "14,55", "UMPP": "37,11", "Efficiency": "23,9"
    },
    "LR7-54HVH-485M": {
        "Manufacturer": "LONGi", "Type": "LR7-54HVH-485M", "Dimensions_WxHxD": "1800x1134x30mm",
        "Weight": "21,6", "IP_Ingress_Protection": "IP68", "Ambient_Temperature": "-40÷85°C",
        "Classification": "Trieda II", "Cell_Type": "6x18 monokryštál", "Design_Load": "5400Pa",
        "Design_Pull": "2400Pa", "Cable_Connection": "MC4 4 mm² 2x1,2m", "UN_MAX": "1500",
        "IREV_MAX": "25", "PMPP": "485", "ISC": "15,23", "UOC": "40,4",
        "IMPP": "14,53", "UMPP": "33,4", "Efficiency": "23,8"
    },
    "TSM-620NEG19RC.20": {
        "Manufacturer": "Trina", "Type": "TSM-620NEG19RC.20", "Dimensions_WxHxD": "2382x1134x30mm",
        "Weight": "33,7", "IP_Ingress_Protection": "IP68", "Ambient_Temperature": "-40÷85°C",
        "Classification": "Trieda II", "Cell_Type": "6x22 monokryštál", "Design_Load": "5400Pa",
        "Design_Pull": "2400Pa", "Cable_Connection": "MC4 4 mm² 2x1,2m", "UN_MAX": "1500",
        "IREV_MAX": "35", "PMPP": "620", "ISC": "15,91", "UOC": "49,6",
        "IMPP": "14,99", "UMPP": "41,4", "Efficiency": "23"
    },
}

# Aliasy pre Notion Panel SELECT hodnoty
PANEL_ALIAS = {
    "LONGi 470 Wp": "LR7-54HVH-485M",  # najbližšia hodnota
    "LONGi 535 Wp": "LR7-60HVH-535M",
    "LONGi 540 Wp": "LR7-60HVH-540M",
    "JA Solar 460 Wp": "JAM72S20 460MR",
}


STRIEDACE = {
    "GW5K-ET": {
        "Manufacturer": "GoodWe", "Type": "GW5K-ET", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "415x516x180mm", "Weight": "24", "IP_Ingress_Protection": "IP66",
        "Ambient_Temperature": "-35÷60°C", "Relative_Humidity": "0÷95%", "Noise_Emission": "<30dB",
        "Efficiency": "97,2", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 180,
        "UMPP": 620, "UMPP_MAX": 1000, "IMPP": 12.5, "ISC": 15.2, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 312, "UN_MAX": 528, "PMAX": 5,
        "I_MAX": 8.5, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 16, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "GW6.5K-ET": {
        "Manufacturer": "GoodWe", "Type": "GW6.5K-ET", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "415x516x180mm", "Weight": "24", "IP_Ingress_Protection": "IP66",
        "Ambient_Temperature": "-35÷60°C", "Relative_Humidity": "0÷95%", "Noise_Emission": "<30dB",
        "Efficiency": "97,2", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 180,
        "UMPP": 620, "UMPP_MAX": 1000, "IMPP": 12.5, "ISC": 15.2, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 312, "UN_MAX": 528, "PMAX": 6.5,
        "I_MAX": 10.8, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 20, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "GW8K-ET": {
        "Manufacturer": "GoodWe", "Type": "GW8K-ET", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "415x516x180mm", "Weight": "24", "IP_Ingress_Protection": "IP66",
        "Ambient_Temperature": "-35÷60°C", "Relative_Humidity": "0÷95%", "Noise_Emission": "<30dB",
        "Efficiency": "97,5", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 180,
        "UMPP": 620, "UMPP_MAX": 1000, "IMPP": 12.5, "ISC": 15.2, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 312, "UN_MAX": 528, "PMAX": 8,
        "I_MAX": 13.5, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 25, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "GW10K-ET": {
        "Manufacturer": "GoodWe", "Type": "GW10K-ET", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "415x516x180mm", "Weight": "24", "IP_Ingress_Protection": "IP66",
        "Ambient_Temperature": "-35÷60°C", "Relative_Humidity": "0÷95%", "Noise_Emission": "<30dB",
        "Efficiency": "97,5", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 180,
        "UMPP": 620, "UMPP_MAX": 1000, "IMPP": 12.5, "ISC": 15.2, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 312, "UN_MAX": 528, "PMAX": 10,
        "I_MAX": 16.5, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 32, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "SUN2000-5K": {
        "Manufacturer": "Huawei", "Type": "SUN2000-5KTL", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "525x470x146,5mm", "Weight": "17", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-25÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<29dB",
        "Efficiency": "97,5", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 200,
        "UMPP": 600, "UMPP_MAX": 1100, "IMPP": 11, "ISC": 15, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 311, "UN_MAX": 478, "PMAX": 5,
        "I_MAX": 8.5, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 16, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "SUN2000-6K": {
        "Manufacturer": "Huawei", "Type": "SUN2000-6KTL", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "525x470x146,5mm", "Weight": "17", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-25÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<29dB",
        "Efficiency": "97,7", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 200,
        "UMPP": 600, "UMPP_MAX": 1100, "IMPP": 11, "ISC": 15, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 311, "UN_MAX": 478, "PMAX": 6,
        "I_MAX": 10.1, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 20, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "SUN2000-8K": {
        "Manufacturer": "Huawei", "Type": "SUN2000-8KTL", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "525x470x146,5mm", "Weight": "17", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-25÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<29dB",
        "Efficiency": "98", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 200,
        "UMPP": 600, "UMPP_MAX": 1080, "IMPP": 11, "ISC": 15, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 311, "UN_MAX": 478, "PMAX": 8,
        "I_MAX": 13.5, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 25, "Cable_Connection_AC_mm2": "16",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "SUN2000-10K": {
        "Manufacturer": "Huawei", "Type": "SUN2000-10KTL", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "525x470x146,5mm", "Weight": "17", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-25÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<29dB",
        "Efficiency": "98", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 200,
        "UMPP": 600, "UMPP_MAX": 1080, "IMPP": 11, "ISC": 15, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 311, "UN_MAX": 478, "PMAX": 10,
        "I_MAX": 16.9, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 32, "Cable_Connection_AC_mm2": "16",
        "RDC": "https://www.dropbox.com/scl/fi/a1nlym1cw2c2fo8j44i0y/2x1.pdf?rlkey=lipnbbka1au33pwzsxihrq1d5&dl=1"
    },
    "MHT-5K-25": {
        "Manufacturer": "Solinteg", "Type": "MHT-5K-25", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "534x418x210mm", "Weight": "26", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-30÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<25dB",
        "Efficiency": "98,1", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 135,
        "UMPP": 120, "UMPP_MAX": 950, "IMPP": 15, "ISC": 20, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 380, "UN_MAX": 415, "PMAX": 5,
        "I_MAX": 8.3, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 16, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/mg3boek83pam2qamw0qze/10-TechList-SOLINTEG-MHT-4-12KW.pdf?rlkey=90p1m47907pqmcwmaov9p4qgc&dl=1"
    },
    "MHT-6K-25": {
        "Manufacturer": "Solinteg", "Type": "MHT-6K-25", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "534x418x210mm", "Weight": "26", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-30÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<25dB",
        "Efficiency": "98,1", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 135,
        "UMPP": 120, "UMPP_MAX": 950, "IMPP": 15, "ISC": 20, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 380, "UN_MAX": 415, "PMAX": 6,
        "I_MAX": 10, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 20, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/mg3boek83pam2qamw0qze/10-TechList-SOLINTEG-MHT-4-12KW.pdf?rlkey=90p1m47907pqmcwmaov9p4qgc&dl=1"
    },
    "MHT-8K-25": {
        "Manufacturer": "Solinteg", "Type": "MHT-8K-25", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "534x418x210mm", "Weight": "26", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-30÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<25dB",
        "Efficiency": "98,2", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 135,
        "UMPP": 200, "UMPP_MAX": 950, "IMPP": 15, "ISC": 20, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 380, "UN_MAX": 415, "PMAX": 8,
        "I_MAX": 13.3, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 25, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/mg3boek83pam2qamw0qze/10-TechList-SOLINTEG-MHT-4-12KW.pdf?rlkey=90p1m47907pqmcwmaov9p4qgc&dl=1"
    },
    "MHT-10K-25": {
        "Manufacturer": "Solinteg", "Type": "MHT-10K-25", "Grid_System": "Hybrid",
        "Dimensions_WxHxD": "534x418x210mm", "Weight": "26", "IP_Ingress_Protection": "IP65",
        "Ambient_Temperature": "-30÷60°C", "Relative_Humidity": "0÷100%", "Noise_Emission": "<25dB",
        "Efficiency": "98,2", "MPPT": 2, "Strings_per_MPPT": 1, "UPV_MIN": 135,
        "UMPP": 200, "UMPP_MAX": 950, "IMPP": 15, "ISC": 20, "SPD_DC": "Type II",
        "Cable_Connection_DC_mm2": 6, "UN": 400, "UN_MIN": 380, "UN_MAX": 415, "PMAX": 10,
        "I_MAX": 16.5, "THD_Total_Harmonic_Distortion": "<3%", "Power_Factor": "0,99",
        "SPD_AC": "Type II", "Protection": 32, "Cable_Connection_AC_mm2": "6",
        "RDC": "https://www.dropbox.com/scl/fi/mg3boek83pam2qamw0qze/10-TechList-SOLINTEG-MHT-4-12KW.pdf?rlkey=90p1m47907pqmcwmaov9p4qgc&dl=1"
    },
}

# Aliasy pre Notion Menič SELECT hodnoty
STRIEDAC_ALIAS = {
    "Solinteg MHT-10K-25": "MHT-10K-25",
    "Huawei SUN2000-5K": "SUN2000-5K",
    "Huawei SUN2000-6K": "SUN2000-6K",
    "Huawei SUN2000-8K": "SUN2000-8K",
    "Huawei SUN2000-10K": "SUN2000-10K",
    "GoodWe GW6000-ET": "GW6.5K-ET",
    "GoodWe GW8000-ET": "GW8K-ET",
    "GoodWe GW10K-ET": "GW10K-ET",
}


# ============================================================
# Distribučné spoločnosti — SR rozdelenie podľa PSČ
# ============================================================
DIS_OFFICIAL = {
    "SSD": "Stredoslovenská distribučná, a.s., Pri Rajčianke 2927/8, 010 47 Žilina",
    "VSD": "Východoslovenská distribučná, a.s., Mlynská 31, 042 91 Košice",
    "ZSDIS": "Západoslovenská distribučná, a.s., Čulenova 6, 816 47 Bratislava",
}


def _resolve_dis_from_psc(psc):
    """Best-effort mapovanie PSČ → DIS (ak nie je v Notione)."""
    if not psc:
        return ""
    digits = re.sub(r'\D', '', str(psc))
    if not digits:
        return ""
    p = int(digits[:2]) if len(digits) >= 2 else 0
    # SSD: 01-03 (ZA), 96-97 (BB), 91-92 čiastočne
    # VSD: 04 (KE), 05-09 (PO + okolie)
    # ZSDIS: 81-85 (BA), 90-94 (TT/NR), 95 (LV)
    if p in (1, 2, 3, 96, 97):
        return "SSD"
    if p in (4, 5, 6, 7, 8, 9) and p != 9:  # 09 → SSD
        # 04 KE, 05 SP, 06 SL, 07 MI, 08 PO
        if p in (4, 5, 6, 7, 8):
            return "VSD"
    if p in (81, 82, 83, 84, 85, 90, 91, 92, 93, 94, 95):
        return "ZSDIS"
    return ""


# ============================================================
# HELPERS
# ============================================================

def _safe(v, fallback=""):
    if v is None or v == "":
        return fallback
    return str(v)


def _resolve_panel(typ_panela):
    """Vyhľadaj panel v cenníku — najprv presný kľúč, potom alias."""
    if not typ_panela:
        return PANELY["LR7-60HVH-535M"]  # default
    if typ_panela in PANELY:
        return PANELY[typ_panela]
    if typ_panela in PANEL_ALIAS:
        return PANELY[PANEL_ALIAS[typ_panela]]
    # Fuzzy match — hľadaj v alias podľa Wp hodnoty
    m = re.search(r'(\d{3})\s*W', typ_panela)
    if m:
        wp = m.group(1)
        for k, v in PANELY.items():
            if v.get("PMPP") == wp:
                return v
    return PANELY["LR7-60HVH-535M"]


def _resolve_striedac(typ_menica):
    """Vyhľadaj striedač v cenníku — najprv presný kľúč, potom alias."""
    if not typ_menica:
        return STRIEDACE["MHT-10K-25"]  # default
    if typ_menica in STRIEDACE:
        return STRIEDACE[typ_menica]
    if typ_menica in STRIEDAC_ALIAS:
        return STRIEDACE[STRIEDAC_ALIAS[typ_menica]]
    # Fuzzy — hľadaj v cenníku podľa substringu
    for k in STRIEDACE.keys():
        if k.lower() in typ_menica.lower() or typ_menica.lower() in k.lower():
            return STRIEDACE[k]
    return STRIEDACE["MHT-10K-25"]


def _set_cell(cell, text, bold=False, size=10, align=None):
    """Helper — nastav text v bunke s formátovaním."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True


def _add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def _add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def _add_p(doc, text, bold=False, size=10, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold:
        r.bold = True
    if align is not None:
        p.alignment = align
    return p


def _build_so01(lead_data):
    """Vygeneruje názov stavebného objektu SO01."""
    vykon = lead_data.get('vykon_kwp', 0)
    bateria = lead_data.get('bateria_kwh', 0)
    txt = f"SO01 — Fotovoltická elektráreň {vykon:.2f} kWp"
    if bateria > 0:
        txt += f" + batériové úložisko {bateria:.2f} kWh"
    return txt


def _build_typ(lead_data):
    """Vygeneruje 'typ' projektu (variant)."""
    variant = lead_data.get('variant', 'B')
    map_typ = {
        "A": f"FVE {lead_data.get('vykon_kwp', 0):.2f} kWp",
        "B": f"FVE {lead_data.get('vykon_kwp', 0):.2f} kWp + BESS {lead_data.get('bateria_kwh', 0):.2f} kWh",
        "C": f"FVE {lead_data.get('vykon_kwp', 0):.2f} kWp + BESS {lead_data.get('bateria_kwh', 0):.2f} kWh + Wallbox",
        "D": f"FVE {lead_data.get('vykon_kwp', 0):.2f} kWp + Wallbox",
    }
    return map_typ.get(variant, map_typ["B"])


def _setup_doc(doc):
    """Spoločné nastavenie margins."""
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)


# ============================================================
# 1. KRYCÍ LIST
# ============================================================

def gen_kryci_list(lead_data, output_path):
    """Krycí list projektu — 1 strana, tabuľka s identifikáciou."""
    doc = Document()
    _setup_doc(doc)

    nazov = _safe(lead_data.get('meno_priezvisko'))
    adresa_klient = _safe(lead_data.get('trvale_bydlisko')) or _safe(lead_data.get('adresa'))
    psc = _safe(lead_data.get('psc'))
    mesto = _safe(lead_data.get('mesto'))
    parcely = _safe(lead_data.get('parcelne_cisla'))
    ev_id = _safe(lead_data.get('ev_id'), "EV-26-XXX")
    datum = _safe(lead_data.get('datum_dnes'), datetime.now().strftime("%d.%m.%Y"))

    # Hlavička – mená firmy
    _add_h1(doc, "Energovision s.r.o.")
    _add_p(doc, "Lamačská cesta 1738/111, 841 03 Bratislava | IČO: 53 036 280", size=10)

    doc.add_paragraph()

    # Titulok
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KRYCÍ LIST PROJEKTOVEJ DOKUMENTÁCIE")
    r.bold = True; r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fotovoltická elektráreň — malý zdroj do 10 kW")
    r.font.size = Pt(11); r.italic = True

    doc.add_paragraph()

    # Hlavná tabuľka identifikácie
    t = doc.add_table(rows=12, cols=2)
    t.style = 'Light Grid Accent 1'
    t.columns[0].width = Cm(6)
    t.columns[1].width = Cm(11)

    rows = [
        ("Názov projektu", "Fotovoltická elektráreň"),
        ("Stavebný objekt", _build_so01(lead_data)),
        ("Investor / Objednávateľ", nazov),
        ("Trvalé bydlisko investora", adresa_klient),
        ("Miesto stavby", f"{adresa_klient}"),
        ("Parcelné čísla", parcely or "—"),
        ("Distribučná spoločnosť", _safe(lead_data.get('dis')) + (" — " + DIS_OFFICIAL[_safe(lead_data.get('dis'))] if _safe(lead_data.get('dis')) in DIS_OFFICIAL else "")),
        ("Číslo zákazky", ev_id),
        ("Stupeň dokumentácie", "DPP — Dokumentácia pre pripojenie"),
        ("Vypracoval", KOMISIA["vypracoval"]),
        ("Kontroloval", KOMISIA["kontroloval"]),
        ("Zodpovedný projektant", KOMISIA["zodpovedny_projektant"]),
    ]
    for i, (label, val) in enumerate(rows):
        _set_cell(t.rows[i].cells[0], label, bold=True, size=10)
        _set_cell(t.rows[i].cells[1], val, size=10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Podpisová tabuľka
    sig = doc.add_table(rows=2, cols=3)
    _set_cell(sig.rows[0].cells[0], "Vypracoval", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[0].cells[1], "Kontroloval", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[0].cells[2], "Zodp. projektant", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[1].cells[0], f"{KOMISIA['vypracoval']}\n\n.................................", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[1].cells[1], f"{KOMISIA['kontroloval']}\n\n.................................", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[1].cells[2], f"{KOMISIA['zodpovedny_projektant']}\n\n.................................", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_p(doc, f"V Bratislave, dňa {datum}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(output_path))
    log.info("[pd-kryci] generovaný pre %s", nazov)
    return output_path


# ============================================================
# 2. ZOZNAM DOKUMENTÁCIE
# ============================================================

def gen_zoznam_dokumentacie(lead_data, output_path):
    """Zoznam priložených dokumentov."""
    doc = Document()
    _setup_doc(doc)

    ev_id = _safe(lead_data.get('ev_id'), "EV-26-XXX")
    panel = _resolve_panel(lead_data.get('panel_typ'))
    striedac = _resolve_striedac(lead_data.get('menic'))
    bateria_typ = _safe(lead_data.get('bateria_typ'))
    ma_bateriu = (lead_data.get('pocet_baterii') or 0) > 0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ZOZNAM DOKUMENTÁCIE")
    r.bold = True; r.font.size = Pt(16)

    _add_p(doc, f"Číslo zákazky: {ev_id}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    items = [
        ("01", "Krycí list", "A4", "1"),
        ("02", "Zoznam dokumentácie", "A4", "1"),
        ("03", "Technická správa", "A4", "3-5"),
        ("04", "Protokol o určení vonkajších vplyvov (PoUVV)", "A4", "2-3"),
        ("05", "Súhrnná technická správa", "A4", "1-2"),
        ("06", f"Technický list panela — {panel['Manufacturer']} {panel['Type']}", "A4", "1"),
        ("07", f"Technický list meniča — {striedac['Manufacturer']} {striedac['Type']}", "A4", "1"),
        ("08", "Schéma zapojenia DC + AC rozvádzača", "A3", "1-2"),
    ]
    if ma_bateriu:
        items.append(("09", f"Technický list batériového úložiska — {bateria_typ}", "A4", "1"))

    t = doc.add_table(rows=len(items) + 1, cols=4)
    t.style = 'Light Grid Accent 1'
    _set_cell(t.rows[0].cells[0], "Č.", bold=True, size=10)
    _set_cell(t.rows[0].cells[1], "Príloha", bold=True, size=10)
    _set_cell(t.rows[0].cells[2], "Formát", bold=True, size=10)
    _set_cell(t.rows[0].cells[3], "Strán", bold=True, size=10)

    for i, (cislo, popis, format, stran) in enumerate(items, start=1):
        _set_cell(t.rows[i].cells[0], cislo, size=9)
        _set_cell(t.rows[i].cells[1], popis, size=9)
        _set_cell(t.rows[i].cells[2], format, size=9)
        _set_cell(t.rows[i].cells[3], stran, size=9)

    doc.save(str(output_path))
    log.info("[pd-zoznam] generovaný")
    return output_path


# ============================================================
# 3. TECHNICKÁ SPRÁVA (zjednodušená pre malé zdroje do 10 kW)
# ============================================================

def gen_technicka_sprava(lead_data, output_path):
    """Technická správa pre malý zdroj FVE do 10 kW."""
    doc = Document()
    _setup_doc(doc)

    nazov = _safe(lead_data.get('meno_priezvisko'))
    adresa = _safe(lead_data.get('trvale_bydlisko')) or _safe(lead_data.get('adresa'))
    ev_id = _safe(lead_data.get('ev_id'), "EV-26-XXX")
    vykon = lead_data.get('vykon_kwp', 0)
    pocet_panelov = lead_data.get('pocet_panelov', 0)
    bateria_kwh = lead_data.get('bateria_kwh', 0)
    pocet_baterii = lead_data.get('pocet_baterii', 0)
    bateria_typ = _safe(lead_data.get('bateria_typ'))
    konstrukcia = _safe(lead_data.get('konstrukcia'), "Šikmá strecha (škridla)")
    hlavny_istic = _safe(lead_data.get('hlavny_istic'), "3x25A")
    ma_wallbox = lead_data.get('ma_wallbox', False)
    wallbox_typ = _safe(lead_data.get('wallbox_typ'))
    eic_odber = _safe(lead_data.get('eic'))
    eic_dodavka = _safe(lead_data.get('eic_dodavka'))
    dis = _safe(lead_data.get('dis'))
    dis_full = DIS_OFFICIAL.get(dis, "—")
    datum = _safe(lead_data.get('datum_dnes'), datetime.now().strftime("%d.%m.%Y"))

    panel = _resolve_panel(lead_data.get('panel_typ'))
    striedac = _resolve_striedac(lead_data.get('menic'))

    # Hlavička
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TECHNICKÁ SPRÁVA")
    r.bold = True; r.font.size = Pt(16)
    _add_p(doc, "Fotovoltická elektráreň — malý zdroj do 10 kW", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _add_p(doc, f"Číslo zákazky: {ev_id}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    # 1. Základné údaje
    _add_h2(doc, "1. Základné údaje")
    t = doc.add_table(rows=6, cols=2)
    t.style = 'Light Grid Accent 1'
    rows = [
        ("Investor", nazov),
        ("Miesto stavby", adresa),
        ("Distribučná spoločnosť", dis_full),
        ("EIC odberného miesta", eic_odber or "—"),
        ("EIC dodávateľa (lokálny zdroj)", eic_dodavka or "—"),
        ("Hlavný istič", hlavny_istic),
    ]
    for i, (l, v) in enumerate(rows):
        _set_cell(t.rows[i].cells[0], l, bold=True, size=10)
        _set_cell(t.rows[i].cells[1], v, size=10)

    # 2. Predpisy a normy
    _add_h2(doc, "2. Predpisy a normy")
    for line in [
        "Zákon č. 251/2012 Z. z. o energetike v znení neskorších predpisov",
        "Vyhláška č. 508/2009 Z. z. — bezpečnosť VTZ elektrických",
        "STN 33 2000-1 — Elektrické inštalácie nízkeho napätia",
        "STN 33 2000-4-41 — Ochrana pred zásahom elektrickým prúdom",
        "STN 33 2000-5-51 — Vonkajšie vplyvy",
        "STN 33 2000-7-712 — Fotovoltické (PV) systémy",
        "STN EN 62446-1 — Skúšky, dokumentácia, údržba FV systémov",
        "Technické podmienky príslušného PDS (SSD/VSD/ZSDIS)",
    ]:
        _add_p(doc, "• " + line, size=10)

    # 3. Technický popis FVE
    _add_h2(doc, "3. Technický popis fotovoltickej elektrárne")
    popis = (
        f"Predmetom dokumentácie je inštalácia fotovoltickej elektrárne (FVE) ON GRID "
        f"s celkovým inštalovaným výkonom {vykon:.2f} kWp"
    )
    if bateria_kwh > 0:
        popis += f" v kombinácii s batériovým úložiskom {bateria_kwh:.2f} kWh"
    popis += (
        f". Systém pozostáva z {pocet_panelov} ks fotovoltických panelov typu "
        f"{panel['Manufacturer']} {panel['Type']} (jednotlivý výkon {panel['PMPP']} Wp), "
        f"meniča (striedača) typu {striedac['Manufacturer']} {striedac['Type']} "
        f"s menovitým výkonom {striedac['PMAX']} kW"
    )
    if pocet_baterii > 0:
        popis += f", batériového úložiska {bateria_typ} v počte {pocet_baterii} ks"
    if ma_wallbox:
        popis += f" a wallboxu {wallbox_typ}"
    popis += "."
    _add_p(doc, popis, size=10)

    _add_p(doc,
        f"FVE bude umiestnená na streche rodinného domu — konštrukcia: {konstrukcia}. "
        f"Pripojenie do AC siete je realizované cez hlavný rozvádzač objektu cez istič "
        f"{hlavny_istic}.", size=10)

    if pocet_baterii > 0:
        _add_p(doc,
            f"Batériové úložisko je pripojené k hybridnému meniču {striedac['Type']} cez "
            "DC port. Slúži na akumuláciu prebytočnej energie z FVE a jej využitie v dobe "
            "nízkej produkcie alebo nepriaznivého počasia.", size=10)

    if ma_wallbox:
        _add_p(doc,
            f"Wallbox {wallbox_typ} je pripojený do AC rozvádzača cez vlastný prúdový "
            "chránič typu B a istič v zmysle STN EN 61851 (nabíjacie stanice elektrických "
            "vozidiel).", size=10)

    # 4. Technické parametre panela
    _add_h2(doc, "4. Technické parametre fotovoltického panela")
    t = doc.add_table(rows=13, cols=2)
    t.style = 'Light Grid Accent 1'
    panel_rows = [
        ("Výrobca", panel["Manufacturer"]),
        ("Typ", panel["Type"]),
        ("Rozmery (š × v × h)", panel["Dimensions_WxHxD"]),
        ("Hmotnosť", f"{panel['Weight']} kg"),
        ("Krytie IP", panel["IP_Ingress_Protection"]),
        ("Rozsah pracovných teplôt", panel["Ambient_Temperature"]),
        ("Trieda ochrany", panel["Classification"]),
        ("Menovitý výkon (Pmpp)", f"{panel['PMPP']} Wp"),
        ("Napätie pri Pmpp (Umpp)", f"{panel['UMPP']} V"),
        ("Prúd pri Pmpp (Impp)", f"{panel['IMPP']} A"),
        ("Napätie naprázdno (Uoc)", f"{panel['UOC']} V"),
        ("Skratový prúd (Isc)", f"{panel['ISC']} A"),
        ("Účinnosť", f"{panel['Efficiency']} %"),
    ]
    for i, (l, v) in enumerate(panel_rows):
        _set_cell(t.rows[i].cells[0], l, bold=True, size=9)
        _set_cell(t.rows[i].cells[1], v, size=9)

    # 5. Technické parametre meniča
    _add_h2(doc, "5. Technické parametre fotovoltického meniča")
    t = doc.add_table(rows=14, cols=2)
    t.style = 'Light Grid Accent 1'
    striedac_rows = [
        ("Výrobca", striedac["Manufacturer"]),
        ("Typ", striedac["Type"]),
        ("Typ systému", striedac["Grid_System"]),
        ("Rozmery (š × v × h)", striedac["Dimensions_WxHxD"]),
        ("Hmotnosť", f"{striedac['Weight']} kg"),
        ("Krytie IP", striedac["IP_Ingress_Protection"]),
        ("Účinnosť", f"{striedac['Efficiency']} %"),
        ("Počet MPPT", str(striedac["MPPT"])),
        ("Reťazcov / MPPT", str(striedac["Strings_per_MPPT"])),
        ("DC napätie min / nominal / max", f"{striedac['UPV_MIN']} / {striedac['UMPP']} / {striedac['UMPP_MAX']} V"),
        ("DC max. prúd (Isc)", f"{striedac['ISC']} A"),
        ("AC menovitý výkon", f"{striedac['PMAX']} kW"),
        ("AC max. prúd", f"{striedac['I_MAX']} A"),
        ("Účinník (cos φ)", striedac["Power_Factor"]),
    ]
    for i, (l, v) in enumerate(striedac_rows):
        _set_cell(t.rows[i].cells[0], l, bold=True, size=9)
        _set_cell(t.rows[i].cells[1], v, size=9)

    # 6. Ochrany
    _add_h2(doc, "6. Ochrany a bezpečnostné prvky")
    for line in [
        "Ochrana pred zásahom elektrickým prúdom (živé časti): izoláciou, krytmi, istiacimi prvkami, prúdovým chráničom",
        "Ochrana pred zásahom (neživé časti): samočinné odpojenie napájania",
        "DC ochrana proti prepätiu (SPD): " + striedac["SPD_DC"] + " — 1100 VDC",
        "AC ochrana proti prepätiu (SPD): " + striedac["SPD_AC"],
        "Núdzové vypnutie FV systému — vypínač DC pri meniči",
        "Sieťová ochrana podľa STN EN 50549-1 (anti-islanding, automatické odpojenie pri výpadku siete)",
        f"Prúdový chránič typu A (FI 30 mA) pre AC stranu, istič {striedac['Protection']} A",
    ]:
        _add_p(doc, "• " + line, size=10)

    # 7. Bezpečnosť pri práci
    _add_h2(doc, "7. Bezpečnosť pri práci a životné prostredie")
    _add_p(doc,
        "Inštaláciu vykonajú pracovníci s príslušnou elektrotechnickou spôsobilosťou podľa "
        "vyhl. č. 508/2009 Z. z., § 23 (samostatný elektrotechnik) alebo vyšší. "
        "Pri prácach vo výške bude použitý certifikovaný kotviaci systém. "
        "Po ukončení realizácie bude vykonaná východisková odborná prehliadka a odborná "
        "skúška (OPaOS) v zmysle STN 33 1500. "
        "Demontovaný a odpadový materiál bude likvidovaný v súlade so zákonom o odpadoch.",
        size=10)

    # 8. Záver
    _add_h2(doc, "8. Záver")
    _add_p(doc,
        "Predmetná dokumentácia opisuje technické riešenie inštalácie fotovoltickej elektrárne "
        f"o výkone {vykon:.2f} kWp na rodinnom dome. Návrh spĺňa všetky platné technické normy "
        "a predpisy. Po realizácii inštalácie bude vykonaná odborná prehliadka a skúška, "
        "vystavená revízna správa a zariadenie odovzdané investorovi do prevádzky.",
        size=10)

    doc.add_paragraph()
    _add_p(doc, f"V Bratislave, dňa {datum}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_p(doc, f"Vypracoval: {KOMISIA['vypracoval']}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_p(doc, f"Zodp. projektant: {KOMISIA['zodpovedny_projektant']}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(output_path))
    log.info("[pd-tech] generovaný pre %s, vykon=%.2f kWp", nazov, vykon)
    return output_path


# ============================================================
# 4. PROTOKOL O URČENÍ VONKAJŠÍCH VPLYVOV (PoUVV)
# ============================================================

def gen_pouvv(lead_data, output_path):
    """Protokol o určení vonkajších vplyvov podľa STN 33 2000-5-51."""
    doc = Document()
    _setup_doc(doc)

    nazov = _safe(lead_data.get('meno_priezvisko'))
    adresa = _safe(lead_data.get('trvale_bydlisko')) or _safe(lead_data.get('adresa'))
    ev_id = _safe(lead_data.get('ev_id'), "EV-26-XXX")
    cislo_pouvv = _safe(lead_data.get('cislo_pouvv'), f"PoUVV-{ev_id}")
    vykon = lead_data.get('vykon_kwp', 0)
    datum = _safe(lead_data.get('datum_dnes'), datetime.now().strftime("%d.%m.%Y"))

    # Hlavička
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROTOKOL O URČENÍ VONKAJŠÍCH VPLYVOV")
    r.bold = True; r.font.size = Pt(14)
    _add_p(doc, f"podľa STN 33 2000-5-51", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _add_p(doc, f"Č. protokolu: {cislo_pouvv}", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True)

    doc.add_paragraph()

    # Identifikácia stavby
    _add_h2(doc, "1. Identifikácia stavby")
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (l, v) in enumerate([
        ("Investor", nazov),
        ("Miesto stavby", adresa),
        ("Predmet", f"Fotovoltická elektráreň {vykon:.2f} kWp"),
        ("Číslo zákazky", ev_id),
    ]):
        _set_cell(t.rows[i].cells[0], l, bold=True, size=10)
        _set_cell(t.rows[i].cells[1], v, size=10)

    # Komisia
    _add_h2(doc, "2. Komisia pre určenie vonkajších vplyvov")
    t = doc.add_table(rows=3, cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (l, v) in enumerate([
        ("Predseda komisie", KOMISIA["zodpovedny_projektant"]),
        ("Člen komisie", KOMISIA["vypracoval"]),
        ("Člen komisie", KOMISIA["kontroloval"]),
    ]):
        _set_cell(t.rows[i].cells[0], l, bold=True, size=10)
        _set_cell(t.rows[i].cells[1], v, size=10)

    # Použité podklady
    _add_h2(doc, "3. Použité podklady")
    for line in [
        "STN 33 2000-5-51 — Vonkajšie vplyvy",
        "STN 33 2000-1 — Elektrické inštalácie nízkeho napätia",
        "STN 33 2000-7-712 — Fotovoltické (PV) systémy",
        "Mapa povodňových oblastí SR (Ministerstvo životného prostredia SR)",
        "Mapa zemetrasných oblastí SR",
        "Klimatické údaje SHMÚ pre danú lokalitu",
    ]:
        _add_p(doc, "• " + line, size=10)

    # Tabuľka vonkajších vplyvov — pre rodinný dom
    _add_h2(doc, "4. Tabuľka vonkajších vplyvov — vonkajšia inštalácia (strecha + fasáda)")
    t = doc.add_table(rows=9, cols=3)
    t.style = 'Light Grid Accent 1'
    _set_cell(t.rows[0].cells[0], "Kód", bold=True, size=10)
    _set_cell(t.rows[0].cells[1], "Vplyv", bold=True, size=10)
    _set_cell(t.rows[0].cells[2], "Stupeň", bold=True, size=10)
    vplyvy_vonk = [
        ("AA", "Teplota okolia", "AA8 (-50÷+40 °C)"),
        ("AB", "Vlhkosť vzduchu", "AB8"),
        ("AD", "Voda", "AD4 (striekajúca voda)"),
        ("AE", "Cudzie pevné telesá", "AE2 (malé predmety)"),
        ("AF", "Korózia", "AF1 (zanedbateľná)"),
        ("AG", "Mechanické namáhanie", "AG1"),
        ("AN", "Slnečné žiarenie", "AN3 (vysoké)"),
        ("BA", "Schopnosť osôb", "BA5 (kvalifikovaní pracovníci)"),
    ]
    for i, (kod, vplyv, stupen) in enumerate(vplyvy_vonk, start=1):
        _set_cell(t.rows[i].cells[0], kod, size=9)
        _set_cell(t.rows[i].cells[1], vplyv, size=9)
        _set_cell(t.rows[i].cells[2], stupen, size=9)

    doc.add_paragraph()

    # Tabuľka vonkajších vplyvov — vnútorná
    _add_h2(doc, "5. Tabuľka vonkajších vplyvov — vnútorná inštalácia (rozvádzač, menič)")
    t = doc.add_table(rows=8, cols=3)
    t.style = 'Light Grid Accent 1'
    _set_cell(t.rows[0].cells[0], "Kód", bold=True, size=10)
    _set_cell(t.rows[0].cells[1], "Vplyv", bold=True, size=10)
    _set_cell(t.rows[0].cells[2], "Stupeň", bold=True, size=10)
    vplyvy_vnut = [
        ("AA", "Teplota okolia", "AA5 (-5÷+40 °C)"),
        ("AB", "Vlhkosť vzduchu", "AB5"),
        ("AD", "Voda", "AD1 (zanedbateľná)"),
        ("AE", "Cudzie pevné telesá", "AE1"),
        ("AF", "Korózia", "AF1"),
        ("AG", "Mechanické namáhanie", "AG1"),
        ("BA", "Schopnosť osôb", "BA1 (laici)"),
    ]
    for i, (kod, vplyv, stupen) in enumerate(vplyvy_vnut, start=1):
        _set_cell(t.rows[i].cells[0], kod, size=9)
        _set_cell(t.rows[i].cells[1], vplyv, size=9)
        _set_cell(t.rows[i].cells[2], stupen, size=9)

    # Záver
    _add_h2(doc, "6. Záver")
    _add_p(doc,
        "Komisia preskúmala podklady a stanovila vonkajšie vplyvy pôsobiace na elektrickú "
        "inštaláciu predmetnej fotovoltickej elektrárne. Použité zariadenia a materiály "
        "spĺňajú stanovené požiadavky pre uvedené stupne vplyvov. "
        "Inštalácia bude realizovaná v zmysle STN 33 2000-5-51.", size=10)

    doc.add_paragraph()

    # Podpisová tabuľka
    sig = doc.add_table(rows=2, cols=3)
    _set_cell(sig.rows[0].cells[0], "Predseda komisie", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[0].cells[1], "Člen komisie", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[0].cells[2], "Člen komisie", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[1].cells[0], f"{KOMISIA['zodpovedny_projektant']}\n\n....................", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[1].cells[1], f"{KOMISIA['vypracoval']}\n\n....................", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(sig.rows[1].cells[2], f"{KOMISIA['kontroloval']}\n\n....................", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_p(doc, f"V Bratislave, dňa {datum}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(output_path))
    log.info("[pd-pouvv] generovaný pre %s", nazov)
    return output_path


# ============================================================
# 5. SÚHRNNÁ TECHNICKÁ SPRÁVA
# ============================================================

def gen_suhrnna_sprava(lead_data, output_path):
    """Súhrnná technická správa — krátka verzia pre stavebné konanie / klienta."""
    doc = Document()
    _setup_doc(doc)

    nazov = _safe(lead_data.get('meno_priezvisko'))
    adresa = _safe(lead_data.get('trvale_bydlisko')) or _safe(lead_data.get('adresa'))
    ev_id = _safe(lead_data.get('ev_id'), "EV-26-XXX")
    vykon = lead_data.get('vykon_kwp', 0)
    pocet_panelov = lead_data.get('pocet_panelov', 0)
    bateria_kwh = lead_data.get('bateria_kwh', 0)
    pocet_baterii = lead_data.get('pocet_baterii', 0)
    konstrukcia = _safe(lead_data.get('konstrukcia'), "Šikmá strecha")
    parcely = _safe(lead_data.get('parcelne_cisla'))
    datum = _safe(lead_data.get('datum_dnes'), datetime.now().strftime("%d.%m.%Y"))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SÚHRNNÁ TECHNICKÁ SPRÁVA")
    r.bold = True; r.font.size = Pt(16)
    _add_p(doc, _build_so01(lead_data), align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True)
    _add_p(doc, f"Číslo zákazky: {ev_id}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    doc.add_paragraph()

    _add_h2(doc, "1. Základné údaje o stavbe")
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (l, v) in enumerate([
        ("Investor", nazov),
        ("Miesto stavby", adresa),
        ("Parcelné čísla", parcely or "—"),
        ("Druh stavby", "Stavebná úprava — inštalácia fotovoltického systému"),
        ("Charakter prevádzky", "Lokálny zdroj elektrickej energie (vlastná spotreba + prebytky do siete)"),
    ]):
        _set_cell(t.rows[i].cells[0], l, bold=True, size=10)
        _set_cell(t.rows[i].cells[1], v, size=10)

    _add_h2(doc, "2. Stručný popis riešenia")
    _add_p(doc,
        f"Predmetom stavby je inštalácia fotovoltickej elektrárne s celkovým výkonom "
        f"{vykon:.2f} kWp na streche existujúceho rodinného domu. Inštalácia pozostáva z "
        f"{pocet_panelov} kusov fotovoltických panelov ukotvených na konštrukcii ({konstrukcia.lower()}). "
        + (f"Súčasťou systému je batériové úložisko s kapacitou {bateria_kwh:.2f} kWh "
           f"({pocet_baterii} ks). " if pocet_baterii > 0 else "")
        + "FVE bude pripojená do existujúcej elektroinštalácie objektu cez menič (striedač) a hlavný rozvádzač.",
        size=10)

    _add_h2(doc, "3. Územné a stavebné podmienky")
    _add_p(doc,
        "Inštalácia FVE je stavebnou úpravou existujúcej stavby — nemení tvar ani objem "
        "existujúceho rodinného domu. Vzhľadom na rozsah (do 10 kW) ide o malý zdroj "
        "elektriny, pre ktorý stavebné povolenie nie je potrebné — postačí ohlásenie "
        "stavebného úradu v zmysle § 57 stavebného zákona.", size=10)

    _add_h2(doc, "4. Vplyv na životné prostredie")
    _add_p(doc,
        "Inštalácia FVE má pozitívny vplyv na životné prostredie — produkuje elektrickú "
        "energiu z obnoviteľného zdroja bez emisií CO₂. Nepoužívajú sa žiadne nebezpečné "
        "látky. Po skončení životnosti panelov bude zabezpečená ich ekologická recyklácia "
        "v zmysle zákona o odpadoch a smernice WEEE.", size=10)

    _add_h2(doc, "5. Bezpečnosť pri realizácii")
    _add_p(doc,
        "Inštalácia bude vykonaná odborne spôsobilou osobou v zmysle vyhl. č. 508/2009 Z. z. "
        "V priebehu realizácie budú dodržané všetky predpisy BOZP, najmä pre práce vo výške "
        "(zaistenie certifikovaným kotviacim systémom) a pre prácu na elektrických zariadeniach. "
        "Po inštalácii bude vykonaná východisková odborná prehliadka a skúška (OPaOS).", size=10)

    doc.add_paragraph()
    _add_p(doc, f"V Bratislave, dňa {datum}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_p(doc, f"Vypracoval: {KOMISIA['vypracoval']}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(output_path))
    log.info("[pd-suhrnna] generovaný pre %s", nazov)
    return output_path


# ============================================================
# MASTER ENTRY POINT
# ============================================================

def vygeneruj_projektovu_dokumentaciu(lead_data, out_dir, solaredge_pdf_bytes=None):
    """
    Vyrobí kompletný balík PD pre malý zdroj do 10 kW.
    Returns: dict {kluc: path} s 5-6 dokumentmi (s/bez technického výkresu).

    solaredge_pdf_bytes: voliteľné, ak je dodané, vygeneruje sa aj technický výkres A3.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    priezvisko = lead_data.get('meno_priezvisko', 'Klient').split()[-1] if lead_data.get('meno_priezvisko') else 'Klient'
    base = re.sub(r'[^A-Za-zÁ-ž0-9]+', '_', priezvisko).strip('_') or 'Klient'
    ev_id = lead_data.get('ev_id', 'EV-XX')

    # Doplniť DIS ak nie je
    if not lead_data.get('dis'):
        psc_guess = _resolve_dis_from_psc(lead_data.get('psc'))
        if psc_guess:
            lead_data['dis'] = psc_guess

    out = {}
    out['kryci'] = gen_kryci_list(lead_data, out_dir / f"{ev_id}_PD_01_Kryci_list_{base}.docx")
    out['zoznam'] = gen_zoznam_dokumentacie(lead_data, out_dir / f"{ev_id}_PD_02_Zoznam_dokumentacie_{base}.docx")
    out['technicka'] = gen_technicka_sprava(lead_data, out_dir / f"{ev_id}_PD_03_Technicka_sprava_{base}.docx")
    out['pouvv'] = gen_pouvv(lead_data, out_dir / f"{ev_id}_PD_04_PoUVV_{base}.docx")
    out['suhrnna'] = gen_suhrnna_sprava(lead_data, out_dir / f"{ev_id}_PD_05_Suhrnna_sprava_{base}.docx")

    # 6. Technický výkres (ak je SolarEdge PDF k dispozícii)
    if solaredge_pdf_bytes:
        try:
            from solar_vykres import vyrob_z_bytes
            vykres_path = out_dir / f"{ev_id}_PD_06_Vykres_FVE_{base}.pdf"
            vyrob_z_bytes(solaredge_pdf_bytes, lead_data, vykres_path)
            out['vykres'] = vykres_path
            log.info("[pd] technický výkres pridaný")
        except Exception as e:
            log.warning("[pd] technický výkres zlyhal: %s", e)

    log.info("[pd] vygenerované %d dokumentov pre %s", len(out), priezvisko)
    return out
