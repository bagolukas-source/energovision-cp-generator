# -*- coding: utf-8 -*-
"""ChocoSuc-grade posudok ako natívny DOCX (python-docx, funguje na Render bez libreoffice).
Mirror PDF: identifikácia, profil (grafy), technika+bilancia, ekonomika (tarif, 3 scenáre, grafy),
skladba+riziko (tornado/MC), cena nečinnosti, odporúčania, záver, predpoklady. Číta ten istý ctx."""
import io, base64, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN=RGBColor(0x16,0xA3,0x4A); LIME=RGBColor(0x92,0xD0,0x50); DARK=RGBColor(0x1A,0x1A,0x1A); GREY=RGBColor(0x6B,0x72,0x80)

def eur(v):
    try: return f"{float(v):,.0f} €".replace(","," ")
    except Exception: return "—"
def num(v,d=0):
    try: return (f"{float(v):,.{d}f}").replace(","," ").replace(".",",")
    except Exception: return "—"
def _png(data_uri):
    if not data_uri or "," not in data_uri: return None
    return io.BytesIO(base64.b64decode(data_uri.split(",",1)[1]))
def _strip(html):
    return re.sub(r"<[^>]+>","",html or "").strip()

def _shade(cell,hexc):
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); cell._tc.get_or_add_tcPr().append(sh)
def _kick(doc,t):
    p=doc.add_paragraph(); r=p.add_run(t.upper()); r.bold=True; r.font.size=Pt(8); r.font.color.rgb=GREEN
    r.font.name="Arial"; pf=p.paragraph_format; pf.space_after=Pt(2); pf.space_before=Pt(10)
    return p
def _h(doc,t,size=14):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=DARK; r.font.name="Arial"
    p.paragraph_format.space_after=Pt(6); return p
def _para(doc,t,size=10,color=DARK,italic=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.italic=italic; r.font.name="Arial"
    p.paragraph_format.space_after=Pt(4); return p
def _table(doc,headers,rows,aligns=None,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    try: t.style="Table Grid"
    except Exception: pass
    for i,hd in enumerate(headers):
        c=t.rows[0].cells[i]; _shade(c,"F0F7F0"); p=c.paragraphs[0]; r=p.add_run(str(hd)); r.bold=True; r.font.size=Pt(8.5); r.font.name="Arial"; r.font.color.rgb=GREY
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            p=cells[i].paragraphs[0]; r=p.add_run(str(v)); r.font.size=Pt(9); r.font.name="Arial"
            if aligns and i<len(aligns) and aligns[i]=="r": p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    return t
def _img(doc,data_uri,width_cm=16.5,cap=None):
    b=_png(data_uri)
    if b:
        doc.add_picture(b,width=Cm(width_cm))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if cap: _para(doc,cap,size=8,color=GREY,italic=True)

def generate_chocosuc_docx(ctx: dict) -> bytes:
    import posudok_chocosuc.charts as C
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.7); sec.right_margin=Cm(1.7)
    S=ctx["scenarios3"]; bza=S[0]; full=next((x for x in S if x.get("recommended")), bza); opti=S[-1]; pm=ctx.get("profile_metrics",{})

    # COVER
    _kick(doc, f"Technicko-ekonomický posudok · {ctx.get('posudok_number','')}" + ((" (k "+ctx['pon_number']+")") if ctx.get('pon_number') else ""))
    _h(doc, ctx.get("client_name",""), size=22)
    _para(doc, f"FVE {num(ctx.get('fve_kwp'))} kWp" + (f" + BESS {num(ctx.get('bess_kwh'))} kWh" if ctx.get('bess_kwh') else ""), size=11)
    _para(doc, ctx.get("cover_subtitle",""), size=9, color=GREY, italic=True)

    # EXEC SUMMARY
    _kick(doc,"Manažérske zhrnutie"); _h(doc, ctx.get("summary_headline",""))
    _para(doc, "Odporúčanie: "+ctx.get("recommendation_line",""), size=10.5)
    _table(doc,["Ukazovateľ","Hodnota"],[
        ["Investícia (Net CAPEX)",eur(ctx.get("net_capex_eur") or ctx.get("capex_total_eur"))],
        ["Ročný prínos (rok 1)",eur(full["save_total"])],
        ["Návratnosť",f"{num(full['payback'],1)} r"],
        ["NPV 20 r. / IRR",f"{eur(full['npv'])} / {num(full['irr'],1)} %"],
        ["Pokrytie spotreby / samospotreba",f"{num(ctx.get('coverage_pct'),1)} % / {num(ctx.get('samosp_pct'),1)} %"],
    ],aligns=["l","r"])
    _summary_html = ctx.get("ai_summary_html") or ctx.get("ai_commentary_html") or ""
    if _strip(_summary_html):
        for chunk in re.split(r"</p>|\n\n", _summary_html):
            txt=_strip(chunk)
            if txt: _para(doc,txt,size=10)

    # 1 IDENTIFIKÁCIA
    _kick(doc,"1 — Východiská a identifikácia OM"); _h(doc,"Stav odberného miesta a vstupné dáta")
    idr=[["Klient",ctx.get("client_name","")]]
    for lab,val in [("Adresa OM",ctx.get("om_address")),("EIC OM / č. miesta", (f"{ctx.get('eic_om')} / {ctx.get('cislo_om')}" if (ctx.get('eic_om') or ctx.get('cislo_om')) else None)),
                    ("Distribučná oblasť",ctx.get("distrib_oblast")),("Sadzba / tarif",ctx.get("om_sadzba"))]:
        if val and str(val).strip() not in ("—","None","None / None"): idr.append([lab,val])
    idr+=[["Ročná spotreba",f"{num(ctx.get('year_mwh'))} MWh"],["Max 15-min odber",f"{num(ctx.get('max15_kw'))} kW"],
          ["Priemerný odber / load factor",f"{num(pm.get('avg_kw'))} kW / {num(pm.get('load_factor'),2)}"],
          ["MRK / RK",f"{num(ctx.get('om_mrk_kw'))} / {num(ctx.get('om_rk_kw'))} kW"]]
    _table(doc,["Údaj","Hodnota"],idr,aligns=["l","l"])
    _para(doc,"Charakteristika prevádzky (odvodená z dát): "+ctx.get("profile_sentence",""),size=9.5,italic=True)

    # 2 PROFIL
    doc.add_page_break(); _kick(doc,"2 — Profil odberu"); _h(doc,"Charakteristika spotreby")
    _img(doc,C.chart_daily(ctx),cap="Graf 1: Denný profil odberu — pracovný deň vs víkend, s PV produkciou.")
    _img(doc,C.chart_monthly(ctx),cap="Graf 2: Mesačná spotreba.")

    # 3 TECHNIKA + BILANCIA
    doc.add_page_break(); _kick(doc,"3 — Technické riešenie"); _h(doc, ctx.get("variant_title","Navrhnutý variant"))
    comp=ctx.get("components") or {}
    _table(doc,["Parameter","Hodnota"],[
        ["Inštalovaný výkon FVE",f"{num(ctx.get('fve_kwp'),2)} kWp DC"],
        ["Moduly",comp.get("panel","—")],["Meniče / optimizéry",comp.get("inverter","—")],
        ["Batériové úložisko",comp.get("battery","—") if ctx.get("bess_kwh") else "—"],
        ["Konštrukcia",comp.get("konstrukcia","—")],
        ["Predpokladaná ročná výroba",f"{num(ctx.get('fve_prod_mwh'))} MWh"],
    ],aligns=["l","l"])
    _kick(doc,"Energetická bilancia")
    _table(doc,["Veličina","Hodnota","Podiel"],[
        ["Ročná výroba FVE",f"{num(ctx.get('fve_prod_mwh'))} MWh","100 %"],
        ["Samospotreba",f"{num(ctx.get('self_use_mwh'))} MWh",f"{num(ctx.get('samosp_pct'),1)} %"],
        ["Export prebytkov",f"{num(ctx.get('export_mwh'))} MWh",f"{num(ctx.get('export_pct') or 0,1)} %"],
        ["Import zo siete",f"{num(ctx.get('grid_import_mwh'))} MWh","—"],
        ["Pokrytie spotreby OM",f"{num(ctx.get('coverage_pct'),1)} %","FVE vs spotreba"],
    ],aligns=["l","r","r"])
    _img(doc,C.chart_energy_balance(ctx),cap="Graf 3: Energetická bilancia.")

    # 4 EKONOMIKA
    doc.add_page_break(); _kick(doc,"4 — Ekonomické posúdenie"); _h(doc,"Cenové predpoklady")
    _table(doc,["Zložka ceny","Hodnota","Zdroj"],[
        ["Silová zložka",f"{num(ctx['p_silova']*1000,2)} €/MWh",ctx.get("tarif_source","")],
        ["Variabilná distribúcia",f"{num(ctx['p_dist_var']*1000,2)} €/MWh","VSD tarif"],
        ["TPS + systémové služby",f"{num((ctx['p_tps']+ctx['p_so'])*1000,2)} €/MWh","pásmo 2"],
        ["Avoided cost samospotreby",f"{num(ctx['p_avoided']*1000,1)} €/MWh","kompozit"],
        ["Pevná zložka distribúcie",f"{num(ctx['p_dist_pevna'],2)} €/kW/mes",f"× RK {num(ctx.get('om_rk_kw'))} kW"],
    ],aligns=["l","r","l"])
    _h(doc,"Tri scenáre prínosu",size=12)
    _table(doc,["Scenár","Úspora €/r","Návratnosť","NPV 20 r.","IRR"],
        [[s["name"],eur(s["save_total"]),f"{num(s['payback'],1)} r",eur(s["npv"]),f"{num(s['irr'],1)} %"] for s in S],
        aligns=["l","r","r","r","r"])
    _img(doc,C.chart_scenarios(ctx),cap="Graf 4: Porovnanie 3 scenárov.")
    _img(doc,C.chart_cumcf(ctx),cap="Graf 5: Kumulatívny cashflow 20 rokov.")

    # 5 SKLADBA + RIZIKO
    doc.add_page_break(); _kick(doc,"5 — Skladba prínosu a riziková analýza"); _h(doc,"Skladba ročného prínosu")
    _table(doc,["Zdroj prínosu","Detail","Hodnota"],
        [[n,f,eur(v)] for n,f,v in ctx.get("benefit_rows",[])]+[["Ročná úspora SPOLU","",eur(full["save_total"])]],
        aligns=["l","l","r"])
    _img(doc,C.chart_benefit(ctx),cap="Graf 6: Skladba ročného prínosu.")
    mc_img=C.chart_montecarlo(ctx)
    _img(doc,C.chart_tornado(ctx),cap="Graf 7: Tornado — citlivosť NPV.")
    _img(doc,mc_img,cap=f"Graf 8: Monte Carlo — P10 {eur(ctx['mc_p10'])}, medián {eur(ctx['mc_p50'])}, P90 {eur(ctx['mc_p90'])}; pravdepodobnosť kladného NPV {num(ctx.get('mc_prob_pos',0)*100,0)} %.")

    # 6 CENA NEČINNOSTI
    doc.add_page_break(); _kick(doc,"6 — Cena nečinnosti"); _h(doc,"Koľko stojí ponechať súčasný stav")
    _table(doc,["Scenár cien energie","Náklad za 20 rokov","Rok 1"],[
        ["Konštantné ceny",eur(ctx.get("inaction_flat_20y")),eur(ctx.get("inaction_y1"))],
        ["Rast cien o 2,5 %/r",eur(ctx.get("inaction_infl_20y")),eur(ctx.get("inaction_y1"))],
    ],aligns=["l","r","r"])

    # 7 ODPORÚČANIA
    _kick(doc,"7 — Expert posúdenie a odporúčania")
    _expert_html = ctx.get("ai_expert_html") or ""
    if _strip(_expert_html):
        for chunk in re.split(r"</p>|\n\n", _expert_html):
            txt=_strip(chunk)
            if txt: _para(doc,txt,size=10)
    _h(doc,"Odporúčané kroky")
    for i,(t,d) in enumerate(ctx.get("recommendations",[])):
        p=doc.add_paragraph(); r=p.add_run(f"{i+1:02d}  {t}"); r.bold=True; r.font.size=Pt(10); r.font.name="Arial"
        _para(doc,d,size=9,color=GREY)

    # 8 ZÁVER
    doc.add_page_break(); _kick(doc,"8 — Záver"); _h(doc, ctx.get("zaver_headline",""))
    _para(doc, ctx.get("zaver_text",""), size=10.5)
    _kick(doc,"Záruky a istoty")
    for li in ["Výkonová záruka na panely (25–30 r.) a záruka výrobcu na meniče a batériu.",
               "Realizácia, revízie a uvedenie do prevádzky pod jednou zodpovednosťou Energovision.",
               "Monitoring výroby a úspor po spustení; servis počas celej životnosti.",
               "Skúsenosť s riešeniami pre porovnateľné priemyselné prevádzky."]:
        doc.add_paragraph(li, style="List Bullet")

    # 9 PREDPOKLADY
    _kick(doc,"9 — Predpoklady a zdroje"); _h(doc,"Metodika a východiská")
    _table(doc,["Predpoklad / vstup","Hodnota","Zdroj"],[
        ["Bilančná simulácia","8 760 h",ctx.get("consumption_source","")],
        ["Špecifický výnos FVE",f"{num(ctx.get('yield'))} kWh/kWp","PVGIS"],
        ["Diskont / horizont / odpis","6 % / 20 r / 6 r","WACC / životnosť / zákon"],
        ["Tarify",ctx.get("tarif_source","") if ctx.get("tarif_real") else "orientačné","faktúra / ÚRSO 2026"],
    ],aligns=["l","r","l"])
    _para(doc,"Model má neistotu ±10–15 %. Čísla pochádzajú zo simulácie dispatchu (energetická bilancia validovaná). Posudok je nezáväzný odborný odhad.",size=8.5,color=GREY)
    _para(doc,f"Kontakt: {ctx.get('prepared_by_name','')} · {ctx.get('prepared_by_email','')} · {ctx.get('prepared_by_phone','')}",size=9)

    out=io.BytesIO(); doc.save(out); return out.getvalue()
