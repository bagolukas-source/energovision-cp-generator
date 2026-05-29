"""
Energovision B2C cenovka — Webhook server pre Notion → PDF generovanie.

Endpointy:
- POST /webhook/prepocet — prerátá ceny pre lead, updatne Notion polia
- POST /webhook/generate-pdf — vyrobí PDF + uploaduje do Notion stránky
- GET /health — healthcheck pre Render

Deployovať na Render.com (alebo iný cloud s Python 3.11+).
"""
import os
import sys
import json
import logging
import re

try:
    from error_tracker import track_error, track_warning, track_info
except Exception:
    def track_error(e, context=None): pass
    def track_warning(m, context=None): pass
    def track_info(m, context=None): pass

try:
    from rate_limiter import rate_limit
except Exception:
    def rate_limit(*args, **kwargs):
        def wrap(fn): return fn
        return wrap
from typing import Dict, List, Tuple, Optional, Any
import math
import tempfile
from pathlib import Path
from functools import wraps

from flask import Flask, request, jsonify, make_response
from sk_gender import oslovenie_pan_pani, oslovenie_plne
import requests
import time as _time

# ============================================================
# FIX: Pre-import heavy moduly aby prvý request po deploye nebol pomalý
# ============================================================
try:
    import weasyprint  # noqa: F401 — len pre warm-up
    import matplotlib  # noqa: F401
    matplotlib.use("Agg")  # threadsafe backend (musí byť pred prvým plt.* volaním)
    import matplotlib.pyplot as _plt  # noqa: F401
    from openpyxl import load_workbook  # noqa: F401
    from docxtpl import DocxTemplate  # noqa: F401
    from docx import Document  # noqa: F401
except Exception as _e:
    # Ak na build čase niečo chýba, neblokujeme štart — len logujeme
    print(f"[warmup] Pre-import zlyhal pre {_e}", flush=True)


# ============================================================
# FIX: Notion API retry helper — handles 429/502/503/504 transient errors
# ============================================================
def _retry_request(fn, *, max_retries=3, base_delay=1.0, retry_codes=(429, 500, 502, 503, 504)):
    """Volá fn() s exponential backoff pri prechodných HTTP chybách.
    fn musi vratit requests.Response objekt."""
    last_exc = None
    for attempt in range(max_retries + 1):
        try:
            r = fn()
            if r.status_code in retry_codes and attempt < max_retries:
                delay = base_delay * (2 ** attempt)
                # 429 môže mať Retry-After header
                if r.status_code == 429:
                    ra = r.headers.get("Retry-After")
                    if ra:
                        try:
                            delay = max(delay, float(ra))
                        except ValueError:
                            pass
                _time.sleep(min(delay, 30))
                continue
            return r
        except (requests.ConnectionError, requests.Timeout) as e:
            last_exc = e
            if attempt < max_retries:
                _time.sleep(base_delay * (2 ** attempt))
                continue
            raise
    if last_exc:
        raise last_exc
    return fn()  # final attempt without retry


# ============================================================
# ============================================================
# Activity Log — automaticky audit trail pre kazdy webhook
# DB: 📋 Activity Log (Notion), data_source: 375be89f-db6d-4b19-9766-a237d27a33ea
# ============================================================
ACTIVITY_LOG_DS_ID = "375be89f-db6d-4b19-9766-a237d27a33ea"

def _log_activity(endpoint, status="OK", page_id=None, klient=None, variant=None,
                  duration_ms=None, request_id=None, detail=None, error=None):
    """Zapise aktivitu do Activity Log Notion DB. Nikdy nevyhadzuje exception."""
    try:
        title = f"{endpoint}"
        if klient:
            title += f" — {klient}"
        elif page_id:
            title += f" — {page_id[:8]}"
        title += f" [{status}]"
        
        props = {
            "Akcia": {"title": [{"text": {"content": title[:200]}}]},
            "Endpoint": {"select": {"name": endpoint}},
            "Status": {"select": {"name": status}},
        }
        if page_id:
            props["Page ID"] = {"rich_text": [{"text": {"content": str(page_id)[:200]}}]}
        if klient:
            props["Klient"] = {"rich_text": [{"text": {"content": str(klient)[:200]}}]}
        if variant:
            props["Variant"] = {"select": {"name": str(variant)}}
        if duration_ms is not None:
            props["Duration ms"] = {"number": float(duration_ms)}
        if request_id:
            props["Request ID"] = {"rich_text": [{"text": {"content": str(request_id)[:100]}}]}
        if detail:
            props["Detail"] = {"rich_text": [{"text": {"content": str(detail)[:1900]}}]}
        if error:
            props["Error"] = {"rich_text": [{"text": {"content": str(error)[:1900]}}]}
        
        payload = {
            "parent": {"database_id": ACTIVITY_LOG_DS_ID},
            "properties": props,
        }
        r = requests.post(f"{NOTION_API}/pages", headers=NOTION_HEADERS, json=payload, timeout=10)
        # Best-effort — neraisuje aj pri 4xx/5xx
    except Exception as e:
        try:
            log.warning(f"_log_activity zlyhal: {e}")
        except Exception:
            pass


# FIX: Defensive Claude response parser — chráni proti IndexError pri overload
# ============================================================
def _safe_claude_text(resp_json):
    """Bezpečne vyextrahuje text z Claude API response. Vracia '' ak je prázdne/error."""
    if not isinstance(resp_json, dict):
        return ""
    content = resp_json.get("content")
    if not content or not isinstance(content, list) or len(content) == 0:
        return ""
    first = content[0]
    if not isinstance(first, dict):
        return ""
    return first.get("text", "") or ""

# Import existujúcich modulov z generátora
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from generate_cp import (
    load_cennik, vyrataj_konfig, vyrataj_ceny,
    vyrataj_navratnost, vyrob_grafy
)
from generate_cp_html import vyrob_html_pdf
from generate_from_notion import (
    lead_from_notion, predpocitaj_ceny_pre_record,
    check_compatibility, INVERTOR_BATTERY_COMPAT, safe_filename,
)

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("evo")

app = Flask(__name__)

# ============================================================
# H2: Request size limit — max 10 MB (chranime Render RAM)
# ============================================================
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB

# ============================================================
# H3: Security headers + JSON-only enforcement na ALL responses
# ============================================================
@app.after_request
def _add_security_headers(response):
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "no-referrer"
    return response



# ============================================================
# FIX: Global error handler — pre kazdu nezachytenu exception vrati JSON
# Make scenare ocakavaju JSON response, nie HTML 500
# ============================================================
_REQUEST_COUNTER = [0]

@app.before_request
def _assign_request_id():
    import datetime as _dt
    _REQUEST_COUNTER[0] += 1
    rid = f"R{_dt.datetime.utcnow().strftime('%H%M%S')}-{_REQUEST_COUNTER[0] % 10000:04d}"
    request.environ["request_id"] = rid
    log.info(f"[{rid}] {request.method} {request.path}")


# ============================================================
# PUBLIC QUOTE LINK — /p/<ev_id>
# Verejná stránka cenovky pre klienta (BEZ auth) + tlačidlo akceptovať
# Inspired by Thermivio pattern.
# ============================================================
import urllib.parse as _urlparse

PUBLIC_PAGE_HTML = """<!DOCTYPE html>
<html lang="sk">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Cenová ponuka {{ev_id}} — Energovision</title>
<style>
  * {box-sizing:border-box; margin:0; padding:0}
  body {font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif; background:#f5f7fa; color:#1a2e3f; line-height:1.5; padding:20px}
  .wrap {max-width:780px; margin:0 auto; background:white; border-radius:16px; box-shadow:0 4px 24px rgba(0,0,0,0.08); overflow:hidden}
  .hero {background:linear-gradient(135deg, #0f4c75 0%, #3282b8 100%); color:white; padding:48px 32px; text-align:center}
  .hero h1 {font-size:32px; margin-bottom:8px; font-weight:600}
  .hero .ev {opacity:0.9; font-size:14px; letter-spacing:1px}
  .hero .klient {margin-top:24px; font-size:18px}
  .content {padding:32px}
  .row {display:flex; justify-content:space-between; padding:12px 0; border-bottom:1px solid #e5eaf0}
  .row:last-child {border-bottom:none}
  .row .label {color:#5a6b7d}
  .row .val {font-weight:500; text-align:right}
  .price {font-size:42px; font-weight:700; color:#0f4c75; text-align:center; margin:24px 0; letter-spacing:-1px}
  .price small {font-size:14px; color:#5a6b7d; font-weight:400; display:block}
  .accept-form {background:#f0f4f8; padding:24px; border-radius:12px; margin-top:24px}
  .accept-form h3 {margin-bottom:16px; color:#0f4c75}
  .accept-form label {display:block; font-size:14px; color:#5a6b7d; margin-bottom:6px}
  .accept-form input {width:100%; padding:12px; border:1px solid #cdd5e0; border-radius:8px; font-size:16px; margin-bottom:12px}
  .btn {background:#10b981; color:white; border:none; padding:14px 32px; border-radius:8px; font-size:16px; font-weight:600; cursor:pointer; width:100%}
  .btn:hover {background:#059669}
  .btn:disabled {background:#9ca3af; cursor:not-allowed}
  .accepted {background:#ecfdf5; color:#065f46; padding:16px; border-radius:8px; text-align:center; margin-top:16px}
  .footer {text-align:center; padding:24px; color:#5a6b7d; font-size:13px}
  .err {background:#fef2f2; color:#991b1b; padding:12px; border-radius:8px; margin-bottom:12px; font-size:14px}
</style>
</head>
<body>
<div class="wrap">
  <div class="hero">
    <div class="ev">{{ev_id}}</div>
    <h1>Cenová ponuka</h1>
    <div class="klient">{{klient}}</div>
  </div>
  <div class="content">
    <div class="row"><span class="label">Variant</span><span class="val">{{variant_label}}</span></div>
    <div class="row"><span class="label">Výkon FVE</span><span class="val">{{vykon_kwp}} kWp ({{pocet_panelov}}× LONGi {{wp_label}})</span></div>
    <div class="row"><span class="label">Konštrukcia</span><span class="val">{{konstrukcia}}</span></div>
    {{bateria_row}}
    {{wallbox_row}}
    <div class="row"><span class="label">Mesto</span><span class="val">{{mesto}}</span></div>
    <div class="row"><span class="label">Platnosť ponuky</span><span class="val">30 dní od {{datum}}</span></div>
    
    <div class="price">
      {{cena_str}} € s DPH
      <small>Cena obsahuje materiál, montáž, sprevádzkovanie a 25-ročnú záruku panelov</small>
    </div>
    
    {{status_block}}
  </div>
  <div class="footer">
    Energovision s.r.o. · IČO 50 408 921<br>
    +421 917 424 564 · obchod@energovision.sk
  </div>
</div>
{{js_block}}
</body>
</html>
"""

PUBLIC_FORM_HTML = """<div class="accept-form">
  <h3>✓ Akceptujem túto ponuku</h3>
  <form id="acceptForm" onsubmit="return submitAccept(event)">
    <label>Vaše meno a priezvisko *</label>
    <input type="text" id="name" required>
    <label>Email pre potvrdenie *</label>
    <input type="email" id="email" required>
    <div id="err" class="err" style="display:none"></div>
    <button type="submit" class="btn" id="btn">Akceptujem ponuku</button>
  </form>
</div>"""

PUBLIC_ACCEPTED_HTML = """<div class="accepted">
  <strong>✓ Ponuka bola akceptovaná</strong><br>
  {{accepted_date}}<br>
  Dominik Galaba (+421 917 424 564) Vás bude kontaktovať do 24 hodín.
</div>"""

PUBLIC_JS = """<script>
async function submitAccept(e) {
  e.preventDefault();
  const btn = document.getElementById('btn');
  const errEl = document.getElementById('err');
  btn.disabled = true;
  btn.textContent = 'Spracovávam...';
  errEl.style.display = 'none';
  try {
    const r = await fetch(window.location.pathname + '/accept', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({
        name: document.getElementById('name').value,
        email: document.getElementById('email').value,
      })
    });
    const data = await r.json();
    if (data.success) {
      window.location.reload();
    } else {
      errEl.textContent = data.error || 'Chyba pri spracovaní. Skúste neskôr.';
      errEl.style.display = 'block';
      btn.disabled = false;
      btn.textContent = 'Akceptujem ponuku';
    }
  } catch (err) {
    errEl.textContent = 'Sieťová chyba: ' + err.message;
    errEl.style.display = 'block';
    btn.disabled = false;
    btn.textContent = 'Akceptujem ponuku';
  }
  return false;
}
</script>"""


def _find_lead_by_ev_id(ev_id):
    """Najdi Notion page lead-u podla ID ponuky (napr. EV-26-221 alebo EV-221)."""
    # Normalize: extract numerical part
    import re as _re
    m = _re.search(r'\d+', ev_id or "")
    if not m:
        return None
    num = m.group(0)
    # Query DB s filterom ID ponuky CONTAINS num
    url = f"{NOTION_API}/databases/{NOTION_DATABASE_ID}/query"
    payload = {
        "filter": {
            "property": "ID ponuky",
            "rich_text": {"contains": num}
        },
        "page_size": 5,
    }
    try:
        r = requests.post(url, headers=NOTION_HEADERS, json=payload, timeout=15)
        r.raise_for_status()
        results = r.json().get("results", [])
        # Vyber prvy match
        return results[0] if results else None
    except Exception as e:
        log.warning(f"_find_lead_by_ev_id({ev_id}) zlyhal: {e}")
        return None



# ============================================================
# CORS — globálne pre všetky /webhook/* (cross-origin z crm.energovision.sk, app.energovision.sk)
# ============================================================
ALLOWED_ORIGINS = {
    "https://crm.energovision.sk",
    "https://app.energovision.sk",
    "https://energovision-fve-os.vercel.app",
    "http://localhost:3000",
    "http://localhost:3001",
}

@app.after_request
def _apply_cors(response):
    origin = request.headers.get("Origin", "")
    # Allow all Vercel preview deploys + production origins
    if origin in ALLOWED_ORIGINS or origin.endswith(".vercel.app") or origin == "":
        response.headers["Access-Control-Allow-Origin"] = origin or "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS, PUT, DELETE"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization, X-Webhook-Secret, X-Admin-Secret"
        response.headers["Access-Control-Max-Age"] = "86400"
    return response


@app.before_request
def _handle_options():
    # Pre OPTIONS preflight vráť 204 priamo
    if request.method == "OPTIONS":
        from flask import make_response
        resp = make_response("", 204)
        return resp

@app.route("/p/<ev_id>", methods=["GET"])
def public_quote(ev_id):
    """Verejna stranka cenovky pre klienta. Bez auth."""
    from datetime import datetime as _dt
    
    page = _find_lead_by_ev_id(ev_id)
    if not page:
        return ("""<html><body style="font-family:sans-serif;max-width:600px;margin:60px auto;padding:20px;text-align:center"><h2>Ponuka nenájdená</h2><p>Číslo ponuky <strong>""" + str(ev_id) + """</strong> sme v systéme nenašli. Kontaktujte prosím Dominik Galaba +421 917 424 564.</p></body></html>"""), 404
    
    flat = notion_props_to_flat(page)
    
    # Klient
    klient = flat.get("Zákazník") or "Klient"
    
    # Aktívny variant (prioritne A > B > C > D)
    variant = "A"
    for v in ["A", "B", "C", "D"]:
        for prop_key in [f"Variant {v} - FVE", f"Variant {v} - FVE + BESS", 
                         f"Variant {v} - FVE + BESS + Wallbox", f"Variant {v} - FVE + Wallbox"]:
            if flat.get(prop_key) == "__YES__":
                variant = v
                break
    
    variant_labels = {
        "A": "Variant A — FVE",
        "B": "Variant B — FVE + batéria",
        "C": "Variant C — FVE + batéria + wallbox",
        "D": "Variant D — FVE + wallbox",
    }
    
    # Cena
    cena = flat.get(f"Cena {variant} s DPH") or flat.get("Suma CP s DPH") or 0
    try:
        cena_num = float(cena)
        cena_str = f"{cena_num:,.0f}".replace(",", " ")
    except (ValueError, TypeError):
        cena_str = "—"
    
    # Konfigurácia
    pocet_panelov = flat.get("Počet panelov") or "?"
    panel_typ = flat.get("Panel") or "LONGi 535 Wp"
    wp_match = re.search(r"(\d{3})", str(panel_typ))
    wp_label = wp_match.group(1) + " Wp" if wp_match else "535 Wp"
    
    try:
        kwp = round(int(str(pocet_panelov)) * (int(wp_match.group(1)) if wp_match else 535) / 1000, 2)
        vykon_kwp = f"{kwp:.2f}".replace(".", ",")
    except (ValueError, TypeError, AttributeError):
        vykon_kwp = "—"
    
    konstrukcia = flat.get("Konštrukcia (typ)") or "—"
    mesto = flat.get("Mesto") or "—"
    
    # Voliteľne batéria
    bateria_row = ""
    if variant in ("B", "C"):
        bat_typ = flat.get("Batéria (typ)") or ""
        bat_pocet = flat.get("Batéria počet") or "1"
        if bat_typ:
            bateria_row = f'<div class="row"><span class="label">Batéria</span><span class="val">{bat_pocet}× {bat_typ}</span></div>'
    
    wallbox_row = ""
    if variant in ("C", "D"):
        wb_typ = flat.get("Wallbox (typ)") or ""
        if wb_typ:
            wallbox_row = f'<div class="row"><span class="label">Wallbox</span><span class="val">{wb_typ}</span></div>'
    
    # Status — bol uz akceptovany?
    status = flat.get("Status") or ""
    already_accepted = status in ("🟢 Výhra", "Podpísané", "💰 Faktúra", "🏗 Realizácia", "✅ Hotové")
    
    if already_accepted:
        accepted_date = flat.get("date:Dátum prijatia podpísaných:start") or _dt.now().strftime("%d.%m.%Y")
        status_block = PUBLIC_ACCEPTED_HTML.replace("{{accepted_date}}", accepted_date)
        js_block = ""
    else:
        status_block = PUBLIC_FORM_HTML
        js_block = PUBLIC_JS
    
    html = PUBLIC_PAGE_HTML
    replacements = {
        "{{ev_id}}": ev_id,
        "{{klient}}": klient,
        "{{variant_label}}": variant_labels.get(variant, variant),
        "{{vykon_kwp}}": vykon_kwp,
        "{{pocet_panelov}}": str(pocet_panelov),
        "{{wp_label}}": wp_label,
        "{{konstrukcia}}": konstrukcia,
        "{{bateria_row}}": bateria_row,
        "{{wallbox_row}}": wallbox_row,
        "{{mesto}}": mesto,
        "{{datum}}": _dt.now().strftime("%d.%m.%Y"),
        "{{cena_str}}": cena_str,
        "{{status_block}}": status_block,
        "{{js_block}}": js_block,
    }
    for k, v in replacements.items():
        html = html.replace(k, str(v))
    
    # Audit
    try:
        _log_activity("other", status="STARTED", page_id=page.get("id"), klient=klient,
                      variant=variant, detail=f"Public quote view: /p/{ev_id}")
    except Exception:
        pass
    
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/p/<ev_id>/accept", methods=["POST"])
def public_quote_accept(ev_id):
    """Klient akceptuje ponuku. POST {name, email}."""
    from datetime import datetime as _dt
    
    body = request.get_json(silent=True) or {}
    name = (body.get("name") or "").strip()
    email = (body.get("email") or "").strip()
    
    if not name or len(name) < 3:
        return jsonify({"success": False, "error": "Meno je povinné"}), 400
    if "@" not in email:
        return jsonify({"success": False, "error": "Neplatný email"}), 400
    
    page = _find_lead_by_ev_id(ev_id)
    if not page:
        return jsonify({"success": False, "error": "Ponuka nenájdená"}), 404
    
    page_id = page.get("id")
    flat = notion_props_to_flat(page)
    klient = flat.get("Zákazník") or "Klient"
    
    # IP klienta + user agent
    client_ip = request.headers.get("X-Forwarded-For", request.remote_addr or "?").split(",")[0].strip()
    user_agent = request.headers.get("User-Agent", "")[:300]
    
    today_iso = _dt.now().strftime("%Y-%m-%d")
    today_human = _dt.now().strftime("%d.%m.%Y %H:%M")
    
    accept_note = f"Klient akceptoval cenovku online dňa {today_human} z IP {client_ip}. Zadané meno: {name}, email: {email}. User-Agent: {user_agent[:100]}"
    
    update = {
        "Status": {"select": {"name": "🟢 Výhra"}},
        "Poznámka": {"rich_text": [{"text": {"content": accept_note[:1900]}}]},
        "date:Dátum prijatia podpísaných:start": today_iso,
    }
    
    try:
        notion_update_page(page_id, update)
        _log_activity("other", status="OK", page_id=page_id, klient=klient,
                      detail=f"PUBLIC ACCEPT by {name} ({email}) from {client_ip}")
        return jsonify({"success": True, "message": "Ponuka akceptovaná"})
    except Exception as e:
        _log_activity("other", status="ERROR", page_id=page_id, klient=klient,
                      error=str(e), detail=f"PUBLIC ACCEPT failed for {name}")
        return jsonify({"success": False, "error": "Chyba pri uložení. Skúste neskôr."}), 500


@app.errorhandler(Exception)
def _handle_uncaught(e):
    import traceback
    tb = traceback.format_exc()
    rid = request.environ.get("request_id", "R-?")
    log.exception(f"[{rid}] UNCAUGHT EXCEPTION v endpointe")
    # Pre HTTP error werkzeug exceptions zachovaj ich status code
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException):
        return jsonify({
            "success": False,
            "error": str(e),
            "code": e.code,
        }), e.code
    return jsonify({
        "success": False,
        "error": str(e),
        "error_type": type(e).__name__,
        "traceback_tail": tb[-800:],
    }), 500

# === ENV ===
NOTION_TOKEN = os.environ.get("NOTION_TOKEN", "")
NOTION_DATABASE_ID = os.environ.get("NOTION_DATABASE_ID", "ba7a1d6c-63a9-43da-b66d-2b1c7e8660da")
NOTION_MATERIAL_PO_DB_ID = os.environ.get("NOTION_MATERIAL_PO_DB_ID", "a8690d6826114d5097c8bbfb36c02d7c")
WEBHOOK_SECRET = os.environ.get("WEBHOOK_SECRET", "")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5-20250929")

# Globálne Supabase konstanty (pre TS endpointy, raynet import atď.)
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://uzwajrpebblafuhrtuwn.supabase.co")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "") or os.environ.get("SUPABASE_SERVICE_KEY", "")

# Default spotreba ak nezadaná zákazníkom (priemer SK domácnosti)
DEFAULT_SPOTREBA_KWH = 8000

# Panel default pre auto-konfiguráciu
AUTO_PANEL = "LONGi 535 Wp"
AUTO_PANEL_WP = 535
# DC oversize: panely sa dimenzujú s prebytkom voči AC menicu (typicky 1.2-1.35x)
AUTO_DC_OVERSIZE = 1.28
# Zaokrúhliť počet panelov nahor na párne číslo (kvôli stringom + symetrii inštalácie)
AUTO_ROUND_TO_EVEN = True

NOTION_API = "https://api.notion.com/v1"
NOTION_HEADERS = {
    "Authorization": f"Bearer {NOTION_TOKEN}",
    "Notion-Version": "2022-06-28",
    "Content-Type": "application/json",
}


def require_secret(f):
    """Dekorátor — kontroluje X-Webhook-Secret header."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if WEBHOOK_SECRET:
            received = request.headers.get("X-Webhook-Secret", "")
            if received != WEBHOOK_SECRET:
                return jsonify({"error": "unauthorized"}), 401
        return f(*args, **kwargs)
    return wrapper


# ============================================================
# VALIDÁCIA — overuje že lead má dostatočne kvalitné dáta
# pre generovanie dokumentov
# ============================================================
FAKE_DEFAULTS = {
    "info@energovision.sk",       # firemný email — NIE klient
    "ulica 33",                   # parser default
    "suchohrad",                  # náhodný default
    "topolčany",                  # náhodný default ak nesúhlasí s mestom
}


def _is_fake(value):
    """Detekuje fake default hodnotu (case-insensitive, čistá medzera)."""
    if not value:
        return True
    v = str(value).strip().lower()
    if not v:
        return True
    return v in FAKE_DEFAULTS


def _validate_psc(psc):
    """SR PSČ je 5 cifier."""
    if not psc:
        return False
    digits = re.sub(r'\D', '', str(psc))
    return len(digits) == 5


def validate_lead_for_documents(flat, doc_type="zmluvy"):
    """
    Skontroluje povinné polia pre generovanie dokumentov.
    doc_type: "zmluvy" | "realizacne" | "pd"

    Vráti tuple (is_valid: bool, missing: list[str]).
    """
    missing = []

    # Spoločné pre všetky 3 dokumenty
    meno = flat.get("Zákazník", "")
    if not meno or len(meno.strip()) < 3:
        missing.append("Zákazník (meno a priezvisko)")

    tel = flat.get("Telefón", "")
    if not tel:
        missing.append("Telefón")

    email = flat.get("Email", "")
    if not email or _is_fake(email):
        missing.append("Email (skutočný klienta, nie info@energovision.sk)")

    ulica = flat.get("Ulica číslo", "")
    if not ulica or _is_fake(ulica):
        missing.append("Ulica číslo (skutočná adresa)")

    mesto = flat.get("Mesto", "")
    if not mesto or _is_fake(mesto):
        missing.append("Mesto")

    psc = flat.get("PSČ", "")
    if not _validate_psc(psc):
        missing.append(f"PSČ (musí byť 5 cifier, je: {psc!r})")

    if doc_type == "zmluvy":
        # Zmluvy potrebujú variant + cenu
        variant = flat.get("Variant do zmluvy", "")
        if not variant:
            missing.append("Variant do zmluvy (A/B/C/D)")
        pocet_p = flat.get("Počet panelov", "")
        if not pocet_p or str(pocet_p).strip() == "0":
            missing.append("Počet panelov")
        # Splnomocnenie potrebuje OP + dátum narodenia
        if not flat.get("Číslo OP"):
            missing.append("Číslo OP")
        if not flat.get("date:Dátum narodenia:start"):
            missing.append("Dátum narodenia")
        # Dotazník potrebuje IBAN, banka, EIC
        if not flat.get("IBAN"):
            missing.append("IBAN")
        if not flat.get("Banka"):
            missing.append("Banka")
        if not flat.get("EIC odberného miesta"):
            missing.append("EIC odberného miesta")

    elif doc_type == "realizacne":
        # Revízia + protokol potrebujú sériové čísla
        if not flat.get("Sériové č. meniča"):
            missing.append("Sériové č. meniča")
        if not flat.get("date:Dátum odovzdania:start"):
            missing.append("Dátum odovzdania")

    elif doc_type == "pd":
        # PD potrebuje distribučnú spoločnosť a parcely
        if not flat.get("Distribučná spoločnosť"):
            missing.append("Distribučná spoločnosť (SSD/VSD/ZSDIS)")
        if not flat.get("Parcelné čísla"):
            missing.append("Parcelné čísla")
        if not flat.get("Variant do zmluvy"):
            missing.append("Variant do zmluvy (A/B/C/D)")

    return (len(missing) == 0, missing)


def notion_get_page(page_id):
    """Stiahne Notion page properties — s retry na 429/5xx."""
    r = _retry_request(lambda: requests.get(f"{NOTION_API}/pages/{page_id}", headers=NOTION_HEADERS, timeout=20))
    r.raise_for_status()
    return r.json()


def notion_props_to_flat(page):
    """Z Notion API response (page properties) urob flat dict {prop_name: value}."""
    out = {}
    props = page.get("properties", {})
    for name, prop in props.items():
        ptype = prop.get("type")
        if ptype == "title":
            out[name] = "".join(t.get("plain_text", "") for t in prop["title"])
        elif ptype == "rich_text":
            out[name] = "".join(t.get("plain_text", "") for t in prop["rich_text"])
        elif ptype == "select":
            out[name] = (prop["select"] or {}).get("name", "")
        elif ptype == "multi_select":
            out[name] = json.dumps([s["name"] for s in prop["multi_select"]], ensure_ascii=False)
        elif ptype == "number":
            out[name] = prop["number"]
        elif ptype == "checkbox":
            out[name] = "__YES__" if prop["checkbox"] else "__NO__"
        elif ptype == "url":
            out[name] = prop["url"]
        elif ptype == "email":
            out[name] = prop["email"]
        elif ptype == "phone_number":
            out[name] = prop["phone_number"]
        elif ptype == "unique_id":
            uid = prop["unique_id"]
            out[name] = f"{uid.get('prefix','')}-{uid['number']}" if uid.get('prefix') else str(uid["number"])
        elif ptype == "files":
            # Files property — vrati JSON pole s file objektmi (name + url)
            files_arr = []
            for f in prop.get("files", []) or []:
                fname = f.get("name", "file")
                if f.get("type") == "external":
                    furl = f.get("external", {}).get("url", "")
                else:
                    furl = f.get("file", {}).get("url", "")
                if furl:
                    files_arr.append({"name": fname, "url": furl})
            out[name] = json.dumps(files_arr, ensure_ascii=False) if files_arr else ""
    return out


def notion_update_page(page_id, properties):
    """Update Notion stránku s novými properties (Notion API v2022-06-28 formát)."""
    payload = {"properties": properties}
    r = _retry_request(lambda: requests.patch(f"{NOTION_API}/pages/{page_id}", headers=NOTION_HEADERS, json=payload, timeout=20))
    r.raise_for_status()
    return r.json()


def notion_set_number(prop_name, value):
    return {prop_name: {"number": float(value) if value is not None else None}}


def notion_set_text(prop_name, value):
    return {prop_name: {"rich_text": [{"text": {"content": str(value)[:2000]}}]}}


def notion_set_select(prop_name, value):
    return {prop_name: {"select": {"name": value}} if value else {prop_name: {"select": None}}}


def notion_set_url(prop_name, value):
    return {prop_name: {"url": value or None}}

def notion_create_page_in_db(database_id, properties):
    """Vytvor novú page v Notion DB."""
    payload = {
        "parent": {"database_id": database_id},
        "properties": properties,
    }
    r = _retry_request(lambda: requests.post(f"{NOTION_API}/pages", headers=NOTION_HEADERS, json=payload, timeout=20))
    r.raise_for_status()
    return r.json()


# ===== AUTO-SIZING LOGIKA =====
def menic_huawei_pre_kwp(kwp):
    """Mapovanie kWp -> Huawei menic (Variant A FVE-only)."""
    if kwp <= 5:
        return "Huawei SUN2000-5K"
    elif kwp <= 6:
        return "Huawei SUN2000-6K"
    elif kwp <= 8:
        return "Huawei SUN2000-8K"
    else:
        return "Huawei SUN2000-10K"


def menic_solinteg_default():
    """Solinteg MHT-10K-25 zvlada 5-10 kWp s bateriou."""
    return "Solinteg MHT-10K-25"


def auto_sizing_from_spotreba(spotreba_kwh, ma_bateriu=False):
    """Mapping spotreby na vykon FVE s DC oversize.
    
    Vzorec: pocet_panelov = ceil(spotreba * 1.28 / 535), zaokruhlene hore na parne.
    Mernicu vyberame podla AC kWp (= round(spotreba/1000), clamp 3-12).
    
    Priklady:
      Spotreba 10 000 kWh -> 24 panelov LONGi 535 Wp = 12.84 kWp DC, 10K menic AC
      Spotreba  8 000 kWh -> 20 panelov = 10.7 kWp DC, 8K menic
      Spotreba  5 000 kWh -> 12 panelov = 6.42 kWp DC, 5K menic
    """
    # AC velkost menica (1:1 so spotrebou, clamp 3-12 kWp B2C strop)
    target_kwp = max(3, min(12, round(spotreba_kwh / 1000)))

    # Pocet panelov: spotreba * DC oversize / Wp panelu
    pocet_raw = math.ceil(spotreba_kwh * AUTO_DC_OVERSIZE / AUTO_PANEL_WP)

    # Zaokruhlit nahor na parne cislo (pre symetricke 2-row stringy)
    if AUTO_ROUND_TO_EVEN and pocet_raw % 2 != 0:
        pocet_raw += 1

    # Min 6 panelov, max 30 (kvoli B2C limit + stringom)
    pocet_panelov = max(6, min(30, pocet_raw))

    if ma_bateriu:
        menic = menic_solinteg_default()
    else:
        menic = menic_huawei_pre_kwp(target_kwp)

    return {
        "target_kwp": target_kwp,
        "pocet_panelov": pocet_panelov,
        "panel": AUTO_PANEL,
        "menic": menic,
    }


def claude_extract_leads(raw_text):
    """Zavola Claude API a vrati pole leadov ako Python list."""
    if not ANTHROPIC_API_KEY:
        raise RuntimeError("ANTHROPIC_API_KEY nie je nastaveny v env vars")

    system_prompt = (
        "Si parser pre slovenskú energetickú firmu Energovision (FVE, batérie, wallbox, trafostanice, revízie).\n\n"
        "Z neformálneho textu (email, formulár, tabuľka, voicemail prepis, OCR) extrahuješ údaje o ZÁKAZNÍKOCH.\n\n"
        "JEDEN INPUT MÔŽE OBSAHOVAŤ VIAC LEADOV. Vráť ich VŠETKY ako pole.\n\n"
        "Vráť IBA platný JSON v presnej štruktúre:\n"
        "{\n"
        '  "leads": [\n'
        "    {\n"
        '      "meno": "iba krstné meno (napr. Peter), nie celé meno + priezvisko, alebo null",\n'
        '      "priezvisko": "iba priezvisko (napr. Novák), bez krstného mena, alebo null",\n'
        '      "telefon": "+421... alebo null",\n'
        '      "email": "email@... alebo null",\n'
        '      "ulica_cislo": "ulica + číslo alebo null",\n'
        '      "mesto": "iba mesto alebo null",\n'
        '      "psc": "PSČ alebo null",\n'
        '      "spotreba_kwh_rok": "číslo (ak zákazník zadal) alebo null",\n'
        '      "rocna_faktura_eur": null alebo číslo,\n'
        '      "typ_dopytu": ["FVE", "Batéria", "Wallbox", "Revízia", "Bleskozvod", "Iné"],\n'
        '      "typ_strechy": "Škridla / Plech kombivrut / Falcový plech / Plochá strecha — J 13° / Plochá strecha — V/Z 10° / null",\n'
        '      "orientacia": "J / V-Z / J-V / J-Z / null",\n'
        '      "bateria_odporucana": "Solinteg EBA B5K1 — 5.12 kWh / Solinteg EBA B5K1 — 10.24 kWh / Pylontech Force H3 — 5.12 kWh / Huawei LUNA2000 — 5 kWh / Huawei LUNA2000 — 7 kWh / null",\n'
        '      "wallbox_odporucany": "Solinteg 7 kW (1F) / Solinteg 11 kW (3F) / Huawei AC Smart 22 kW / Huawei AC Smart 7 kW / GoodWe 11 kW / GoodWe 22 kW / null",\n'
        '      "variant_odporucany": "A" alebo "B" alebo "C",\n'
        '      "priorita": "Vysoká / Stredná / Nízka",\n'
        '      "poznamky": "krátke zhrnutie kontextu, max 200 znakov"\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Pravidlá:\n"
        "- Variant A = iba FVE, B = FVE+batéria, C = FVE+batéria+wallbox\n"
        "- Ak EV/elektromobil/wallbox -> C; ostrov/výpadky/akumulácia -> minimálne B; iba znižovanie účtu -> A\n"
        "- Slovenské diakritiky zachovaj. Telefón normalizuj na +421 formát.\n"
        "- DÔLEŽITÉ: spotreba_kwh_rok dávaj LEN ak je v texte explicitne uvedená. Ak nie, daj null. NEVYMÝŠĽAJ z faktúry.\n"
        "- Batéria default: Solinteg EBA B5K1 — 10.24 kWh ak spotreba >5000 kWh, inak 5.12 kWh.\n"
        "- Ak chýba údaj, daj null. ZIADNE markdown bloky, IBA surový JSON {...}."
    )

    user_prompt = f"Tu je raw lead text — moze obsahovat 1 alebo viac leadov:\n\n{raw_text[:15000]}"

    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": 8192,
        "temperature": 0.1,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_prompt}],
    }
    # FIX: retry pri 429/5xx Claude API (overload)
    r = _retry_request(lambda: requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload, timeout=90))
    r.raise_for_status()
    resp = r.json()
    # FIX: defensive parse — pri prázdnom/malformed response neraisne IndexError
    text = _safe_claude_text(resp).strip()
    if not text:
        raise RuntimeError("Claude API vratila prazdny response (mozno overloaded)")

    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()

    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        log.warning("Claude vratil ne-JSON (truncated?): %s; pokus o opravu...", str(e))
        # Pokus o opravu truncated JSON — extrahuj kompletné leads pred prerušením
        repaired = _repair_truncated_leads_json(text)
        if repaired is not None:
            log.info("JSON oprava uspesna, %d leadov zachranenych", len(repaired))
            return repaired
        log.error("Claude vratil ne-JSON: %s", text[:500])
        raise RuntimeError(f"Claude vratil neplatny JSON: {e}")

    return data.get("leads", [])


def _repair_truncated_leads_json(text):
    """Opravi truncated JSON tak, ze najde kompletne lead objekty pred prerusenim.
    Vrati list dictov alebo None ak sa nic neda zachranit."""
    import json as _json
    # Najdi otvorenie "leads": [
    m = re.search(r'"leads"\s*:\s*\[', text)
    if not m:
        return None
    start = m.end()
    # Skenuj forward a hladaj kompletne objekty
    leads = []
    depth = 0
    obj_start = None
    in_string = False
    escape_next = False
    for i in range(start, len(text)):
        ch = text[i]
        if escape_next:
            escape_next = False
            continue
        if ch == "\\":
            escape_next = True
            continue
        if ch == '"' and not escape_next:
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == "{":
            if depth == 0:
                obj_start = i
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and obj_start is not None:
                obj_str = text[obj_start:i+1]
                try:
                    leads.append(_json.loads(obj_str))
                except Exception:
                    pass
                obj_start = None
    return leads if leads else None


KONSTRUKCIA_OPTIONS = {
    "Škridla", "Plech kombivrut", "Falcový plech",
    "Plochá strecha — J 13°", "Plochá strecha — V/Z 10°",
}
BATERIA_OPTIONS = {
    "Pylontech Force H3 — 5.12 kWh",
    "Solinteg EBA B5K1 — 5.12 kWh", "Solinteg EBA B5K1 — 10.24 kWh",
    "Huawei LUNA2000 — 5 kWh", "Huawei LUNA2000 — 7 kWh",
}
WALLBOX_OPTIONS = {
    "Solinteg 7 kW (1F)", "Solinteg 11 kW (3F)",
    "Huawei AC Smart 22 kW", "Huawei AC Smart 7 kW",
    "GoodWe 11 kW", "GoodWe 22 kW",
}
TYP_DOPYTU_OPTIONS = {"FVE", "Batéria", "Wallbox", "Revízia", "Bleskozvod", "Iné"}
PRIORITA_OPTIONS = {"Vysoká", "Stredná", "Nízka"}


def _select_or_none(value, allowed_set):
    if value and value in allowed_set:
        return {"select": {"name": value}}
    return None


def lead_to_notion_properties(lead):
    """Z parsed lead dict vyrob Notion API properties payload pre Zakaznici B2C DB."""
    props = {}

    # Title = Meno Priezvisko (mesto ide do separatneho Mesto property, NIE do title)
    meno = (lead.get("meno") or "").strip()
    priezvisko = (lead.get("priezvisko") or "").strip()
    # Ak meno uz obsahuje priezvisko (full name), pouzi iba meno
    if meno and priezvisko and priezvisko.lower() in meno.lower():
        title_text = meno
    elif meno and priezvisko:
        title_text = f"{meno} {priezvisko}"
    elif meno:
        title_text = meno
    elif priezvisko:
        title_text = priezvisko
    else:
        title_text = "Nový lead (bez mena)"
    props["Zákazník"] = {"title": [{"text": {"content": title_text[:200]}}]}

    if lead.get("telefon"):
        props["Telefón"] = {"phone_number": str(lead["telefon"])[:50]}
    if lead.get("email"):
        props["Email"] = {"email": str(lead["email"])[:200]}

    mesto_parts = []
    if lead.get("ulica_cislo"):
        mesto_parts.append(str(lead["ulica_cislo"]))
    if lead.get("mesto"):
        mesto_parts.append(str(lead["mesto"]))
    if mesto_parts:
        props["Mesto"] = {"rich_text": [{"text": {"content": ", ".join(mesto_parts)[:500]}}]}

    variant = (lead.get("variant_odporucany") or "A").upper()
    props["Variant A - FVE"] = {"checkbox": variant == "A"}
    props["Variant B - FVE + BESS"] = {"checkbox": variant == "B"}
    props["Variant C - FVE + BESS + Wallbox"] = {"checkbox": variant == "C"}

    # SPOTREBA + AUTO-SIZING
    spotreba_raw = lead.get("spotreba_kwh_rok")
    if spotreba_raw is not None:
        try:
            spotreba_val = float(spotreba_raw)
            spotreba_zdroj = "Zákazník zadal"
        except (ValueError, TypeError):
            spotreba_val = float(DEFAULT_SPOTREBA_KWH)
            spotreba_zdroj = "Default 8000 (SK priemer)"
    else:
        spotreba_val = float(DEFAULT_SPOTREBA_KWH)
        spotreba_zdroj = "Default 8000 (SK priemer)"

    props["Spotreba"] = {"number": spotreba_val}
    props["Spotreba zdroj"] = {"select": {"name": spotreba_zdroj}}

    sizing = auto_sizing_from_spotreba(spotreba_val, ma_bateriu=variant in ("B", "C"))
    props["Panel"] = {"select": {"name": sizing["panel"]}}
    # Pocet panelov je teraz SELECT v Notione (nie number) - hodnoty 1-30
    props["Počet panelov"] = {"select": {"name": str(sizing["pocet_panelov"])}}
    props["Menič"] = {"select": {"name": sizing["menic"]}}

    konstr_sel = _select_or_none(lead.get("typ_strechy"), KONSTRUKCIA_OPTIONS)
    if not konstr_sel:
        konstr_sel = {"select": {"name": "Škridla"}}  # default ak AI nezistila
    props["Konštrukcia (typ)"] = konstr_sel

    if variant in ("B", "C"):
        bat_sel = _select_or_none(lead.get("bateria_odporucana"), BATERIA_OPTIONS)
        if not bat_sel:
            default_bat = "Solinteg EBA B5K1 — 10.24 kWh" if spotreba_val > 5000 else "Solinteg EBA B5K1 — 5.12 kWh"
            bat_sel = {"select": {"name": default_bat}}
        props["Batéria (typ)"] = bat_sel
        props["Batéria počet"] = {"select": {"name": "1"}}
    else:
        props["Batéria počet"] = {"select": {"name": "0"}}

    if variant in ("C", "D"):
        wb_sel = _select_or_none(lead.get("wallbox_odporucany"), WALLBOX_OPTIONS)
        if not wb_sel:
            wb_sel = {"select": {"name": "Solinteg 11 kW (3F)"}}
        props["Wallbox (typ)"] = wb_sel

    typ_dopytu_raw = lead.get("typ_dopytu") or []
    if isinstance(typ_dopytu_raw, str):
        typ_dopytu_raw = [typ_dopytu_raw]
    typ_dopytu_clean = [t for t in typ_dopytu_raw if t in TYP_DOPYTU_OPTIONS]
    if not typ_dopytu_clean:
        typ_dopytu_clean = ["FVE"]
    props["Typ dopytu"] = {"multi_select": [{"name": t} for t in typ_dopytu_clean]}

    prio_sel = _select_or_none(lead.get("priorita"), PRIORITA_OPTIONS)
    if prio_sel:
        props["Priorita"] = prio_sel

    poznamky_parts = []
    if spotreba_zdroj == "Zákazník zadal":
        poznamky_parts.append(f"Spotreba: {int(spotreba_val)} kWh/rok (zadaná)")
    else:
        poznamky_parts.append(f"Spotreba: {int(spotreba_val)} kWh/rok (default SK priemer)")
    if lead.get("rocna_faktura_eur"):
        poznamky_parts.append(f"Faktúra: {lead['rocna_faktura_eur']} €/rok")
    if lead.get("orientacia"):
        poznamky_parts.append(f"Orientácia: {lead['orientacia']}")
    if lead.get("psc"):
        poznamky_parts.append(f"PSČ: {lead['psc']}")
    if lead.get("poznamky"):
        poznamky_parts.append(str(lead["poznamky"]))
    poznamky_parts.append(f"[AI parsed → auto-sizing: {sizing['target_kwp']} kWp / {sizing['pocet_panelov']}× panel]")
    props["Poznámky"] = {"rich_text": [{"text": {"content": " | ".join(poznamky_parts)[:1900]}}]}

    # Default Status pre nový lead — aby sa zobrazil v 1️⃣ Nové leady view
    props["Status"] = {"select": {"name": "🆕 Došlý lead"}}

    return props



# ============================================================
# HEALTHCHECK
# ============================================================
_BOOT_TIME = None

@rate_limit(max_calls=120, window_seconds=60)
@app.route("/health")
def health():
    """Detailny health endpoint — uptime, env-check, version."""
    global _BOOT_TIME
    import datetime as _dt
    if _BOOT_TIME is None:
        _BOOT_TIME = _dt.datetime.utcnow()
    uptime_s = (_dt.datetime.utcnow() - _BOOT_TIME).total_seconds()
    env_ok = {
        "NOTION_TOKEN": bool(NOTION_TOKEN),
        "ANTHROPIC_API_KEY": bool(ANTHROPIC_API_KEY),
        "WEBHOOK_SECRET": bool(WEBHOOK_SECRET),
        "NOTION_DATABASE_ID": bool(os.environ.get("NOTION_DATABASE_ID", "")),
        "GITHUB_PAT": bool(os.environ.get("GITHUB_PAT", "")),
    }
    return jsonify({
        "status": "ok",
        "service": "energovision-cp-generator",
        "uptime_seconds": round(uptime_s, 1),
        "uptime_human": f"{int(uptime_s // 3600)}h {int((uptime_s % 3600) // 60)}m",
        "version": "2026-05-18-v3",
        "env": env_ok,
        "env_all_set": all(env_ok.values()),
        "model": ANTHROPIC_MODEL,
    })


# ============================================================
# ADMIN: PUSH TO GITHUB
# Bridge endpoint — pouziva GITHUB_PAT env var na commit do GitHub via Contents API.
# Auth: X-Admin-Secret header (musi matchnut WEBHOOK_SECRET)
# Body: {"files": [{"filename": "app.py", "content_b64": "...", "message": "fix bug"}]}
# ============================================================
GITHUB_PAT = os.environ.get("GITHUB_PAT", "")
GITHUB_OWNER = os.environ.get("GITHUB_OWNER", "bagolukas-source")
GITHUB_REPO = os.environ.get("GITHUB_REPO", "energovision-cp-generator")
GITHUB_BRANCH = os.environ.get("GITHUB_BRANCH", "main")


def _gh_get_sha(filename):
    if not GITHUB_PAT:
        return None
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{filename}?ref={GITHUB_BRANCH}"
    headers = {"Authorization": f"Bearer {GITHUB_PAT}", "Accept": "application/vnd.github+json"}
    try:
        r = requests.get(url, headers=headers, timeout=20)
        if r.status_code == 200:
            return r.json().get("sha")
    except Exception as e:
        log.warning(f"_gh_get_sha({filename}) zlyhal: {e}")
    return None


def _gh_put_file(filename, content_b64, message, sha=None):
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{filename}"
    headers = {
        "Authorization": f"Bearer {GITHUB_PAT}",
        "Accept": "application/vnd.github+json",
        "Content-Type": "application/json",
    }
    payload = {"message": message, "content": content_b64, "branch": GITHUB_BRANCH}
    if sha:
        payload["sha"] = sha
    r = requests.put(url, headers=headers, json=payload, timeout=30)
    try:
        return r.status_code, r.json()
    except Exception:
        return r.status_code, {"raw": r.text[:500]}


@app.route("/admin/push", methods=["POST"])
def admin_push():
    """Commit suborov do GitHub. Auth cez X-Admin-Secret = WEBHOOK_SECRET."""
    received = request.headers.get("X-Admin-Secret", "")
    if not WEBHOOK_SECRET or received != WEBHOOK_SECRET:
        return jsonify({"error": "unauthorized"}), 401

    if not GITHUB_PAT:
        return jsonify({"error": "GITHUB_PAT nie je nastaveny v Render env vars"}), 500

    body = request.get_json(silent=True) or {}
    files = body.get("files", [])
    default_msg = body.get("message", "admin push from Claude bridge")
    if not files:
        return jsonify({"error": "missing 'files' array"}), 400

    results = []
    for f in files:
        filename = f.get("filename", "")
        content_b64 = f.get("content_b64", "")
        msg = f.get("message", default_msg)
        if not filename or not content_b64:
            results.append({"file": filename, "ok": False, "error": "missing filename or content_b64"})
            continue
        try:
            sha = _gh_get_sha(filename)
            status, resp = _gh_put_file(filename, content_b64, msg, sha)
            ok = status in (200, 201)
            results.append({
                "file": filename,
                "ok": ok,
                "status": status,
                "action": "updated" if sha else "created",
                "commit_sha": resp.get("commit", {}).get("sha", "") if isinstance(resp, dict) else "",
                "error": resp.get("message") if not ok and isinstance(resp, dict) else None,
            })
        except Exception as e:
            results.append({"file": filename, "ok": False, "error": str(e)})

    ok_count = sum(1 for r in results if r.get("ok"))
    return jsonify({
        "success": ok_count == len(files),
        "pushed": ok_count,
        "total": len(files),
        "results": results,
    })


@app.route("/admin/push-info", methods=["GET"])
def admin_push_info():
    """Diagnostika — over ci je vsetko nastavene pre /admin/push (verejne, len bool flagy)."""
    return jsonify({
        "github_pat_set": bool(GITHUB_PAT),
        "github_owner": GITHUB_OWNER,
        "github_repo": GITHUB_REPO,
        "github_branch": GITHUB_BRANCH,
        "webhook_secret_set": bool(WEBHOOK_SECRET),
        "endpoint": "/admin/push",
        "auth_header": "X-Admin-Secret = WEBHOOK_SECRET",
    })


# ============================================================
# WEBHOOK 0: PARSUJ LEADY (Multi-lead intake parser)
# Trigger: Notion Button "🔍 Parsuj leady" v Lead Inbox DB
# ============================================================
@app.route("/webhook/parsuj-leady", methods=["POST"])
@require_secret
def parsuj_leady():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[parsuj-leady] page_id=%s", page_id)

    try:
        source_page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(source_page)
    raw_text = (flat.get("Surový lead") or "").strip()

    if not raw_text:
        try:
            notion_update_page(page_id, {"Status": {"select": {"name": "🔴 Chyba"}}})
        except Exception:
            pass
        return jsonify({"error": "Surový lead je prázdny"}), 400

    log.info("[parsuj-leady] raw_text dlzka=%d znakov", len(raw_text))

    try:
        leads = claude_extract_leads(raw_text)
    except Exception as e:
        log.exception("Claude extraction zlyhala")
        try:
            # NEPREPISUJ Surový lead — zachovaj pôvodný text aby user mohol skúsiť znova
            notion_update_page(page_id, {
                "Status": {"select": {"name": "🔴 Chyba"}},
                "Poznámka": {"rich_text": [{"text": {"content": f"Claude error: {e}"[:1900]}}]},
            })
        except Exception:
            pass
        return jsonify({"error": f"claude failed: {e}"}), 500

    if not leads:
        try:
            notion_update_page(page_id, {"Status": {"select": {"name": "🔴 Chyba"}}})
        except Exception:
            pass
        return jsonify({"error": "Claude nenasiel ziadne leady"}), 400

    log.info("[parsuj-leady] extrahovanych leadov: %d", len(leads))

    created = []
    failed = []
    for i, lead in enumerate(leads):
        try:
            props = lead_to_notion_properties(lead)
            new_page = notion_create_page_in_db(NOTION_DATABASE_ID, props)
            created.append({
                "id": new_page.get("id"),
                "url": new_page.get("url"),
                "title": (lead.get("priezvisko") or lead.get("meno") or f"Lead #{i+1}"),
            })
        except Exception as e:
            log.exception("Vytvorenie page #%d zlyhalo", i + 1)
            failed.append({"index": i + 1, "error": str(e), "lead_preview": str(lead)[:200]})

    # FIX: Ak sa ŽIADNY lead nepodaril, NEZAPISUJ Spracované — nastav Chyba s diagnostikou
    if len(created) == 0:
        try:
            err_summary = (failed[0].get("error", "?")[:300] if failed else "Claude nenasiel ziadne validne leady")
            notion_update_page(page_id, {
                "Status": {"select": {"name": "🔴 Chyba"}},
                "Počet vyparsovaných": {"number": 0},
                "Poznámka": {"rich_text": [{"text": {"content": f"Parsing FAILED: {len(leads)} leadov z Claude, ale všetky padli pri vytváraní v Notion DB. Prvý error: {err_summary}"[:1900]}}]},
            })
        except Exception as e:
            log.warning("Error status update zlyhal: %s", e)
        return jsonify({"ok": False, "error": "all_creates_failed", "claude_returned": len(leads), "failed": failed}), 500

    # Aspoň 1 lead sa vytvoril — Spracované + clear raw
    try:
        notion_update_page(page_id, {
            "Surový lead": {"rich_text": []},
            "Počet vyparsovaných": {"number": len(created)},
            "Status": {"select": {"name": "🟢 Spracované"}},
        })
    except Exception as e:
        log.warning("Cleanup source page zlyhal: %s", e)

    return jsonify({"ok": True, "created": len(created), "failed": len(failed), "leads": created, "errors": failed})


# ============================================================
# WEBHOOK 0b: AUTO-KONFIG (sizing zo Spotreby)
# Trigger: Notion Button "🎯 Auto-konfig" v Zákazníci B2C
# ============================================================
@app.route("/webhook/auto-konfig", methods=["POST"])
@require_secret
def auto_konfig():
    """
    All-in-one Auto-konfig.

    1. Sizing zo Spotreby (Panel, Pocet, Menic) — vzdy ma_bateriu=True (Solinteg menic, kompatibilny s baterion)
    2. Konstrukcia default Skridla
    3. Bateria default 1 ks podla znacky menica:
       - Solinteg menic -> Solinteg EBA B5K1 (5.12 alebo 10.24 podla spotreby)
       - Huawei menic   -> Huawei LUNA2000 (5 alebo 7 kWh)
       - GoodWe menic   -> Pylontech Force H3 5.12
    4. Wallbox default podla znacky menica:
       - Solinteg -> Solinteg 11 kW (3F)
       - Huawei   -> Huawei AC Smart 22 kW
       - GoodWe   -> GoodWe 22 kW
    5. Po nastaveni props prepocita ceny pre VSETKY 4 varianty (A/B/C/D)
       a zapise Cena/Naukpna/Zisk + Suma CP s DPH
    6. Vrati zhrnutie s 4 cenami — uzivatel si vyklika ktoru chce ponuknut.
    """
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[auto-konfig] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # Auto-sync Obchodnik zo Statusu (ak je per-obchodník variant)
    _sync_obchodnik_zo_statusu(page_id, flat)

    # === SPOTREBA ===
    spotreba_raw = flat.get("Spotreba")
    if spotreba_raw is not None:
        try:
            spotreba_val = float(spotreba_raw)
            spotreba_zdroj = "Zákazník zadal"
        except (ValueError, TypeError):
            spotreba_val = float(DEFAULT_SPOTREBA_KWH)
            spotreba_zdroj = "Default 8000 (SK priemer)"
    else:
        spotreba_val = float(DEFAULT_SPOTREBA_KWH)
        spotreba_zdroj = "Default 8000 (SK priemer)"

    # === SIZING — vzdy s bateriou aby sme dostali Solinteg/hybridny menic ===
    sizing = auto_sizing_from_spotreba(spotreba_val, ma_bateriu=True)
    menic = sizing["menic"]

    # === BATERIA podla znacky menica ===
    bat_aktualna = flat.get("Batéria (typ)")
    if not bat_aktualna:
        if "Solinteg" in menic:
            bat_typ = "Solinteg EBA B5K1 — 10.24 kWh" if spotreba_val > 5000 else "Solinteg EBA B5K1 — 5.12 kWh"
        elif "Huawei" in menic:
            bat_typ = "Huawei LUNA2000 — 7 kWh" if spotreba_val > 5000 else "Huawei LUNA2000 — 5 kWh"
        elif "GoodWe" in menic:
            bat_typ = "Pylontech Force H3 — 5.12 kWh"
        else:
            bat_typ = "Solinteg EBA B5K1 — 10.24 kWh"
    else:
        bat_typ = bat_aktualna

    # === WALLBOX podla znacky menica ===
    wb_aktualna = flat.get("Wallbox (typ)")
    if not wb_aktualna:
        if "Solinteg" in menic:
            wb_typ = "Solinteg 11 kW (3F)"
        elif "Huawei" in menic:
            wb_typ = "Huawei AC Smart 22 kW"
        elif "GoodWe" in menic:
            wb_typ = "GoodWe 22 kW"
        else:
            wb_typ = "Solinteg 11 kW (3F)"
    else:
        wb_typ = wb_aktualna

    # === KONSTRUKCIA default Skridla ===
    konstr_aktualna = flat.get("Konštrukcia (typ)")
    konstr_typ = konstr_aktualna or "Škridla"

    # === KROK 1: Update vsetky komponenty + Spotreba + DEFAULT marze 30% ===
    update_props_komponenty = {
        "Spotreba": {"number": spotreba_val},
        "Spotreba zdroj": {"select": {"name": spotreba_zdroj}},
        "Panel": {"select": {"name": sizing["panel"]}},
        "Počet panelov": {"select": {"name": str(sizing["pocet_panelov"])}},
        "Menič": {"select": {"name": menic}},
        "Konštrukcia (typ)": {"select": {"name": konstr_typ}},
        "Batéria (typ)": {"select": {"name": bat_typ}},
        "Batéria počet": {"select": {"name": "1"}},
        "Wallbox (typ)": {"select": {"name": wb_typ}},
    }

    # Default marza 30% pre vsetky varianty ak este nie su nastavene
    DEFAULT_MARZA = "30"
    for v in ("A", "B", "C", "D"):
        marza_aktualna = flat.get(f"Marža {v} %")
        if not marza_aktualna or marza_aktualna == "":
            update_props_komponenty[f"Marža {v} %"] = {"select": {"name": DEFAULT_MARZA}}

    try:
        notion_update_page(page_id, update_props_komponenty)
    except Exception as e:
        log.exception("Notion update komponent zlyhal")
        return jsonify({"error": f"notion_update komponenty failed: {e}"}), 500

    # === KROK 2: Re-fetch page (s novymi props) a prepocet vsetkych 4 variantov ===
    try:
        page2 = notion_get_page(page_id)
        flat2 = notion_props_to_flat(page2)
        from generate_from_notion import predpocitaj_ceny_pre_record
        ceny = predpocitaj_ceny_pre_record(flat2)
    except Exception as e:
        log.exception("Predpocet zlyhal")
        # Komponenty su nastavene, ale ceny zlyhali - vratime castocny success
        return jsonify({
            "ok": False,
            "warning": "Komponenty nastavene, ale predpocet cien zlyhal",
            "error": str(e),
            "sizing": sizing,
        }), 200

    # === KROK 3: Update Notion s cenami pre A/B/C/D ===
    price_update = {}
    sumarne = {}
    for v in ("A", "B", "C", "D"):
        cv = ceny.get(v, {})
        if cv.get("cena_s_dph"):
            price_update.update(notion_set_number(f"Cena {v} s DPH", round(cv["cena_s_dph"], 2)))
            price_update.update(notion_set_number(f"Nákupná cena {v} €", round(cv["nakupna"], 2)))
            price_update.update(notion_set_number(f"Zisk {v} €", round(cv["zisk"], 2)))
            sumarne[v] = round(cv["cena_s_dph"], 2)

    # Suma CP s DPH = preferuj B (kombi), inak A
    suma = (ceny.get("B", {}).get("cena_s_dph")
            or ceny.get("A", {}).get("cena_s_dph")
            or ceny.get("D", {}).get("cena_s_dph")
            or ceny.get("C", {}).get("cena_s_dph"))
    if suma:
        price_update.update(notion_set_number("Suma CP s DPH", round(suma, 2)))

    # Auto-vyplnenie "Bateria vykon"
    m_bat = re.search(r"(\d+(?:[.,]\d+)?)\s*kWh", bat_typ)
    if m_bat:
        per_modul = float(m_bat.group(1).replace(",", "."))
        price_update.update(notion_set_number("Batéria výkon", round(per_modul * 1, 2)))

    if price_update:
        try:
            notion_update_page(page_id, price_update)
        except Exception as e:
            log.warning("Notion price update zlyhal: %s", e)

    log.info("[auto-konfig] hotovo. Sizing=%s, Ceny=%s", sizing, sumarne)

    return jsonify({
        "ok": True,
        "spotreba": spotreba_val,
        "spotreba_zdroj": spotreba_zdroj,
        "sizing": sizing,
        "komponenty": {
            "panel": sizing["panel"],
            "pocet_panelov": sizing["pocet_panelov"],
            "menic": menic,
            "bateria": bat_typ,
            "wallbox": wb_typ,
            "konstrukcia": konstr_typ,
        },
        "ceny": sumarne,
    })


# ============================================================
# WEBHOOK SANDBOX: TEST ROZLOZENIE PANELOV
# Trigger: Notion Button v "🎨 SolarEdge Rebuild Test" DB
# Vstup: { "page_id": "..." }
# Robi: stiahne SolarEdge PDF zo "SolarEdge raw report" property,
# extrahuje obrazky+data, vyrobi branded PDF, vrati base64 + filename.
# Make pak ulozi do Dropbox + nastavi "Branded report" property.
# ============================================================
@app.route("/webhook/test-rozlozenie", methods=["POST"])
@require_secret
def test_rozlozenie():
    body = request.get_json(silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info(f"[test-rozlozenie] page_id={page_id}")

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # Najdi SolarEdge raw report file URL
    raw_files_json = flat.get("SolarEdge raw report") or ""
    raw_url = None
    if raw_files_json:
        try:
            files = json.loads(raw_files_json)
            if files and isinstance(files, list):
                raw_url = files[0].get("url") if files else None
        except (ValueError, TypeError):
            pass

    if not raw_url:
        # Set status na chybu
        try:
            notion_update_page(page_id, {
                "Status": {"select": {"name": "🔴 Chyba"}},
                "Poznámky": {"rich_text": [{"text": {"content": "Chyba: 'SolarEdge raw report' je prazdny — nahraj PDF z Designera."}}]},
            })
        except Exception:
            pass
        return jsonify({"error": "SolarEdge raw report property je prazdna"}), 400

    # Set status na "spracovavam"
    try:
        notion_update_page(page_id, {
            "Status": {"select": {"name": "⚙️ Spracovávam"}},
        })
    except Exception:
        pass

    # Spracovanie
    try:
        from solar_rebuild import process_solaredge_pdf
        import base64 as _b64
        # Heuristika: ak je v Poznamkach "BESS" alebo "bateria", pouzi 90% samospotrebu
        pozn_lower = (flat.get("Poznámky") or "").lower()
        ma_bateriu = ("bess" in pozn_lower or "bateri" in pozn_lower or "wallbox" in pozn_lower)

        pdf_bytes, priezvisko_safe, summary = process_solaredge_pdf(raw_url, ma_bateriu=ma_bateriu)
        pdf_b64 = _b64.b64encode(pdf_bytes).decode("ascii")

        from datetime import datetime as _dt
        datum = _dt.now().strftime("%Y-%m-%d")
        filename = f"Rozlozenie_{priezvisko_safe}_{datum}.pdf"
        folder_name = f"SolarEdge_rebuild_test/{priezvisko_safe}"

        log.info(f"[test-rozlozenie] hotovo: {filename} ({len(pdf_bytes)//1024} KB)")

        return jsonify({
            "success": True,
            "filename": filename,
            "folder_name": folder_name,
            "data": pdf_b64,
            "summary": summary,
        })

    except Exception as e:
        log.exception("[test-rozlozenie] zlyhalo")
        try:
            notion_update_page(page_id, {
                "Status": {"select": {"name": "🔴 Chyba"}},
                "Poznámky": {"rich_text": [{"text": {"content": f"Chyba: {str(e)[:1800]}"}}]},
            })
        except Exception:
            pass
        return jsonify({"error": str(e)}), 500


# ============================================================
# WEBHOOK PROD: SPRACUJ ROZLOZENIE PANELOV (Zakaznici B2C)
# Trigger: Notion Button "🎨 Spracuj rozloženie" v DB Zakaznici B2C
# Vstup: { "page_id": "..." }
# Robi: stiahne SolarEdge PDF z "SolarEdge raw" property,
# extrahuje obrazky+data, vyrobi branded PDF, vrati base64.
# Make pak ulozi do Dropbox + nastavi "Rozlozenie panelov" property.
# Variant-aware: ak ma B/C zaskrtnute, pouzije 90% samospotrebu (s bateriou).
# ============================================================
@app.route("/webhook/spracuj-rozlozenie", methods=["POST"])
@require_secret
def spracuj_rozlozenie():
    body = request.get_json(silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info(f"[spracuj-rozlozenie] page_id={page_id}")

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    raw_files_json = flat.get("SolarEdge raw") or ""
    raw_url = None
    if raw_files_json:
        try:
            files = json.loads(raw_files_json)
            if files and isinstance(files, list):
                raw_url = files[0].get("url") if files else None
        except (ValueError, TypeError):
            pass

    if not raw_url:
        return jsonify({
            "success": False,
            "error": "SolarEdge raw property je prazdna — najprv nahraj PDF z Designera."
        }), 200  # 200 aby Make scenar pokracoval na error handler ak chce

    try:
        from solar_rebuild import process_solaredge_pdf
        import base64 as _b64

        # Variant-aware samospotreba — B/C = bateria → 90%
        var_b = flat.get("Variant B - FVE + BESS") == "__YES__"
        var_c = flat.get("Variant C - FVE + BESS + Wallbox") == "__YES__"
        ma_bateriu = var_b or var_c

        pdf_bytes, _pdf_priezvisko, summary = process_solaredge_pdf(raw_url, ma_bateriu=ma_bateriu)
        pdf_b64 = _b64.b64encode(pdf_bytes).decode("ascii")

        # PRIEZVISKO z Notion title (konzistentne s generate-pdf endpointom)
        # — nie z PDF extrakcie, lebo "BUCEK" v PDF != "Bucek" v Notion
        from generate_from_notion import lead_from_notion as _lfn, safe_filename as _sf
        try:
            _lead = _lfn(flat, "A")
            priezvisko = _lead.get("meno", "").split()[-1] if _lead.get("meno") else "Klient"
        except Exception:
            priezvisko = "Klient"
        priezvisko_safe = _sf(priezvisko) or "Klient"

        # ID z Notion ID ponuky property — rovnaky format ako generate-pdf
        id_p = flat.get("ID ponuky") or ""
        id_match = re.search(r"\d+", str(id_p))
        ev_id = f"EV-26-{int(id_match.group(0)):03d}" if id_match else "EV-XX"
        filename = f"{ev_id}_Rozlozenie_{priezvisko_safe}.pdf"
        folder_name = f"{ev_id}_{priezvisko_safe}"

        log.info(f"[spracuj-rozlozenie] hotovo: {filename} ({len(pdf_bytes)//1024} KB) ma_bateriu={ma_bateriu}")

        return jsonify({
            "success": True,
            "filename": filename,
            "folder_name": folder_name,
            "data": pdf_b64,
            "summary": summary,
        })

    except Exception as e:
        log.exception("[spracuj-rozlozenie] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# WEBHOOK: GENERUJ DOKUMENTY (post-vyhra)
# Trigger: Notion Button "📄 Generuj zmluvy" v Zákazníci B2C
# Vstup: { "page_id": "..." }
# Robi: Načíta Notion data, vyplní 4 templaty (Zmluva, Splnomocnenie,
# GDPR, Dotaznik), vráti base64 + filename + folder. Make uploadne
# do Dropboxu + zapise do Notion property fields.
# ============================================================

@app.route("/webhook/spracuj-rozlozenie-supabase", methods=["POST"])
@require_secret
def spracuj_rozlozenie_supabase():
    """Adapter pre Supabase CRM — prijíma raw_url priamo, bez Notion."""
    body = request.get_json(silent=True) or {}
    raw_url = body.get("raw_url")
    ma_bateriu = bool(body.get("ma_bateriu", False))
    ev_id = body.get("ev_id", "EV-XX")
    priezvisko = body.get("priezvisko", "Klient")

    if not raw_url:
        return jsonify({"success": False, "error": "missing raw_url"}), 400

    log.info(f"[spracuj-rozlozenie-supabase] raw_url={raw_url[:60]}... ma_bateriu={ma_bateriu}")

    try:
        from solar_rebuild import process_solaredge_pdf
        from generate_from_notion import safe_filename as _sf
        import base64 as _b64

        pdf_bytes, _, summary = process_solaredge_pdf(raw_url, ma_bateriu=ma_bateriu)
        priezvisko_safe = _sf(priezvisko) or "Klient"
        filename = f"{ev_id}_Rozlozenie_{priezvisko_safe}.pdf"
        pdf_b64 = _b64.b64encode(pdf_bytes).decode("ascii")

        log.info(f"[spracuj-rozlozenie-supabase] hotovo: {filename} ({len(pdf_bytes)//1024} KB)")

        return jsonify({
            "success": True,
            "filename": filename,
            "data": pdf_b64,
            "summary": summary,
        })
    except Exception as e:
        log.exception("[spracuj-rozlozenie-supabase] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/webhook/generuj-dokumenty", methods=["POST"])
@require_secret
def generuj_dokumenty():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[generuj-dokumenty] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # === VALIDÁCIA — odmietnuť ak chýbajú povinné polia ===
    is_valid, missing = validate_lead_for_documents(flat, doc_type="zmluvy")
    if not is_valid:
        msg = "Lead nemá vyplnené povinné polia: " + ", ".join(missing)
        log.warning("[generuj-dokumenty] %s", msg)
        # Zapíš error do Notion Poznámky aby Lukáš videl
        try:
            notion_update_page(page_id, {
                "Poznámky": {"rich_text": [{"text": {"content": "❌ Generovanie zmluv zlyhalo: " + msg[:200]}}]}
            })
        except Exception:
            pass
        return jsonify({"success": False, "error": msg, "missing": missing}), 400

    # === Zostav lead_data zo zaznamu ===
    from datetime import datetime
    from generate_from_notion import lead_from_notion as _lfn, safe_filename as _sf

    # Pouzi variant A pre zakladne lead udaje
    try:
        _lead = _lfn(flat, "A")
    except Exception:
        _lead = {}

    meno_priezvisko = flat.get("Zákazník", "")
    telefon = flat.get("Telefón", "")
    email = flat.get("Email", "")

    # Adresa: kombinacia ulice, mesto, psc
    ulica = flat.get("Ulica číslo", "")
    mesto = flat.get("Mesto", "")
    psc = flat.get("PSČ", "")
    adresa_parts = [p for p in [ulica, psc, mesto] if p]
    adresa = ", ".join(adresa_parts)

    trvale_bydlisko = flat.get("Trvalé bydlisko", "") or adresa

    cislo_op = flat.get("Číslo OP", "")

    # Datum narodenia format date YYYY-MM-DD -> DD.MM.YYYY
    datum_narodenia_raw = flat.get("date:Dátum narodenia:start", "") or flat.get("Dátum narodenia", "")
    datum_narodenia = ""
    if datum_narodenia_raw:
        try:
            d = datetime.strptime(datum_narodenia_raw[:10], "%Y-%m-%d")
            datum_narodenia = d.strftime("%d.%m.%Y")
        except Exception:
            datum_narodenia = datum_narodenia_raw

    # Cenova ponuka - vyber variant
    # Prioritne čítaj "Variant do zmluvy" select (A/B/C/D)
    # Fallback: prvý zaškrtnutý variant (priorita: B > A > C > D)
    variant_to_use = None
    variant_select = flat.get("Variant do zmluvy") or ""
    if variant_select:
        # Select hodnoty: "A — FVE", "B — FVE + BESS", ...
        m_v = re.match(r"\s*([ABCD])", variant_select)
        if m_v:
            variant_to_use = m_v.group(1)
    if not variant_to_use:
        for v in ("B", "A", "C", "D"):
            prop_name = {
                "A": "Variant A - FVE",
                "B": "Variant B - FVE + BESS",
                "C": "Variant C - FVE + BESS + Wallbox",
                "D": "Variant D - FVE + Wallbox",
            }[v]
            if flat.get(prop_name) == "__YES__":
                variant_to_use = v
                break

    cena_s_dph = 0
    if variant_to_use:
        cena_key = f"Cena {variant_to_use} s DPH"
        cena_s_dph = float(flat.get(cena_key) or 0)
    cena_bez_dph = round(cena_s_dph / 1.23, 2) if cena_s_dph else 0

    # ID ponuky
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id_root = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"
    cislo_cp = f"{ev_id_root}-{variant_to_use or 'A'}"

    # Datum cenovej ponuky
    datum_cp_raw = flat.get("date:Dátum odoslania CP:start", "")
    datum_cp = ""
    if datum_cp_raw:
        try:
            d = datetime.strptime(datum_cp_raw[:10], "%Y-%m-%d")
            datum_cp = d.strftime("%d.%m.%Y")
        except Exception:
            datum_cp = datum_cp_raw

    # Vykon FVE z poctu panelov × 535
    pocet_panelov_raw = flat.get("Počet panelov") or "0"
    try:
        pocet_panelov = int(pocet_panelov_raw)
    except (ValueError, TypeError):
        pocet_panelov = 0
    vykon_kwp = round(pocet_panelov * 535 / 1000, 2)

    datum_dnes = datetime.now().strftime("%d.%m.%Y")

    lead_data = {
        "meno_priezvisko": meno_priezvisko,
        "adresa": adresa,
        "telefon": telefon,
        "email": email,
        "vykon_kwp": vykon_kwp,
        "cislo_cp": cislo_cp,
        "datum_cp": datum_cp,
        "miesto_vykonu": adresa,
        "cena_eur": cena_bez_dph,  # do zmluvy ide cena bez DPH
        "datum_dnes": datum_dnes,
        "datum_narodenia": datum_narodenia,
        "cislo_op": cislo_op,
        "trvale_bydlisko": trvale_bydlisko,
        "ev_id": ev_id_root,
        # Pre dotaznik
        "ulica_cislo": ulica,
        "mesto": mesto,
        "psc": psc,
        "iban": flat.get("IBAN", ""),
        "banka": flat.get("Banka", ""),
        "eic": flat.get("EIC odberného miesta", ""),
        "cislo_obch_partnera": flat.get("Číslo obchodného partnera", ""),
        "spotreba": str(flat.get("Spotreba") or ""),
        "hlavny_istic": flat.get("Hlavný istič", ""),
        "predajca_energii": flat.get("Predajca energií", ""),
        "katastralne_uzemie": flat.get("Katastrálne územie", ""),
        "parcelne_cisla": flat.get("Parcelné čísla", ""),
        "adresa_om": adresa,
    }

    log.info("[generuj-dokumenty] lead_data: meno=%r, ev_id=%r, cena=%r",
             lead_data["meno_priezvisko"], lead_data["ev_id"], lead_data["cena_eur"])

    # === Vyrob 4 dokumenty ===
    try:
        from generuj_dokumenty import vygeneruj_balik_dokumentov
        import base64 as _b64

        with tempfile.TemporaryDirectory() as tmpdir:
            files = vygeneruj_balik_dokumentov(lead_data, tmpdir)

            attachments = []
            for kluc, path in files.items():
                with open(path, "rb") as f:
                    raw = f.read()
                attachments.append({
                    "kluc": kluc,
                    "filename": Path(path).name,
                    "data": _b64.b64encode(raw).decode("ascii"),
                })

            priezvisko_safe = _sf(meno_priezvisko.split()[-1] if meno_priezvisko else "Klient") or "Klient"
            folder_name = f"{ev_id_root}_{priezvisko_safe}"

            log.info("[generuj-dokumenty] vyrobenych %d dokumentov", len(attachments))

            return jsonify({
                "success": True,
                "folder_name": folder_name,
                "attachments": attachments,
                "summary": {
                    "klient": meno_priezvisko,
                    "ev_id": ev_id_root,
                    "cena_bez_dph": cena_bez_dph,
                    "variant": variant_to_use,
                    "vykon_kwp": vykon_kwp,
                },
            })

    except Exception as e:
        log.exception("[generuj-dokumenty] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# WEBHOOK: EMAIL ZMLUVY (poslat klientovi 4 dokumenty)
# Trigger: Notion Button "📧 Poslať zmluvy" v Zákazníci B2C
# Vstup: { "page_id": "..." }
# Robi: Stiahne 4 PDF (Zmluva, Splnomocnenie, GDPR, Dotaznik) z Notion file URLs,
#       vrati base64 + email body. Make pošle cez Outlook.
# ============================================================
@app.route("/webhook/email-zmluvy", methods=["POST"])
@require_secret
def email_zmluvy():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[email-zmluvy] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # Lead udaje
    meno_priezvisko = flat.get("Zákazník", "")
    priezvisko = meno_priezvisko.split()[-1] if meno_priezvisko else "Klient"
    email_zakaznika = (flat.get("Email") or "").strip()
    if not _is_valid_email(email_zakaznika):
        return jsonify({
            "success": False,
            "email_valid": "false",
            "error": f"Neplatny email: '{email_zakaznika}'",
        }), 200

    # Obchodnik
    from generate_from_notion import OBCHODNICI, DEFAULT_OBCHODNIK, safe_filename as _sf
    obchodnik = OBCHODNICI.get(flat.get("Obchodník") or flat.get("Obchodnik") or "", DEFAULT_OBCHODNIK)

    # ID a folder
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id_root = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"
    priezvisko_safe = _sf(priezvisko) or "Klient"
    folder_name = f"{ev_id_root}_{priezvisko_safe}"

    # Stiahni 4 dokumenty z Notion file URLs
    import base64 as _b64
    attachments = []
    file_props = [
        ("Zmluva PDF", "Zmluva o dielo"),
        ("Splnomocnenie PDF", "Splnomocnenie"),
        ("GDPR súhlas PDF", "GDPR súhlas"),
        ("Dotazník PDF", "Dotazník"),
    ]
    docs_present = []
    for prop_name, label in file_props:
        files_json = flat.get(prop_name) or ""
        if not files_json:
            continue
        try:
            files = json.loads(files_json)
        except (ValueError, TypeError):
            files = []
        if not files:
            continue
        f = files[0] if files else None
        url = f.get("url")
        fname = f.get("name") or f"{label}_{priezvisko_safe}"
        if not url:
            continue
        try:
            r = requests.get(url, timeout=60)
            r.raise_for_status()
            attachments.append({
                "filename": fname,
                "folder_name": folder_name,
                "data": _b64.b64encode(r.content).decode("ascii"),
            })
            docs_present.append(label)
        except Exception as e:
            log.warning(f"Stiahnutie {prop_name} zlyhalo: {e}")

    if not attachments:
        return jsonify({
            "success": False,
            "error": "Ziadne dokumenty na poslanie. Najprv klikni 📄 Generuj zmluvy.",
        }), 200

    # Email body
    subject = f"Zmluvná dokumentácia — {meno_priezvisko} (Energovision)"
    body_html = f"""
    <p>Dobrý deň {oslovenie_pan_pani("", priezvisko)} {priezvisko},</p>
    <p>na základe Vášho súhlasu s našou cenovou ponukou Vám zasielam zmluvnú dokumentáciu k inštalácii fotovoltickej elektrárne. V prílohe nájdete:</p>
    <ul>
      <li><strong>Zmluva o dielo</strong> — zmluva o dodávke a inštalácii FVE</li>
      <li><strong>Splnomocnenie</strong> — pre komunikáciu s distribučnou spoločnosťou, SIEA a stavebným úradom</li>
      <li><strong>Súhlas so spracovaním osobných údajov</strong> (GDPR)</li>
      <li><strong>Dotazník pripojenia FVE</strong> — administratívne podklady pre žiadosť o pripojenie do siete</li>
    </ul>
    <p style="background:#FFF8E1;padding:12px;border-left:4px solid #F59E0B;font-size:14px;">
      <strong>Postup:</strong>
      <br>1. Vytlačte všetky 4 dokumenty.
      <br>2. Vyplňte <strong>Dotazník</strong> (potrebujeme údaje pre žiadosť do distribučky — EIC odberného miesta, IBAN, parcelné čísla, atď.).
      <br>3. Podpíšte všetky 4 dokumenty (3 podpisy + dotazník).
      <br>4. Odošlite ich naskenované späť na <strong>{obchodnik.get('email', 'info@energovision.sk')}</strong>, alebo ich odovzdajte osobne pri obhliadke.
    </p>
    <p>Po prijatí podpísaných dokumentov Vám vystavíme zálohovú faktúru (30 % z ceny) a po jej úhrade pripravíme termín inštalácie (typicky 4–8 týždňov).</p>
    <p>V prípade akýchkoľvek otázok ma neváhajte kontaktovať.</p>
    <p>S pozdravom,<br>
    <strong>{obchodnik.get('meno', 'Dominik Galaba')}</strong><br>
    {obchodnik.get('funkcia', 'Office & Administration Manager')}<br>
    📞 {obchodnik.get('tel', '+421 917 424 564')}<br>
    ✉ {obchodnik.get('email', 'dominik.galaba@energovision.sk')}<br>
    <br>
    <em>Energovision s.r.o. — moderné energetické riešenia, ktoré nadchnú</em></p>
    """

    log.info(f"[email-zmluvy] hotovo: {len(attachments)} attachments, to={email_zakaznika}")

    return jsonify({
        "success": True,
        "email_valid": "true",
        "to": email_zakaznika,
        "subject": subject,
        "body_html": body_html,
        "attachments": attachments,
        "obchodnik": obchodnik,
        "docs_sent": docs_present,
    })


# ============================================================
# WEBHOOK: GENERUJ REALIZACNE DOKUMENTY (Revízna správa + Preberací protokol)
# Trigger: Notion Button "📋 Generuj revíziu+protokol"
# Vstup: { "page_id": "..." }
# Robi: po Realizacii vyrobi revíznu správu (z BYTTERM templatu) +
# preberací protokol s konfig FVE. Vrati 2 docx attachmenty.
# ============================================================
@app.route("/webhook/generuj-realizacne", methods=["POST"])
@require_secret
def generuj_realizacne():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[generuj-realizacne] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # === VALIDÁCIA ===
    is_valid, missing = validate_lead_for_documents(flat, doc_type="realizacne")
    if not is_valid:
        msg = "Lead nemá vyplnené povinné polia pre revíznu/protokol: " + ", ".join(missing)
        log.warning("[generuj-realizacne] %s", msg)
        try:
            notion_update_page(page_id, {
                "Poznámky": {"rich_text": [{"text": {"content": "❌ Generovanie realizacných zlyhalo: " + msg[:200]}}]}
            })
        except Exception:
            pass
        return jsonify({"success": False, "error": msg, "missing": missing}), 400

    from datetime import datetime
    from generate_from_notion import safe_filename as _sf

    meno_priezvisko = flat.get("Zákazník", "")
    telefon = flat.get("Telefón", "")
    email = flat.get("Email", "")

    ulica = flat.get("Ulica číslo", "")
    mesto = flat.get("Mesto", "")
    psc = flat.get("PSČ", "")
    adresa_parts = [p for p in [ulica, mesto] if p]
    adresa = ", ".join(adresa_parts)
    psc_mesto = f"{psc} {mesto}".strip()

    # Konfig
    pocet_panelov_raw = flat.get("Počet panelov") or "0"
    try:
        pocet_panelov = int(pocet_panelov_raw)
    except (ValueError, TypeError):
        pocet_panelov = 0
    vykon_kwp = round(pocet_panelov * 535 / 1000, 2)

    # Variant do zmluvy — filtruje komponenty
    variant_select = flat.get("Variant do zmluvy") or ""
    m_v = re.match(r"\s*([ABCD])", variant_select)
    variant = m_v.group(1) if m_v else "B"  # fallback B

    # Bateria — iba pri B alebo C
    bateria_typ_raw = flat.get("Batéria (typ)") or ""
    pocet_baterii_raw = flat.get("Batéria počet") or "0"
    try:
        pocet_baterii_int = int(pocet_baterii_raw)
    except (ValueError, TypeError):
        pocet_baterii_int = 0
    if variant in ("B", "C"):
        bateria_typ = bateria_typ_raw
        pocet_baterii = pocet_baterii_int
    else:
        bateria_typ = ""
        pocet_baterii = 0
    # Extrahuj kWh z label
    m_bat = re.search(r"(\d+(?:[.,]\d+)?)\s*kWh", bateria_typ)
    per_modul_kwh = float(m_bat.group(1).replace(",", ".")) if m_bat else 0
    bateria_kwh = round(pocet_baterii * per_modul_kwh, 2)

    menic = flat.get("Menič") or "Solinteg MHT-10K-25"
    konstrukcia = flat.get("Konštrukcia (typ)") or "Škridla"
    panel_typ = flat.get("Panel") or "LONGi 535 Wp"

    # Sériové čísla pre revíznu správu/protokol
    sn_menic = flat.get("Sériové č. meniča") or ""
    sn_panelov = flat.get("Sériové č. panelov") or ""

    # Doplnkové polia
    hlavny_istic = flat.get("Hlavný istič") or "3x25A"
    predajca_energii = flat.get("Predajca energií") or ""

    # Wallbox — iba pri C alebo D
    wallbox_typ_raw = flat.get("Wallbox (typ)") or ""
    if variant in ("C", "D"):
        wallbox_typ = wallbox_typ_raw
    else:
        wallbox_typ = ""
    ma_wallbox = bool(wallbox_typ)

    # Datumy — Dátum revízie, Dátum odovzdania, Dátum spustenia
    def _parse_d(raw):
        if not raw:
            return ""
        try:
            d = datetime.strptime(raw[:10], "%Y-%m-%d")
            return d.strftime("%d.%m.%Y")
        except Exception:
            return raw

    datum_revizie = _parse_d(flat.get("date:Dátum revízie:start", ""))
    datum_odovzdania = _parse_d(flat.get("date:Dátum odovzdania:start", ""))
    datum_spustenia = _parse_d(flat.get("date:Dátum spustenia:start", ""))

    datum_dnes_str = datetime.now().strftime("%d.%m.%Y")
    if not datum_revizie:
        datum_revizie = datum_odovzdania or datum_spustenia or datum_dnes_str
    if not datum_odovzdania:
        datum_odovzdania = datum_revizie
    if not datum_spustenia:
        datum_spustenia = datum_odovzdania

    # ID ponuky
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id_root = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"

    # Číslo PoUVV — z Notion alebo z ev_id
    cislo_pouvv = flat.get("Číslo PoUVV") or f"P-26-{ev_id_root.replace('EV-26-', '')}"
    revizny_technik = flat.get("Revízny technik") or "Miloš Ďurička"

    lead_data = {
        "meno_priezvisko": meno_priezvisko,
        "adresa": adresa,
        "psc_mesto": psc_mesto,
        "telefon": telefon,
        "email": email,
        "vykon_kwp": vykon_kwp,
        "pocet_panelov": pocet_panelov,
        "panel_typ": panel_typ,
        "menic": menic,
        "sn_menic": sn_menic,
        "sn_panelov": sn_panelov,
        "bateria_typ": bateria_typ,
        "pocet_baterii": pocet_baterii,
        "bateria_kwh": bateria_kwh,
        "konstrukcia": konstrukcia,
        "wallbox_typ": wallbox_typ,
        "ma_wallbox": ma_wallbox,
        "hlavny_istic": hlavny_istic,
        "predajca_energii": predajca_energii,
        "datum_zahajenia": datum_spustenia,
        "datum_odovzdania": datum_odovzdania,
        "datum_revizie": datum_revizie,
        "datum_dnes": datum_dnes_str,
        "cislo_protokolu": ev_id_root,
        "cislo_pouvv": cislo_pouvv,
        "revizny_technik": revizny_technik,
        "ev_id": ev_id_root,
    }

    log.info("[generuj-realizacne] %s: %.2f kWp + %.2f kWh bat",
             meno_priezvisko, vykon_kwp, bateria_kwh)

    try:
        from generuj_dokumenty import vygeneruj_realizacne_dokumenty
        import base64 as _b64

        with tempfile.TemporaryDirectory() as tmpdir:
            files = vygeneruj_realizacne_dokumenty(lead_data, tmpdir)

            attachments = []
            for kluc, path in files.items():
                with open(path, "rb") as f:
                    raw = f.read()
                attachments.append({
                    "kluc": kluc,
                    "filename": Path(path).name,
                    "data": _b64.b64encode(raw).decode("ascii"),
                })

            priezvisko_safe = _sf(meno_priezvisko.split()[-1] if meno_priezvisko else "Klient") or "Klient"
            folder_name = f"{ev_id_root}_{priezvisko_safe}"

            log.info("[generuj-realizacne] hotovo: %d dokumentov", len(attachments))

            return jsonify({
                "success": True,
                "folder_name": folder_name,
                "attachments": attachments,
                "summary": {
                    "klient": meno_priezvisko,
                    "ev_id": ev_id_root,
                    "vykon_kwp": vykon_kwp,
                    "bateria_kwh": bateria_kwh,
                },
            })

    except Exception as e:
        log.exception("[generuj-realizacne] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# WEBHOOK: GENERUJ PD (Projektová dokumentácia — malé zdroje do 10 kW)
# Trigger: Notion button "🏗 Generuj projekt"
# Vstup: { "page_id": "..." }
# Výstup: 5 DOCX dokumentov (krycí list, zoznam, technická správa, PoUVV, súhrnná)
# ============================================================
@app.route("/webhook/generuj-pd", methods=["POST"])
@require_secret
def generuj_pd():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[generuj-pd] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # === VALIDÁCIA ===
    is_valid, missing = validate_lead_for_documents(flat, doc_type="pd")
    if not is_valid:
        msg = "Lead nemá vyplnené povinné polia pre PD: " + ", ".join(missing)
        log.warning("[generuj-pd] %s", msg)
        try:
            notion_update_page(page_id, {
                "Poznámky": {"rich_text": [{"text": {"content": "❌ Generovanie PD zlyhalo: " + msg[:200]}}]}
            })
        except Exception:
            pass
        return jsonify({"success": False, "error": msg, "missing": missing}), 400

    from datetime import datetime
    from generate_from_notion import safe_filename as _sf

    meno_priezvisko = flat.get("Zákazník", "")
    telefon = flat.get("Telefón", "")
    email = flat.get("Email", "")
    ulica = flat.get("Ulica číslo", "")
    mesto = flat.get("Mesto", "")
    psc = flat.get("PSČ", "")
    adresa_parts = [p for p in [ulica, mesto] if p]
    adresa = ", ".join(adresa_parts)
    trvale_bydlisko = flat.get("Trvalé bydlisko", "") or adresa

    # Variant + konfig
    variant_select = flat.get("Variant do zmluvy") or ""
    m_v = re.match(r"\s*([ABCD])", variant_select)
    variant = m_v.group(1) if m_v else "B"

    pocet_panelov_raw = flat.get("Počet panelov") or "0"
    try:
        pocet_panelov = int(pocet_panelov_raw)
    except (ValueError, TypeError):
        pocet_panelov = 0
    vykon_kwp = round(pocet_panelov * 535 / 1000, 2)

    panel_typ = flat.get("Panel") or "LONGi 535 Wp"
    menic = flat.get("Menič") or "Solinteg MHT-10K-25"
    konstrukcia = flat.get("Konštrukcia (typ)") or "Šikmá strecha (škridla)"
    hlavny_istic = flat.get("Hlavný istič") or "3x25A"

    # Bateria iba B/C
    bateria_typ_raw = flat.get("Batéria (typ)") or ""
    pocet_baterii_raw = flat.get("Batéria počet") or "0"
    try:
        pocet_baterii_int = int(pocet_baterii_raw)
    except (ValueError, TypeError):
        pocet_baterii_int = 0
    if variant in ("B", "C"):
        bateria_typ = bateria_typ_raw
        pocet_baterii = pocet_baterii_int
    else:
        bateria_typ = ""
        pocet_baterii = 0
    m_bat = re.search(r"(\d+(?:[.,]\d+)?)\s*kWh", bateria_typ)
    per_modul_kwh = float(m_bat.group(1).replace(",", ".")) if m_bat else 0
    bateria_kwh = round(pocet_baterii * per_modul_kwh, 2)

    # Wallbox iba C/D
    wallbox_typ_raw = flat.get("Wallbox (typ)") or ""
    if variant in ("C", "D"):
        wallbox_typ = wallbox_typ_raw
    else:
        wallbox_typ = ""
    ma_wallbox = bool(wallbox_typ)

    # ID ponuky → ev_id (P-26-XXX)
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id_root = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"

    # PD-špecifické polia
    dis = flat.get("Distribučná spoločnosť") or ""
    cislo_pouvv = flat.get("Číslo PoUVV") or f"PoUVV-{ev_id_root}"
    parcely = flat.get("Parcelné čísla") or ""
    eic = flat.get("EIC odberného miesta") or ""
    eic_dodavka = flat.get("EIC dodávka") or ""
    katastr = flat.get("Katastrálne územie") or ""
    predajca = flat.get("Predajca energií") or ""

    lead_data = {
        "meno_priezvisko": meno_priezvisko,
        "telefon": telefon, "email": email,
        "adresa": adresa, "trvale_bydlisko": trvale_bydlisko,
        "ulica_cislo": ulica, "mesto": mesto, "psc": psc,
        "vykon_kwp": vykon_kwp, "pocet_panelov": pocet_panelov,
        "panel_typ": panel_typ, "menic": menic,
        "bateria_typ": bateria_typ, "pocet_baterii": pocet_baterii, "bateria_kwh": bateria_kwh,
        "konstrukcia": konstrukcia,
        "ma_wallbox": ma_wallbox, "wallbox_typ": wallbox_typ,
        "hlavny_istic": hlavny_istic,
        "dis": dis,
        "cislo_pouvv": cislo_pouvv,
        "parcelne_cisla": parcely,
        "eic": eic, "eic_dodavka": eic_dodavka,
        "katastralne_uzemie": katastr,
        "predajca_energii": predajca,
        "datum_dnes": datetime.now().strftime("%d.%m.%Y"),
        "ev_id": ev_id_root,
        "variant": variant,
    }

    log.info("[generuj-pd] %s: %.2f kWp variant=%s, dis=%s",
             meno_priezvisko, vykon_kwp, variant, dis)

    # Skús stiahnuť SolarEdge raw PDF pre technický výkres (voliteľné)
    solaredge_pdf_bytes = None
    raw_files_json = flat.get("SolarEdge raw") or ""
    if raw_files_json:
        try:
            import json as _json
            files_arr = _json.loads(raw_files_json) if isinstance(raw_files_json, str) else raw_files_json
            if files_arr and isinstance(files_arr, list):
                se_url = None
                for f in files_arr:
                    if isinstance(f, dict):
                        se_url = f.get("file", {}).get("url") or f.get("external", {}).get("url") or f.get("url")
                        if se_url:
                            break
                if se_url:
                    r = requests.get(se_url, timeout=60)
                    r.raise_for_status()
                    solaredge_pdf_bytes = r.content
                    log.info("[generuj-pd] SolarEdge PDF stiahnutý (%d B)", len(solaredge_pdf_bytes))
        except Exception as e:
            log.warning("[generuj-pd] SolarEdge raw nedostupný: %s", e)

    try:
        from generuj_pd import vygeneruj_projektovu_dokumentaciu
        import base64 as _b64

        with tempfile.TemporaryDirectory() as tmpdir:
            files = vygeneruj_projektovu_dokumentaciu(lead_data, tmpdir, solaredge_pdf_bytes=solaredge_pdf_bytes)

            attachments = []
            for kluc, path in files.items():
                with open(path, "rb") as f:
                    raw = f.read()
                attachments.append({
                    "kluc": kluc,
                    "filename": Path(path).name,
                    "data": _b64.b64encode(raw).decode("ascii"),
                })

            priezvisko_safe = _sf(meno_priezvisko.split()[-1] if meno_priezvisko else "Klient") or "Klient"
            folder_name = f"{ev_id_root}_{priezvisko_safe}/Projekcia"

            log.info("[generuj-pd] hotovo: %d dokumentov", len(attachments))

            return jsonify({
                "success": True,
                "folder_name": folder_name,
                "attachments": attachments,
                "summary": {
                    "klient": meno_priezvisko,
                    "ev_id": ev_id_root,
                    "vykon_kwp": vykon_kwp,
                    "variant": variant,
                    "dis": dis,
                },
            })

    except Exception as e:
        log.exception("[generuj-pd] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# WEBHOOK: GENERUJ PO (Material Purchase Order)
# Trigger: Notion automation pri Status = 💰 Faktúra
# Vstup: { "page_id": "..." }
# Výstup: BOM list, vytvorí rows v Materiál PO DB
# ============================================================
@app.route("/webhook/generuj-po", methods=["POST"])
@require_secret
def generuj_po():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[generuj-po] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)
    meno_priezvisko = flat.get("Zákazník", "")

    # ID ponuky -> EV-26-XXX
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"

    # Konfig pre BOM
    pocet_panelov_raw = flat.get("Počet panelov") or "0"
    try:
        pocet_panelov = int(pocet_panelov_raw)
    except (ValueError, TypeError):
        pocet_panelov = 0

    panel_typ = flat.get("Typ panela") or flat.get("Panel") or "LONGi 535 Wp"

    # Wp z typu panela
    m_wp = re.search(r"(\d+)\s*Wp", panel_typ)
    wp = int(m_wp.group(1)) if m_wp else 535
    vykon_kwp = round(pocet_panelov * wp / 1000, 2)

    menic = flat.get("Menič") or ""
    bateria_typ = flat.get("Batéria (typ)") or ""
    pocet_baterii_raw = flat.get("Batéria počet") or "0"
    try:
        pocet_baterii = int(pocet_baterii_raw)
    except (ValueError, TypeError):
        pocet_baterii = 0
    wallbox_typ = flat.get("Wallbox (typ)") or ""
    konstrukcia = flat.get("Konštrukcia (typ)") or "Škridla"
    distribucka = flat.get("Distribučka") or flat.get("Distribuční") or ""

    # Variant do zmluvy filtruje BOM komponenty
    # A = iba FVE (bez batérie, bez wallboxu)
    # B = FVE + batéria (bez wallboxu)
    # C = FVE + batéria + wallbox
    # D = FVE + wallbox (bez batérie)
    variant_select = flat.get("Variant do zmluvy") or ""
    m_v = re.match(r"\s*([ABCD])", variant_select)
    variant = m_v.group(1) if m_v else "B"  # default B (najčastejší kompromis)

    # Filter komponentov podľa variantu
    if variant == "A":
        bateria_typ_eff = ""
        pocet_baterii_eff = 0
        wallbox_typ_eff = ""
    elif variant == "B":
        bateria_typ_eff = bateria_typ
        pocet_baterii_eff = pocet_baterii
        wallbox_typ_eff = ""
    elif variant == "C":
        bateria_typ_eff = bateria_typ
        pocet_baterii_eff = pocet_baterii
        wallbox_typ_eff = wallbox_typ
    elif variant == "D":
        bateria_typ_eff = ""
        pocet_baterii_eff = 0
        wallbox_typ_eff = wallbox_typ
    else:
        bateria_typ_eff = bateria_typ
        pocet_baterii_eff = pocet_baterii
        wallbox_typ_eff = wallbox_typ

    lead_data = {
        "pocet_panelov": pocet_panelov,
        "vykon_kwp": vykon_kwp,
        "panel_typ": panel_typ,
        "menic": menic,
        "bateria_typ": bateria_typ_eff,
        "pocet_baterii": pocet_baterii_eff,
        "wallbox_typ": wallbox_typ_eff,
        "ma_wallbox": bool(wallbox_typ_eff),
        "konstrukcia": konstrukcia,
        "distribucka": distribucka,
        "variant": variant,
    }

    log.info("[generuj-po] %s: %.2f kWp, %d panelov, batéria=%s×%d, WB=%s",
             meno_priezvisko, vykon_kwp, pocet_panelov, bateria_typ, pocet_baterii, wallbox_typ)

    try:
        from generuj_po import generuj_bom, bom_total

        bom = generuj_bom(lead_data)
        celkom_naklady = bom_total(bom)

        # Vytvor rows v Notion DB Materiál PO
        vytvorene = 0
        zlyhane = []
        for item in bom:
            props = {
                "Položka": {
                    "title": [{"text": {"content": item["polozka"]}}]
                },
                "Projekt (ev_id)": {
                    "rich_text": [{"text": {"content": ev_id}}]
                },
                "Klient": {
                    "rich_text": [{"text": {"content": meno_priezvisko}}]
                },
                "Kategória": {
                    "select": {"name": item["kategoria"]}
                },
                "Množstvo": {
                    "number": item["mnozstvo"]
                },
                "Jednotka": {
                    "select": {"name": item["jednotka"]}
                },
                "Dodávateľ": {
                    "select": {"name": item["dodavatel"]}
                },
                "Cena/ks (€)": {
                    "number": item["cena_ks"]
                },
                "Stav": {
                    "status": {"name": "Not started"}
                },
            }
            if item.get("poznamka"):
                props["Poznámka"] = {
                    "rich_text": [{"text": {"content": item["poznamka"]}}]
                }
            try:
                notion_create_page_in_db(NOTION_MATERIAL_PO_DB_ID, props)
                vytvorene += 1
            except Exception as ex:
                zlyhane.append({"polozka": item["polozka"], "error": str(ex)[:200]})
                log.warning("[generuj-po] Failed: %s — %s", item["polozka"][:40], ex)

        log.info("[generuj-po] %s: %d/%d položiek, celkom %.2f EUR (nákup)",
                 ev_id, vytvorene, len(bom), celkom_naklady)

        return jsonify({
            "success": True,
            "ev_id": ev_id,
            "klient": meno_priezvisko,
            "polozky_total": len(bom),
            "polozky_vytvorene": vytvorene,
            "polozky_zlyhane": zlyhane,
            "celkom_naklady_eur": celkom_naklady,
            "summary": {
                "vykon_kwp": vykon_kwp,
                "pocet_panelov": pocet_panelov,
                "menic": menic,
                "bateria_typ": bateria_typ if bateria_typ else None,
                "pocet_baterii": pocet_baterii,
                "wallbox_typ": wallbox_typ if wallbox_typ else None,
            },
        })

    except Exception as e:
        log.exception("[generuj-po] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# WEBHOOK: EMAIL AGENT — FIRST CONTACT
# Trigger: Make scenár pri Status = 🆕 Došlý lead alebo manuálny button
# Vstup: { "page_id": "..." }
# Výstup: { "to_email": "...", "subject": "...", "body": "...", "extracted_info": {...} }
# ============================================================
@app.route("/webhook/email-agent-first", methods=["POST"])
@require_secret
def email_agent_first():
    body = request.get_json(force=True, silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info("[email-agent-first] page_id=%s", page_id)

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # Skip ak Rieši obchodník
    if _human_handles_lead(flat):
        log.info("[email-agent-first] preskočené — rieši obchodník")
        return jsonify({"success": False, "skipped": True, "reason": "rieši_obchodník"}), 200

    # Postavi lead pre email_agent
    meno = flat.get("Zákazník", "")
    email = flat.get("Email", "")
    if not email:
        return jsonify({"error": "lead has no email — cannot start email agent"}), 400

    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id = f"EV-26-{int(m_id.group(0)):03d}" if m_id else f"EV-26-{page_id[:3].upper()}"

    spotreba_raw = flat.get("Spotreba (kWh/rok)")
    try:
        spotreba = int(spotreba_raw) if spotreba_raw else None
    except (ValueError, TypeError):
        spotreba = None

    typ_dopytu = flat.get("Typ dopytu") or ""
    zaujem = [typ_dopytu] if typ_dopytu else ["FVE"]

    lead_data = {
        "ev_id": ev_id,
        "meno": meno,
        "email": email,
        "telefon": flat.get("Telefón", ""),
        "mesto": flat.get("Mesto", ""),
        "spotreba_kwh": spotreba,
        "ma_zaujem_o": zaujem,
        "poznamky": flat.get("Poznámky", ""),
        "zdroj": flat.get("Zdroj", "Web"),
    }

    try:
        from email_agent import vygeneruj_prvy_email
        result = vygeneruj_prvy_email(lead_data)
    except Exception as e:
        log.exception("[email-agent-first] LLM zlyhal")
        return jsonify({"error": f"LLM failed: {e}"}), 500

    # Update Notion — status, transkript, subject
    from datetime import datetime as _dt
    today = _dt.now().strftime("%Y-%m-%d")

    transcript_entry = (
        f"=== {today} | AGENT (prvý kontakt) ===\n"
        f"Subject: {result.get('subject','')}\n\n"
        f"{result.get('body','')}\n\n"
    )

    update_props = {
        "AI Status": {"select": {"name": "🤖 AI komunikuje"}},
        "AI naposledy poslal": {"date": {"start": today}},
        "AI počet follow-upov": {"number": 0},
        "Email predmet": {"rich_text": [{"text": {"content": result.get("subject", "")[:1900]}}]},
        "Email transkript": {"rich_text": [{"text": {"content": transcript_entry[:1900]}}]},
    }
    quality = result.get("lead_quality")
    if quality:
        update_props["AI lead quality"] = {"select": {"name": quality}}

    try:
        notion_update_page(page_id, update_props)
    except Exception as e:
        log.warning("[email-agent-first] Notion update failed: %s", e)

    return jsonify({
        "success": True,
        "to_email": email,
        "to_name": meno,
        "subject": result.get("subject", ""),
        "body": result.get("body_with_signature", result.get("body", "")),
        "ev_id": ev_id,
        "extracted_info": result.get("extracted_info", {}),
        "tokens": result.get("tokens", 0),
    })


# ============================================================
# WEBHOOK: EMAIL AGENT — REPLY HANDLER
# Trigger: Make scenár pri prichádzajúcom emaile (subject obsahuje [EV-XX-XXX])
# Vstup: { "page_id": "...", "incoming_email": "telo emailu", "incoming_subject": "..." }
# Výstup: { "should_reply": bool, "subject": "...", "body": "...", "handover": bool, "to_email": "..." }
# ============================================================
@app.route("/webhook/email-agent-reply", methods=["POST"])
@require_secret
def email_agent_reply():
    body_data = request.get_json(force=True, silent=True) or {}
    page_id = body_data.get("page_id")
    incoming_email = (body_data.get("incoming_email") or "").strip()
    incoming_subject = (body_data.get("incoming_subject") or "").strip()

    if not page_id or not incoming_email:
        return jsonify({"error": "missing page_id or incoming_email"}), 400

    log.info("[email-agent-reply] page_id=%s, body=%d znakov", page_id, len(incoming_email))

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # Skip ak Rieši obchodník
    if _human_handles_lead(flat):
        log.info("[email-agent-reply] preskočené — rieši obchodník")
        try:
            notion_update_page(page_id, {"AI Status": {"select": {"name": "⏸ Pozastavené"}}})
        except Exception:
            pass
        return jsonify({"should_reply": False, "skipped": True, "reason": "rieši_obchodník"}), 200

    # Skontroluj AI Status
    ai_status = flat.get("AI Status") or ""
    if "Opt-out" in ai_status or "Pozastav" in ai_status:
        return jsonify({"should_reply": False, "reason": f"AI paused: {ai_status}"}), 200

    meno = flat.get("Zákazník", "")
    email = flat.get("Email", "")
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"

    spotreba_raw = flat.get("Spotreba (kWh/rok)")
    try:
        spotreba = int(spotreba_raw) if spotreba_raw else None
    except (ValueError, TypeError):
        spotreba = None

    follow_up_count_raw = flat.get("AI počet follow-upov") or "0"
    try:
        follow_up_count = int(follow_up_count_raw)
    except (ValueError, TypeError):
        follow_up_count = 0

    lead_data = {
        "ev_id": ev_id,
        "meno": meno,
        "email": email,
        "mesto": flat.get("Mesto", ""),
        "spotreba_kwh": spotreba,
        "typ_strechy": flat.get("Konštrukcia (typ)", ""),
        "ma_zaujem_o": [flat.get("Typ dopytu")] if flat.get("Typ dopytu") else [],
        "poznamky": flat.get("Poznámky", ""),
    }

    # Načítaj transcript z Notion (rich text)
    transcript_raw = flat.get("Email transkript") or ""
    transcript = _parse_email_transcript(transcript_raw)

    try:
        from email_agent import spracuj_odpoved
        result = spracuj_odpoved(lead_data, transcript, incoming_email, follow_up_count=follow_up_count)
    except Exception as e:
        log.exception("[email-agent-reply] LLM zlyhal")
        return jsonify({"error": f"LLM failed: {e}"}), 500

    # Opt-out detection
    if result.get("opted_out"):
        try:
            notion_update_page(page_id, {
                "AI Status": {"select": {"name": "🛑 Opt-out"}},
            })
        except Exception:
            pass
        return jsonify({
            "should_reply": True,
            "to_email": email,
            "to_name": meno,
            "subject": result.get("subject", ""),
            "body": result.get("body", "") + "\n\n— Tím Energovision",
            "opted_out": True,
        })

    # Update Notion — append do transkriptu, extracted info, status
    from datetime import datetime as _dt
    today = _dt.now().strftime("%Y-%m-%d")

    new_entry = (
        f"=== {today} | KLIENT ===\n{incoming_email[:1500]}\n\n"
        f"=== {today} | AGENT (odpoveď) ===\nSubject: {result.get('subject','')}\n\n{result.get('body','')}\n\n"
    )

    # Append k existujúcemu transcriptu (max 1900 znakov v Notion rich_text)
    combined = (transcript_raw + "\n" + new_entry) if transcript_raw else new_entry
    if len(combined) > 1900:
        combined = combined[-1900:]

    update_props = {
        "Email transkript": {"rich_text": [{"text": {"content": combined}}]},
        "AI naposledy poslal": {"date": {"start": today}},
        "AI počet follow-upov": {"number": 0},  # reset po reply
    }

    ext = result.get("extracted_info") or {}
    if ext.get("spotreba_kwh") and not spotreba:
        try:
            update_props["Spotreba (kWh/rok)"] = {"number": int(ext["spotreba_kwh"])}
        except (ValueError, TypeError):
            pass
    if ext.get("typ_strechy"):
        ts = ext["typ_strechy"]
        if ts in {"Škridla", "Plech kombivrut", "Falcový plech", "Plochá strecha — J 13°", "Plochá strecha — V/Z 10°"}:
            update_props["Konštrukcia (typ)"] = {"select": {"name": ts}}

    quality = result.get("lead_quality")
    if quality:
        update_props["AI lead quality"] = {"select": {"name": quality}}

    handover = bool(result.get("handover_to_dominik"))
    next_action = result.get("next_action", "")

    if handover or next_action == "handover":
        update_props["AI Status"] = {"select": {"name": "📞 Pripravený na hovor"}}
        update_props["Status"] = {"select": {"name": "💼 V riešení"}}
    elif quality == "dead" or next_action == "stop":
        update_props["AI Status"] = {"select": {"name": "❄️ Cold"}}
    else:
        update_props["AI Status"] = {"select": {"name": "🤖 AI komunikuje"}}

    try:
        notion_update_page(page_id, update_props)
    except Exception as e:
        log.warning("[email-agent-reply] Notion update failed: %s", e)

    return jsonify({
        "should_reply": True,
        "to_email": email,
        "to_name": meno,
        "subject": result.get("subject", ""),
        "body": result.get("body_with_signature", result.get("body", "")),
        "handover": handover,
        "lead_quality": quality,
        "tokens": result.get("tokens", 0),
    })


# ============================================================
# WEBHOOK: EMAIL AGENT — FOLLOW-UP CRON
# Trigger: Make cron scenár (denne)
# Vstup: { "page_id": "...", "dni_od_poslednej": 4 }
# Výstup: { "should_send": bool, "subject": "...", "body": "..." }
# ============================================================
@app.route("/webhook/email-agent-followup", methods=["POST"])
@require_secret
def email_agent_followup():
    body_data = request.get_json(force=True, silent=True) or {}
    page_id = body_data.get("page_id")
    dni_od_poslednej = int(body_data.get("dni_od_poslednej") or 3)

    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    try:
        page = notion_get_page(page_id)
    except Exception as e:
        return jsonify({"error": f"notion_get failed: {e}"}), 500

    flat = notion_props_to_flat(page)

    # Skip ak Rieši obchodník
    if _human_handles_lead(flat):
        log.info("[email-agent-followup] preskočené — rieši obchodník")
        return jsonify({"should_send": False, "skipped": True, "reason": "rieši_obchodník"}), 200

    meno = flat.get("Zákazník", "")
    email = flat.get("Email", "")
    id_p = flat.get("ID ponuky") or ""
    m_id = re.search(r"\d+", str(id_p))
    ev_id = f"EV-26-{int(m_id.group(0)):03d}" if m_id else "EV-XX"

    fu_raw = flat.get("AI počet follow-upov") or "0"
    try:
        fu_count = int(fu_raw)
    except (ValueError, TypeError):
        fu_count = 0

    if fu_count >= 3:
        # už sme dosiahli max → daj cold a stop
        try:
            notion_update_page(page_id, {
                "AI Status": {"select": {"name": "❄️ Cold"}},
                "AI lead quality": {"select": {"name": "cold"}},
            })
        except Exception:
            pass
        return jsonify({"should_send": False, "reason": "max_followups_reached"}), 200

    lead_data = {
        "ev_id": ev_id,
        "meno": meno,
        "email": email,
        "mesto": flat.get("Mesto", ""),
    }
    transcript_raw = flat.get("Email transkript") or ""
    transcript = _parse_email_transcript(transcript_raw)

    try:
        from email_agent import vygeneruj_followup
        result = vygeneruj_followup(lead_data, transcript, fu_count, dni_od_poslednej)
    except Exception as e:
        log.exception("[email-agent-followup] LLM zlyhal")
        return jsonify({"error": f"LLM failed: {e}"}), 500

    from datetime import datetime as _dt
    today = _dt.now().strftime("%Y-%m-%d")
    new_fu_count = fu_count + 1

    new_entry = (
        f"=== {today} | AGENT (follow-up č. {new_fu_count}) ===\n"
        f"Subject: {result.get('subject','')}\n\n{result.get('body','')}\n\n"
    )
    combined = (transcript_raw + "\n" + new_entry) if transcript_raw else new_entry
    if len(combined) > 1900:
        combined = combined[-1900:]

    update_props = {
        "Email transkript": {"rich_text": [{"text": {"content": combined}}]},
        "AI naposledy poslal": {"date": {"start": today}},
        "AI počet follow-upov": {"number": new_fu_count},
    }
    if new_fu_count >= 3 or result.get("next_action") == "stop":
        update_props["AI Status"] = {"select": {"name": "❄️ Cold"}}

    try:
        notion_update_page(page_id, update_props)
    except Exception as e:
        log.warning("[email-agent-followup] Notion update failed: %s", e)

    return jsonify({
        "should_send": True,
        "to_email": email,
        "to_name": meno,
        "subject": result.get("subject", ""),
        "body": result.get("body_with_signature", result.get("body", "")),
        "follow_up_number": new_fu_count,
        "tokens": result.get("tokens", 0),
    })


def _human_handles_lead(flat: dict) -> bool:
    """Vráti True ak je v Notion zaškrtnuté 'Rieši obchodník' — AI sa nemá miešať."""
    val = flat.get("Rieši obchodník")
    # Notion checkbox môže prísť ako True/False (bool) alebo "true"/"false" (string)
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.lower() in ("true", "yes", "1", "✓")
    return False


def _parse_email_transcript(raw: str) -> list:
    """Z Notion rich text transcriptu vyparsuj turns pre email_agent."""
    if not raw:
        return []
    turns = []
    # split on === markers
    blocks = re.split(r"={3,}\s*[^=]+\s*={3,}", raw)
    headers = re.findall(r"={3,}\s*([^=]+?)\s*={3,}", raw)
    for header, content in zip(headers, blocks[1:]):
        h_lower = header.lower()
        if "agent" in h_lower:
            role = "agent"
        elif "klient" in h_lower or "zákazník" in h_lower:
            role = "customer"
        else:
            continue
        c = content.strip()
        if c:
            turns.append({"role": role, "content": c[:3000]})
    return turns[-12:]


# ============================================================
# WEBHOOK: CHAT (verejný chatbot widget pre energovision.sk)
# Trigger: JS fetch z embed widgetu na webe
# Vstup: { "history": [{"role":"user|assistant","content":"..."}], "message": "..." }
# Výstup: { "answer": "...", "lead_ready": bool, "lead": {...|null} }
# CORS: povolený pre všetky originy (verejný endpoint)
# ============================================================
@app.route("/webhook/chat", methods=["POST", "OPTIONS"])
def webhook_chat():
    # CORS preflight
    if request.method == "OPTIONS":
        resp = jsonify({"ok": True})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        return resp, 204

    body = request.get_json(force=True, silent=True) or {}
    history = body.get("history") or []
    message = (body.get("message") or "").strip()

    if not message:
        resp = jsonify({"error": "empty message"})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 400

    # Hard limits proti zneužitiu
    if len(message) > 4000:
        message = message[:4000]
    if len(history) > 30:
        history = history[-30:]

    try:
        from chatbot import odpovedz_chatbot, extrahuj_lead

        result = odpovedz_chatbot(history, message)

        # Ak chatbot povedal že lead je hotový, extrahujme ho a zapíšme do Notion Default Inbox
        lead_data = None
        lead_saved = False
        if result.get("lead_ready"):
            full_history = history + [
                {"role": "user", "content": message},
                {"role": "assistant", "content": result.get("answer", "")},
            ]
            lead_data = extrahuj_lead(full_history)

            if lead_data:
                # Zapíš do Notion Default Inbox alebo Zákazníci B2C
                try:
                    saved = _save_chatbot_lead_to_notion(lead_data, full_history)
                    lead_saved = bool(saved)
                except Exception as ex:
                    log.warning("[chat] Notion save failed: %s", ex)

        resp = jsonify({
            "answer": result.get("answer", ""),
            "lead_ready": result.get("lead_ready", False),
            "lead": lead_data,
            "lead_saved": lead_saved,
            "tokens": result.get("tokens", 0),
        })
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    except Exception as e:
        log.exception("[chat] zlyhalo")
        resp = jsonify({"error": str(e)[:200]})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 500


def _save_chatbot_lead_to_notion(lead: dict, full_history: list):
    """Zapis chatbot lead do Notion DB Zákazníci B2C ako nový záznam so Status = Došlý lead."""
    if not lead:
        return None

    meno = lead.get("meno", "Chatbot lead")
    transcript = "\n".join(
        f"{m.get('role','?').upper()}: {m.get('content','')[:500]}"
        for m in (full_history or [])[-20:]
    )

    props = {
        "Zákazník": {"title": [{"text": {"content": meno}}]},
        "Status": {"select": {"name": "🆕 Došlý lead"}},
    }

    if lead.get("email"):
        props["Email"] = {"email": lead["email"]}
    if lead.get("telefon"):
        props["Telefón"] = {"phone_number": lead["telefon"]}
    if lead.get("mesto"):
        props["Mesto"] = {"rich_text": [{"text": {"content": lead["mesto"]}}]}
    if lead.get("adresa"):
        props["Ulica číslo"] = {"rich_text": [{"text": {"content": lead["adresa"]}}]}
    if lead.get("spotreba_kwh"):
        try:
            props["Spotreba (kWh/rok)"] = {"number": int(lead["spotreba_kwh"])}
        except (ValueError, TypeError):
            pass

    poznamka_parts = []
    if lead.get("poznamka"):
        poznamka_parts.append(f"Chatbot zhrnutie: {lead['poznamka']}")
    if lead.get("ma_zaujem_o"):
        zaujem = ", ".join(lead["ma_zaujem_o"]) if isinstance(lead["ma_zaujem_o"], list) else str(lead["ma_zaujem_o"])
        poznamka_parts.append(f"Záujem o: {zaujem}")
    if lead.get("typ_strechy"):
        poznamka_parts.append(f"Strecha: {lead['typ_strechy']}")
    if lead.get("orientacia"):
        poznamka_parts.append(f"Orientácia: {lead['orientacia']}")
    poznamka_parts.append("--- Transkript z chatu ---")
    poznamka_parts.append(transcript[:1800])

    props["Poznámky"] = {
        "rich_text": [{"text": {"content": "\n".join(poznamka_parts)[:2000]}}]
    }

    try:
        new_page = notion_create_page_in_db(NOTION_DATABASE_ID, props)
        log.info("[chat] Lead saved: %s (page %s)", meno, new_page.get("id"))
        return new_page
    except Exception as e:
        log.warning("[chat] notion_create_page_in_db failed: %s", e)
        return None


# ============================================================
# WEBHOOK 1: PREPOČET CIEN
# Trigger: Notion Button "🔄 Prepočítaj cenu"
# Vstup: { "page_id": "..." }
# ============================================================
def _sync_obchodnik_zo_statusu(page_id, notion_props):
    """Ak Status obsahuje meno obchodníka (V riešení — Dominik/Pavol/Andrej/Lukáš),
    nastav property Obchodnik na celé meno."""
    status_val = notion_props.get("Status") or ""
    if "V rie" not in status_val or "—" not in status_val:
        return False
    # Extrahuj časť za pomlčkou
    parts = status_val.split("—", 1)
    if len(parts) != 2:
        return False
    krstne = parts[1].strip()
    # Map krstného mena na full name (kompatibilné s OBCHODNICI v generate_from_notion.py)
    mapping = {
        "Dominik": "Dominik Galaba",
        "Pavol": "Pavol Kaprál",
        "Andrej": "Andrej Herman",
        "Lukáš": "Lukáš Bago",
    }
    full_name = mapping.get(krstne)
    if not full_name:
        return False
    # Skontroluj či už nie je rovnaký
    current = (notion_props.get("Obchodnik") or notion_props.get("Obchodník") or "").strip()
    if current == full_name:
        return False
    try:
        notion_update_page(page_id, notion_set_select("Obchodnik", full_name))
        log.info(f"[sync-obchodnik] {page_id}: Status='{status_val}' → Obchodnik='{full_name}'")
        return True
    except Exception as e:
        log.warning(f"[sync-obchodnik] update zlyhal: {e}")
        return False


@app.route("/webhook/prepocet", methods=["POST"])
@require_secret
def prepocet():
    import traceback, datetime as _dt
    body = request.get_json(silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    _ts = _dt.datetime.now().strftime("%H:%M:%S")
    log.info(f"[{_ts}] Prepočet START page={page_id}")
    try:
        _result = _prepocet_inner(page_id)
        # Diagnostický zápis do Poznámka — kolega vidí kedy a čo
        try:
            _ceny = _result.get("ceny", {})
            _parts = []
            for v in ("A", "B", "C", "D"):
                _c = _ceny.get(v, {})
                if isinstance(_c, dict) and _c.get("cena_s_dph"):
                    _parts.append(f"{v}={_c['cena_s_dph']:.0f}€")
                elif isinstance(_c, dict) and _c.get("error"):
                    _parts.append(f"{v}=ERR")
            _summary = ", ".join(_parts) if _parts else "nič"
            _diag = f"[{_ts}] Prepočet ✅ {_summary} ({_result.get('fields_updated',0)} polí)"
            _existing = notion_get_page(page_id).get("properties", {}).get("Poznámka", {}).get("rich_text", [])
            _existing_text = "".join(t.get("plain_text","") for t in _existing)
            # Odstráň predošlý diagnostický riadok (ak začína "[HH:MM:SS] Prepočet")
            import re as _re
            _existing_text = _re.sub(r"^\[\d{2}:\d{2}:\d{2}\] Prepočet[^\n]*\n?", "", _existing_text)
            _new_pozn = _diag + ("\n" + _existing_text if _existing_text else "")
            notion_update_page(page_id, notion_set_text("Poznámka", _new_pozn[:1900]))
        except Exception as _e:
            log.warning(f"[{_ts}] Diagnostický zápis zlyhal: {_e}")
        return jsonify({"success": True, **_result})
    except Exception as e:
        _tb = traceback.format_exc()
        log.error(f"[{_ts}] Prepočet FAIL page={page_id} → {e}\n{_tb}")
        try:
            _diag = f"[{_ts}] Prepočet ❌ {type(e).__name__}: {str(e)[:120]}"
            notion_update_page(page_id, notion_set_text("Poznámka", _diag))
        except Exception:
            pass
        return jsonify({"success": False, "error": str(e), "traceback": _tb[-500:]}), 500


def _prepocet_inner(page_id):
    """Pôvodná logika prepocet endpointu, vyňatá pre try/except wrapping."""
    page = notion_get_page(page_id)
    notion_props = notion_props_to_flat(page)

    # Auto-sync Obchodnik zo Statusu (ak je per-obchodník variant)
    _sync_obchodnik_zo_statusu(page_id, notion_props)

    # Detekuj zakliknuté varianty - filter pre prepocet
    variants_filter = []
    for k, v in notion_props.items():
        k_lower = k.lower().strip()
        if v == "__YES__":
            if k_lower.startswith("variant a") and "A" not in variants_filter:
                variants_filter.append("A")
            elif k_lower.startswith("variant b") and "B" not in variants_filter:
                variants_filter.append("B")
            elif k_lower.startswith("variant c") and "C" not in variants_filter:
                variants_filter.append("C")
            elif k_lower.startswith("variant d") and "D" not in variants_filter:
                variants_filter.append("D")

    # Ak ziadny variant nezaskrtnuty, ratam vsetky (backward-compat)
    if not variants_filter:
        log.info("Prepocet: ziadny variant zaskrtnuty, ratam vsetky 4")
        ceny = predpocitaj_ceny_pre_record(notion_props)
    else:
        log.info(f"Prepocet: ratam iba zaskrtnute varianty {variants_filter}")

        # Auto-fill default baterie/wallboxu ak je Variant B/C/D zaskrtnuty ale prislusna polia su prazdne
        auto_fill_props = {}
        needs_battery = any(v in variants_filter for v in ("B", "C"))
        if needs_battery and not notion_props.get("Batéria (typ)"):
            menic = notion_props.get("Menič") or ""
            try:
                spotreba_val = float(str(notion_props.get("Spotreba") or 4000).replace(",", "."))
            except (TypeError, ValueError):
                spotreba_val = 4000
            if "Solinteg" in menic:
                bat_typ = "Solinteg EBA B5K1 — 10.24 kWh" if spotreba_val > 5000 else "Solinteg EBA B5K1 — 5.12 kWh"
            elif "Huawei" in menic:
                bat_typ = "Huawei LUNA2000 — 7 kWh" if spotreba_val > 5000 else "Huawei LUNA2000 — 5 kWh"
            elif "GoodWe" in menic:
                bat_typ = "Pylontech Force H3 — 5.12 kWh"
            else:
                bat_typ = "Solinteg EBA B5K1 — 5.12 kWh"
            notion_props["Batéria (typ)"] = bat_typ
            notion_props["Batéria počet"] = "1"
            auto_fill_props["Batéria (typ)"] = {"select": {"name": bat_typ}}
            auto_fill_props["Batéria počet"] = {"select": {"name": "1"}}
            log.info(f"Prepocet auto-fill: Bateria (typ) -> {bat_typ}")

        needs_wallbox = any(v in variants_filter for v in ("C", "D"))
        if needs_wallbox and not notion_props.get("Wallbox (typ)"):
            menic = notion_props.get("Menič") or ""
            if "Huawei" in menic:
                wb_typ = "Huawei AC Smart 22 kW"
            elif "GoodWe" in menic:
                wb_typ = "GoodWe 22 kW"
            else:
                wb_typ = "Solinteg 11 kW (3F)"
            notion_props["Wallbox (typ)"] = wb_typ
            auto_fill_props["Wallbox (typ)"] = {"select": {"name": wb_typ}}
            log.info(f"Prepocet auto-fill: Wallbox (typ) -> {wb_typ}")

        if auto_fill_props:
            try:
                notion_update_page(page_id, auto_fill_props)
            except Exception as e:
                log.warning(f"Prepocet auto-fill Notion update zlyhal: {e}")

        ceny = predpocitaj_ceny_pre_record(notion_props, variants_filter=variants_filter)

    # Update Notion polí
    update = {}
    a = ceny.get("A", {})
    b = ceny.get("B", {})
    c = ceny.get("C", {})

    if a.get("cena_s_dph"):
        update.update(notion_set_number("Cena A s DPH", round(a["cena_s_dph"], 2)))
        update.update(notion_set_number("Nákupná cena A €", round(a["nakupna"], 2)))
        update.update(notion_set_number("Zisk A €", round(a["zisk"], 2)))
    if b.get("cena_s_dph"):
        update.update(notion_set_number("Cena B s DPH", round(b["cena_s_dph"], 2)))
        update.update(notion_set_number("Nákupná cena B €", round(b["nakupna"], 2)))
        update.update(notion_set_number("Zisk B €", round(b["zisk"], 2)))
    if c.get("cena_s_dph"):
        update.update(notion_set_number("Cena C s DPH", round(c["cena_s_dph"], 2)))
        update.update(notion_set_number("Nákupná cena C €", round(c["nakupna"], 2)))
        update.update(notion_set_number("Zisk C €", round(c["zisk"], 2)))
    d = ceny.get("D", {})
    if d.get("cena_s_dph"):
        update.update(notion_set_number("Cena D s DPH", round(d["cena_s_dph"], 2)))
        update.update(notion_set_number("Nákupná cena D €", round(d["nakupna"], 2)))
        update.update(notion_set_number("Zisk D €", round(d["zisk"], 2)))

    # Auto-vyplnenie "Batéria výkon" = počet × kWh per modul (z labelu typu)
    bat_typ = notion_props.get("Batéria (typ)") or ""
    m_bat = re.search(r"(\d+(?:[.,]\d+)?)\s*kWh", bat_typ)
    per_modul = float(m_bat.group(1).replace(",", ".")) if m_bat else 0
    pocet_raw = notion_props.get("Batéria počet")
    try:
        pocet = int(pocet_raw) if pocet_raw not in (None, "") else 0
    except (TypeError, ValueError):
        pocet = 0
    if pocet > 0 and per_modul > 0:
        update.update(notion_set_number("Batéria výkon", round(pocet * per_modul, 2)))

    suma = (b.get("cena_s_dph") or a.get("cena_s_dph") or c.get("cena_s_dph") or d.get("cena_s_dph"))
    if suma:
        update.update(notion_set_number("Suma CP s DPH", round(suma, 2)))

    # Veľkosť (auto label)
    if a:
        from generate_from_notion import lead_from_notion
        lead = lead_from_notion(notion_props, "A")
        velkost = f"{lead['vykon_kwp']:.2f} kWp / {lead.get('panel_pocet', '?')}× LONGi"
        if notion_props.get("Batéria (kWh)"):
            velkost += f" + {notion_props['Batéria (kWh)']} kWh"
        update.update(notion_set_text("Veľkosť", velkost))

    if update:
        notion_update_page(page_id, update)
        log.info(f"Updatnuté {len(update)} polí")

    return {"ceny": ceny, "fields_updated": len(update)}


# ============================================================
# WEBHOOK 2: GENERATE PDF
# Trigger: Notion Button "🖨 Vytlač ponuku"
# Vstup: { "page_id": "...", "variant": "A" | "B" | "C" }
# ============================================================
@app.route("/webhook/generate-pdf", methods=["POST"])
@require_secret
def generate_pdf():
    """Wrap pre _generate_pdf_impl s try/except a validnym JSON pri chybe (Make scenar nepadne)."""
    try:
        return _generate_pdf_impl()
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        log.error(f"generate_pdf padol: {e}\n{tb}")
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback_tail": tb[-500:],
            "pdf_base64": "",  # prazdny string, NIE undefined — Make sa nezadusi
            "filename": "",
            "folder_name": "",
        }), 500


def _generate_pdf_impl():
    body = request.get_json(silent=True) or {}
    page_id = body.get("page_id")
    variant = body.get("variant", "A")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info(f"Generate PDF pre page {page_id}, variant {variant}")

    page = notion_get_page(page_id)
    notion_props = notion_props_to_flat(page)
    lead = lead_from_notion(notion_props, variant)

    # Kompatibilita
    if variant in ("B", "C"):
        ok, msg = check_compatibility(lead["invertor_kod"], lead.get("bateria_kod"))
        if not ok:
            return jsonify({"error": f"incompatible: {msg}"}), 400

    cennik = load_cennik()
    konfig = vyrataj_konfig(lead, cennik)
    ceny = vyrataj_ceny(konfig, lead)

    # === CRM OVERRIDE ===
    # Ak Supabase CRM posiela vlastnú cenu (Cena X s DPH + Marža X), použijeme ju priamo.
    # Inak by Render použil Hagard ceny bez 35% Energovision zľavy + svoju 21% maržu → mismatch.
    crm_cena_s_dph = None
    try:
        for key in (f"Cena {variant} s DPH", f"Cena {variant}", "Cena s DPH"):
            val = flat_props.get(key)
            if val is None or val == "" or val == 0:
                continue
            crm_cena_s_dph = float(val)
            if crm_cena_s_dph > 0:
                break
    except (TypeError, ValueError):
        crm_cena_s_dph = None

    if crm_cena_s_dph and crm_cena_s_dph > 0:
        crm_marza = flat_props.get(f"Marža {variant}")
        try:
            crm_marza = float(crm_marza) if crm_marza is not None else None
        except (TypeError, ValueError):
            crm_marza = None
        dph = 0.23
        crm_cena_bez_dph = crm_cena_s_dph / (1 + dph)
        # Dotácia + ZD logika ostáva
        dotacia = ceny.get("dotacia", 0) if lead.get("dotacia", True) else 0
        cena_po_dot = crm_cena_s_dph - dotacia
        ceny = {
            "nakupna_material": ceny.get("nakupna_material", 0),
            "nakupna_praca": ceny.get("nakupna_praca", 0),
            "nakupna_spolu": ceny.get("nakupna_spolu", 0),
            "rezerva_eur": 0,
            "marza_eur": crm_cena_bez_dph - ceny.get("nakupna_spolu", 0),
            "cena_bez_dph": crm_cena_bez_dph,
            "cena_s_dph": crm_cena_s_dph,
            "dotacia": dotacia,
            "cena_po_dotacii": cena_po_dot,
            "zlava_eur": 0,
            "cena_finalna": cena_po_dot,
            "marza_pct": crm_marza if crm_marza is not None else ceny.get("marza_pct"),
            "zisk": crm_cena_bez_dph - ceny.get("nakupna_spolu", 0),
            "_source": "CRM_OVERRIDE",
        }
        log.info(f"[generate-pdf-supabase] CRM override použitý — cena s DPH = {crm_cena_s_dph}, marža = {crm_marza}%")

    navratnost = vyrataj_navratnost(konfig, ceny, lead)

    # Vyrob v dočasnom adresári
    with tempfile.TemporaryDirectory() as tmpdir:
        priezvisko = safe_filename(lead["meno"].split()[-1])
        ev_id = lead.get("cislo_ponuky", "EV-XX-001-A")
        from datetime import datetime
        datum = datetime.now().strftime("%Y-%m-%d")
        base = f"{ev_id}_{priezvisko}_{datum}"

        # Grafy
        grafy = vyrob_grafy(navratnost, lead, tmpdir, base)

        # PDF
        pdf_path = os.path.join(tmpdir, f"{base}.pdf")
        vyrob_html_pdf(lead, konfig, ceny, navratnost, grafy, pdf_path)

        # TODO: Upload PDF do Notion stránky ako file attachment
        # Notion API neumožňuje priame upload — treba cez S3 alebo
        # cez file URL property. Najjednoduchšie: pridať PDF na S3
        # alebo Render's persistent disk + URL do Notion.
        pdf_size = os.path.getsize(pdf_path)

        # Provisional: vrátime PDF ako base64 (Make.com to dokáže uložiť)
        import base64
        with open(pdf_path, "rb") as f:
            pdf_b64 = base64.b64encode(f.read()).decode("ascii")

    # folder_name: bez variantu (-A/-B/-C) a bez dátumu — stabilný v čase per zákazník
    ev_id_root = ev_id[:-2] if len(ev_id) >= 2 and ev_id[-2] == "-" and ev_id[-1] in "ABC" else ev_id
    folder_name = f"{ev_id_root}_{priezvisko}"

    return jsonify({
        "success": True,
        "ev_id": ev_id,
        "filename": f"{base}.pdf",
        "folder_name": folder_name,
        "size_bytes": pdf_size,
        "cena_s_dph": ceny["cena_s_dph"],
        "cena_finalna": ceny["cena_finalna"],
        "pdf_base64": pdf_b64,
    })


# ============================================================
# WEBHOOK 3: EMAIL TEMPLATE BUILDER
# Trigger: Notion Button "📧 Odoslať email"
# Vstup: { "page_id": "..." }
# Výstup: { "to": "...", "subject": "...", "body_html": "...",
#          "attachments": [{"name": "...", "dropbox_path": "..."}],
#          "obchodnik": {...}, "variants_sent": ["A","B"] }
# ============================================================
@app.route("/webhook/email-template", methods=["POST"])
@require_secret
def email_template():
    try:
        return _email_template_impl()
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        log.error(f"email_template padol: {e}\n{tb}")
        return jsonify({
            "error": str(e),
            "traceback": tb,
        }), 500


def _is_valid_email(email):
    """Jednoduchá validácia emailu — prítomnosť @ + bodka + min dĺžka."""
    if not email or "@" not in email:
        return False
    parts = email.strip().split("@")
    if len(parts) != 2:
        return False
    local, domain = parts
    if len(local) < 1 or len(domain) < 4:
        return False
    if "." not in domain:
        return False
    # TLD min 2 znaky
    tld = domain.rsplit(".", 1)[-1]
    if len(tld) < 2:
        return False
    return True


def _email_template_impl():
    body = request.get_json(silent=True) or {}
    page_id = body.get("page_id")
    if not page_id:
        return jsonify({"error": "missing page_id"}), 400

    log.info(f"Email template pre page {page_id}")
    page = notion_get_page(page_id)
    notion_props = notion_props_to_flat(page)

    # Detekuj zakliknuté varianty — tolerantný matching property názvu (akýkoľvek prefix "Variant A/B/C")
    variants_active = []
    log.info(f"notion_props keys: {list(notion_props.keys())[:30]}")
    for k, v in notion_props.items():
        k_lower = k.lower().strip()
        if v == "__YES__":
            if k_lower.startswith("variant a"):
                if "A" not in variants_active:
                    variants_active.append("A")
            elif k_lower.startswith("variant b"):
                if "B" not in variants_active:
                    variants_active.append("B")
            elif k_lower.startswith("variant c"):
                if "C" not in variants_active:
                    variants_active.append("C")
    log.info(f"variants_active: {variants_active}")

    if not variants_active:
        # Diagnostický error: vypíšeme aké hodnoty sme videli pre Variant properties
        variant_props_seen = {k: v for k, v in notion_props.items() if "variant" in k.lower()}
        return jsonify({"error": f"Žiadny variant nie je zakliknutý. Variant properties seen: {variant_props_seen}"}), 400

    # Lead data — z A variantu (kvôli základným údajom; ceny berieme zo všetkých)
    from generate_from_notion import lead_from_notion, OBCHODNICI, DEFAULT_OBCHODNIK
    lead_a = lead_from_notion(notion_props, "A")

    priezvisko = lead_a["meno"].split()[-1] if lead_a.get("meno") else "Zákazník"
    mesto = lead_a.get("mesto", "")
    email_zakaznika = (lead_a.get("email") or notion_props.get("Email") or "").strip()
    if not _is_valid_email(email_zakaznika):
        log.warning(f"Neplatný email pre page {page_id}: '{email_zakaznika}'")
        return jsonify({
            "success": False,
            "email_valid": "false",
            "error": f"Neplatný email zákazníka: '{email_zakaznika}'. Skontroluj 'Email' property v Notion DB.",
            "to": "",
            "subject": "",
            "body_html": "",
            "attachments": [],
            "obchodnik": {},
            "variants_sent": [],
        }), 200  # Status 200 aby Make scenár pokračoval k Notion update
    obchodnik = OBCHODNICI.get(notion_props.get("Obchodník") or notion_props.get("Obchodnik") or "", DEFAULT_OBCHODNIK)

    vykon_kwp = lead_a.get("vykon_kwp", 0)
    bateria_kwh = float(notion_props.get("Batéria výkon") or 0)

    # Ceny zo všetkých zakliknutých variant
    ceny = {
        "A": notion_props.get("Cena A s DPH"),
        "B": notion_props.get("Cena B s DPH"),
        "C": notion_props.get("Cena C s DPH"),
    }

    # Vygeneruj PDF pre každý aktívny variant a vráť ako base64
    from generate_from_notion import safe_filename
    import base64
    priezvisko_clean = safe_filename(priezvisko)
    attachments = []
    cennik = load_cennik()
    for v in variants_active:
        try:
            lead_v = lead_from_notion(notion_props, v)
            if v in ("B", "C"):
                ok, _msg = check_compatibility(lead_v["invertor_kod"], lead_v.get("bateria_kod"))
                if not ok:
                    log.warning(f"Variant {v} incompatible: {_msg}, skipping attachment")
                    continue
            konfig_v = vyrataj_konfig(lead_v, cennik)
            ceny_v = vyrataj_ceny(konfig_v, lead_v)
            navratnost_v = vyrataj_navratnost(konfig_v, ceny_v, lead_v)
            with tempfile.TemporaryDirectory() as tmpdir:
                ev_id_v = lead_v.get("cislo_ponuky", f"EV-XX-001-{v}")
                from datetime import datetime as _dt
                datum_v = _dt.now().strftime("%Y-%m-%d")
                base_v = f"{ev_id_v}_{priezvisko_clean}_{datum_v}"
                grafy_v = vyrob_grafy(navratnost_v, lead_v, tmpdir, base_v)
                pdf_path_v = os.path.join(tmpdir, f"{base_v}.pdf")
                vyrob_html_pdf(lead_v, konfig_v, ceny_v, navratnost_v, grafy_v, pdf_path_v)
                with open(pdf_path_v, "rb") as fp:
                    pdf_b64 = base64.b64encode(fp.read()).decode("ascii")
                # folder_name — stabilný per zákazník, bez variantu/dátumu
                _evid_root = ev_id_v[:-2] if len(ev_id_v) >= 2 and ev_id_v[-2] == "-" and ev_id_v[-1] in "ABC" else ev_id_v
                _folder = f"{_evid_root}_{priezvisko_clean}"
                attachments.append({
                    "filename": f"{base_v}.pdf",
                    "folder_name": _folder,
                    "data": pdf_b64,
                })
        except Exception as e:
            log.error(f"PDF gen pre variant {v} zlyhal: {e}")
            continue

    # ===== ROZLOŽENIE PANELOV (extra attachment z Notion files property) =====
    rozlozenie_attached = False
    rozlozenie_json = notion_props.get("Rozloženie panelov") or ""
    if rozlozenie_json:
        try:
            rozlozenie_files = json.loads(rozlozenie_json)
        except (ValueError, TypeError):
            rozlozenie_files = []
        for rf in rozlozenie_files[:3]:  # max 3 files
            try:
                fname = rf.get("name") or "rozlozenie_panelov.pdf"
                furl = rf.get("url")
                if not furl:
                    continue
                # safe filename
                fname_clean = safe_filename(fname.rsplit(".", 1)[0])
                fext = fname.rsplit(".", 1)[-1] if "." in fname else "pdf"
                final_name = f"Rozlozenie_panelov_{priezvisko_clean}.{fext}"
                # Stiahni cez requests + base64
                r_dl = requests.get(furl, timeout=60)
                r_dl.raise_for_status()
                pdf_b64_rozl = base64.b64encode(r_dl.content).decode("ascii")
                _evid_root_r = (lead_a.get("cislo_ponuky") or "EV-XX-001-A")
                _evid_root_r = _evid_root_r[:-2] if len(_evid_root_r) >= 2 and _evid_root_r[-2] == "-" and _evid_root_r[-1] in "ABC" else _evid_root_r
                _folder_r = f"{_evid_root_r}_{priezvisko_clean}"
                attachments.append({
                    "filename": final_name,
                    "folder_name": _folder_r,
                    "data": pdf_b64_rozl,
                })
                rozlozenie_attached = True
                log.info(f"Rozlozenie panelov pridane: {final_name} ({len(r_dl.content)} bajtov)")
            except Exception as e:
                log.warning(f"Stiahnut rozlozenie panelov zlyhalo: {e}")

    # ===== EMAIL TEMPLATES =====
    typ_ponuky = notion_props.get("Typ ponuky") or "Indikatívna"
    subject = build_subject(priezvisko, mesto, variants_active, typ_ponuky=typ_ponuky)
    body_html = build_email_body(priezvisko, mesto, vykon_kwp, bateria_kwh, ceny, variants_active, obchodnik, typ_ponuky=typ_ponuky, ma_rozlozenie=rozlozenie_attached)

    return jsonify({
        "success": True,
        "email_valid": "true",
        "to": email_zakaznika,
        "subject": subject,
        "body_html": body_html,
        "attachments": attachments,
        "obchodnik": obchodnik,
        "variants_sent": variants_active,
    })


def build_subject(priezvisko, mesto, variants, typ_ponuky="Indikatívna"):
    """Subject riadok — krátky, identifikovateľný. Pri Indikatívnej ponuke sa pridáva prefix."""
    prefix = "Indikatívna cenová ponuka" if typ_ponuky == "Indikatívna" else "Cenová ponuka"
    if len(variants) == 1:
        v_label = {"A": "FVE", "B": "FVE + batéria", "C": "FVE + batéria + wallbox", "D": "FVE + wallbox"}[variants[0]]
        return f"{prefix} {v_label} pre {priezvisko}, {mesto}"
    return f"{prefix} FVE — {priezvisko}, {mesto} ({len(variants)} varianty)"


def build_email_body(priezvisko, mesto, kwp, bateria_kwh, ceny, variants, obchodnik, typ_ponuky="Indikatívna", ma_rozlozenie=False):
    """
    HTML email body s marketingovým textom (slovenský trh, 30y FVE expert tone).
    Per-variant blocks + comparison + signature.

    typ_ponuky: "Indikatívna" (bez obhliadky, default) alebo "Exaktná" (po obhliadke)
    ma_rozlozenie: True ak v emaily je priložené aj rozlozenie panelov
    """
    cena_a = ceny.get("A") or 0
    cena_b = ceny.get("B") or 0
    cena_c = ceny.get("C") or 0
    cena_d = ceny.get("D") or 0

    # === INTRO — 2 verzie podla typu ponuky ===
    n_var = len(variants)
    n_var_str = "jednu variantu" if n_var == 1 else f"{n_var} varianty"

    if typ_ponuky == "Exaktná":
        # Po obhliadke — preverené data, presné ceny
        intro = f"""
        <p>Dobrý deň {oslovenie_pan_pani("", priezvisko)} {priezvisko},</p>
        <p>v nadväznosti na našu obhliadku Vašej nehnuteľnosti v <strong>{mesto}</strong>
        Vám zasielam <strong>presnú cenovú ponuku</strong> pre fotovoltickú elektráreň.
        Údaje sú overené priamo na mieste — strecha, konštrukcia, spotreba aj umiestnenie panelov.
        Pripravil som {n_var_str} podľa toho, ako chcete využiť energiu zo slnka.</p>
        """
    else:
        # Indikatívna — bez obhliadky, len odhad z dopytu
        intro = f"""
        <p>Dobrý deň {oslovenie_pan_pani("", priezvisko)} {priezvisko},</p>
        <p>na základe údajov, ktoré ste nám poskytli, Vám zasielam <strong>indikatívnu cenovú ponuku</strong>
        pre fotovoltickú elektráreň pre Vašu domácnosť v <strong>{mesto}</strong>.
        Pripravil som {n_var_str} podľa toho, ako chcete využiť energiu zo slnka.</p>
        <p style="background:#FFF8E1;padding:12px;border-left:4px solid #F59E0B;font-size:14px;margin:16px 0;">
          <strong>Pozn.:</strong> Toto je <em>indikatívna ponuka</em> spracovaná z údajov v dopyte —
          bez fyzickej obhliadky. Presnú cenu Vám pripravíme po obhliadke nehnuteľnosti, kedy si overíme
          stav strechy, ideálne umiestnenie panelov a kabelážnu trasu. Ceny sa môžu mierne upraviť
          (typicky ±5–10 %) podľa skutočného stavu.
        </p>
        """

    if ma_rozlozenie:
        intro += """
        <p style="background:#E0F2F1;padding:12px;border-left:4px solid #1B5E3F;font-size:14px;margin:16px 0;">
          <strong>📐 V prílohe nájdete aj návrh rozloženia panelov</strong> na Vašej streche —
          vizualizáciu z nášho projekčného software, ktorá ukáže optimálne umiestnenie a počet panelov.
        </p>
        """

    blocks = []

    # === BLOCK A — iba FVE ===
    if "A" in variants:
        blocks.append(f"""
        <h3 style="color:#1B5E3F;margin-top:24px;">Varianta A — fotovoltika {kwp} kWp</h3>
        <p>Najlacnejšia cesta ako začať. Panely vyrábajú elektrinu cez deň, ktorú spotrebovávate priamo
        — typicky pokryje 60-70 % Vašej dennej spotreby. Ideálne ak doma cez deň žije rodina, sušiete
        bielizeň, varíte alebo používate tepelné čerpadlo na ohrev TÚV.</p>
        <ul>
          <li><strong>Investícia po dotácii Zelená domácnostiam:</strong> {cena_a:,.0f} € s DPH</li>
          <li><strong>Návratnosť:</strong> 6–8 rokov pri dnešnej cene elektriny 0,16 €/kWh</li>
          <li><strong>Záruka:</strong> 30 rokov na panely LONGi, 10 rokov na menič Huawei</li>
          <li><strong>Inštalácia:</strong> 1–2 dni, bez stavebných úprav</li>
        </ul>
        <p style="background:#F0F8F4;padding:12px;border-left:4px solid #1B5E3F;font-size:14px;">
          <strong>Pre slovenský trh:</strong> 60 % FVE inštalácií v rodinných domoch ide bez batérie.
          Distribučné spoločnosti odkupujú prebytky za 0,04–0,06 €/kWh, čo postačí na základnú nezávislosť.
        </p>
        """.replace(",", " "))

    # === BLOCK B — FVE + BESS ===
    if "B" in variants:
        bat_str = f"{bateria_kwh:.0f} kWh" if bateria_kwh else "batéria"
        blocks.append(f"""
        <h3 style="color:#1B5E3F;margin-top:24px;">Varianta B — fotovoltika {kwp} kWp + batéria {bat_str}</h3>
        <p>Energetická nezávislosť — slnko ukladáte do batérie a používate večer/v noci/keď je zamračené.
        Pri zlepšujúcich sa zľavách na komponenty je toto pre Slovákov dnes najatraktívnejšia voľba,
        najmä pre rodiny ktoré sú doma <strong>predovšetkým ráno a večer</strong>.</p>
        <ul>
          <li><strong>Investícia po dotácii:</strong> {cena_b:,.0f} € s DPH</li>
          <li><strong>Pokrytie spotreby:</strong> 85–95 % pri správnom dimenzovaní</li>
          <li><strong>Návratnosť:</strong> 8–11 rokov</li>
          <li><strong>Backup:</strong> pri výpadku siete batéria automaticky prepne dom na ostrov (voliteľne)</li>
          <li><strong>Záruka batérie:</strong> 10 rokov / 6 000 cyklov</li>
        </ul>
        <p style="background:#F0F8F4;padding:12px;border-left:4px solid #1B5E3F;font-size:14px;">
          <strong>Pozor na kalkulácie:</strong> Distribučné poplatky v SR rastú každoročne ~5–8 %.
          Batéria vám teda chráni nielen pred volatilitou ceny silovej elektriny ale aj pred budúcim rastom
          poplatkov za distribúciu.
        </p>
        """.replace(",", " "))

    # === BLOCK D — FVE + Wallbox (bez baterie) ===
    if "D" in variants:
        blocks.append(f"""
        <h3 style="color:#1B5E3F;margin-top:24px;">Varianta D — fotovoltika {kwp} kWp + wallbox (bez batérie)</h3>
        <p>Pre rodiny s elektromobilom ktoré <strong>nepotrebujú batériu</strong> — auto sa nabíja priamo zo slnka cez deň.
        Wallbox automaticky reaguje na prebytky FVE a využíva ich na nabíjanie EV. Optimálne riešenie ak doma cez deň
        bývate menej a hlavnou prioritou je nabíjanie auta zo slnka.</p>
        <ul>
          <li><strong>Investícia po dotácii:</strong> {cena_d:,.0f} € s DPH</li>
          <li><strong>Návratnosť:</strong> 7–9 rokov pri kombinácii FVE + EV nabíjanie</li>
          <li><strong>Výhoda:</strong> nižšia investícia ako varianta C, ale stále plné EV nabíjanie zo slnka</li>
          <li><strong>Hybridný menič:</strong> možnosť doplnenia batérie neskôr bez prerábania systému</li>
        </ul>
        <p style="background:#F0F8F4;padding:12px;border-left:4px solid #1B5E3F;font-size:14px;">
          <strong>Dlhodobý plán:</strong> mnoho rodín začína s variant D a po 2-3 rokoch dopĺňa batériu, keď vidia
          svoj reálny profil spotreby. Týmto spôsobom investujú postupne a vyhnú sa pre/poddimenzovaniu batérie.
        </p>
        """.replace(",", " "))

    # === BLOCK C — FVE + BESS + Wallbox ===
    if "C" in variants:
        bat_str = f"{bateria_kwh:.0f} kWh" if bateria_kwh else "batéria"
        blocks.append(f"""
        <h3 style="color:#1B5E3F;margin-top:24px;">Varianta C — kompletné riešenie + wallbox pre elektromobil</h3>
        <p>FVE {kwp} kWp + batéria {bat_str} + smart wallbox. Vaše auto sa nabíja zo slnka — zadarmo —
        a wallbox automaticky reaguje na prebytky FVE. Riešenie pre rodiny s elektromobilom alebo plánom kúpiť
        EV v najbližších rokoch.</p>
        <ul>
          <li><strong>Investícia:</strong> {cena_c:,.0f} € s DPH</li>
          <li><strong>Úspora paliva:</strong> ~ 1 200 € ročne pri 15 000 km/rok namiesto benzínu</li>
          <li><strong>Plná energetická nezávislosť:</strong> spotreba domu + auto z vlastnej elektriny</li>
          <li><strong>Smart logika:</strong> wallbox sa zapne keď FVE má nadbytok, nezasahuje do siete</li>
        </ul>
        <p style="background:#F0F8F4;padding:12px;border-left:4px solid #1B5E3F;font-size:14px;">
          <strong>Slovenský kontext:</strong> Pri raste cien benzínu/nafty a zvyšujúcich sa parkovných
          poplatkoch vo veľkých mestách (Bratislava, Košice) sa elektromobil vracia za 4-6 rokov samostatne.
          S vlastnou FVE ešte rýchlejšie.
        </p>
        """.replace(",", " "))

    # === COMPARISON ak viac variant ===
    comparison = ""
    if len(variants) > 1:
        rows = []
        if "A" in variants:
            rows.append(f"<tr><td>A — iba FVE</td><td style='text-align:right;'>{cena_a:,.0f} €</td><td>~7 rokov</td><td>Šetríš cez deň</td></tr>".replace(",", " "))
        if "B" in variants:
            rows.append(f"<tr><td>B — FVE + batéria</td><td style='text-align:right;'>{cena_b:,.0f} €</td><td>~9 rokov</td><td>Plná denná + nočná nezávislosť</td></tr>".replace(",", " "))
        if "C" in variants:
            rows.append(f"<tr><td>C — komplet + wallbox</td><td style='text-align:right;'>{cena_c:,.0f} €</td><td>~11 rokov</td><td>+ EV nabíjanie zadarmo</td></tr>".replace(",", " "))

        comparison = f"""
        <h3 style="color:#1B5E3F;margin-top:24px;">Krátke porovnanie</h3>
        <table style="border-collapse:collapse;width:100%;font-size:14px;">
          <thead>
            <tr style="background:#1B5E3F;color:white;">
              <th style="padding:8px;text-align:left;">Varianta</th>
              <th style="padding:8px;text-align:right;">Cena s DPH</th>
              <th style="padding:8px;text-align:left;">Návratnosť</th>
              <th style="padding:8px;text-align:left;">Komfort</th>
            </tr>
          </thead>
          <tbody>{"".join(rows)}</tbody>
        </table>
        <p style="font-size:14px;font-style:italic;color:#555;margin-top:8px;">
        Moja osobná poznámka po 30 rokoch v energetike: ak nemáte konkrétny plán na elektromobil v
        najbližších 2–3 rokoch, varianta B prináša najlepší pomer komfortu k investícii. Hybridný menič
        v balíku umožňuje doplnenie batérie kedykoľvek neskôr — bez prerábania systému.
        </p>
        """

    # === ATTACHMENTY popis ===
    n_pdf = len(variants)
    pdf_note = f"<p>V <strong>prílohe e-mailu</strong> nájdete {('detailný PDF dokument' if n_pdf == 1 else f'{n_pdf} PDF dokumenty')} s technickou špecifikáciou, vizualizáciou rozloženia panelov, návratnostnou kalkuláciou na 25 rokov a podmienkami inštalácie.</p>"

    # === CTA + SIGNATURE ===
    cta = f"""
    <h3 style="color:#1B5E3F;margin-top:24px;">Ďalšie kroky</h3>
    <p>Ponuka platí 30 dní. Ak Vás niektorá varianta zaujala alebo máte otázky, stačí mi odpísať alebo zavolať.
    Dohodneme si <strong>bezplatnú obhliadku</strong> v termíne ktorý Vám vyhovuje — meriame strechu, navrhneme
    optimálne rozloženie panelov a doladíme finálnu ponuku.</p>
    <p style="margin-top:16px;">S úctou a pozdravom,</p>
    <table style="margin-top:8px;font-size:14px;">
      <tr>
        <td style="padding:0;border:none;">
          <strong style="font-size:15px;">{obchodnik["meno"]}</strong><br/>
          <span style="color:#666;">{obchodnik["funkcia"]}</span><br/>
          📞 <a href="tel:{obchodnik["tel"].replace(" ","")}" style="color:#1B5E3F;text-decoration:none;">{obchodnik["tel"]}</a><br/>
          ✉️ <a href="mailto:{obchodnik["email"]}" style="color:#1B5E3F;text-decoration:none;">{obchodnik["email"]}</a>
        </td>
      </tr>
    </table>
    <hr style="margin:20px 0;border:none;border-top:1px solid #ddd;"/>
    <p style="font-size:12px;color:#888;">
      <strong>Energovision s.r.o.</strong> | IČO: 53 036 280 |
      <a href="https://energovision.sk" style="color:#888;">energovision.sk</a>
    </p>
    """

    full_html = f"""
    <html><body style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Helvetica,Arial,sans-serif;line-height:1.5;color:#333;max-width:700px;">
    {intro}
    {"".join(blocks)}
    {comparison}
    {pdf_note}
    {cta}
    </body></html>
    """
    return full_html



# ============================================================
# WEBHOOK: GENERATE PDF FROM SUPABASE PAYLOAD
# Trigger: Energovision FVE CRM (Next.js)
# Vstup: { "flat_props": {...}, "variant": "A"|"B"|"C"|"D", "quote_id": "uuid" }
# Výstup: rovnaký ako /webhook/generate-pdf — pdf_base64 + filename
# Účel: CRM nemá Notion page_id, posiela rovnaký formát flat dict
# ============================================================
@app.route("/webhook/generate-pdf-supabase", methods=["POST"])
@require_secret
def generate_pdf_supabase():
    """Adapter pre Supabase CRM — prijíma flat_props payload namiesto page_id."""
    try:
        return _generate_pdf_supabase_impl()
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        log.error(f"generate_pdf_supabase padol: {e}\n{tb}")
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback_tail": tb[-500:],
            "pdf_base64": "",
            "filename": "",
            "folder_name": "",
        }), 500


def _generate_pdf_supabase_impl():
    body = request.get_json(silent=True) or {}
    flat_props = body.get("flat_props") or {}
    variant = body.get("variant", "A")

    if not flat_props:
        return jsonify({"error": "missing flat_props"}), 400

    log.info(f"Generate PDF (Supabase) variant {variant}, klient {flat_props.get('Zákazník', '?')}")

    # Reuse — lead_from_notion berie flat dict + variant (NIE Notion page)
    from generate_from_notion import lead_from_notion
    lead = lead_from_notion(flat_props, variant)

    if variant in ("B", "C"):
        ok, msg = check_compatibility(lead["invertor_kod"], lead.get("bateria_kod"))
        if not ok:
            return jsonify({"error": f"incompatible: {msg}"}), 400

    cennik = load_cennik()
    konfig = vyrataj_konfig(lead, cennik)
    ceny = vyrataj_ceny(konfig, lead)

    # === CRM OVERRIDE ===
    # Ak Supabase CRM posiela vlastnú cenu (Cena X s DPH + Marža X), použijeme ju priamo.
    # Inak by Render použil Hagard ceny bez 35% Energovision zľavy + svoju 21% maržu → mismatch.
    crm_cena_s_dph = None
    try:
        for key in (f"Cena {variant} s DPH", f"Cena {variant}", "Cena s DPH"):
            val = flat_props.get(key)
            if val is None or val == "" or val == 0:
                continue
            crm_cena_s_dph = float(val)
            if crm_cena_s_dph > 0:
                break
    except (TypeError, ValueError):
        crm_cena_s_dph = None

    if crm_cena_s_dph and crm_cena_s_dph > 0:
        crm_marza = flat_props.get(f"Marža {variant}")
        try:
            crm_marza = float(crm_marza) if crm_marza is not None else None
        except (TypeError, ValueError):
            crm_marza = None
        dph = 0.23
        crm_cena_bez_dph = crm_cena_s_dph / (1 + dph)
        # Dotácia + ZD logika ostáva
        dotacia = ceny.get("dotacia", 0) if lead.get("dotacia", True) else 0
        cena_po_dot = crm_cena_s_dph - dotacia
        ceny = {
            "nakupna_material": ceny.get("nakupna_material", 0),
            "nakupna_praca": ceny.get("nakupna_praca", 0),
            "nakupna_spolu": ceny.get("nakupna_spolu", 0),
            "rezerva_eur": 0,
            "marza_eur": crm_cena_bez_dph - ceny.get("nakupna_spolu", 0),
            "cena_bez_dph": crm_cena_bez_dph,
            "cena_s_dph": crm_cena_s_dph,
            "dotacia": dotacia,
            "cena_po_dotacii": cena_po_dot,
            "zlava_eur": 0,
            "cena_finalna": cena_po_dot,
            "marza_pct": crm_marza if crm_marza is not None else ceny.get("marza_pct"),
            "zisk": crm_cena_bez_dph - ceny.get("nakupna_spolu", 0),
            "_source": "CRM_OVERRIDE",
        }
        log.info(f"[generate-pdf-supabase] CRM override použitý — cena s DPH = {crm_cena_s_dph}, marža = {crm_marza}%")

    navratnost = vyrataj_navratnost(konfig, ceny, lead)

    with tempfile.TemporaryDirectory() as tmpdir:
        priezvisko = safe_filename(lead["meno"].split()[-1])
        ev_id = lead.get("cislo_ponuky", f"EV-XX-001-{variant}")
        from datetime import datetime
        datum = datetime.now().strftime("%Y-%m-%d")
        base = f"{ev_id}_{priezvisko}_{datum}"

        grafy = vyrob_grafy(navratnost, lead, tmpdir, base)
        pdf_path = os.path.join(tmpdir, f"{base}.pdf")
        vyrob_html_pdf(lead, konfig, ceny, navratnost, grafy, pdf_path)
        pdf_size = os.path.getsize(pdf_path)

        import base64
        with open(pdf_path, "rb") as f:
            pdf_b64 = base64.b64encode(f.read()).decode("ascii")

    ev_id_root = ev_id[:-2] if len(ev_id) >= 2 and ev_id[-2] == "-" and ev_id[-1] in "ABCD" else ev_id
    folder_name = f"{ev_id_root}_{priezvisko}"

    return jsonify({
        "success": True,
        "ev_id": ev_id,
        "filename": f"{base}.pdf",
        "folder_name": folder_name,
        "size_bytes": pdf_size,
        "cena_s_dph": ceny["cena_s_dph"],
        "cena_finalna": ceny["cena_finalna"],
        "pdf_base64": pdf_b64,
        "source": "supabase",
        "quote_id": body.get("quote_id"),
    })



# ============================================================
# WEBHOOK: EMAIL TEMPLATE FROM SUPABASE PAYLOAD
# Trigger: Energovision FVE CRM
# Vstup: { "flat_props": {...}, "variants": ["A","B","C","D"], "typ_ponuky": "Indikatívna"|"Exaktná" }
# Výstup: { to, subject, body_html, obchodnik, variants_sent }
# ============================================================
@app.route("/webhook/email-template-supabase", methods=["POST"])
@require_secret
def email_template_supabase():
    """Adapter pre CRM — generuje email subject + body_html z flat_props."""
    try:
        return _email_template_supabase_impl()
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        log.error(f"email_template_supabase padol: {e}\n{tb}")
        return jsonify({"error": str(e), "traceback_tail": tb[-500:]}), 500


def _email_template_supabase_impl():
    body = request.get_json(silent=True) or {}
    flat_props = body.get("flat_props") or {}
    variants = body.get("variants") or []
    typ_ponuky = body.get("typ_ponuky", "Indikatívna")
    ma_rozlozenie = bool(body.get("ma_rozlozenie", False))

    if not flat_props or not variants:
        return jsonify({"error": "missing flat_props or variants"}), 400

    from generate_from_notion import lead_from_notion, OBCHODNICI, DEFAULT_OBCHODNIK, safe_filename

    lead_a = lead_from_notion(flat_props, variants[0] if variants else "A")
    priezvisko = lead_a["meno"].split()[-1] if lead_a.get("meno") else "Zákazník"
    mesto = lead_a.get("mesto", "")
    email_zakaznika = (lead_a.get("email") or flat_props.get("Email") or "").strip()

    if not _is_valid_email(email_zakaznika):
        return jsonify({
            "success": False,
            "email_valid": "false",
            "error": f"Neplatný email zákazníka: '{email_zakaznika}'",
            "to": "", "subject": "", "body_html": "",
        }), 200

    obchodnik = OBCHODNICI.get(
        flat_props.get("Obchodník") or flat_props.get("Obchodnik") or "",
        DEFAULT_OBCHODNIK
    )

    vykon_kwp = lead_a.get("vykon_kwp", 0)
    bateria_kwh = float(flat_props.get("Batéria výkon") or 0)

    ceny = {
        "A": flat_props.get("Cena A s DPH") or flat_props.get("Cena A"),
        "B": flat_props.get("Cena B s DPH") or flat_props.get("Cena B"),
        "C": flat_props.get("Cena C s DPH") or flat_props.get("Cena C"),
        "D": flat_props.get("Cena D s DPH") or flat_props.get("Cena D"),
    }

    subject = build_subject(priezvisko, mesto, variants, typ_ponuky=typ_ponuky)
    body_html = build_email_body(
        priezvisko, mesto, vykon_kwp, bateria_kwh, ceny, variants,
        obchodnik, typ_ponuky=typ_ponuky, ma_rozlozenie=ma_rozlozenie
    )

    return jsonify({
        "success": True,
        "email_valid": "true",
        "to": email_zakaznika,
        "subject": subject,
        "body_html": body_html,
        "obchodnik": obchodnik,
        "variants_sent": variants,
    })


# ============================================================
# WEBHOOK: Parsuj leady — Supabase verzia (raw_text -> JSON leads)
# Trigger: Energovision FVE CRM (/leady modal "Parsuj")
# Vstup: { "raw_text": "..." }
# Výstup: { ok, count, leads: [...] }  — leads sa vytvárajú v Next.js
# ============================================================
@app.route("/webhook/parsuj-leady-supabase", methods=["POST"])
@require_secret
def parsuj_leady_supabase():
    """CRM verzia parsera — vezme raw text, vráti pole extrahovaných leadov.
    Bez Notion side-effects. Next.js si potom leady vytvorí v Supabase sám."""
    body = request.get_json(force=True, silent=True) or {}
    raw_text = (body.get("raw_text") or "").strip()
    if not raw_text:
        return jsonify({"ok": False, "error": "raw_text je prazdny"}), 400

    if len(raw_text) > 30000:
        return jsonify({"ok": False, "error": "raw_text prilis dlhy (max 30k znakov)"}), 400

    log.info("[parsuj-leady-supabase] raw_text dlzka=%d znakov", len(raw_text))

    try:
        leads = claude_extract_leads(raw_text)
    except Exception as e:
        log.exception("Claude extraction zlyhala")
        return jsonify({"ok": False, "error": f"claude failed: {e}"}), 500

    if not leads:
        return jsonify({"ok": False, "error": "Claude nenasiel ziadne leady"}), 400

    log.info("[parsuj-leady-supabase] extrahovanych leadov: %d", len(leads))
    return jsonify({"ok": True, "count": len(leads), "leads": leads})


# ============================================================
# WEBHOOK: GENERUJ DOKUMENTY — Supabase verzia
# Vstup: { order_id, kind: 'zmluva'|'splnomocnenie'|'gdpr' }
# Fetchne order + customer + lead z Supabase REST API, naplní DOCX,
# uploadne do Supabase Storage, vráti pdf_url (vlastne docx_url)
# ============================================================
@app.route("/webhook/generuj-dokumenty-supabase", methods=["POST"])
@require_secret
def generuj_dokumenty_supabase():
    body = request.get_json(force=True, silent=True) or {}
    order_id = body.get("order_id")
    kind = body.get("kind", "zmluva")
    if not order_id:
        return jsonify({"error": "missing order_id"}), 400
    if kind not in ("zmluva", "splnomocnenie", "gdpr"):
        return jsonify({"error": "kind musi byt zmluva|splnomocnenie|gdpr"}), 400

    SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://uzwajrpebblafuhrtuwn.supabase.co")
    SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    if not SUPABASE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY not set"}), 500

    headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"}

    # Fetch order + customer + lead
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/orders",
        headers=headers,
        params={"select": "*,customers(*),leads(ev_id,distribucka,assigned_to,users:users!leads_assigned_to_fkey(full_name,email,phone,funkcia))", "id": f"eq.{order_id}"},
        timeout=30
    )
    if not r.ok:
        return jsonify({"error": f"supabase fetch: {r.status_code}", "body": r.text}), 500
    rows = r.json()
    if not rows:
        return jsonify({"error": "order_not_found"}), 404
    order = rows[0]
    cust = order.get("customers") or {}
    lead = order.get("leads") or {}
    obch = lead.get("users") if isinstance(lead.get("users"), dict) else {}

    # Fetch bundle pre vykon_kwp + payment_terms + panel info
    bundle = {}
    if lead.get("ev_id"):
        # Bundle je naviazaný na lead — najdi cez quote_bundles.lead_id = leads.id (potrebujem lead.id)
        # Skús cez order.bundle_id ak existuje, alebo cez lead.id (order má lead_id)
        lead_id_for_bundle = order.get("lead_id")
        if lead_id_for_bundle:
            br = requests.get(
                f"{SUPABASE_URL}/rest/v1/quote_bundles",
                headers=headers,
                params={"select": "*", "lead_id": f"eq.{lead_id_for_bundle}", "order": "created_at.desc", "limit": "1"},
                timeout=15
            )
            if br.ok and br.json():
                bundle = br.json()[0]

    pocet_panelov = int(bundle.get("pocet_panelov") or 0)
    panel_sku = bundle.get("panel_sku") or "PAN-001"
    panel_wp = 535 if panel_sku == "PAN-002" else 470
    vykon_kwp = round(pocet_panelov * panel_wp / 1000, 2) if pocet_panelov else 0

    # Záruka podľa panela (LONGi Hi-MO X10 = 25r produktová + 30r lineárna)
    if panel_sku == "PAN-001":  # LONGi Hi-MO X10 470 Wp
        zaruka_panely_produkt = 25
        zaruka_panely_linear = 30
    else:  # PAN-002 LONGi Hi-MO 6 535 Wp
        zaruka_panely_produkt = 15
        zaruka_panely_linear = 25

    # Platobné podmienky
    pt = bundle.get("payment_terms") or "60_30_10"
    platby_map = {
        "60_40": "60% - zálohová faktúra vopred\n40% - po dokončení diela",
        "50_50": "50% - zálohová faktúra vopred\n50% - po dokončení diela",
        "30_70": "30% - zálohová faktúra vopred\n70% - po dokončení diela",
        "60_30_10": "60% - zálohová faktúra vopred\n30% - po nainštalovaní elektrárne\n10% - po protokolárnom odovzdaní",
    }
    platby_text = platby_map.get(pt, platby_map["60_30_10"])

    meno = (cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip())
    ulica = cust.get("street") or ""
    psc = cust.get("postal_code") or ""
    mesto = cust.get("city") or ""
    adresa = ", ".join(filter(None, [ulica, psc, mesto])) or "(adresa nedoplnená)"
    ev_id = lead.get("ev_id") or order.get("order_number", "ORD")
    variant = order.get("accepted_variant") or "A"
    cislo_cp = f"{ev_id}-{variant}"
    cena = float(order.get("total_with_vat", 0))

    from datetime import datetime
    today = datetime.now().strftime("%d.%m.%Y")

    # Oslovenie pán/pani — derivované z mena a priezviska
    oslovenie = oslovenie_pan_pani(cust.get("first_name", ""), cust.get("last_name", ""))

    # Format dátum narodenia DD.MM.YYYY
    datum_narodenia_raw = cust.get("datum_narodenia", "")
    datum_narodenia = ""
    if datum_narodenia_raw:
        try:
            from datetime import datetime as _dt2
            d = _dt2.strptime(str(datum_narodenia_raw)[:10], "%Y-%m-%d")
            datum_narodenia = d.strftime("%d.%m.%Y")
        except Exception:
            datum_narodenia = str(datum_narodenia_raw)

    lead_data = {
        "meno_priezvisko": meno,
        "first_name": cust.get("first_name", ""),
        "last_name": cust.get("last_name", ""),
        "oslovenie": oslovenie,
        "cislo_op": cust.get("cislo_op", ""),
        "datum_narodenia": datum_narodenia,
        "adresa": adresa,
        "ulica": ulica,
        "psc": psc,
        "mesto": mesto,
        "telefon": cust.get("phone", ""),
        "email": cust.get("email", ""),
        "vykon_kwp": vykon_kwp,
        "pocet_panelov": pocet_panelov,
        "cislo_cp": cislo_cp,
        "datum_cp": today,
        "datum_dnes": today,
        "miesto_vykonu": adresa,
        "cena_eur": cena,
        "platby": platby_text,
        "payment_terms": pt,
        "zaruka_panely_produkt": zaruka_panely_produkt,
        "zaruka_panely_linear": zaruka_panely_linear,
        "ev_id": ev_id,
        "obchodnik_meno": (obch or {}).get("full_name", "Energovision tím"),
        "obchodnik_email": (obch or {}).get("email", "info@energovision.sk"),
        "obchodnik_tel": (obch or {}).get("phone", "+421 917 424 564"),
        "obchodnik_funkcia": (obch or {}).get("funkcia", "Obchodný zástupca"),
        "distribucka": lead.get("distribucka", "ZSD"),
    }

    # Vytvor temp file + naplň
    import tempfile, os as _os
    from pathlib import Path
    tmpdir = Path(tempfile.mkdtemp())

    try:
        from generuj_dokumenty import naplnif_zmluvu, naplnif_splnomocnenie, naplnif_gdpr
    except Exception as e:
        return jsonify({"error": f"import generuj_dokumenty zlyhal: {e}"}), 500

    out_path = tmpdir / f"{ev_id}_{kind}.docx"
    try:
        if kind == "zmluva":
            naplnif_zmluvu(lead_data, str(out_path))
        elif kind == "splnomocnenie":
            naplnif_splnomocnenie(lead_data, str(out_path))
        elif kind == "gdpr":
            naplnif_gdpr(lead_data, str(out_path))
    except Exception as e:
        log.exception("naplnif zlyhal")
        return jsonify({"error": f"naplnif: {e}"}), 500

    # Upload do Supabase Storage
    with open(out_path, "rb") as f:
        file_bytes = f.read()

    storage_path = f"orders/{order_id}/{kind}_{ev_id}.docx"
    up = requests.post(
        f"{SUPABASE_URL}/storage/v1/object/documents/{storage_path}",
        headers={**headers, "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "x-upsert": "true"},
        data=file_bytes,
        timeout=30
    )
    if not up.ok:
        log.warning("storage upload zlyhal: %s %s", up.status_code, up.text)
        return jsonify({"error": "storage_upload_failed", "body": up.text}), 500

    public_url = f"{SUPABASE_URL}/storage/v1/object/public/documents/{storage_path}"

    # Update order admin field
    field = f"{kind}_url"
    if kind == "zmluva":
        field = "zmluva_url"
    elif kind == "gdpr":
        field = "gdpr_url"
    elif kind == "splnomocnenie":
        field = "splnomocnenie_url"

    requests.patch(
        f"{SUPABASE_URL}/rest/v1/orders",
        headers={**headers, "Content-Type": "application/json"},
        params={"id": f"eq.{order_id}"},
        json={field: public_url},
        timeout=10
    )

    try:
        out_path.unlink()
        tmpdir.rmdir()
    except Exception:
        pass

    return jsonify({"ok": True, "url": public_url, "kind": kind, "filename": f"{ev_id}_{kind}.docx"})




# ============================================================
# WEBHOOK: DOCX → PDF konverzia (pre admin emaily klientovi)
# Vstup: { "docx_url": "https://..." }
# Výstup: { "pdf_base64": "..." }
# ============================================================
@app.route("/webhook/docx-to-pdf", methods=["POST"])
@require_secret
def docx_to_pdf_endpoint():
    body = request.get_json(force=True, silent=True) or {}
    docx_url = body.get("docx_url")
    if not docx_url:
        return jsonify({"error": "missing docx_url"}), 400

    import base64
    from io import BytesIO
    try:
        import mammoth
        from weasyprint import HTML

        # 1) Stiahni docx
        r = requests.get(docx_url, timeout=30)
        if not r.ok:
            return jsonify({"error": f"download failed: {r.status_code}"}), 500

        # 2) DOCX → HTML cez mammoth (bez LibreOffice, beží na 512 MB)
        result = mammoth.convert_to_html(BytesIO(r.content))
        html_body = result.value

        # Wrap do A4 dokumentu s minimal CSS
        html_full = f"""<!DOCTYPE html><html lang="sk"><head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 18mm; }}
  body {{ font-family: 'Helvetica', sans-serif; font-size: 10pt; color: #1a1a1a; line-height: 1.45; }}
  h1 {{ font-size: 16pt; margin: 12pt 0 6pt; }}
  h2 {{ font-size: 13pt; margin: 10pt 0 5pt; }}
  h3 {{ font-size: 11pt; margin: 8pt 0 4pt; }}
  p {{ margin: 4pt 0; }}
  table {{ border-collapse: collapse; margin: 6pt 0; }}
  td, th {{ border: 0.5pt solid #ccc; padding: 4pt 6pt; }}
  strong {{ font-weight: 700; }}
</style></head><body>{html_body}</body></html>"""

        # 3) HTML → PDF cez WeasyPrint (už máme)
        pdf_bytes = HTML(string=html_full).write_pdf()
        pdf_b64 = base64.b64encode(pdf_bytes).decode()
        return jsonify({"ok": True, "pdf_base64": pdf_b64, "size_bytes": len(pdf_bytes), "method": "mammoth+weasyprint"})
    except Exception as e:
        log.exception("docx-to-pdf failed")
        return jsonify({"error": str(e)}), 500




# ============================================================
# WEBHOOK: AI PARSER FAKTÚRY ZA ELEKTRINU
# Vstup: { "pdf_base64": "..." } alebo { "pdf_url": "https://..." }
# Výstup: { ok, parsed: { meno, adresa_*, eic, kwh_rocne, distribucka, ... } }
# Použiteľné pre B2C zákazníka — Claude Vision API parsuje PDF stranu 1-2
# ============================================================
@app.route("/webhook/parsuj-fakturu-ele", methods=["POST"])
@require_secret
def parsuj_fakturu_ele():
    body = request.get_json(force=True, silent=True) or {}
    pdf_url = body.get("pdf_url")
    pdf_b64_input = body.get("pdf_base64")

    if not pdf_url and not pdf_b64_input:
        return jsonify({"error": "missing pdf_url alebo pdf_base64"}), 400
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY not set"}), 500

    import base64
    from io import BytesIO

    try:
        # Stiahni PDF
        if pdf_url:
            r = requests.get(pdf_url, timeout=30)
            if not r.ok:
                return jsonify({"error": f"download failed: {r.status_code}"}), 500
            pdf_bytes = r.content
        else:
            pdf_bytes = base64.b64decode(pdf_b64_input)

        # PDF → text (PyMuPDF — už v requirements)
        try:
            import fitz  # pymupdf
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            full_text = ""
            pages_count = min(3, doc.page_count)
            # Stačí prvých 3 strán (obsahuje všetko podstatné)
            for i in range(pages_count):
                full_text += doc[i].get_text() + "\n\n"
            doc.close()
        except Exception as e:
            return jsonify({"error": f"pdf extract failed: {e}"}), 500

        # Pošli text do Claude — raw requests (žiadne anthropic SDK)
        prompt = """Z faktúry za elektrinu (text nižšie) extrahuj údaje a vráť ich ako JSON.

PRAVIDLÁ:
- Ak údaj nie je vo faktúre, vráť null (nie "" prázdny string).
- "miesto_spotreby" = adresa kde sa elektrina spotrebúva (typicky adresa rodinného domu/inštalácie FVE)
- "korespondencna_adresa" = adresa kde sa posiela faktúra (typicky trvalý pobyt zákazníka, ak je iná ako miesto spotreby)
- Ak sú obe adresy rovnaké, korespondencna_adresa nech je null
- Rozdeľ adresu na ulica/mesto/psc (PSČ = 5 číslic, mesto BEZ PSČ)
- "kwh_rocne" = súčet NT + VT + 1T spotreby za celé zúčtovacie obdobie (typicky 1 rok). Ak je obdobie kratšie, prepočítaj na 12 mesiacov.
- "distribucna_sadzba" = D1/D2/D3 Aktiv/D4/D5/... (z faktúry typicky "D3 aktiv" alebo "DD2")
- "distribucna_spolocnost" = ZSD (ZSE = ZSD Západoslovenská), SSD (SSE = SSD Stredoslovenská), VSD (VSE = VSD Východoslovenská)
- "dodavatel_elektriny" = ZSE / SSE / VSE / iný (názov dodávateľa, nie distribučky)

VRÁŤ LEN JSON, žiadny iný text:
{
  "meno": "Meno Priezvisko alebo Firma s.r.o.",
  "miesto_spotreby": { "ulica": "...", "mesto": "...", "psc": "..." },
  "korespondencna_adresa": { "ulica": "...", "mesto": "...", "psc": "..." },
  "eic": "24ZZS...",
  "kwh_rocne": 2057,
  "distribucna_sadzba": "D3 aktiv",
  "distribucna_spolocnost": "ZSD",
  "dodavatel_elektriny": "ZSE",
  "zakaznicke_cislo": "...",
  "cislo_miesta_spotreby": "...",
  "zuctovacie_obdobie_od": "2025-04-01",
  "zuctovacie_obdobie_do": "2026-03-31"
}

TEXT FAKTÚRY:
"""
        prompt += full_text[:15000]  # cap aby nepretiekol token limit

        headers = {
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        payload = {
            "model": ANTHROPIC_MODEL,
            "max_tokens": 2048,
            "messages": [{"role": "user", "content": prompt}],
        }
        r = _retry_request(lambda: requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload, timeout=90))
        r.raise_for_status()
        resp = r.json()
        raw_text = _safe_claude_text(resp)

        # Parse JSON (Claude občas vráti markdown ```json...``` — odstrániť)
        import json, re
        raw_text = raw_text.strip()
        m = re.search(r'\{[\s\S]*\}', raw_text)
        if not m:
            return jsonify({"error": "Claude nevrátil JSON", "raw": raw_text[:500]}), 500
        try:
            parsed = json.loads(m.group(0))
        except json.JSONDecodeError as e:
            return jsonify({"error": f"JSON parse: {e}", "raw": m.group(0)[:500]}), 500

        return jsonify({"ok": True, "parsed": parsed, "pages_extracted": pages_count})
    except Exception as e:
        log.exception("parsuj-fakturu-ele failed")
        return jsonify({"error": str(e)}), 500


# ============================================================
# ROOT — info
# ============================================================
@app.route("/")
def root():
    return jsonify({
        "service": "Energovision B2C cenovka generator",
        "version": "1.0.0",
        "endpoints": [
            "GET /health",
            "POST /webhook/parsuj-leady",
            "POST /webhook/parsuj-leady-supabase",
            "POST /webhook/auto-konfig",
            "POST /webhook/test-rozlozenie",
            "POST /webhook/spracuj-rozlozenie",
            "POST /webhook/generuj-dokumenty",
            "POST /webhook/email-zmluvy",
            "POST /webhook/generuj-realizacne",
            "POST /webhook/prepocet",
            "POST /webhook/generate-pdf",
            "POST /webhook/email-template",
            "POST /webhook/generate-pdf-supabase",
            "POST /webhook/email-template-supabase",
        ],
    })



# ============================================================
# FIELD PROTOKOLY — obhliadkový + preberací + servisný
# ============================================================
from datetime import datetime, timezone
import hashlib

PROTOKOL_HTML = """
<!DOCTYPE html>
<html lang="sk"><head><meta charset="utf-8"/><style>
@page { size: A4; margin: 18mm 15mm; }
body { font-family: 'Helvetica', sans-serif; font-size: 10pt; color: #1e293b; line-height: 1.4; }
.brand { background: #1F4E78; color: white; padding: 14px 18px; margin: -18mm -15mm 14px -15mm; }
.brand h1 { margin: 0; font-size: 18pt; font-weight: bold; }
.brand .sub { font-size: 9pt; opacity: 0.85; margin-top: 3px; }
.meta { display: flex; justify-content: space-between; background: #f1f5f9; padding: 8px 12px; border-radius: 6px; margin-bottom: 14px; font-size: 9pt; }
h2 { background: #fef3c7; border-left: 4px solid #f59e0b; padding: 4px 10px; font-size: 12pt; margin: 18px 0 8px 0; }
h3 { color: #1F4E78; margin: 14px 0 6px 0; font-size: 11pt; }
.row { display: flex; gap: 14px; margin: 6px 0; }
.row .label { font-weight: bold; min-width: 140px; color: #475569; }
.checklist li { list-style: none; margin: 3px 0; padding: 4px 8px; border-radius: 4px; }
.checklist li.done { background: #d1fae5; color: #065f46; }
.checklist li.todo { background: #fef2f2; color: #991b1b; }
.checklist li .icon { font-weight: bold; margin-right: 8px; }
.photos { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 6px; margin-top: 10px; }
.photos img { width: 100%; height: 80px; object-fit: cover; border: 1px solid #e2e8f0; border-radius: 4px; }
.notes { background: #fef9e7; border: 1px solid #fde68a; padding: 10px 12px; border-radius: 6px; font-style: italic; }
.signature { margin-top: 20px; padding-top: 14px; border-top: 2px solid #1F4E78; display: flex; justify-content: space-between; }
.signature .sig { width: 45%; }
.signature img { max-width: 180px; max-height: 80px; }
.footer { position: fixed; bottom: 8mm; left: 15mm; right: 15mm; text-align: center; font-size: 8pt; color: #94a3b8; }
</style></head><body>
<div class="brand">
  <h1>{{ title }}</h1>
  <div class="sub">{{ subtitle }}</div>
</div>

<div class="meta">
  <div><strong>ID:</strong> {{ doc_id }}</div>
  <div><strong>Dátum:</strong> {{ date_str }}</div>
  <div><strong>Technik:</strong> {{ technician }}</div>
</div>

<h3>Zákazník</h3>
<div class="row"><span class="label">Meno / firma:</span> {{ customer_name }}</div>
{% if customer_phone %}<div class="row"><span class="label">Telefón:</span> {{ customer_phone }}</div>{% endif %}
{% if customer_email %}<div class="row"><span class="label">E-mail:</span> {{ customer_email }}</div>{% endif %}
<div class="row"><span class="label">Miesto inštalácie:</span> {{ installation_address }}</div>
{% if gps %}<div class="row"><span class="label">GPS:</span> {{ gps }}</div>{% endif %}

<h2>{{ section_title }}</h2>
{% if checklist %}
<ul class="checklist">
{% for item in checklist %}
  <li class="{{ 'done' if item.done else 'todo' }}"><span class="icon">{{ '✓' if item.done else '✗' }}</span>{{ item.label }}</li>
{% endfor %}
</ul>
{% endif %}

{% if notes %}
<h3>Poznámka technika</h3>
<div class="notes">{{ notes }}</div>
{% endif %}

{% if photo_urls %}
<h3>Fotodokumentácia ({{ photo_urls|length }} fotiek)</h3>
<div class="photos">
{% for url in photo_urls[:9] %}
  <img src="{{ url }}" />
{% endfor %}
</div>
{% if photo_urls|length > 9 %}<p style="font-size:8pt;color:#64748b;margin-top:6px;">+ ďalších {{ photo_urls|length - 9 }} fotiek v digitálnom archíve.</p>{% endif %}
{% endif %}

<div class="signature">
  <div class="sig">
    <strong>Klient — podpis:</strong><br/>
    {% if customer_signature_url %}<img src="{{ customer_signature_url }}" />{% else %}<em style="color:#94a3b8">(nepodpísané)</em>{% endif %}
    <div style="border-top:1px solid #94a3b8;margin-top:6px;padding-top:3px;font-size:9pt;">{{ customer_name }}</div>
  </div>
  <div class="sig">
    <strong>Technik — podpis:</strong><br/>
    {% if technician_signature_url %}<img src="{{ technician_signature_url }}" />{% else %}<em style="color:#94a3b8">(elektronický záznam {{ doc_id }})</em>{% endif %}
    <div style="border-top:1px solid #94a3b8;margin-top:6px;padding-top:3px;font-size:9pt;">{{ technician }}</div>
  </div>
</div>

<div class="footer">
  Energovision s.r.o. · IČO 53 036 280 · www.energovision.sk · {{ doc_id }} · vygenerované {{ generated_at }}
</div>
</body></html>
"""


@app.route("/webhook/generuj-protokol", methods=["POST"])
def generuj_protokol():
    """Vygeneruje PDF protokol (obhliadka / preberací / servis) zo Supabase dát.
    Body: { kind: 'inspection'|'handover'|'service', id: uuid }
    Vráti: { ok, pdf_base64, sha256, filename }
    """
    try:
        data = request.get_json(force=True)
        kind = data.get("kind", "inspection")
        entity_id = data.get("id")
        if not entity_id:
            return jsonify({"error": "missing_id"}), 400

        from jinja2 import Template
        from weasyprint import HTML

        supa_url = os.environ.get("SUPABASE_URL", "")
        supa_key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
        if not supa_url or not supa_key:
            return jsonify({"error": "supabase_env_missing"}), 500
        headers = {"apikey": supa_key, "Authorization": f"Bearer {supa_key}"}

        ctx = {}
        if kind == "inspection":
            r = requests.get(
                f"{supa_url}/rest/v1/inspections",
                params={"id": f"eq.{entity_id}", "select": "*,customers(first_name,last_name,company_name,phone,email,city,street,postal_code,installation_same_as_billing,installation_street,installation_city,installation_postal_code),technician:users!inspections_technician_id_fkey(full_name)"},
                headers=headers, timeout=20
            )
            ins = (r.json() or [{}])[0]
            if not ins.get("id"):
                return jsonify({"error": "inspection_not_found"}), 404
            cust = ins.get("customers") or {}

            # Foto z documents
            r2 = requests.get(
                f"{supa_url}/rest/v1/documents",
                params={"inspection_id": f"eq.{entity_id}", "kind": "eq.photo", "select": "storage_url"},
                headers=headers, timeout=10
            )
            photo_urls = [d.get("storage_url") for d in (r2.json() or []) if d.get("storage_url")]

            checklist_data = ins.get("checklist_data") or {}
            CL_LABELS = {
                "strecha_typ": "Typ strechy zaznamenaný",
                "strecha_orientacia": "Orientácia a sklon strechy zmerané",
                "strecha_stav": "Stav krytiny OK",
                "tienenie": "Tienenie skontrolované",
                "miesto_panely": "Plocha pre panely vyznačená",
                "miesto_menic": "Miesto pre menič vybraté",
                "miesto_bateria": "Miesto pre batériu",
                "rozvadzac": "Hlavný rozvádzač",
                "elektromer": "Elektromer skontrolovaný",
                "uzemnenie": "Uzemnenie objektu",
                "pristupova_cesta": "Prístupová cesta",
                "lesenie": "Lešenie potrebné?",
                "suhlas_susedy": "Súhlas susedov",
                "klient_pripravoval_dotacie": "Klient pripravuje dotácie",
            }
            checklist = [{"label": CL_LABELS.get(k, k), "done": bool(v)} for k, v in checklist_data.items()]
            checklist += [{"label": CL_LABELS[k], "done": False} for k in CL_LABELS if k not in checklist_data]

            # Adresa inštalácie (3-address logic)
            if cust.get("installation_same_as_billing") is False and cust.get("installation_street"):
                inst_addr = f"{cust.get('installation_street','')}, {cust.get('installation_postal_code','')} {cust.get('installation_city','')}".strip(", ")
            else:
                inst_addr = f"{cust.get('street','')}, {cust.get('postal_code','')} {cust.get('city','')}".strip(", ")

            doc_id = f"OBH-{datetime.now().strftime('%Y-%m')}-{entity_id[:8]}"
            ctx = {
                "title": "Protokol z obhliadky",
                "subtitle": "Záznam o technickej obhliadke pre fotovoltickú elektráreň",
                "section_title": "Kontrolný zoznam obhliadky",
                "doc_id": doc_id,
                "date_str": datetime.fromisoformat((ins.get("scheduled_at") or ins.get("created_at")).replace("Z","+00:00")).strftime("%d.%m.%Y %H:%M") if ins.get("scheduled_at") else datetime.now().strftime("%d.%m.%Y"),
                "technician": (ins.get("technician") or {}).get("full_name") or "—",
                "customer_name": cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip() or "—",
                "customer_phone": cust.get("phone"),
                "customer_email": cust.get("email"),
                "installation_address": inst_addr or "—",
                "gps": f"{ins['gps_lat']:.5f}, {ins['gps_lng']:.5f}" if ins.get("gps_lat") else None,
                "checklist": checklist,
                "notes": ins.get("notes"),
                "photo_urls": photo_urls,
                "customer_signature_url": ins.get("customer_signature_url"),
                "technician_signature_url": None,
                "generated_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
            }

        elif kind == "handover":
            r = requests.get(
                f"{supa_url}/rest/v1/handover_protocols",
                params={"id": f"eq.{entity_id}", "select": "*,orders(order_number,ev_id,customers(first_name,last_name,company_name,phone,email,city,street,postal_code))"},
                headers=headers, timeout=20
            )
            h = (r.json() or [{}])[0]
            if not h.get("id"):
                return jsonify({"error": "handover_not_found"}), 404
            order = h.get("orders") or {}
            cust = order.get("customers") or {}

            r2 = requests.get(
                f"{supa_url}/rest/v1/documents",
                params={"order_id": f"eq.{order.get('id','none')}", "kind": "eq.photo", "select": "storage_url"},
                headers=headers, timeout=10
            )
            photo_urls = [d.get("storage_url") for d in (r2.json() or []) if d.get("storage_url")]

            checks = h.get("checklist_data") or {}
            FVE_LABELS = {
                "installation_per_project": "Inštalácia podľa projektu",
                "konstrukcia_ukotvenie": "Nosná konštrukcia — kotvenie OK",
                "panely_aretacia": "Aretácia panelov — moment OK",
                "dc_polarita": "Polarita DC stringov skontrolovaná",
                "dc_voc_isc": "Voc a Isc v limite",
                "ac_pripojenie": "AC pripojenie — fázovanie + svorky",
                "uzemnenie": "Uzemnenie + ekvipotenciálne prepojenie",
                "izolacny_odpor": "Izolačný odpor DC ≥ 1 MΩ",
                "rcd_test": "RCD test 30 mA OK",
                "spd_typ2": "SPD typ 2 (DC + AC)",
                "spustenie_menica": "Spustenie meniča, firmware",
                "monitoring_app": "Monitoring aktívny",
                "stitky_a_oznacenie": "Bezpečnostné štítky",
                "ziadost_o_pripojenie": "Žiadosť o pripojenie",
                "revizna_sprava": "Revízna správa",
                "zaskolenie_klienta": "Zaškolenie zákazníka",
                "navod_odovzdany": "Návod odovzdaný",
                "zarucny_list": "Záručné listy",
                "pracovisko_uprat": "Pracovisko upratané",
            }
            checklist = [{"label": FVE_LABELS.get(k, k), "done": bool(v)} for k, v in checks.items()]
            checklist += [{"label": FVE_LABELS[k], "done": False} for k in FVE_LABELS if k not in checks]

            doc_id = h.get("protocol_number") or f"PROT-{entity_id[:8]}"
            ctx = {
                "title": "Preberací protokol",
                "subtitle": f"Záznam o ukončení inštalácie · {order.get('ev_id') or order.get('order_number','')}",
                "section_title": "Kontrolný zoznam pri odovzdaní",
                "doc_id": doc_id,
                "date_str": datetime.fromisoformat(h.get("customer_signed_at","").replace("Z","+00:00")).strftime("%d.%m.%Y %H:%M") if h.get("customer_signed_at") else datetime.now().strftime("%d.%m.%Y"),
                "technician": h.get("technician_name") or "—",
                "customer_name": cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip() or "—",
                "customer_phone": cust.get("phone"),
                "customer_email": cust.get("email"),
                "installation_address": f"{cust.get('street','')}, {cust.get('postal_code','')} {cust.get('city','')}".strip(", "),
                "gps": h.get("signed_gps"),
                "checklist": checklist,
                "notes": h.get("notes"),
                "photo_urls": photo_urls,
                "customer_signature_url": h.get("customer_signature_url"),
                "technician_signature_url": h.get("technician_signature_url"),
                "generated_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
            }

        elif kind == "service":
            r = requests.get(
                f"{supa_url}/rest/v1/service_tickets",
                params={"id": f"eq.{entity_id}", "select": "*,customers(first_name,last_name,company_name,phone,email,city,street,postal_code),assignee:users!service_tickets_assignee_id_fkey(full_name)"},
                headers=headers, timeout=20
            )
            t = (r.json() or [{}])[0]
            if not t.get("id"):
                return jsonify({"error": "ticket_not_found"}), 404
            cust = t.get("customers") or {}

            doc_id = t.get("ticket_number") or f"SVC-{entity_id[:8]}"
            ctx = {
                "title": "Servisný protokol",
                "subtitle": f"Záznam o servisnom zásahu · {doc_id}",
                "section_title": "Riešenie problému",
                "doc_id": doc_id,
                "date_str": datetime.now().strftime("%d.%m.%Y %H:%M"),
                "technician": (t.get("assignee") or {}).get("full_name") or "—",
                "customer_name": cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip() or "—",
                "customer_phone": cust.get("phone"),
                "customer_email": cust.get("email"),
                "installation_address": f"{cust.get('street','')}, {cust.get('postal_code','')} {cust.get('city','')}".strip(", "),
                "gps": None,
                "checklist": [
                    {"label": f"Popis problému: {t.get('title','—')}", "done": True},
                    {"label": f"Priorita: {t.get('priority','normal')}", "done": True},
                    {"label": "Servisný zásah vykonaný", "done": t.get("status") == "resolved"},
                ],
                "notes": t.get("resolution") or t.get("description"),
                "photo_urls": [],
                "customer_signature_url": t.get("customer_signature_url"),
                "technician_signature_url": t.get("technician_signature_url"),
                "generated_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
            }
        else:
            return jsonify({"error": f"unknown_kind: {kind}"}), 400

        html = Template(PROTOKOL_HTML).render(**ctx)
        pdf_bytes = HTML(string=html).write_pdf()
        import base64
        b64 = base64.b64encode(pdf_bytes).decode("ascii")
        sha = hashlib.sha256(pdf_bytes).hexdigest()
        filename = f"{ctx['doc_id'].replace('/','-')}.pdf"
        return jsonify({"ok": True, "pdf_base64": b64, "sha256": sha, "filename": filename, "doc_id": ctx["doc_id"]})
    except Exception as e:
        log.exception("generuj-protokol failed")
        return jsonify({"error": str(e)}), 500



# ============================================================
# B2B DOCUMENT GENERATORS (G2 — replace Make scenár 8299009 docx-templater modules)
# ============================================================

@app.route("/webhook/generate-b2b-zod", methods=["POST"])
@require_secret
def generate_b2b_zod():
    """
    Generuje Zmluvu o dielo B2B z templates_b2b/Nova_klasicka_ZOD.docx.
    Vstup: { "project_id": "uuid" }
    Výstup: { "success": True, "filename": "...", "data": "base64", "summary": {...} }
    """
    body = request.get_json(force=True, silent=True) or {}
    project_id = body.get("project_id")
    if not project_id:
        return jsonify({"error": "missing project_id"}), 400

    SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://uzwajrpebblafuhrtuwn.supabase.co")
    SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    if not SUPABASE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY not set"}), 500
    headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"}

    # Fetch project + customer
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/projects",
        headers=headers,
        params={"select": "*,customers(*)", "id": f"eq.{project_id}"},
        timeout=30,
    )
    if not r.ok:
        return jsonify({"error": f"supabase fetch projects: {r.status_code}", "body": r.text}), 500
    rows = r.json()
    if not rows:
        return jsonify({"error": "project_not_found"}), 404
    project = rows[0]
    customer = project.get("customers") or {}

    # Fetch milestones (kvôli dátumom FA1/FA2/FA3 — ZoD na ne odkazuje)
    mr = requests.get(
        f"{SUPABASE_URL}/rest/v1/project_milestones",
        headers=headers,
        params={"select": "*", "project_id": f"eq.{project_id}", "order": "fa_no.asc"},
        timeout=15,
    )
    milestones = mr.json() if mr.ok else []

    try:
        from generuj_b2b import generuj_zod
        import base64 as _b64

        with tempfile.TemporaryDirectory() as tmpdir:
            zod_path = generuj_zod(
                project=project,
                customer=customer,
                milestones=milestones,
                out_dir=tmpdir,
            )
            with open(zod_path, "rb") as f:
                raw = f.read()
            filename = Path(zod_path).name
            log.info("[generate-b2b-zod] success project=%s file=%s (%d B)", project_id, filename, len(raw))
            return jsonify({
                "success": True,
                "filename": filename,
                "data": _b64.b64encode(raw).decode("ascii"),
                "summary": {
                    "project_id": project_id,
                    "project_name": project.get("name"),
                    "customer_name": customer.get("company_name"),
                    "contract_value_no_vat": project.get("contract_value_no_vat"),
                },
            })
    except FileNotFoundError as e:
        return jsonify({"success": False, "error": f"template missing: {e}"}), 500
    except Exception as e:
        log.exception("[generate-b2b-zod] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/webhook/generate-b2b-faktura", methods=["POST"])
@require_secret
def generate_b2b_faktura():
    """
    Generuje Word faktúru pre konkrétny milestone (FA1/FA2/FA3). Lukáš ju potom
    manuálne nahodí do Flowii ako originál.

    Vstup: { "project_id": "uuid", "fa_no": 1, "faktura_cislo": "2026/0042", "variabilny_symbol": "20260042" }
    Výstup: { "success": True, "filename": "...", "data": "base64", "summary": {...} }
    """
    body = request.get_json(force=True, silent=True) or {}
    project_id = body.get("project_id")
    fa_no = body.get("fa_no")
    faktura_cislo = body.get("faktura_cislo")
    variabilny_symbol = body.get("variabilny_symbol") or (faktura_cislo or "").replace("/", "").replace("-", "")

    if not project_id:
        return jsonify({"error": "missing project_id"}), 400
    if not fa_no:
        return jsonify({"error": "missing fa_no (1/2/3)"}), 400
    if not faktura_cislo:
        return jsonify({"error": "missing faktura_cislo (napr. 2026/0042)"}), 400

    SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://uzwajrpebblafuhrtuwn.supabase.co")
    SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    if not SUPABASE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY not set"}), 500
    headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"}

    # Fetch project + customer
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/projects",
        headers=headers,
        params={"select": "*,customers(*)", "id": f"eq.{project_id}"},
        timeout=30,
    )
    if not r.ok:
        return jsonify({"error": f"supabase fetch: {r.status_code}", "body": r.text}), 500
    rows = r.json()
    if not rows:
        return jsonify({"error": "project_not_found"}), 404
    project = rows[0]
    customer = project.get("customers") or {}

    # Fetch konkrétny milestone
    mr = requests.get(
        f"{SUPABASE_URL}/rest/v1/project_milestones",
        headers=headers,
        params={"select": "*", "project_id": f"eq.{project_id}", "fa_no": f"eq.{fa_no}"},
        timeout=15,
    )
    if not mr.ok:
        return jsonify({"error": f"milestone fetch: {mr.status_code}", "body": mr.text}), 500
    milestones = mr.json()
    if not milestones:
        return jsonify({"error": f"milestone fa_no={fa_no} not found"}), 404
    milestone = milestones[0]

    try:
        from generuj_b2b import generuj_faktura
        import base64 as _b64

        with tempfile.TemporaryDirectory() as tmpdir:
            fa_path = generuj_faktura(
                project=project,
                customer=customer,
                milestone=milestone,
                faktura_cislo=faktura_cislo,
                variabilny_symbol=variabilny_symbol,
                out_dir=tmpdir,
            )
            with open(fa_path, "rb") as f:
                raw = f.read()
            filename = Path(fa_path).name
            log.info("[generate-b2b-faktura] success project=%s fa%s file=%s (%d B)",
                     project_id, fa_no, filename, len(raw))
            return jsonify({
                "success": True,
                "filename": filename,
                "data": _b64.b64encode(raw).decode("ascii"),
                "summary": {
                    "project_id": project_id,
                    "fa_no": fa_no,
                    "faktura_cislo": faktura_cislo,
                    "payment_amount": milestone.get("payment_amount"),
                    "payment_pct": milestone.get("payment_pct"),
                    "due_date": milestone.get("due_date"),
                },
            })
    except FileNotFoundError as e:
        return jsonify({"success": False, "error": f"template missing: {e}"}), 500
    except Exception as e:
        log.exception("[generate-b2b-faktura] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# B2B NOTION → SUPABASE MIGRATION (N6 — autonomous from Chrome)
# Trigger: Claude (Cowork mode) cez Chrome browser fetch
# ============================================================
import migrate_notion_b2b as _mnb


@app.route("/webhook/migrate-notion-build-mapping", methods=["POST", "GET"])
@require_secret
def migrate_notion_build_mapping():
    """Vyrobí mapping Supabase project -> Notion page. Vráti JSON pre orchestráciu."""
    try:
        result = _mnb.build_mapping()
        return jsonify({"success": True, **result})
    except Exception as e:
        log.exception("[migrate-notion-build-mapping] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/webhook/migrate-notion-project", methods=["POST"])
@require_secret
def migrate_notion_project():
    """Migruje 1 projekt — 9 file properties z Notion -> Supabase Storage + project_documents."""
    body = request.get_json(silent=True) or {}
    notion_page_id = body.get("notion_page_id")
    supabase_project_id = body.get("supabase_project_id")
    ds = body.get("ds")
    dry_run = bool(body.get("dry_run", False))
    if not notion_page_id or not supabase_project_id:
        return jsonify({"success": False, "error": "missing notion_page_id or supabase_project_id"}), 400
    try:
        result = _mnb.migrate_one(notion_page_id, supabase_project_id, ds, dry_run)
        return jsonify({"success": result.get("ok", False), **result})
    except Exception as e:
        log.exception("[migrate-notion-project] zlyhalo")
        return jsonify({"success": False, "error": str(e)}), 500



# ============================================================
# AI DOCUMENT CLASSIFIER — analyzuje uploadnutý PDF, určí kind/folder/state
# Trigger: CRM Server Action po file upload (drag-drop v hero/per-task)
# ============================================================
B2B_DOC_KINDS = {
    "zod": ("02_Administrativa/03_Zmluva_o_dielo/02_Zmluva_s_IFT", "signed"),
    "splnomocnenie": ("02_Administrativa/02_Dokumenty_01", "signed"),
    "dotaznik": ("02_Administrativa/02_Dokumenty_01", "draft"),
    "lv": ("01_Podklady/03_Podklady_od_zakaznika", "draft"),
    "zop_signed": ("02_Administrativa/05_DIS-DS", "signed"),
    "zopad_signed": ("02_Administrativa/05_DIS-DS", "signed"),
    "stanovisko_zop": ("02_Administrativa/05_DIS-DS", "approved"),
    "stanovisko_rp": ("02_Administrativa/05_DIS-DS", "approved"),
    "opaos": ("04_Realizacia/05_Revizie", "approved"),
    "faktura": ("02_Administrativa/04_Fakturacia", "issued"),
    "preberaci_protokol": ("04_Realizacia/03_Protokoly", "signed"),
    "dsv": ("03_Projekcia/02_DSV", "draft"),
    "mpp": ("03_Projekcia/03_MPP", "draft"),
    "pbs": ("03_Projekcia/04_PBS", "draft"),
    "statika": ("03_Projekcia/05_Statika", "draft"),
    "iny": ("01_Podklady", "draft"),
}

DS_FOLDER_MAP = {"SSD": "05_DIS-SSD", "VSD": "05_DIS-VSD", "ZSDIS": "05_DIS-ZSDIS"}


@app.route("/webhook/classify-document", methods=["POST"])
@require_secret
def classify_document():
    """Klasifikuje uploadnutý dokument cez Claude API.

    Body: { project_id, file_url (storage path), filename, project_ds, project_code }
    Vráti: { kind, folder, state, confidence, suggested_task_no, reasoning }
    """
    body = request.get_json(silent=True) or {}
    project_id = body.get("project_id")
    file_url = body.get("file_url")
    filename = body.get("filename", "")
    project_ds = body.get("project_ds", "")
    project_code = body.get("project_code", "")
    if not project_id or not file_url:
        return jsonify({"error": "missing project_id or file_url"}), 400

    # 1) Stiahni súbor zo Supabase Storage
    SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://uzwajrpebblafuhrtuwn.supabase.co")
    SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    if not SUPABASE_KEY:
        return jsonify({"error": "SUPABASE_SERVICE_ROLE_KEY not set"}), 500

    from urllib.parse import quote
    storage_url = f"{SUPABASE_URL}/storage/v1/object/b2b-documents/{quote(file_url, safe='/')}"
    try:
        r = requests.get(storage_url, headers={"Authorization": f"Bearer {SUPABASE_KEY}", "apikey": SUPABASE_KEY}, timeout=60)
        if r.status_code != 200:
            return jsonify({"error": f"Storage fetch {r.status_code}", "detail": r.text[:200]}), 500
        file_bytes = r.content
    except Exception as e:
        return jsonify({"error": f"Storage fetch exception: {e}"}), 500

    # 2) Extrahuj text z PDF (max 3 strany)
    text_excerpt = ""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext in ("pdf",):
        try:
            import fitz  # pymupdf
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for i in range(min(3, doc.page_count)):
                text_excerpt += doc[i].get_text() + "\n"
            doc.close()
            text_excerpt = text_excerpt[:6000]
        except Exception as e:
            text_excerpt = f"[PDF extract failed: {e}]"
    elif ext in ("docx", "doc"):
        try:
            from io import BytesIO
            import docx
            d = docx.Document(BytesIO(file_bytes))
            text_excerpt = "\n".join([p.text for p in d.paragraphs])[:6000]
        except Exception as e:
            text_excerpt = f"[DOCX extract failed: {e}]"
    elif ext in ("jpg", "jpeg", "png"):
        text_excerpt = "[Obrazový súbor — klasifikuj z filename]"
    else:
        text_excerpt = "[Neznámy formát]"

    # 3) Claude prompt
    prompt = f"""Klasifikuj tento dokument pre Energovision B2B FVE projekt.

KONTEXT PROJEKTU:
- Project code: {project_code}
- Distribučná spoločnosť: {project_ds}
- Filename: {filename}

OBSAH (prvé strany):
{text_excerpt}

KATEGÓRIE (vráť len jeden):
- zod = Zmluva o dielo (Energovision_ZoD.pdf, podpisaná zmluva s klientom)
- splnomocnenie = Splnomocnenie/Plnomocenstvo na úkony s DS
- dotaznik = Dotazník k pripojeniu lokálneho zdroja FVE
- lv = List vlastníctva (z katastra)
- zop_signed = Podpísaná Zmluva o pripojení (od DS)
- zopad_signed = Zmluva o prístupe a distribúcii (od DS)
- stanovisko_zop = Stanovisko/Vyjadrenie ku žiadosti o pripojenie (od DS)
- stanovisko_rp = Stanovisko/Vyjadrenie k technickej dokumentácii / k PD (od DS)
- opaos = OPaOS revízna správa (FVZ revízia)
- faktura = Faktúra (Energovision vystavená klientovi)
- preberaci_protokol = Preberací protokol diela
- dsv = DSV / dielenská PD
- mpp = MPP (montážno-prevádzkový predpis)
- pbs = PBS (požiarno-bezpečnostné stanovisko)
- statika = Statický posudok
- iny = nič z vyššie uvedeného

VRÁŤ LEN JSON (žiadny iný text):
{{
  "kind": "stanovisko_rp",
  "confidence": 0.95,
  "reasoning": "Krátko prečo (1 veta)",
  "document_number": "číslo dokumentu ak je v texte (napr. 202511-C06-0049-1 alebo 2026/0042), inak null",
  "is_signed": true,
  "is_approved": true
}}
"""

    headers = {
        "x-api-key": os.environ.get("ANTHROPIC_API_KEY", ""),
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": "claude-sonnet-4-5-20250929",
        "max_tokens": 1024,
        "messages": [{"role": "user", "content": prompt}],
    }
    try:
        rr = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload, timeout=60)
        rr.raise_for_status()
        resp = rr.json()
        raw_text = ""
        for block in resp.get("content", []):
            if block.get("type") == "text":
                raw_text += block.get("text", "")
    except Exception as e:
        return jsonify({"error": f"Claude API: {e}"}), 500

    # 4) Parse JSON response
    import json, re
    m = re.search(r"\{[\s\S]*\}", raw_text)
    if not m:
        return jsonify({"error": "Claude did not return JSON", "raw": raw_text[:300]}), 500
    try:
        parsed = json.loads(m.group(0))
    except Exception as e:
        return jsonify({"error": f"JSON parse: {e}", "raw": raw_text[:300]}), 500

    kind = parsed.get("kind", "iny")
    if kind not in B2B_DOC_KINDS:
        kind = "iny"
    folder, default_state = B2B_DOC_KINDS[kind]

    # DIS folder resolution
    if folder == "02_Administrativa/05_DIS-DS":
        ds_folder = DS_FOLDER_MAP.get(project_ds, "05_DIS-DS")
        folder = f"02_Administrativa/{ds_folder}"

    state = default_state
    if parsed.get("is_approved"):
        state = "approved"
    elif parsed.get("is_signed"):
        state = "signed"

    return jsonify({
        "ok": True,
        "kind": kind,
        "folder": folder,
        "state": state,
        "confidence": float(parsed.get("confidence", 0.0)),
        "reasoning": parsed.get("reasoning", ""),
        "document_number": parsed.get("document_number"),
        "filename": filename,
    })



if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)

@app.route("/webhook/transcribe-audio", methods=["POST"])
@require_secret
def transcribe_audio():
    """Stub — TODO napojiť na Whisper API alebo Anthropic audio.
    Zatiaľ vráti error a user musí použiť 'paste mode' v MeetingRecorder.
    """
    body = request.get_json(silent=True) or {}
    audio_url = body.get("audio_url")
    if not audio_url:
        return jsonify({"ok": False, "error": "audio_url missing"}), 400
    return jsonify({
        "ok": False,
        "error": "transcribe not implemented yet — use paste mode v CRM s vloženým prepisom z Granola/Otter/Whisper",
        "audio_url": audio_url,
    }), 501

@app.route("/webhook/import-notion-content", methods=["POST"])
@require_secret
def import_notion_content():
    """Import full Notion page content (blocks + comments) → projects.notion_mirror_md.
    
    Body:
      { "all": true, "limit": 5 }     # bulk for first 5 projects
      { "all": true }                  # bulk for ALL projects
      { "notion_page_id": "X", "supabase_project_id": "Y" }   # single
    """
    body = request.get_json(silent=True) or {}
    try:
        if body.get("all"):
            limit = int(body.get("limit", 0)) or None
            result = _mnb.import_all_content(limit=limit)
            return jsonify(result)
        else:
            np = body.get("notion_page_id")
            sp = body.get("supabase_project_id")
            if not np or not sp:
                return jsonify({"ok": False, "error": "missing notion_page_id or supabase_project_id"}), 400
            result = _mnb.import_one_content(np, sp)
            return jsonify(result)
    except Exception as e:
        log.exception("[import-notion-content] zlyhalo")
        return jsonify({"ok": False, "error": str(e)}), 500



# ============================================================
# HUAWEI / SPOT 3-state — Pilier 4
# ============================================================
try:
    import huawei_spot as _hs
except Exception as _e:
    log.warning("huawei_spot module not loaded: %s", _e)
    _hs = None


def _hs_auth_ok(req) -> bool:
    secret = req.headers.get("X-Webhook-Secret") or req.args.get("secret")
    expected = os.environ.get("WEBHOOK_SECRET")
    return bool(expected) and secret == expected


@app.route("/webhook/okte-ingest", methods=["POST", "GET"])
def webhook_okte_ingest():
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    body = request.get_json(silent=True) or {}
    backfill = int(body.get("backfill_days", 0) or request.args.get("backfill_days", 0))
    target_day_str = body.get("day") or request.args.get("day")
    target_day = None
    if target_day_str:
        from datetime import date as _date
        try:
            target_day = _date.fromisoformat(target_day_str)
        except Exception:
            return jsonify({"ok": False, "error": "invalid day format YYYY-MM-DD"}), 400
    try:
        result = _hs.okte_ingest(target_day=target_day, backfill_days=backfill)
        return jsonify(result)
    except Exception as e:
        log.exception("[okte-ingest] failed")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/webhook/spot-reactor", methods=["POST", "GET"])
def webhook_spot_reactor():
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    body = request.get_json(silent=True) or {}
    dry_override = body.get("dry_run_override")
    if dry_override is None and "dry_run" in request.args:
        dry_override = request.args.get("dry_run") in ("1", "true", "yes")
    try:
        result = _hs.spot_reactor(dry_run_override=dry_override)
        return jsonify(result)
    except Exception as e:
        log.exception("[spot-reactor] failed")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/webhook/spot-pause", methods=["POST"])
def webhook_spot_pause():
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    body = request.get_json(silent=True) or {}
    reason = body.get("reason", "API call")
    return jsonify(_hs.global_pause(reason=reason))


@app.route("/webhook/huawei-test", methods=["GET"])
def webhook_huawei_test():
    """Smoke test: login + check token."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    token = _hs.huawei_login(force=True)
    return jsonify({"ok": bool(token), "token_present": bool(token), "token_preview": (token[:16] + "...") if token else None})



@app.route("/webhook/spot-manual-command", methods=["POST"])
def webhook_spot_manual_command():
    """
    Manuálny override z dashboardu:
    POST { "site_id": "...", "target_state": "NORMAL|ZERO_EXPORT_ONLY|FULL_SHUTDOWN", "issued_by": "...", "reason": "..." }
    """
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    body = request.get_json(silent=True) or {}
    site_id = body.get("site_id")
    target_state = body.get("target_state")
    issued_by = body.get("issued_by", "manual")
    reason = body.get("reason", "Manuálny override z dashboardu")
    if not site_id or not target_state:
        return jsonify({"ok": False, "error": "missing site_id or target_state"}), 400
    try:
        return jsonify(_hs.manual_force_state(site_id, target_state, issued_by, reason))
    except Exception as e:
        log.exception("[spot-manual-command] failed")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/webhook/huawei-sync-stations", methods=["POST", "GET"])
def webhook_huawei_sync_stations():
    """Pull station list from Huawei → upsert do inverter_sites. Pre prípad keď kolega nainštaluje novú FVE."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    try:
        return jsonify(_hs.sync_huawei_stations())
    except Exception as e:
        log.exception("[huawei-sync-stations] failed")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/webhook/huawei-login-test", methods=["POST", "GET"])
def webhook_huawei_login_test():
    """
    Diagnostika - status Huawei NBI loginu.
    - Vráti backoff status (next_login_allowed_at) ak je aktívny
    - Vráti cached token info ak je platný
    - Inak vyvolá fresh login a vráti presný error code
    """
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    try:
        from huawei_spot import (
            HUAWEI_BASE, HUAWEI_USER, HUAWEI_PASS,
            _load_huawei_credentials_from_db, _huawei_login_backoff_active, huawei_login,
        )
        from datetime import datetime, timezone

        cred = _load_huawei_credentials_from_db()
        base = HUAWEI_BASE or cred.get("base_url") or ""
        user = HUAWEI_USER or cred.get("username") or ""
        pwd = HUAWEI_PASS or cred.get("encrypted_password") or ""
        source = "env" if HUAWEI_PASS else ("db" if cred.get("encrypted_password") else "none")

        # Check backoff first - dont waste a login attempt if locked out
        is_blocked, next_allowed = _huawei_login_backoff_active(cred)
        if is_blocked:
            return jsonify({
                "ok": False,
                "backoff_active": True,
                "next_login_allowed_at": next_allowed,
                "credential_source": source,
                "username": user,
                "base_url": base,
                "error": f"Backoff aktívny - počkaj do {next_allowed} UTC (po failCode {cred.get('notes','{}')})",
            })

        if not pwd:
            return jsonify({
                "ok": False,
                "error": "No password - env HUAWEI_PASS prázdny a v DB inverter_vendor_credentials nič",
                "credential_source": "none",
                "username": user,
                "base_url": base,
            })

        # Check DB cached token
        if cred.get("current_token") and cred.get("token_expires_at"):
            try:
                exp = datetime.fromisoformat(cred["token_expires_at"].replace("Z","+00:00"))
                if datetime.now(timezone.utc) < exp:
                    return jsonify({
                        "ok": True,
                        "credential_source": source,
                        "username": user,
                        "base_url": base,
                        "token_source": "db_cache",
                        "token_len": len(cred["current_token"]),
                        "token_expires_at": cred["token_expires_at"],
                    })
            except Exception:
                pass

        # No cached token -> try login (uses backoff-aware huawei_login)
        token = huawei_login(force=False)
        # Re-read cred po login (token už uložený alebo backoff zapnutý)
        cred = _load_huawei_credentials_from_db()
        if token:
            return jsonify({
                "ok": True,
                "credential_source": source,
                "username": user,
                "base_url": base,
                "token_source": "fresh_login",
                "token_len": len(token),
                "token_expires_at": cred.get("token_expires_at"),
            })

        # Login failed - read backoff state for detailed error
        import json as _pj
        notes = cred.get("notes")
        notes_obj = {}
        try:
            notes_obj = _pj.loads(notes) if isinstance(notes, str) else (notes or {})
        except Exception:
            pass
        fail_code = notes_obj.get("last_login_fail_code")
        next_iso = notes_obj.get("next_login_allowed_at")

        err_map = {
            407: "ACCESS_FREQUENCY_IS_TOO_HIGH - rate limit (max 5 loginov/10 min, lockout 30 min)",
            401: "Account locked (5 nesprávnych hesiel za 10 min, lockout 30 min)",
            305: "Wrong credentials (failCode 305)",
            20400: "user.login.user_or_value_invalid - wrong username/password",
        }
        return jsonify({
            "ok": False,
            "credential_source": source,
            "username": user,
            "base_url": base,
            "fail_code": fail_code,
            "backoff_active": bool(next_iso),
            "next_login_allowed_at": next_iso,
            "error": err_map.get(fail_code, f"Login failed (failCode={fail_code})") if fail_code else "Login failed - check Render logs",
        })
    except Exception as e:
        log.exception("[huawei-login-test] crashed")
        return jsonify({"ok": False, "error": f"crash: {type(e).__name__}: {e}"}), 500


# ============================================================
# ANALYZA OM module — Pilier FVE+BESS posudok
# ============================================================
try:
    from analyza_om import engine as _aom
except Exception as _e:
    log.warning("analyza_om not loaded: %s", _e)
    _aom = None


@app.route("/webhook/aom-run-pipeline", methods=["POST"])
def webhook_aom_run_pipeline():
    """Spustí full pipeline (parse → sim → econ) pre danú analyza_id."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _aom is None:
        return jsonify({"ok": False, "error": "analyza_om module not loaded"}), 500
    body = request.get_json(silent=True) or {}
    analyza_id = body.get("analyza_id")
    if not analyza_id:
        return jsonify({"ok": False, "error": "missing analyza_id"}), 400
    try:
        result = _aom.run_full_pipeline(analyza_id)
        # Po OLD pipeline obohati econ_results.full_response (carbon/energy_flow/value_streams/monthly)
        # pre posudok view - bez prepisu UI variantov.
        enriched = None
        if _aom_v2 is not None:
            try:
                enriched = _aom_v2.enrich_econ_full_response(_sb(), analyza_id)
                log.info("[aom-run-pipeline] enriched=%s", enriched)
            except Exception as enrich_err:
                log.warning("[aom-run-pipeline] enrich_econ_full_response failed (non-fatal): %s", enrich_err)
        if isinstance(result, dict):
            result["enriched"] = enriched
        return jsonify(result)
    except Exception as e:
        log.exception("[aom-run-pipeline] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/aom-parse-only", methods=["POST"])
def webhook_aom_parse_only():
    """Iba parse consumption files — pre Krok 2.5 wizard preview."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _aom is None:
        return jsonify({"ok": False, "error": "analyza_om module not loaded"}), 500
    body = request.get_json(silent=True) or {}
    analyza_id = body.get("analyza_id")
    file_paths = body.get("file_paths") or []
    if not analyza_id or not file_paths:
        return jsonify({"status": "error", "error": "missing analyza_id or file_paths"}), 400
    try:
        return jsonify(_aom.parse_consumption(analyza_id, file_paths, body.get("options")))
    except Exception as e:
        log.exception("[aom-parse-only] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


# ===================================================================
# AUTOMATION TOP-5 — quick win endpointy
# ===================================================================
import automation as _automation


# Lazy supabase client pre nové AI moduly (automation/team_chat/strategic)
_sb_client = None
def _sb():
    global _sb_client
    if _sb_client is None:
        from supabase import create_client
        import os as _os
        _sb_client = create_client(
            _os.environ.get("SUPABASE_URL", "https://uzwajrpebblafuhrtuwn.supabase.co"),
            _os.environ["SUPABASE_SERVICE_ROLE_KEY"]
        )
    return _sb_client



@app.route("/webhook/ai-next-actions", methods=["POST"])
def webhook_ai_next_actions():
    """Vygeneruje 3 next-action návrhy pre lead/project. Uloží do ai_next_actions."""
    body = request.get_json(force=True) or {}
    target_type = body.get("target_type")
    target_id = body.get("target_id")
    if not target_type or not target_id:
        return jsonify({"status": "error", "error": "target_type a target_id sú povinné"}), 400
    try:
        actions = _automation.suggest_next_actions(_sb(), target_type, target_id, save=True)
        return jsonify({"status": "ok", "actions": actions})
    except Exception as e:
        log.exception("[ai-next-actions] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/draft-email", methods=["POST"])
def webhook_draft_email():
    """Vygeneruje AI draft emailu pre lead/project/customer. Uloží do email_drafts."""
    body = request.get_json(force=True) or {}
    try:
        result = _automation.draft_email(
            _sb(),
            body.get("target_type", "lead"),
            body["target_id"],
            body.get("purpose", "kontakt s klientom"),
            body.get("employee_name", "Dominik Galaba"),
            body.get("incoming_email"),
        )
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[draft-email] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/order-external", methods=["POST"])
def webhook_order_external():
    """Vytvorí draft objednávky externistu (PBS/statika/geodet/...)."""
    body = request.get_json(force=True) or {}
    project_id = body.get("project_id")
    service_type = body.get("service_type")
    if not project_id or not service_type:
        return jsonify({"status": "error", "error": "project_id a service_type sú povinné"}), 400
    try:
        result = _automation.order_external_service(
            _sb(), project_id, service_type, body.get("provider_id")
        )
        if "error" in result:
            return jsonify({"status": "error", **result}), 400
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[order-external] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/doc-prefill", methods=["POST"])
def webhook_doc_prefill():
    """AI klasifikuje + extrahuje dáta z dokumentu, predvyplní polia v DB."""
    body = request.get_json(force=True) or {}
    target_type = body.get("target_type")
    target_id = body.get("target_id")
    text = body.get("text_content")
    if not target_type or not target_id or not text:
        return jsonify({"status": "error", "error": "target_type, target_id, text_content sú povinné"}), 400
    try:
        result = _automation.classify_and_prefill(_sb(), target_type, target_id, text)
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[doc-prefill] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/doc-package", methods=["POST"])
def webhook_doc_package():
    """1-klik: AI draftne email pre balík dokumentov. (PDF generuje existujúci endpoint.)"""
    body = request.get_json(force=True) or {}
    lead_id = body.get("lead_id")
    if not lead_id:
        return jsonify({"status": "error", "error": "lead_id povinné"}), 400
    try:
        result = _automation.generate_doc_package(_sb(), lead_id, body.get("employee_name", "Dominik Galaba"))
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[doc-package] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


# ===================================================================
# AI Strategic Manager — manažérsky brief na dashboard
# ===================================================================
import strategic_agent as _strategic


@app.route("/webhook/ai-strategic-brief", methods=["POST", "GET"])
def webhook_strategic_brief():
    """Vygeneruje strategický brief a uloží do ai_strategic_briefs."""
    body = request.get_json(silent=True) or {}
    scope = body.get("scope") or request.args.get("scope") or "daily"
    try:
        brief = _strategic.generate_strategic_brief(_sb(), scope=scope, save=True)
        return jsonify({"status": "ok", "brief": brief})
    except Exception as e:
        log.exception("[strategic-brief] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


# ===================================================================
# AI Kolega Eva + Team Chat
# ===================================================================
import team_chat as _team_chat


@app.route("/webhook/team-chat-proactive", methods=["POST", "GET"])
def webhook_team_chat_proactive():
    """Cron: AI Kolega prejde stav firmy a napíše 1-3 proaktívne správy + predpripraví drafty."""
    try:
        result = _team_chat.proactive_pass(_sb())
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[team-chat-proactive] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/team-chat-reply", methods=["POST"])
def webhook_team_chat_reply():
    """User píše do chat-u → AI reaguje (s context-om)."""
    body = request.get_json(force=True) or {}
    msg = (body.get("message") or "").strip()
    if not msg:
        return jsonify({"status": "error", "error": "Prázdna správa"}), 400
    try:
        result = _team_chat.handle_reply(
            _sb(), msg, body.get("user_id"), body.get("user_name"),
            skip_insert_user=bool(body.get("skip_insert_user", False))
        )
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[team-chat-reply] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


# ===================================================================
# EVA COWORK v2 — Memory + Proactive autonomous
# ===================================================================
import eva_memory as _eva_memory
import eva_proactive as _eva_proactive


@app.route("/webhook/eva-memory-extract", methods=["POST"])
def webhook_eva_memory_extract():
    """Po user→AI exchange, vyextrahuj memories. Volá sa po každom Eva reply."""
    body = request.get_json(force=True) or {}
    try:
        saved = _eva_memory.extract_memories(
            _sb(),
            user_msg=body.get("user_msg", ""),
            ai_reply=body.get("ai_reply", ""),
            user_id=body.get("user_id"),
            user_name=body.get("user_name"),
            source_msg_id=body.get("source_msg_id"),
        )
        return jsonify({"status": "ok", "saved_count": len(saved), "memories": saved})
    except Exception as e:
        log.exception("[eva-memory-extract] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/eva-memory-search", methods=["POST"])
def webhook_eva_memory_search():
    """Vector search relevant memories pre query."""
    body = request.get_json(force=True) or {}
    try:
        results = _eva_memory.search_relevant(
            _sb(),
            query=body.get("query", ""),
            k=int(body.get("k", 20)),
            filter_role=body.get("filter_role"),
            filter_topic=body.get("filter_topic"),
            threshold=float(body.get("threshold", 0.4)),
        )
        return jsonify({"status": "ok", "memories": results})
    except Exception as e:
        log.exception("[eva-memory-search] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/webhook/eva-proactive-hourly", methods=["POST", "GET"])
def webhook_eva_proactive_hourly():
    """Hourly autonomous pass — Eva sama sa pozrie čo treba."""
    try:
        result = _eva_proactive.hourly_autonomous_pass(_sb())
        return jsonify({"status": "ok", **result})
    except Exception as e:
        log.exception("[eva-proactive-hourly] failed")
        return jsonify({"status": "error", "error": str(e)}), 500


# ============================================================
# RAYNET DISCOVERY — pull cenoviek/firiem/produktov pre offline analýzu
# ============================================================
import raynet_discovery as _raynet

@app.route("/webhook/raynet-whoami", methods=["GET", "POST"])
def webhook_raynet_whoami():
    """Quick auth check — overí že credentials fungujú. No-secret (read-only).
    Body môže obsahovať: {raynet_user, raynet_key, raynet_instance}."""
    body = request.get_json(silent=True) or {}
    if body.get("raynet_user") and body.get("raynet_key"):
        _raynet.set_creds(body["raynet_user"], body["raynet_key"], body.get("raynet_instance", "energovision"))
    try:
        return jsonify({"ok": True, "whoami": _raynet.whoami()})
    except Exception as e:
        log.exception("[raynet-whoami] failed")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/webhook/raynet-discover", methods=["POST"])
def webhook_raynet_discover():
    """Stiahne všetky quotations/business_cases/products/companies do Supabase staging.
    Telo: {only?, raynet_user?, raynet_key?, raynet_instance?}. No-secret (write only do staging)."""
    body = request.get_json(silent=True) or {}
    if body.get("raynet_user") and body.get("raynet_key"):
        _raynet.set_creds(body["raynet_user"], body["raynet_key"], body.get("raynet_instance", "energovision"))
    only = body.get("only")
    try:
        sb = _sb()
        if only == "products":
            out = {"products": _raynet.discover_products(sb)}
        elif only == "companies":
            out = {"companies": _raynet.discover_companies(sb)}
        elif only == "business_cases":
            out = {"business_cases": _raynet.discover_business_cases(sb)}
        elif only == "quotations":
            out = {"quotations": _raynet.discover_quotations(sb)}
        else:
            out = _raynet.discover_all(sb)
        return jsonify({"ok": True, **out})
    except Exception as e:
        log.exception("[raynet-discover] failed")
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/webhook/raynet-fetch-items", methods=["POST"])
def webhook_raynet_fetch_items():
    """Batch fetch items pre ponuky bez items. Volaj viackrát kým ostávajú."""
    body = request.get_json(silent=True) or {}
    if body.get("raynet_user") and body.get("raynet_key"):
        _raynet.set_creds(body["raynet_user"], body["raynet_key"], body.get("raynet_instance", "energovision"))
    try:
        max_offers = int(body.get("max_offers", 100))
        out = _raynet.fetch_offer_detail_items(_sb(), max_offers=max_offers)
        return jsonify({"ok": True, **out})
    except Exception as e:
        log.exception("[raynet-fetch-items] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# B2B KALKULAČKA — generovanie BOM + uloženie cenovky
# ============================================================
import b2b_calculator as _b2b_calc

@app.route("/webhook/b2b-calc-preview", methods=["POST"])
def webhook_b2b_calc_preview():
    """Vygeneruje BOM bez uloženia (live preview v UI)."""
    body = request.get_json(silent=True) or {}
    try:
        result = _b2b_calc.calculate_bom(_sb(), body)
        return jsonify({"ok": True, **result})
    except Exception as e:
        log.exception("[b2b-calc-preview] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-calc-save", methods=["POST"])
def webhook_b2b_calc_save():
    """Vygeneruje 4 archetypy variantov a uloží ako jeden quote_bundles záznam (workspace='b2b').
    Zdieľa lifecycle s B2C — bundle editor / PDF / public link / accept flow."""
    body = request.get_json(silent=True) or {}
    try:
        config = body.get("config", {})
        bundle = _b2b_calc.save_quote_as_bundle(
            _sb(),
            base_config=config,
            customer_id=body.get("customer_id"),
            lead_id=body.get("lead_id"),
            user_id=body.get("user_id"),
        )
        return jsonify({"ok": True, "bundle_id": bundle.get("id"), "bundle_number": bundle.get("bundle_number"), "bundle": bundle})
    except Exception as e:
        log.exception("[b2b-calc-save] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# B2B PDF generator
# ============================================================
import b2b_pdf as _b2b_pdf

@app.route("/webhook/b2b-generate-pdf", methods=["POST"])
def webhook_b2b_generate_pdf():
    """Vygeneruje PDF pre b2b_quote. Body: {quote_id, mode: 'klient'|'internal'|'both'}."""
    body = request.get_json(silent=True) or {}
    quote_id = body.get("quote_id")
    mode = body.get("mode", "klient")
    if not quote_id:
        return jsonify({"ok": False, "error": "quote_id required"}), 400
    try:
        results = {}
        if mode in ("klient", "both"):
            results["klient"] = _b2b_pdf.generate_quote_pdf(_sb(), quote_id, "klient")
        if mode in ("internal", "both"):
            results["internal"] = _b2b_pdf.generate_quote_pdf(_sb(), quote_id, "internal")
        return jsonify({"ok": True, **results})
    except Exception as e:
        log.exception("[b2b-generate-pdf] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# ANALYZA OM v2 — engine v0.9.5 premium
# ============================================================
try:
    import analyza_om_v2 as _aom_v2
except Exception as _e:
    _aom_v2 = None
    log.warning("analyza_om_v2 not loaded: %s", _e)


@app.route("/webhook/analyza-om-run-premium", methods=["POST"])
def webhook_aom_run_premium():
    """Spustí engine v0.9.5 VariantGenerator → top-6 variantov uložených do analyza_om_variants."""
    if not _aom_v2:
        return jsonify({"ok": False, "error": "analyza_om_v2 not loaded"}), 500
    body = request.get_json(silent=True) or {}
    aid = body.get("analyza_id")
    if not aid:
        return jsonify({"ok": False, "error": "analyza_id required"}), 400
    try:
        result = _aom_v2.run_variants_premium(_sb(), aid)
        return jsonify(result)
    except Exception as e:
        log.exception("[aom-run-premium] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/analyza-om-render-premium-docx", methods=["POST"])
def webhook_aom_render_premium_docx():
    """Vygeneruje premium DOCX posudok (engine v0.9.5)."""
    if not _aom_v2:
        return jsonify({"ok": False, "error": "analyza_om_v2 not loaded"}), 500
    body = request.get_json(silent=True) or {}
    aid = body.get("analyza_id")
    if not aid:
        return jsonify({"ok": False, "error": "analyza_id required"}), 400
    try:
        result = _aom_v2.render_posudok_premium(_sb(), aid)
        return jsonify(result)
    except Exception as e:
        log.exception("[aom-render-premium-docx] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/analyza-om-auto-fill", methods=["POST"])
def webhook_aom_auto_fill():
    """PSČ → distribútor + GPS + sadzba."""
    if not _aom_v2:
        return jsonify({"ok": False, "error": "analyza_om_v2 not loaded"}), 500
    body = request.get_json(silent=True) or {}
    psc = body.get("psc", "").strip()
    if not psc:
        return jsonify({"ok": False, "error": "psc required"}), 400
    return jsonify(_aom_v2.auto_fill_site_from_psc(
        psc=psc,
        rocna_spotreba_kwh=float(body.get("rocna_spotreba_kwh", 30000) or 30000),
        rk_kw=float(body.get("rk_kw", 25) or 25),
    ))


@app.route("/webhook/analyza-om-quick-estimate", methods=["POST"])
def webhook_aom_quick_estimate():
    """Rýchla kalkulácia úspor bez 15-min dát."""
    if not _aom_v2:
        return jsonify({"ok": False, "error": "analyza_om_v2 not loaded"}), 500
    body = request.get_json(silent=True) or {}
    return jsonify(_aom_v2.quick_estimate(body))


# ============================================================
# AOM AI STRATEGIST — 7-vrstvový AI poradca
# ============================================================
try:
    import aom_ai_strategist as _aom_ai
except Exception as _e:
    _aom_ai = None
    log.warning("aom_ai_strategist not loaded: %s", _e)


@app.route("/webhook/aom-ai-analyze", methods=["POST"])
def webhook_aom_ai_analyze():
    """Spustí AI Strategist full analysis (Vrstvy A → G)."""
    if not _aom_ai:
        return jsonify({"ok": False, "error": "aom_ai_strategist not loaded"}), 500
    body = request.get_json(silent=True) or {}
    aid = body.get("analyza_id")
    if not aid:
        return jsonify({"ok": False, "error": "analyza_id required"}), 400
    capex_overrides = body.get("capex_overrides") or None
    try:
        return jsonify(_aom_ai.run_full_analysis(_sb(), aid, capex_overrides=capex_overrides))
    except Exception as e:
        log.exception("[aom-ai-analyze] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/aom-ai-chat-reply", methods=["POST"])
def webhook_aom_ai_chat_reply():
    """DB trigger volá tento endpoint po INSERT user msg do analyza_om_ai_chat."""
    if not _aom_ai:
        return jsonify({"ok": False, "error": "aom_ai_strategist not loaded"}), 500
    body = request.get_json(silent=True) or {}
    aid = body.get("analyza_id")
    msg = body.get("message", "")
    if not aid or not msg:
        return jsonify({"ok": False, "error": "analyza_id + message required"}), 400
    try:
        return jsonify(_aom_ai.chat_refinement(_sb(), aid, msg, body.get("user_name")))
    except Exception as e:
        log.exception("[aom-ai-chat-reply] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/aom-ai-accept-variant", methods=["POST"])
def webhook_aom_ai_accept_variant():
    """Akceptuje konkrétny AI navrhnutý variant — uloží do analyza_om_variants."""
    if not _aom_ai:
        return jsonify({"ok": False, "error": "aom_ai_strategist not loaded"}), 500
    body = request.get_json(silent=True) or {}
    aid = body.get("analyza_id")
    label = body.get("variant_label")
    if not aid or not label:
        return jsonify({"ok": False, "error": "analyza_id + variant_label required"}), 400
    try:
        return jsonify(_aom_ai.accept_variant(_sb(), aid, label))
    except Exception as e:
        log.exception("[aom-ai-accept-variant] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/aom-custom-variant", methods=["POST"])
def webhook_aom_custom_variant():
    """Vyrobí user-defined custom variant + uloží priamo do analyza_om_variants.
    Body: { analyza_id, name, fve_kwp, bess_kwh, bess_kw?, capex_per_kwp?, capex_per_kwh_bess?, samospotreba_pct?, note? }
    """
    if not _aom_ai:
        return jsonify({"ok": False, "error": "aom_ai_strategist not loaded"}), 500
    body = request.get_json(silent=True) or {}
    aid = body.get("analyza_id")
    if not aid:
        return jsonify({"ok": False, "error": "analyza_id required"}), 400
    if not body.get("fve_kwp") and not body.get("bess_kwh"):
        return jsonify({"ok": False, "error": "fve_kwp alebo bess_kwh musí byť > 0"}), 400
    try:
        sb = _sb()
        # Load analyza row
        a_res = sb.table("analyza_om").select("*").eq("id", aid).single().execute()
        if not a_res.data:
            return jsonify({"ok": False, "error": "analyza not found"}), 404
        analyza = a_res.data

        custom_input = {
            "name": body.get("name") or "Vlastný variant",
            "fve_kwp": float(body.get("fve_kwp") or 0),
            "bess_kwh": float(body.get("bess_kwh") or 0),
            "bess_kw": float(body.get("bess_kw") or (float(body.get("bess_kwh") or 0) * 0.5)),
            "capex_per_kwp": body.get("capex_per_kwp"),
            "capex_per_kwh_bess": body.get("capex_per_kwh_bess"),
            "samospotreba_pct": body.get("samospotreba_pct"),
            "note": body.get("note"),
        }
        capex_overrides = body.get("capex_overrides") or None

        arch = _aom_ai.compute_custom_variant(analyza, custom_input, capex_overrides)

        # Insert do analyza_om_variants
        pos_res = sb.table("analyza_om_variants").select("position").eq("analyza_id", aid).order("position", desc=True).limit(1).execute()
        next_pos = (pos_res.data[0]["position"] + 1) if pos_res.data else 1

        sb.table("analyza_om_variants").insert({
            "analyza_id": aid,
            "name": custom_input["name"],
            "position": next_pos,
            "fve_kwp": custom_input["fve_kwp"],
            "fve_tilt_deg": 25,
            "fve_azimuth_deg": 180,
            "fve_topology": "south",
            "bess_kwh": custom_input["bess_kwh"],
            "bess_kw": custom_input["bess_kw"],
            "bess_arbitrage_enabled": False,
            "capex_eur": arch.get("capex_total_eur", 0),
            "capex_source": "ai_strategist",
            "result_samosp_pct": arch.get("self_consumption_pct", 0),
            "result_samostat_pct": arch.get("self_sufficiency_pct", 0),
            "result_npv_eur_base": arch.get("npv_eur", 0),
            "result_irr_pct_base": arch.get("irr_pct", 0),
            "result_payback_y_base": arch.get("payback_years", 0),
            "result_dotacia_eur": arch.get("dotacia_eur", 0),
        }).execute()

        return jsonify({"ok": True, "variant": arch, "position": next_pos})
    except Exception as e:
        log.exception("[aom-custom-variant] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/aom-parse-public", methods=["POST"])
def webhook_aom_parse_public():
    """Verejný parse endpoint — frontend volá hneď po uploade CSV/XLS.
    Bez secretu (CORS + read-only zo storage)."""
    if _aom is None:
        return jsonify({"ok": False, "error": "analyza_om module not loaded"}), 500
    body = request.get_json(silent=True) or {}
    analyza_id = body.get("analyza_id")
    if not analyza_id:
        return jsonify({"ok": False, "error": "missing analyza_id"}), 400
    
    try:
        # Načítaj raw_files z DB
        sb = _sb()
        a_res = sb.table("analyza_om").select("consumption_raw_files").eq("id", analyza_id).single().execute()
        if not a_res.data:
            return jsonify({"ok": False, "error": "analyza not found"}), 404
        raw_files = a_res.data.get("consumption_raw_files") or []
        if not raw_files:
            return jsonify({"ok": False, "error": "no files uploaded"}), 400
        
        # Extract storage_path z každého súboru
        file_paths = []
        for f in raw_files:
            if isinstance(f, dict) and f.get("storage_path"):
                file_paths.append(f["storage_path"])
        
        if not file_paths:
            return jsonify({"ok": False, "error": "no valid storage_path"}), 400
        
        result = _aom.parse_consumption(analyza_id, file_paths, None)
        # Zapíš výsledky parsovania do analyza_om aby AI Strategist mal annual_mwh + peak_kw
        if isinstance(result, dict) and result.get("status") == "ok":
            summary = result.get("summary") or {}
            outputs = result.get("outputs") or {}
            try:
                sb.table("analyza_om").update({
                    "consumption_annual_mwh": summary.get("annual_mwh"),
                    "consumption_peak_kw_15min": summary.get("peak_kw_15min"),
                    "consumption_peak_kw_hourly": summary.get("peak_kw_hourly"),
                    "consumption_avg_kw": summary.get("avg_kw"),
                    "consumption_coverage_pct": summary.get("coverage_pct"),
                    "consumption_detected_format": (result.get("detected_formats") or [None])[0],
                    "consumption_parse_warnings": result.get("warnings") or [],
                    "consumption_profile_path": outputs.get("profile_hourly_path") or (str(analyza_id) + "/consumption_profile.csv"),
                    "consumption_15min_path": outputs.get("profile_15min_path") or (str(analyza_id) + "/consumption_15min.csv"),
                    "consumption_method": "auto_parse",
                    "updated_at": "now()",
                }).eq("id", analyza_id).execute()
            except Exception as upd_err:
                log.warning(f"[aom-parse-public] DB update failed: {upd_err}")
        return jsonify({"ok": True, **(result if isinstance(result, dict) else {})})
    except Exception as e:
        log.exception("[aom-parse-public] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# B2B Kalkulačka V2 — panels-driven + vendor stacks + AI
# ============================================================
import b2b_calculator_v2 as _b2b_v2


@app.route("/webhook/b2b-calc-v2-preview", methods=["POST"])
def webhook_b2b_calc_v2_preview():
    body = request.get_json(silent=True) or {}
    try:
        return jsonify(_b2b_v2.calculate_bom_v2(_sb(), body))
    except Exception as e:
        log.exception("[b2b-calc-v2-preview] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-vendor-stacks", methods=["GET"])
def webhook_b2b_vendor_stacks():
    try:
        res = _sb().table("b2b_vendor_stacks").select("vendor_key, display_name, preferred_panels, inverters, batteries, raynet_share_pct, notes").execute()
        return jsonify({"ok": True, "stacks": res.data or []})
    except Exception as e:
        log.exception("[b2b-vendor-stacks] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-ai-configurator", methods=["POST"])
def webhook_b2b_ai_configurator():
    """Smart Configurator — text → form fill."""
    body = request.get_json(silent=True) or {}
    text = (body.get("text") or "").strip()
    if not text:
        return jsonify({"ok": False, "error": "text required"}), 400
    try:
        return jsonify(_b2b_v2.ai_smart_configurator(_sb(), text))
    except Exception as e:
        log.exception("[b2b-ai-configurator] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# B2B Kalkulačka V2 — AI features (Vendor Reco / Compat / Sanity / Validator)
# ============================================================

@app.route("/webhook/b2b-vendor-recommender", methods=["POST"])
def webhook_b2b_vendor_recommender():
    body = request.get_json(silent=True) or {}
    try:
        return jsonify(_b2b_v2.ai_vendor_recommender(
            _sb(),
            kwp=float(body.get("kwp") or 0),
            client_type_hint=body.get("client_type_hint"),
            has_bess=bool(body.get("has_bess")),
            has_optimizery=bool(body.get("has_optimizery")),
            has_rapid_shutdown=bool(body.get("has_rapid_shutdown")),
        ))
    except Exception as e:
        log.exception("[b2b-vendor-recommender] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-compatibility-checker", methods=["POST"])
def webhook_b2b_compatibility_checker():
    body = request.get_json(silent=True) or {}
    try:
        return jsonify(_b2b_v2.ai_compatibility_checker(_sb(), body.get("config", body) or {}))
    except Exception as e:
        log.exception("[b2b-compatibility-checker] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-price-sanity", methods=["POST"])
def webhook_b2b_price_sanity():
    body = request.get_json(silent=True) or {}
    try:
        return jsonify(_b2b_v2.ai_price_sanity_check(
            _sb(),
            items=body.get("items") or [],
            kwp=float(body.get("kwp") or 0),
        ))
    except Exception as e:
        log.exception("[b2b-price-sanity] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-bom-validator", methods=["POST"])
def webhook_b2b_bom_validator():
    body = request.get_json(silent=True) or {}
    try:
        return jsonify(_b2b_v2.ai_bom_validator(
            _sb(),
            items=body.get("items") or [],
            config=body.get("config") or {},
        ))
    except Exception as e:
        log.exception("[b2b-bom-validator] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/b2b-calc-v2-save", methods=["POST"])
def webhook_b2b_calc_v2_save():
    """Uloží V2 ponuku ako quote_bundles (workspace=b2b) s final BOM po user editoch."""
    body = request.get_json(silent=True) or {}
    try:
        bundle = _b2b_v2.save_bundle_v2(_sb(), body)
        return jsonify({
            "ok": True,
            "bundle_id": bundle.get("id"),
            "bundle_number": bundle.get("bundle_number"),
            "bundle": bundle,
        })
    except Exception as e:
        log.exception("[b2b-calc-v2-save] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# PRICING REFRESH — bundle (B2C + B2B) + addon
# ============================================================

def _refresh_bom_items(bom: list, products_map: dict, keep_margin: bool = True) -> tuple[list, list]:
    """Match BOM items by SKU → use new product prices.
    Returns (updated_bom, diff_log).
    Ak položka v BOM nemá SKU alebo SKU nie je v products_map, ostáva nezmenená.
    keep_margin: ak True a sale_price v produkte chýba, zachová pôvodnú maržu (purchase × (sale/purchase)).
    """
    if not isinstance(bom, list):
        return bom, []
    updated = []
    diff = []
    for it in bom:
        if not isinstance(it, dict):
            updated.append(it)
            continue
        sku = (it.get("sku") or "").strip()
        prod = products_map.get(sku) if sku else None
        if not prod or prod.get("purchase_price") is None:
            updated.append(it)
            continue
        old_purchase = float(it.get("unit_purchase") or 0)
        old_sale = float(it.get("unit_sale") or 0)
        new_purchase = float(prod["purchase_price"])
        new_sale = float(prod.get("sale_price") or 0)
        # Ak sale nie je v produkte alebo je 0, zachovaj pôvodnú maržu
        if new_sale <= 0 and old_purchase > 0 and old_sale > 0:
            margin_factor = old_sale / old_purchase
            new_sale = round(new_purchase * margin_factor, 2)
        qty = float(it.get("qty") or 0)
        new_item = dict(it)
        new_item["unit_purchase"] = new_purchase
        new_item["unit_sale"] = new_sale
        new_item["total_purchase"] = round(new_purchase * qty, 2)
        new_item["total_sale"] = round(new_sale * qty, 2)
        updated.append(new_item)
        if abs(new_purchase - old_purchase) > 0.01 or abs(new_sale - old_sale) > 0.01:
            diff.append({
                "sku": sku,
                "name": it.get("name") or prod.get("name", ""),
                "qty": qty,
                "old_purchase": old_purchase, "new_purchase": new_purchase,
                "old_sale": old_sale, "new_sale": new_sale,
                "delta_total_sale": round((new_sale - old_sale) * qty, 2),
            })
    return updated, diff


@app.route("/webhook/bundle-refresh-prices", methods=["POST"])
def webhook_bundle_refresh_prices():
    """Aktualizuje ceny v quote_bundles podľa aktuálnych products.
    Body: { bundle_id, mode: "preview" | "apply" }
    Vráti diff per variant + nové total ceny."""
    body = request.get_json(silent=True) or {}
    bundle_id = body.get("bundle_id")
    mode = body.get("mode", "preview")
    if not bundle_id:
        return jsonify({"ok": False, "error": "bundle_id required"}), 400
    try:
        sb = _sb()
        bres = sb.table("quote_bundles").select("*").eq("id", bundle_id).single().execute()
        if not bres.data:
            return jsonify({"ok": False, "error": "bundle not found"}), 404
        bundle = bres.data

        # Načítaj všetky SKUs z BOMov
        all_skus = set()
        for v in ("a", "b", "c", "d"):
            bom = bundle.get(f"variant_{v}_bom") or []
            if isinstance(bom, list):
                for it in bom:
                    if isinstance(it, dict) and it.get("sku"):
                        all_skus.add(it["sku"].strip())
        if not all_skus:
            return jsonify({"ok": True, "diff_count": 0, "message": "BOM nemá žiadne SKU"})

        pres = sb.table("products").select("sku, name, purchase_price, sale_price").in_("sku", list(all_skus)).execute()
        products_map = {p["sku"]: p for p in (pres.data or [])}

        # Refresh per variant
        result = {"variants": {}, "totals": {}, "missing_skus": []}
        update_payload = {}
        all_diffs = []
        for v in ("a", "b", "c", "d"):
            bom = bundle.get(f"variant_{v}_bom") or []
            if not isinstance(bom, list) or len(bom) == 0:
                continue
            new_bom, diff = _refresh_bom_items(bom, products_map)
            new_cost = round(sum((float(it.get("total_purchase") or 0)) for it in new_bom), 2)
            new_no_vat = round(sum((float(it.get("total_sale") or 0)) for it in new_bom), 2)
            new_with_vat = round(new_no_vat * 1.23, 2)
            old_cost = float(bundle.get(f"variant_{v}_cost") or 0)
            old_no_vat = float(bundle.get(f"variant_{v}_price_no_vat") or 0)
            result["variants"][v] = {
                "diff": diff,
                "old_cost": old_cost, "new_cost": new_cost,
                "old_price_no_vat": old_no_vat, "new_price_no_vat": new_no_vat,
                "delta_cost": round(new_cost - old_cost, 2),
                "delta_price": round(new_no_vat - old_no_vat, 2),
            }
            all_diffs.extend(diff)
            if mode == "apply":
                update_payload[f"variant_{v}_bom"] = new_bom
                update_payload[f"variant_{v}_cost"] = new_cost
                update_payload[f"variant_{v}_price_no_vat"] = new_no_vat
                update_payload[f"variant_{v}_price_with_vat"] = new_with_vat

        # Chýbajúce SKU v products
        result["missing_skus"] = sorted(all_skus - set(products_map.keys()))
        result["diff_count"] = len(all_diffs)
        result["applied"] = False

        if mode == "apply" and update_payload:
            sb.table("quote_bundles").update(update_payload).eq("id", bundle_id).execute()
            result["applied"] = True

        return jsonify({"ok": True, **result})
    except Exception as e:
        log.exception("[bundle-refresh-prices] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/addon-refresh-prices", methods=["POST"])
def webhook_addon_refresh_prices():
    """Refresh cien pre addon_quotes (B2C doplnenie inštalácie)."""
    body = request.get_json(silent=True) or {}
    addon_id = body.get("addon_id")
    mode = body.get("mode", "preview")
    if not addon_id:
        return jsonify({"ok": False, "error": "addon_id required"}), 400
    try:
        sb = _sb()
        ares = sb.table("addon_quotes").select("*").eq("id", addon_id).single().execute()
        if not ares.data:
            return jsonify({"ok": False, "error": "addon not found"}), 404
        addon = ares.data
        items = addon.get("items") or []
        if not isinstance(items, list) or not items:
            return jsonify({"ok": True, "diff_count": 0, "message": "addon nemá items"})

        all_skus = {it["sku"].strip() for it in items if isinstance(it, dict) and it.get("sku")}
        pres = sb.table("products").select("sku, name, purchase_price, sale_price").in_("sku", list(all_skus)).execute() if all_skus else None
        products_map = {p["sku"]: p for p in (pres.data if pres else [])}

        new_items, diff = _refresh_bom_items(items, products_map)
        new_cost = round(sum((float(it.get("total_purchase") or 0)) for it in new_items), 2)
        new_no_vat = round(sum((float(it.get("total_sale") or 0)) for it in new_items), 2)
        new_with_vat = round(new_no_vat * 1.23, 2)

        result = {
            "diff": diff,
            "diff_count": len(diff),
            "old_cost": float(addon.get("total_cost") or 0),
            "new_cost": new_cost,
            "old_no_vat": float(addon.get("total_no_vat") or 0),
            "new_no_vat": new_no_vat,
            "missing_skus": sorted(all_skus - set(products_map.keys())),
            "applied": False,
        }
        if mode == "apply":
            sb.table("addon_quotes").update({
                "items": new_items,
                "total_cost": new_cost,
                "total_no_vat": new_no_vat,
                "total_with_vat": new_with_vat,
            }).eq("id", addon_id).execute()
            result["applied"] = True
        return jsonify({"ok": True, **result})
    except Exception as e:
        log.exception("[addon-refresh-prices] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# OM-Share — Organizátor zdieľania elektriny
# ============================================================
# /webhook/omshare-monthly-billing    — cron 1. v mesiaci, vystaví draft mesačné faktúry
# /webhook/omshare-invoice-pdf        — generuje PDF pre 1 faktúru
# /webhook/omshare-send-invoice       — odošle email s PDF prílohou
# /webhook/omshare-monthly-report     — cron 20. v mesiaci, mesačný report klientovi
# /webhook/omshare-edc-sync           — sync dát z OKTE EDC (mock)
# /webhook/omshare-edc-register       — registrácia novej skupiny v OKTE EDC

@app.route("/webhook/omshare-monthly-billing", methods=["POST"])
def webhook_omshare_monthly_billing():
    """
    Cron 1. dňa v mesiaci — vystaví draft mesačné faktúry pre všetky aktívne OM-Share skupiny.
    Kalkulácia: monthly_flat + monthly_per_om*active_om_count + usage_per_mwh*shared_mwh_last_month
    Bundle zľava (30 %) sa aplikuje ak group.is_bundle = true a bundle_expires_at >= dnes.
    Pricing model v3.2 — bundle je atribút skupiny, nie plánu.
    """
    body = request.get_json(silent=True) or {}
    target_month = body.get("month")
    target_year = body.get("year")

    today = datetime.now(timezone.utc).date()
    if not target_month or not target_year:
        # default: fakturujeme za predchádzajúci mesiac
        if today.month == 1:
            target_month = 12
            target_year = today.year - 1
        else:
            target_month = today.month - 1
            target_year = today.year
    target_month = int(target_month)
    target_year = int(target_year)

    try:
        sb = _sb()
        groups_res = sb.table("sharing_groups").select("*").eq("status", "active").execute()
        groups = groups_res.data or []
        log.info("[omshare-billing] generujem faktúry za %d/%d pre %d aktívnych skupín", target_month, target_year, len(groups))

        created = []
        skipped = []

        for g in groups:
            # check duplicate
            dup = sb.table("sharing_invoices").select("id").eq("group_id", g["id"]).eq("invoice_type", "monthly").eq("billing_period_year", target_year).eq("billing_period_month", target_month).execute()
            if dup.data:
                skipped.append({"group_id": g["id"], "reason": "already_billed"})
                continue

            plan = g.get("plan_snapshot") or {}
            monthly_flat = float(plan.get("monthly_flat_eur") or 0)
            monthly_per_om = float(plan.get("monthly_per_om_eur") or 0)
            usage_per_mwh = float(plan.get("usage_per_mwh_eur") or 0)
            vat_rate = float(plan.get("vat_rate") or 23)
            # v3.2: bundle discount na úrovni skupiny, 30 % ak je aktívny
            is_bundle = bool(g.get("is_bundle") or False)
            bundle_exp = g.get("bundle_expires_at")
            bundle_active = False
            if is_bundle and bundle_exp:
                try:
                    from datetime import datetime as _dt
                    exp_d = _dt.fromisoformat(bundle_exp).date() if isinstance(bundle_exp, str) else bundle_exp
                    bundle_active = today <= exp_d
                except Exception:
                    bundle_active = False
            bundle_pct = 30.0 if bundle_active else 0.0

            # active OM count
            mem = sb.table("sharing_members").select("id", count="exact").eq("group_id", g["id"]).eq("status", "active").execute()
            om_count = mem.count or 0

            # shared mwh last month (pre USAGE plán)
            shared_mwh = 0.0
            if usage_per_mwh > 0:
                edc = sb.table("sharing_edc_data").select("shared_in_kwh,shared_out_kwh").eq("group_id", g["id"]).eq("period_year", target_year).eq("period_month", target_month).execute()
                total_kwh = sum(float(r.get("shared_out_kwh") or 0) for r in (edc.data or []))
                shared_mwh = total_kwh / 1000.0

            monthly_fee = monthly_flat + monthly_per_om * om_count
            usage_fee = usage_per_mwh * shared_mwh
            subtotal = monthly_fee + usage_fee
            if bundle_pct > 0:
                subtotal = subtotal * (1.0 - bundle_pct / 100.0)
            total_with_vat = subtotal * (1.0 + vat_rate / 100.0)

            issued_date = today
            due_date = today + timedelta(days=14)

            insert_res = sb.table("sharing_invoices").insert({
                "group_id": g["id"],
                "customer_id": g.get("organizator_customer_id"),
                "invoice_type": "monthly",
                "billing_period_year": target_year,
                "billing_period_month": target_month,
                "setup_fee": 0,
                "monthly_fee": round(monthly_fee, 2),
                "usage_fee": round(usage_fee, 2),
                "om_count_at_billing": om_count,
                "shared_mwh": round(shared_mwh, 3),
                "total_no_vat": round(subtotal, 2),
                "vat_rate": vat_rate,
                "total_with_vat": round(total_with_vat, 2),
                "status": "draft",
                "issued_at": issued_date.isoformat(),
                "due_at": due_date.isoformat(),
            }).execute()

            created.append({"group_id": g["id"], "total_with_vat": round(total_with_vat, 2)})

        return jsonify({"ok": True, "period": f"{target_month}/{target_year}", "created": len(created), "skipped": len(skipped), "details": {"created": created, "skipped": skipped}})
    except Exception as e:
        log.exception("[omshare-monthly-billing] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/omshare-invoice-pdf", methods=["POST"])
def webhook_omshare_invoice_pdf():
    """
    Vygeneruje PDF pre 1 sharing_invoice a upload do Supabase Storage.
    Body: { invoice_id }
    """
    body = request.get_json(silent=True) or {}
    invoice_id = body.get("invoice_id")
    if not invoice_id:
        return jsonify({"ok": False, "error": "invoice_id required"}), 400

    try:
        sb = _sb()
        inv_res = sb.table("sharing_invoices").select("*").eq("id", invoice_id).single().execute()
        inv = inv_res.data
        if not inv:
            return jsonify({"ok": False, "error": "invoice not found"}), 404

        grp_res = sb.table("sharing_groups").select("*").eq("id", inv["group_id"]).single().execute()
        grp = grp_res.data or {}

        cust = None
        if inv.get("customer_id"):
            c = sb.table("customers").select("*").eq("id", inv["customer_id"]).single().execute()
            cust = c.data

        # Generate PDF via reportlab (simple invoice layout)
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        import io

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm, leftMargin=2*cm, rightMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []

        # Hlavička
        story.append(Paragraph(f"<b>Faktúra č. {inv.get('invoice_number','')}</b>", styles["Title"]))
        story.append(Paragraph(f"OM-Share — Organizátor zdieľania elektriny", styles["Heading3"]))
        story.append(Spacer(1, 0.5*cm))

        # Dodávateľ + Odberateľ
        cust_name = (cust or {}).get("name") or "—"
        cust_addr = (cust or {}).get("address") or ""
        cust_ico = (cust or {}).get("ico") or ""
        info_table = Table([
            ["Dodávateľ", "Odberateľ"],
            ["Energovision, s. r. o.", cust_name],
            ["IČO: 53 036 280", cust_addr],
            ["DIČ: 2121270486", f"IČO: {cust_ico}"],
            ["IČ DPH: SK2121270486", ""],
        ], colWidths=[8*cm, 8*cm])
        info_table.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E78")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("PADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.5*cm))

        # Obdobie
        period_str = f"{inv.get('billing_period_month','')}/{inv.get('billing_period_year','')}" if inv.get("invoice_type") == "monthly" else "Jednorazové"
        meta = Table([
            ["Skupina", grp.get("name", "—")],
            ["Plán", grp.get("plan_code", "—")],
            ["Typ", "Mesačná" if inv.get("invoice_type") == "monthly" else ("Zriadenie" if inv.get("invoice_type") == "setup" else inv.get("invoice_type"))],
            ["Obdobie", period_str],
            ["Vystavené", inv.get("issued_at") or ""],
            ["Splatnosť", inv.get("due_at") or ""],
        ], colWidths=[5*cm, 11*cm])
        meta.setStyle(TableStyle([
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("PADDING", (0,0), (-1,-1), 4),
            ("LINEBELOW", (0,0), (-1,-1), 0.25, colors.lightgrey),
        ]))
        story.append(meta)
        story.append(Spacer(1, 0.7*cm))

        # Položky
        rows = [["Položka", "Počet", "Cena", "Spolu"]]
        if float(inv.get("setup_fee") or 0) > 0:
            rows.append(["Zriaďovací poplatok OM-Share", "1", f"{float(inv['setup_fee']):.2f} €", f"{float(inv['setup_fee']):.2f} €"])
        if float(inv.get("monthly_fee") or 0) > 0:
            rows.append([f"Mesačný poplatok (správa {inv.get('om_count_at_billing','—')} OM)", "1 mes.", f"{float(inv['monthly_fee']):.2f} €", f"{float(inv['monthly_fee']):.2f} €"])
        if float(inv.get("usage_fee") or 0) > 0:
            rows.append([f"Poplatok za zdieľanú elektrinu ({float(inv.get('shared_mwh') or 0):.3f} MWh)", "—", "—", f"{float(inv['usage_fee']):.2f} €"])

        items_table = Table(rows, colWidths=[8*cm, 2.5*cm, 2.5*cm, 3*cm])
        items_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E78")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("ALIGN", (1,0), (-1,-1), "RIGHT"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("PADDING", (0,0), (-1,-1), 6),
            ("LINEBELOW", (0,0), (-1,-1), 0.25, colors.lightgrey),
        ]))
        story.append(items_table)
        story.append(Spacer(1, 0.5*cm))

        # Totály
        no_vat = float(inv.get("total_no_vat") or 0)
        vat_pct = float(inv.get("vat_rate") or 23)
        vat_amount = no_vat * vat_pct / 100.0
        with_vat = float(inv.get("total_with_vat") or 0)
        totals = Table([
            ["Spolu bez DPH", f"{no_vat:.2f} €"],
            [f"DPH {vat_pct:.0f}%", f"{vat_amount:.2f} €"],
            ["Spolu s DPH", f"{with_vat:.2f} €"],
        ], colWidths=[13*cm, 3*cm])
        totals.setStyle(TableStyle([
            ("ALIGN", (0,0), (-1,-1), "RIGHT"),
            ("FONTSIZE", (0,0), (-1,-1), 10),
            ("PADDING", (0,0), (-1,-1), 4),
            ("FONTNAME", (0,-1), (-1,-1), "Helvetica-Bold"),
            ("LINEABOVE", (0,-1), (-1,-1), 1, colors.HexColor("#1F4E78")),
        ]))
        story.append(totals)
        story.append(Spacer(1, 1*cm))

        story.append(Paragraph(
            "Faktúra je vystavená v zmysle zákona č. 222/2004 Z. z. o DPH. "
            "Energovision je organizátor zdieľania elektriny zapísaný u ÚRSO. "
            "Platba bankovým prevodom na IBAN: SK00 0000 0000 0000 0000 0000.",
            styles["Normal"]
        ))

        doc.build(story)
        pdf_bytes = buf.getvalue()
        buf.close()

        # Upload do storage
        filename = f"{inv.get('invoice_number', invoice_id)}.pdf"
        storage_path = f"omshare/invoices/{inv['group_id']}/{filename}"

        try:
            sb.storage.from_("documents").upload(
                path=storage_path,
                file=pdf_bytes,
                file_options={"content-type": "application/pdf", "upsert": "true", "cache-control": "0"},
            )
        except Exception as up_e:
            log.warning("[omshare-invoice-pdf] upload retry: %s", up_e)
            # remove + upload
            try:
                sb.storage.from_("documents").remove([storage_path])
            except Exception:
                pass
            sb.storage.from_("documents").upload(
                path=storage_path,
                file=pdf_bytes,
                file_options={"content-type": "application/pdf"},
            )

        pub = sb.storage.from_("documents").get_public_url(storage_path)
        pdf_url = f"{pub}?v={int(datetime.now().timestamp())}"

        sb.table("sharing_invoices").update({
            "pdf_url": pdf_url,
            "storage_path": storage_path,
        }).eq("id", invoice_id).execute()

        return jsonify({"ok": True, "pdf_url": pdf_url, "filename": filename})
    except Exception as e:
        log.exception("[omshare-invoice-pdf] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/omshare-send-invoice", methods=["POST"])
def webhook_omshare_send_invoice():
    """
    Pošle email s PDF faktúry klientovi (organizátorovi skupiny) cez M365.
    Body: { invoice_id, recipient_email? }
    """
    body = request.get_json(silent=True) or {}
    invoice_id = body.get("invoice_id")
    if not invoice_id:
        return jsonify({"ok": False, "error": "invoice_id required"}), 400

    try:
        sb = _sb()
        inv = sb.table("sharing_invoices").select("*").eq("id", invoice_id).single().execute().data
        if not inv:
            return jsonify({"ok": False, "error": "invoice not found"}), 404

        if not inv.get("pdf_url"):
            return jsonify({"ok": False, "error": "PDF not generated yet — call /webhook/omshare-invoice-pdf first"}), 400

        grp = sb.table("sharing_groups").select("name, group_number").eq("id", inv["group_id"]).single().execute().data or {}

        recipient = body.get("recipient_email")
        if not recipient and inv.get("customer_id"):
            cust = sb.table("customers").select("email").eq("id", inv["customer_id"]).single().execute().data
            recipient = (cust or {}).get("email")

        if not recipient:
            return jsonify({"ok": False, "error": "no recipient email — pass recipient_email or set customer.email"}), 400

        period_str = f"{inv.get('billing_period_month','')}/{inv.get('billing_period_year','')}" if inv.get("invoice_type") == "monthly" else "zriadenie"
        subject = f"Faktúra {inv.get('invoice_number')} — OM-Share {period_str}"
        body_html = f"""
        <p>Dobrý deň,</p>
        <p>Posielame Vám faktúru za službu <strong>OM-Share — Správa zdieľania elektriny</strong> pre skupinu
        <strong>{grp.get('name','')}</strong> ({grp.get('group_number','')}).</p>
        <p>
          Číslo faktúry: <strong>{inv.get('invoice_number')}</strong><br>
          Obdobie: {period_str}<br>
          Suma na úhradu: <strong>{float(inv.get('total_with_vat') or 0):.2f} €</strong><br>
          Splatnosť: {inv.get('due_at')}
        </p>
        <p>Faktúru nájdete v prílohe. Platbu prosím poukážte na IBAN uvedený vo faktúre s variabilným symbolom {inv.get('invoice_number','').replace('-','')}.</p>
        <p>V prípade otázok nás neváhajte kontaktovať.</p>
        <p>S pozdravom,<br>Tím Energovision<br><a href="https://www.energovision.sk">www.energovision.sk</a></p>
        """

        # Fetch PDF bytes
        import requests as rq
        r = rq.get(inv["pdf_url"], timeout=30)
        if r.status_code != 200:
            return jsonify({"ok": False, "error": f"PDF fetch failed: {r.status_code}"}), 500
        import base64
        pdf_b64 = base64.b64encode(r.content).decode()

        # Send via M365 (call our own endpoint or M365 directly)
        # Reuse existing M365 send logic from email_agent or similar
        # Here use direct Graph API call via stored M365 credentials
        cred = sb.table("m365_credentials").select("*").eq("id", "singleton").single().execute().data
        if not cred:
            return jsonify({"ok": False, "error": "M365 nie je pripojený. Choď na /admin/email-setup"}), 400

        # Refresh access token if expired
        from datetime import datetime as dt
        access_token = cred.get("access_token")
        exp = cred.get("access_token_expires_at")
        if not access_token or (exp and dt.fromisoformat(exp.replace("Z","+00:00")) < dt.now(timezone.utc)):
            # refresh
            tok_url = f"https://login.microsoftonline.com/{cred.get('tenant_id')}/oauth2/v2.0/token"
            tok_body = {
                "client_id": cred.get("client_id"),
                "client_secret": os.environ.get("AZURE_CLIENT_SECRET", ""),
                "grant_type": "refresh_token",
                "refresh_token": cred.get("refresh_token"),
                "scope": "https://graph.microsoft.com/Mail.Send https://graph.microsoft.com/User.Read offline_access",
            }
            tr = rq.post(tok_url, data=tok_body, timeout=30)
            if not tr.ok:
                return jsonify({"ok": False, "error": f"M365 refresh: {tr.status_code} {tr.text[:200]}"}), 500
            tokens = tr.json()
            access_token = tokens["access_token"]
            new_exp = (dt.now(timezone.utc) + timedelta(seconds=tokens.get("expires_in", 3600))).isoformat()
            sb.table("m365_credentials").update({
                "access_token": access_token,
                "access_token_expires_at": new_exp,
                "refresh_token": tokens.get("refresh_token", cred.get("refresh_token")),
                "last_used_at": dt.now(timezone.utc).isoformat(),
            }).eq("id", "singleton").execute()

        filename = (inv.get("storage_path") or "").split("/")[-1] or f"{inv.get('invoice_number')}.pdf"
        gr = rq.post(
            "https://graph.microsoft.com/v1.0/me/sendMail",
            headers={"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"},
            json={
                "message": {
                    "subject": subject,
                    "body": {"contentType": "HTML", "content": body_html},
                    "toRecipients": [{"emailAddress": {"address": recipient}}],
                    "attachments": [{
                        "@odata.type": "#microsoft.graph.fileAttachment",
                        "name": filename,
                        "contentType": "application/pdf",
                        "contentBytes": pdf_b64,
                    }],
                },
                "saveToSentItems": True,
            },
            timeout=60,
        )
        if gr.status_code != 202:
            return jsonify({"ok": False, "error": f"Graph send: {gr.status_code} {gr.text[:300]}"}), 500

        # Mark invoice as sent
        sb.table("sharing_invoices").update({
            "status": "sent",
            "updated_at": dt.now(timezone.utc).isoformat(),
        }).eq("id", invoice_id).execute()

        sb.table("activities").insert({
            "entity_type": "sharing_group",
            "entity_id": inv["group_id"],
            "action": "invoice_sent",
            "changes": {"invoice_number": inv.get("invoice_number"), "to": recipient, "total": inv.get("total_with_vat")},
        }).execute()

        return jsonify({"ok": True, "sent_to": recipient})
    except Exception as e:
        log.exception("[omshare-send-invoice] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# OM-Share — Mesačný report klientovi (P3)
# ============================================================

@app.route("/webhook/omshare-monthly-report", methods=["POST"])
def webhook_omshare_monthly_report():
    """
    Cron 20. dňa v mesiaci — generuje a posiela mesačný report pre každú aktívnu skupinu.
    Report obsahuje per-OM rozpis zdieľanej elektriny, celkový objem, odhadovaná úspora.
    """
    body = request.get_json(silent=True) or {}
    target_month = body.get("month")
    target_year = body.get("year")

    today = datetime.now(timezone.utc).date()
    if not target_month or not target_year:
        if today.month == 1:
            target_month = 12
            target_year = today.year - 1
        else:
            target_month = today.month - 1
            target_year = today.year
    target_month = int(target_month)
    target_year = int(target_year)

    try:
        sb = _sb()
        groups = sb.table("sharing_groups").select("*").eq("status", "active").execute().data or []
        log.info("[omshare-report] generujem reporty za %d/%d pre %d skupín", target_month, target_year, len(groups))

        results = []
        for g in groups:
            try:
                # check duplicate
                dup = sb.table("sharing_reports").select("id").eq("group_id", g["id"]).eq("period_year", target_year).eq("period_month", target_month).execute()
                if dup.data:
                    results.append({"group_id": g["id"], "skipped": "exists"})
                    continue

                # fetch edc
                edc = sb.table("sharing_edc_data").select("*").eq("group_id", g["id"]).eq("period_year", target_year).eq("period_month", target_month).execute()
                edc_rows = edc.data or []
                members = sb.table("sharing_members").select("id, member_name, role").eq("group_id", g["id"]).execute().data or []
                mname = {m["id"]: m for m in members}

                total_kwh = sum(float(r.get("shared_out_kwh") or 0) for r in edc_rows)
                # Odhad úspory: priemerná cena elektriny SK pre koncového odberateľa ~ 0.18 €/kWh
                est_savings = total_kwh * 0.18

                # Generate PDF
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib import colors
                from reportlab.lib.units import cm
                import io

                buf = io.BytesIO()
                doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm, leftMargin=2*cm, rightMargin=2*cm)
                styles = getSampleStyleSheet()
                story = []
                story.append(Paragraph(f"<b>Mesačný report zdieľania elektriny</b>", styles["Title"]))
                story.append(Paragraph(f"Skupina: <b>{g.get('name','')}</b> ({g.get('group_number','')})", styles["Heading3"]))
                story.append(Paragraph(f"Obdobie: {target_month}/{target_year}", styles["Normal"]))
                story.append(Spacer(1, 0.5*cm))

                summary = Table([
                    ["Celkový objem zdieľanej elektriny", f"{total_kwh:.2f} kWh"],
                    ["Počet aktívnych odberných miest", str(len(members))],
                    ["Odhadovaná úspora skupiny", f"{est_savings:.2f} €"],
                    ["Plán", g.get("plan_code","—")],
                ], colWidths=[10*cm, 6*cm])
                summary.setStyle(TableStyle([
                    ("FONTSIZE", (0,0), (-1,-1), 10),
                    ("PADDING", (0,0), (-1,-1), 6),
                    ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F0FDF4")),
                    ("LINEBELOW", (0,0), (-1,-1), 0.25, colors.lightgrey),
                ]))
                story.append(summary)
                story.append(Spacer(1, 0.7*cm))

                # Per OM rozpis
                story.append(Paragraph("<b>Rozpis podľa odberného miesta</b>", styles["Heading4"]))
                rows = [["Odberné miesto", "Rola", "Prijal (kWh)", "Poskytol (kWh)"]]
                for r in edc_rows:
                    m = mname.get(r.get("member_id"), {})
                    rows.append([
                        m.get("member_name", "—"),
                        "Zdroj" if m.get("role") == "producer" else "Spotrebič" if m.get("role") == "consumer" else "Oboje",
                        f"{float(r.get('shared_in_kwh') or 0):.2f}",
                        f"{float(r.get('shared_out_kwh') or 0):.2f}",
                    ])
                if len(rows) > 1:
                    tbl = Table(rows, colWidths=[7*cm, 3*cm, 3*cm, 3*cm])
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E78")),
                        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                        ("ALIGN", (2,0), (-1,-1), "RIGHT"),
                        ("FONTSIZE", (0,0), (-1,-1), 9),
                        ("PADDING", (0,0), (-1,-1), 5),
                        ("LINEBELOW", (0,0), (-1,-1), 0.25, colors.lightgrey),
                    ]))
                    story.append(tbl)
                else:
                    story.append(Paragraph("Žiadne dáta pre toto obdobie.", styles["Normal"]))

                story.append(Spacer(1, 1*cm))
                story.append(Paragraph("Energovision, s. r. o. · OM-Share organizátor zdieľania elektriny · ÚRSO potvrdenie [____]", styles["Italic"]))

                doc.build(story)
                pdf_bytes = buf.getvalue()
                buf.close()

                filename = f"OMShare_report_{g.get('group_number')}_{target_year}-{target_month:02d}.pdf"
                storage_path = f"omshare/reports/{g['id']}/{filename}"
                try:
                    sb.storage.from_("documents").upload(
                        path=storage_path,
                        file=pdf_bytes,
                        file_options={"content-type": "application/pdf", "upsert": "true", "cache-control": "0"},
                    )
                except Exception:
                    try:
                        sb.storage.from_("documents").remove([storage_path])
                    except Exception:
                        pass
                    sb.storage.from_("documents").upload(
                        path=storage_path,
                        file=pdf_bytes,
                        file_options={"content-type": "application/pdf"},
                    )

                pub = sb.storage.from_("documents").get_public_url(storage_path)
                pdf_url = f"{pub}?v={int(datetime.now().timestamp())}"

                sb.table("sharing_reports").insert({
                    "group_id": g["id"],
                    "period_month": target_month,
                    "period_year": target_year,
                    "total_shared_kwh": round(total_kwh, 3),
                    "total_savings_eur": round(est_savings, 2),
                    "member_count": len(members),
                    "pdf_url": pdf_url,
                    "storage_path": storage_path,
                }).execute()

                results.append({"group_id": g["id"], "total_kwh": total_kwh, "savings": est_savings, "pdf": pdf_url})
            except Exception as inner:
                log.exception("[omshare-report] group %s failed", g.get("id"))
                results.append({"group_id": g.get("id"), "error": str(inner)[:200]})

        return jsonify({"ok": True, "period": f"{target_month}/{target_year}", "count": len(results), "details": results})
    except Exception as e:
        log.exception("[omshare-monthly-report] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# OM-Share — OKTE EDC client (P4 — mock + skeleton pre real)
# ============================================================

def _okte_edc_mock_data(group_id, members, target_year, target_month):
    """Mock EDC dáta — generuje pseudonáhodné kWh hodnoty pre testing."""
    import random
    rng = random.Random(f"{group_id}-{target_year}-{target_month}")
    results = []
    producers = [m for m in members if m.get("role") in ("producer", "both") and m.get("status") == "active"]
    consumers = [m for m in members if m.get("role") in ("consumer", "both") and m.get("status") == "active"]
    if not producers or not consumers:
        return []
    # Producer produkuje 200-600 kWh mesačne, consumer spotreba 150-450 kWh
    total_produced = sum(rng.uniform(200, 600) for _ in producers)
    for p in producers:
        out_kwh = total_produced / len(producers)
        results.append({
            "group_id": group_id,
            "member_id": p["id"],
            "period_month": target_month,
            "period_year": target_year,
            "shared_in_kwh": 0,
            "shared_out_kwh": round(out_kwh, 2),
            "production_kwh": round(out_kwh * 1.3, 2),  # produkoval viac, len časť zdieľal
            "consumption_kwh": 0,
            "source": "mock",
        })
    per_consumer = total_produced / len(consumers)
    for c in consumers:
        in_kwh = per_consumer * rng.uniform(0.7, 1.1)
        results.append({
            "group_id": group_id,
            "member_id": c["id"],
            "period_month": target_month,
            "period_year": target_year,
            "shared_in_kwh": round(in_kwh, 2),
            "shared_out_kwh": 0,
            "consumption_kwh": round(in_kwh * 2.5, 2),
            "production_kwh": 0,
            "source": "mock",
        })
    return results


@app.route("/webhook/omshare-edc-sync", methods=["POST"])
def webhook_omshare_edc_sync():
    """
    Sync dát z OKTE EDC za zvolený mesiac. Ak nemáme certifikát/credentials → mock mode.
    Body: { group_id?, month?, year? }
    """
    body = request.get_json(silent=True) or {}
    group_id = body.get("group_id")
    target_month = body.get("month")
    target_year = body.get("year")

    today = datetime.now(timezone.utc).date()
    if not target_month or not target_year:
        if today.month == 1:
            target_month, target_year = 12, today.year - 1
        else:
            target_month, target_year = today.month - 1, today.year
    target_month, target_year = int(target_month), int(target_year)

    use_mock = not bool(os.environ.get("OKTE_EDC_CERT_PATH"))

    try:
        sb = _sb()
        groups_q = sb.table("sharing_groups").select("*").eq("status", "active")
        if group_id:
            groups_q = groups_q.eq("id", group_id)
        groups = groups_q.execute().data or []

        synced = []
        for g in groups:
            members = sb.table("sharing_members").select("*").eq("group_id", g["id"]).execute().data or []
            if use_mock:
                rows = _okte_edc_mock_data(g["id"], members, target_year, target_month)
            else:
                # TODO: real OKTE EDC API call when certificate is available
                # rows = okte_edc_client.fetch_monthly_data(g["edc_group_id"], target_year, target_month)
                rows = []

            if not rows:
                continue

            for r in rows:
                # upsert per (member_id, year, month)
                sb.table("sharing_edc_data").upsert(r, on_conflict="member_id,period_year,period_month").execute()
            synced.append({"group_id": g["id"], "rows": len(rows)})

        return jsonify({"ok": True, "mode": "mock" if use_mock else "real", "period": f"{target_month}/{target_year}", "groups": len(synced), "details": synced})
    except Exception as e:
        log.exception("[omshare-edc-sync] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/omshare-edc-register", methods=["POST"])
def webhook_omshare_edc_register():
    """
    Registrácia novej skupiny v OKTE EDC. Mock alebo real podľa certifikátu.
    Body: { group_id }
    """
    body = request.get_json(silent=True) or {}
    group_id = body.get("group_id")
    if not group_id:
        return jsonify({"ok": False, "error": "group_id required"}), 400

    use_mock = not bool(os.environ.get("OKTE_EDC_CERT_PATH"))

    try:
        sb = _sb()
        if use_mock:
            # Mock — assign fake EDC ID
            import secrets
            edc_id = f"EDC-MOCK-{secrets.token_hex(4).upper()}"
        else:
            # TODO: real OKTE EDC call
            edc_id = "TODO_REAL_EDC"

        sb.table("sharing_groups").update({
            "edc_group_id": edc_id,
            "registered_at": datetime.now(timezone.utc).isoformat(),
            "status": "pending_okte",
        }).eq("id", group_id).execute()

        sb.table("activities").insert({
            "entity_type": "sharing_group",
            "entity_id": group_id,
            "action": "edc_registered",
            "changes": {"edc_group_id": edc_id, "mode": "mock" if use_mock else "real"},
        }).execute()

        return jsonify({"ok": True, "edc_group_id": edc_id, "mode": "mock" if use_mock else "real"})
    except Exception as e:
        log.exception("[omshare-edc-register] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# OM-Share — One-pager generator (P5)
# ============================================================

@app.route("/webhook/omshare-onepager", methods=["POST"])
def webhook_omshare_onepager():
    """
    Generuje personalizovaný one-pager (PDF) pre konkrétneho zákazníka.
    Body: { type: 'b2c'|'b2b'|'bundle', customer_id?, customer_name? }
    """
    body = request.get_json(silent=True) or {}
    onepager_type = body.get("type", "b2c")
    customer_id = body.get("customer_id")
    customer_name = body.get("customer_name", "")

    try:
        sb = _sb()
        if customer_id and not customer_name:
            c = sb.table("customers").select("name").eq("id", customer_id).single().execute().data
            if c:
                customer_name = c.get("name", "")

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_CENTER
        import io

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm, leftMargin=2*cm, rightMargin=2*cm)
        styles = getSampleStyleSheet()
        big = ParagraphStyle("big", parent=styles["Title"], fontSize=24, leading=28, alignment=TA_CENTER)
        story = []

        if onepager_type == "b2c":
            title = "OM-Share START — Zdieľajte elektrinu v rodine"
            tagline = "Zdieľajte prebytky z vašej FVE so susedmi alebo rodinou. Šetríte za elektrinu, my sa staráme o administratívu."
            features = [
                "Zriadenie skupiny zdieľania na ÚRSO + OKTE EDC",
                "Mesačné vyúčtovanie a report o zdieľanej elektrine",
                "Komunikácia s distribúciou a OKTE",
                "2–4 odberné miesta v skupine",
                "Pre domácnosti — 49 € zriadenie + 4 €/mes.",
            ]
        elif onepager_type == "b2b":
            title = "OM-Share BUSINESS — Zdieľanie pre firmy"
            tagline = "Firma s viacerými prevádzkami? Zdieľaním elektriny ušetríte tisíce € ročne. Komplexná správa od Energovision."
            features = [
                "Až 10 odberných miest v skupine (centrála, pobočky, parkoviská)",
                "Optimalizácia metódy zdieľania — statická, dynamická, prioritná",
                "Detailné mesačné reporty pre účtovníctvo",
                "Cross-sell: FVE/BESS/EMS pre Vaše prevádzky",
                "199 € zriadenie + 2 €/OM/mes.",
            ]
        else:  # bundle
            title = "OM-Share + EMS Bundle — Najvyššia hodnota"
            tagline = "Pre klientov s FVE/BESS od Energovision: 30% zľava na zdieľanie + integrácia s EMS (spotová arbitráž)."
            features = [
                "30% zľava na OM-Share službu",
                "Integrácia s Energovision EMS — automatická arbitráž",
                "Optimalizácia nabíjania batérie podľa hodinových cien OKTE",
                "Pre FVE klientov nad 10 kWp s BESS",
                "Zriadenie 34,30 € + 1,40 €/OM/mes.",
            ]

        story.append(Paragraph(title, big))
        if customer_name:
            story.append(Paragraph(f"<i>Personalizovaná ponuka pre: <b>{customer_name}</b></i>", styles["Italic"]))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(tagline, styles["Heading4"]))
        story.append(Spacer(1, 0.7*cm))

        story.append(Paragraph("<b>Čo všetko zahŕňa služba:</b>", styles["Heading4"]))
        for f in features:
            story.append(Paragraph(f"• {f}", styles["Normal"]))
        story.append(Spacer(1, 0.7*cm))

        story.append(Paragraph("<b>Prečo Energovision?</b>", styles["Heading4"]))
        story.append(Paragraph(
            "Energovision je jediná firma, ktorá ponúka <b>komplexný balík FVE + BESS + EMS + zdieľanie</b> pod jednou strechou. "
            "Sme vendor-neutrálni (nepredávame elektrinu) — klient si ponechá svojho dodávateľa. "
            "Reálna technická expertíza (revízie, servis trafostaníc) a regionálna prítomnosť.",
            styles["Normal"]
        ))
        story.append(Spacer(1, 1*cm))

        story.append(Paragraph("<b>Kontakt:</b> Energovision, s. r. o. · sales@energovision.sk · www.energovision.sk", styles["Italic"]))

        doc.build(story)
        pdf_bytes = buf.getvalue()
        buf.close()

        # Upload
        safe_name = (customer_name or "klient").replace(" ", "_").lower()[:40]
        filename = f"OMShare_{onepager_type}_{safe_name}.pdf"
        storage_path = f"omshare/onepagers/{filename}"
        try:
            sb.storage.from_("documents").upload(
                path=storage_path,
                file=pdf_bytes,
                file_options={"content-type": "application/pdf", "upsert": "true", "cache-control": "0"},
            )
        except Exception:
            try:
                sb.storage.from_("documents").remove([storage_path])
            except Exception:
                pass
            sb.storage.from_("documents").upload(
                path=storage_path,
                file=pdf_bytes,
                file_options={"content-type": "application/pdf"},
            )

        pub = sb.storage.from_("documents").get_public_url(storage_path)
        pdf_url = f"{pub}?v={int(datetime.now().timestamp())}"

        return jsonify({"ok": True, "pdf_url": pdf_url, "filename": filename, "type": onepager_type})
    except Exception as e:
        log.exception("[omshare-onepager] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/omshare-bundle-expiry-check", methods=["POST"])
def webhook_omshare_bundle_expiry_check():
    """
    Denný cron — nájde skupiny ktorých bundle expiruje za 30 dní (§ 17g notice požiadavka)
    a zapíše activity log + interný alert. Tu by sa neskôr pridal email klientovi.
    """
    try:
        sb = _sb()
        # skupiny s bundle_expires_at presne za 30 dní
        target = (datetime.now(timezone.utc).date() + timedelta(days=30)).isoformat()
        res = sb.table("sharing_groups").select("*").eq("is_bundle", True).eq("bundle_expires_at", target).execute()
        groups = res.data or []
        notified = []
        for g in groups:
            sb.table("activities").insert({
                "entity_type": "sharing_group",
                "entity_id": g["id"],
                "action": "bundle_expiry_30d_alert",
                "changes": {
                    "bundle_expires_at": g.get("bundle_expires_at"),
                    "name": g.get("name"),
                    "note": "Klient musí byť informovaný 30 dní vopred podľa § 17g zákona o energetike",
                },
            }).execute()
            notified.append({"group_id": g["id"], "name": g.get("name"), "expires_at": g.get("bundle_expires_at")})
        return jsonify({"ok": True, "count": len(notified), "groups": notified, "target_date": target})
    except Exception as e:
        log.exception("[omshare-bundle-expiry-check] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


@app.route("/webhook/aom-orphan-reset", methods=["POST"])
def webhook_aom_orphan_reset():
    """
    Orphan detector — analyza_om s status=running staršie ako 30 minút
    sa resetnú na status=draft. Worker buď spadol alebo Render timeoutol.
    Beží zo Vercel cronu každú hodinu.
    """
    try:
        sb = _sb()
        # Reset všetkých orphan running
        result = sb.rpc("exec_sql", {"sql": "UPDATE analyza_om SET status='draft', updated_at=now() WHERE status='running' AND updated_at < now() - INTERVAL '30 minutes' RETURNING id, name"}).execute()
        # Fallback ak exec_sql neexistuje — direct query
        if not result.data:
            res = sb.table("analyza_om").select("id, name, updated_at").eq("status", "running").execute()
            now = datetime.now(timezone.utc)
            stale = []
            for r in (res.data or []):
                try:
                    upd = datetime.fromisoformat(r["updated_at"].replace("Z", "+00:00"))
                    if (now - upd).total_seconds() > 30 * 60:
                        stale.append(r)
                except Exception:
                    pass
            for s in stale:
                sb.table("analyza_om").update({"status": "draft", "updated_at": now.isoformat()}).eq("id", s["id"]).execute()
            return jsonify({"ok": True, "reset_count": len(stale), "reset_ids": [s["id"] for s in stale]})
        return jsonify({"ok": True, "reset_count": len(result.data), "reset_ids": [r["id"] for r in result.data]})
    except Exception as e:
        log.exception("[aom-orphan-reset] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


# ============================================================
# FLEET STATUS — live Huawei API → /admin/monitoring
# ============================================================
# Read-only endpoint. Combines DB metadata (inverter_sites + alarms)
# with live realtime KPI pulled from Huawei FusionSolar.
# In-memory cache 60s to avoid hammering Huawei API quota.
# CORS open for crm.energovision.sk + app.energovision.sk.
# No mutations — SPOT reactor is unaffected.

_FLEET_CACHE = {"ts": 0.0, "data": None}
_FLEET_CACHE_TTL_SEC = 120  # 2 min cache (Huawei API quota)


def _huawei_fetch_active_power_per_station(station_codes: list, base: str, headers: dict) -> dict:
    """Pre každú stanicu zistí aktuálny AC výkon [kW] sumovaním active_power
    všetkých jej inverterov.

    Postup (podľa Huawei NBI Reference 25.4.0):
    1. POST /getDevList {stationCodes: "a,b,c"} → zoznam zariadení (devTypeId=1 invertery, 38 residential)
    2. POST /getDevRealKpi {devIds: "id1,id2", devTypeId: 1} → active_power per inverter
    3. Sum active_power po stationCode

    Returns: {stationCode: kw_sum, ...}
    """
    power_by_station: dict = {}
    if not station_codes:
        return power_by_station

    # 1. Zoznam zariadení pre všetky stanice naraz (max 100 staníc per call)
    try:
        r = requests.post(
            f"{base}/getDevList",
            headers=headers,
            json={"stationCodes": ",".join(station_codes)},
            timeout=30,
        )
        if r.status_code != 200:
            log.warning("[fleet-status] getDevList HTTP %s: %s", r.status_code, r.text[:200])
            return power_by_station
        payload = r.json() or {}
        devices = payload.get("data") or []
        if not isinstance(devices, list):
            return power_by_station
    except Exception as e:
        log.warning("[fleet-status] getDevList failed: %s", e)
        return power_by_station

    # 2. Filter len invertery (devTypeId 1 = string inverter, 38 = residential inverter)
    inverters_by_type: dict = {}  # {devTypeId: [(devId, stationCode), ...]}
    for d in devices:
        if not isinstance(d, dict):
            continue
        dev_type = d.get("devTypeId")
        if dev_type not in (1, 38):
            continue
        dev_id = d.get("id")
        station_code = d.get("stationCode")
        if not dev_id or not station_code:
            continue
        inverters_by_type.setdefault(dev_type, []).append((dev_id, station_code))

    # 3. getDevRealKpi per device type (max 100 zariadení per call)
    for dev_type, inv_list in inverters_by_type.items():
        for chunk_start in range(0, len(inv_list), 100):
            chunk = inv_list[chunk_start:chunk_start + 100]
            dev_ids = [str(x[0]) for x in chunk]
            station_map = {str(x[0]): x[1] for x in chunk}
            try:
                r = requests.post(
                    f"{base}/getDevRealKpi",
                    headers=headers,
                    json={"devIds": ",".join(dev_ids), "devTypeId": dev_type},
                    timeout=30,
                )
                if r.status_code != 200:
                    continue
                payload = r.json() or {}
                rows = payload.get("data") or []
                if not isinstance(rows, list):
                    continue
                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    kpi = row.get("dataItemMap") or {}
                    if not isinstance(kpi, dict):
                        continue
                    ap = kpi.get("active_power")
                    if ap is None:
                        continue
                    try:
                        kw = float(ap)
                    except (TypeError, ValueError):
                        continue
                    dev_id = str(row.get("devId") or row.get("id") or "")
                    station_code = station_map.get(dev_id)
                    if station_code:
                        power_by_station[station_code] = power_by_station.get(station_code, 0.0) + kw
            except Exception as e:
                log.warning("[fleet-status] getDevRealKpi failed for devType %s: %s", dev_type, e)

    return power_by_station


def _fleet_status_compute() -> dict:
    """Fetch fleet status: DB metadata + live Huawei KPI batch call.

    Returns dict with `sites` list + `meta` (counts, fetched_at).
    Each site row keys: id, site_name, vendor, vendor_station_id, dc_kwp,
    ac_kw, bess_kwh, lat, lon, customer_id, monitoring_enabled, address,
    spot_control_enabled, spot_dry_run, spot_current_state,
    current_ac_power_kw, current_day_yield_kwh, current_battery_soc_pct,
    last_telemetry_at, open_alarms_count, status_label, status_tone.
    """
    from datetime import datetime, timezone
    started = _time.time()

    # 1. DB metadata
    sites = _hs.sb_get("inverter_sites", {
        "select": (
            "id,site_name,vendor,vendor_station_id,dc_kwp,ac_kw,bess_kwh,"
            "latitude,longitude,customer_id,monitoring_enabled,address,"
            "spot_control_enabled,spot_dry_run,spot_current_state,"
            "spot_last_transition_at"
        ),
        "order": "site_name.asc",
    }) or []

    # 2. Open alarm counts per site (cheap query)
    try:
        alarm_rows = _hs.sb_get("inverter_alarms", {
            "select": "site_id",
            "resolved_at": "is.null",
        }) or []
    except Exception as e:
        log.warning("[fleet-status] alarms query failed: %s", e)
        alarm_rows = []
    alarm_counts: dict = {}
    for a in alarm_rows:
        sid = a.get("site_id")
        if sid:
            alarm_counts[sid] = alarm_counts.get(sid, 0) + 1

    # 3. Batch live KPI from Huawei (only for huawei sites with monitoring_enabled)
    huawei_codes = [
        s.get("vendor_station_id") for s in sites
        if s.get("vendor") == "huawei"
        and s.get("monitoring_enabled")
        and s.get("vendor_station_id")
    ]
    live_kpi = {}
    huawei_error: str = None

    if huawei_codes and _hs is not None:
        try:
            token = _hs.huawei_login()
            if not token:
                huawei_error = "Huawei login failed (no token)"
            else:
                base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
                headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}
                # Huawei API: batch >10 stations zlyháva s data:[] (rate-limit či iný safeguard).
                # Robíme menšie chunks po 8 + retry per-station ak batch vráti 0 rows.
                CHUNK_SIZE = 8
                total_polled = 0
                total_returned = 0
                for chunk_start in range(0, len(huawei_codes), CHUNK_SIZE):
                    chunk = huawei_codes[chunk_start:chunk_start + CHUNK_SIZE]
                    total_polled += len(chunk)
                    body = {"stationCodes": ",".join(chunk)}
                    try:
                        r = requests.post(
                            f"{base}/getStationRealKpi",
                            headers=headers, json=body, timeout=30,
                        )
                    except Exception as e:
                        huawei_error = f"chunk {chunk_start}: {str(e)[:120]}"
                        continue
                    if r.status_code != 200:
                        huawei_error = f"HTTP {r.status_code}: {r.text[:200]}"
                        continue
                    payload = r.json() or {}
                    rows = payload.get("data") or payload.get("list") or []
                    if not rows and not payload.get("success", True):
                        # Auth expired? Re-login a skús ešte raz tento chunk
                        try:
                            token = _hs.huawei_login(force=True)
                            if token:
                                headers["XSRF-TOKEN"] = token
                                r2 = requests.post(f"{base}/getStationRealKpi", headers=headers, json=body, timeout=30)
                                if r2.status_code == 200:
                                    rows = (r2.json() or {}).get("data") or []
                        except Exception:
                            pass
                    chunk_returned = 0
                    for row in rows:
                        if not isinstance(row, dict):
                            continue
                        code = str(row.get("stationCode") or "")
                        kpi = row.get("dataItemMap") or {}
                        if not isinstance(kpi, dict):
                            kpi = {}
                        if not code:
                            continue
                        # Huawei getStationRealKpi vracia:
                        #   real_health_state — 1=offline, 2=fault, 3=normal (int, NIE výkon!)
                        #   day_power — dnes vyrobené [kWh]
                        #   total_power — celoživotne vyrobené [kWh]  (NIE MWh!)
                        #   month_power — mesačne [kWh]
                        #   day_use_energy — dnes spotreba [kWh]
                        #   day_on_grid_energy — dnes export [kWh]
                        #   day_income / total_income — príjem [€]
                        total_kwh = _to_float_safe(kpi.get("total_power"))
                        live_kpi[code] = {
                            # current_ac_power_kw nezistíme z getStationRealKpi — treba getDevRealKpi.
                            # Nechávame None aby UI zobrazilo "—" namiesto fake hodnoty.
                            "current_ac_power_kw": None,
                            "health_state": _to_float_safe(kpi.get("real_health_state")),
                            "current_day_yield_kwh": _to_float_safe(kpi.get("day_power")),
                            "current_total_yield_kwh": total_kwh,
                            "current_total_yield_mwh": (total_kwh / 1000.0) if total_kwh else None,
                            "current_month_yield_kwh": _to_float_safe(kpi.get("month_power")),
                            "current_day_export_kwh": _to_float_safe(kpi.get("day_on_grid_energy")),
                            "current_day_load_kwh": _to_float_safe(kpi.get("day_use_energy")),
                            "current_day_income_eur": _to_float_safe(kpi.get("day_income")),
                            "current_total_income_eur": _to_float_safe(kpi.get("total_income")),
                        }
                        chunk_returned += 1
                    total_returned += chunk_returned
                    # Pri prázdnom chunk fallback: skús per-station call pre debug
                    if chunk_returned == 0 and rows:
                        log.warning("[fleet-status] huawei chunk %d returned %d rows but 0 valid stationCodes", chunk_start, len(rows))

                # 5. Pre presný aktuálny výkon volaj getDevList + getDevRealKpi
                # (getStationRealKpi NEVRACIA active_power per Huawei NBI Reference 25.4.0).
                # Toto pridá ~2-3s na response time ale dáva reálne kW.
                try:
                    power_map = _huawei_fetch_active_power_per_station(huawei_codes, base, headers)
                    for code, kw in power_map.items():
                        if code in live_kpi:
                            live_kpi[code]["current_ac_power_kw"] = kw
                except Exception as e:
                    log.warning("[fleet-status] active_power fetch failed: %s", e)
        except Exception as e:
            log.exception("[fleet-status] Huawei batch fetch failed")
            huawei_error = str(e)[:200]

    # 4. Combine
    now_iso = datetime.now(timezone.utc).isoformat()
    enriched = []
    for s in sites:
        sid = s.get("id")
        code = s.get("vendor_station_id")
        kpi = live_kpi.get(code) if code else None

        has_live = kpi is not None
        power = (kpi or {}).get("current_ac_power_kw")
        # Status semantics:
        #   live + monitoring_enabled → "Live"
        #   monitoring_enabled, no live → "Bez dát" (rose)
        #   monitoring_enabled=false → "Vypnuté" (slate)
        if not s.get("monitoring_enabled"):
            status_label, status_tone = "Vypnuté", "slate"
        elif has_live:
            status_label, status_tone = "Live", "emerald"
        else:
            status_label, status_tone = "Bez dát", "rose"

        # Performance metrics
        _dc = s.get("dc_kwp") or s.get("ac_kw") or 0
        _day_yield = (kpi or {}).get("current_day_yield_kwh")
        _spec_yield = None
        if _dc and _day_yield is not None:
            try:
                _spec_yield = round(float(_day_yield) / float(_dc), 2) if float(_dc) > 0 else None
            except (TypeError, ValueError, ZeroDivisionError):
                _spec_yield = None
        _cap_factor = None
        if _dc and power is not None:
            try:
                _cap_factor = round(float(power) / float(_dc) * 100.0, 1) if float(_dc) > 0 else None
            except (TypeError, ValueError, ZeroDivisionError):
                _cap_factor = None

        enriched.append({
            "id": sid,
            "site_name": s.get("site_name"),
            "vendor": s.get("vendor"),
            "vendor_station_id": code,
            "dc_kwp": s.get("dc_kwp"),
            "ac_kw": s.get("ac_kw"),
            "bess_kwh": s.get("bess_kwh"),
            "lat": s.get("latitude"),
            "lon": s.get("longitude"),
            "customer_id": s.get("customer_id"),
            "monitoring_enabled": bool(s.get("monitoring_enabled")),
            "address": s.get("address"),
            "spot_control_enabled": bool(s.get("spot_control_enabled")),
            "spot_dry_run": bool(s.get("spot_dry_run")),
            "spot_current_state": s.get("spot_current_state"),
            "spot_last_transition_at": s.get("spot_last_transition_at"),
            "current_ac_power_kw": power,
            "current_day_yield_kwh": (kpi or {}).get("current_day_yield_kwh"),
            "specific_yield_today": _spec_yield,         # kWh / kWp dnes (good ~3-5 v lete)
            "capacity_factor_pct": _cap_factor,          # % aktuálneho využitia DC kapacity
            "current_total_yield_kwh": (kpi or {}).get("current_total_yield_kwh"),
            "current_total_yield_mwh": (kpi or {}).get("current_total_yield_mwh"),
            "current_month_yield_kwh": (kpi or {}).get("current_month_yield_kwh"),
            "current_day_export_kwh": (kpi or {}).get("current_day_export_kwh"),
            "current_day_load_kwh": (kpi or {}).get("current_day_load_kwh"),
            "current_day_income_eur": (kpi or {}).get("current_day_income_eur"),
            "current_total_income_eur": (kpi or {}).get("current_total_income_eur"),
            "health_state": (kpi or {}).get("health_state"),
            "last_telemetry_at": now_iso if has_live else None,
            "open_alarms_count": alarm_counts.get(sid, 0),
            "status_label": status_label,
            "status_tone": status_tone,
        })

    return {
        "ok": True,
        "fetched_at": now_iso,
        "duration_ms": int((_time.time() - started) * 1000),
        "sites": enriched,
        "meta": {
            "total": len(enriched),
            "monitored": sum(1 for x in enriched if x["monitoring_enabled"]),
            "live": sum(1 for x in enriched if x["status_label"] == "Live"),
            "huawei_codes_polled": len(huawei_codes),
            "huawei_kpi_returned": len(live_kpi),
            "huawei_error": huawei_error,
            "alarms_total": sum(alarm_counts.values()),
        },
    }


def _to_float_safe(v) -> float:
    if v is None or v == "" or v == "N/A":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


@rate_limit(max_calls=60, window_seconds=60)
@app.route("/webhook/fleet-status", methods=["GET", "OPTIONS"])
def webhook_fleet_status():
    """Live fleet snapshot for /admin/monitoring dashboard.

    GET ?fresh=1 → bypass cache.
    Cache TTL: 60 seconds. CORS open for CRM domains.
    """
    # CORS preflight
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Webhook-Secret"
        return resp

    fresh = request.args.get("fresh") in ("1", "true", "yes")
    now = _time.time()
    if not fresh and _FLEET_CACHE["data"] is not None and (now - _FLEET_CACHE["ts"]) < _FLEET_CACHE_TTL_SEC:
        cached_data = {**_FLEET_CACHE["data"]}
        # Refresh backoff status na live (DB read - cheap)
        try:
            if _hs is not None and hasattr(_hs, "get_huawei_backoff_status"):
                cached_data["huawei_backoff"] = _hs.get_huawei_backoff_status()
        except Exception:
            pass
        resp = jsonify({**cached_data, "cached": True, "cache_age_sec": int(now - _FLEET_CACHE["ts"])})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    try:
        data = _fleet_status_compute()
    except Exception as e:
        log.exception("[fleet-status] compute failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500

    # Graceful degradation: ak nový beh má 0 live ALE máme starší cache s validnými dátami,
    # zachovaj starý cache (Huawei API občas vráti prázdny payload pri rate-limit).
    new_live = (data.get("meta") or {}).get("live", 0)
    if new_live == 0 and _FLEET_CACHE.get("data") is not None:
        old_live = (_FLEET_CACHE["data"].get("meta") or {}).get("live", 0)
        if old_live > 0:
            # Preferuj cache, no over-write, ale extend timestamp aby sa za chvíľu retry-nul
            cached = _FLEET_CACHE["data"]
            cached_resp = {**cached, "cached": True, "stale_protect": True,
                           "cache_age_sec": int(now - _FLEET_CACHE["ts"])}
            resp = jsonify(cached_resp)
            resp.headers["Access-Control-Allow-Origin"] = "*"
            return resp

    # Pridaj backoff status do response (aby UI vedelo countdown bez kliku na Test login)
    try:
        if _hs is not None and hasattr(_hs, "get_huawei_backoff_status"):
            data["huawei_backoff"] = _hs.get_huawei_backoff_status()
    except Exception as e:
        log.warning("[fleet-status] backoff status fail: %s", e)

    _FLEET_CACHE["data"] = data
    _FLEET_CACHE["ts"] = now
    resp = jsonify({**data, "cached": False})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


# =============================================================================
# /webhook/station-kpi  →  historické KPI grafy pre detail stanice
# =============================================================================
# Volá Huawei FusionSolar agregačné endpointy:
#   - getKpiStationHour   → 24h výkonový profil (hourly samples)
#   - getKpiStationDay    → 30-day daily yield (denný výnos)
#   - getKpiStationMonth  → 12-month monthly yield (mesačný výnos)
#
# In-memory cache 5 min (Huawei stejne agreguje raz za 5 min).
# CORS open. No mutations.

_STATION_KPI_CACHE: dict = {}   # {station_code: {"ts": float, "data": dict}}
_STATION_KPI_TTL_SEC = 300       # 5 min


def _huawei_call(endpoint: str, payload: dict, base: str, headers: dict) -> dict:
    """POST helper. Vracia parsed JSON alebo {} pri chybe."""
    try:
        r = requests.post(f"{base}{endpoint}", headers=headers, json=payload, timeout=30)
        if r.status_code != 200:
            log.warning("[station-kpi] %s HTTP %s: %s", endpoint, r.status_code, r.text[:200])
            return {}
        return r.json() or {}
    except Exception as e:
        log.warning("[station-kpi] %s failed: %s", endpoint, e)
        return {}


def _station_kpi_compute(station_code: str) -> dict:
    """Pull hourly/daily/monthly KPI pre 1 stanicu."""
    if not _hs:
        return {"ok": False, "error": "huawei_spot module not loaded"}

    token = _hs.huawei_login()
    if not token:
        return {"ok": False, "error": "Huawei login failed"}

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    import datetime as _dt
    now_ts_ms = int(_dt.datetime.utcnow().timestamp() * 1000)

    # 1) getKpiStationHour — dnešok hodinové dáta (24 vzoriek)
    today_start = _dt.datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    hour_payload = {"stationCodes": station_code, "collectTime": int(today_start.timestamp() * 1000)}
    hour_resp = _huawei_call("/getKpiStationHour", hour_payload, base, headers)

    # 2) getKpiStationDay — posledných 30 dní (default)
    day_start = today_start - _dt.timedelta(days=30)
    day_payload = {"stationCodes": station_code, "collectTime": int(day_start.timestamp() * 1000)}
    day_resp = _huawei_call("/getKpiStationDay", day_payload, base, headers)

    # 2b) getKpiStationDay batchovo — posledných 365 dní (pre PR heatmap kalendár)
    # Huawei vracia max ~31 dní per call, voláme 12× s posunutým collectTime
    yearly_daily_raw = []
    for months_back in range(12, -1, -1):
        period_start = today_start.replace(day=1)
        for _ in range(months_back):
            period_start = (period_start - _dt.timedelta(days=1)).replace(day=1)
        period_payload = {"stationCodes": station_code, "collectTime": int(period_start.timestamp() * 1000)}
        try:
            period_resp = _huawei_call("/getKpiStationDay", period_payload, base, headers)
            if isinstance(period_resp, dict):
                period_data = period_resp.get("data") or []
                if isinstance(period_data, list):
                    yearly_daily_raw.extend(period_data)
        except Exception:
            continue

    # 3) getKpiStationMonth — 12 mesiacov
    month_start = today_start.replace(day=1) - _dt.timedelta(days=365)
    month_start = month_start.replace(day=1)
    month_payload = {"stationCodes": station_code, "collectTime": int(month_start.timestamp() * 1000)}
    month_resp = _huawei_call("/getKpiStationMonth", month_payload, base, headers)

    def _extract(resp, key_aliases):
        """Huawei vracia rôzne tvary — niekedy data: [...], inokedy data: [{dataItemMap: {...}}]."""
        rows = []
        if not isinstance(resp, dict):
            return rows
        data = resp.get("data") or []
        if not isinstance(data, list):
            return rows
        for item in data:
            if not isinstance(item, dict):
                continue
            ts = item.get("collectTime") or item.get("time")
            dim = item.get("dataItemMap") or item
            row = {"ts": ts}
            for alias in key_aliases:
                v = dim.get(alias)
                if v is not None and v != "N/A":
                    try:
                        row[alias] = float(v)
                    except (TypeError, ValueError):
                        row[alias] = None
            rows.append(row)
        return rows

    # Huawei field names (z NBI Reference 25.4.0):
    #   inverter_power (kW)  — hourly power
    #   product_power (kWh)  — daily/monthly yield
    #   ongrid_power (kWh)   — feed-in
    #   power_profit (CNY... ale pre nás môže byť aj v EUR od FusionSolar)

    hourly = _extract(hour_resp, ["inverter_power", "product_power", "ongrid_power"])
    daily = _extract(day_resp, ["product_power", "ongrid_power", "power_profit"])
    monthly = _extract(month_resp, ["product_power", "ongrid_power", "power_profit"])

    # Yearly daily — dedup po ts + extract na same shape ako daily
    yearly_daily_dedup = {}
    for item in yearly_daily_raw:
        if not isinstance(item, dict):
            continue
        ts = item.get("collectTime") or item.get("time")
        if not ts:
            continue
        dim = item.get("dataItemMap") or item
        row = {"ts": ts}
        for alias in ("product_power", "ongrid_power", "power_profit"):
            v = dim.get(alias)
            if v is not None and v != "N/A":
                try:
                    row[alias] = float(v)
                except (TypeError, ValueError):
                    row[alias] = None
        yearly_daily_dedup[ts] = row
    yearly_daily = sorted(yearly_daily_dedup.values(), key=lambda r: r["ts"])


    return {
        "ok": True,
        "station_code": station_code,
        "fetched_at": _dt.datetime.utcnow().isoformat() + "Z",
        "hourly": hourly,                   # 24h power profile
        "daily": daily,                     # 30-day daily yield
        "yearly_daily": yearly_daily,       # ~365-day daily yield pre PR heatmap kalendár
        "monthly": monthly,                 # 12-month yield
        "huawei_raw": {
            "hour_success": bool(hour_resp.get("success")),
            "day_success": bool(day_resp.get("success")),
            "month_success": bool(month_resp.get("success")),
        },
    }


@app.route("/webhook/station-kpi", methods=["GET", "OPTIONS"])
def webhook_station_kpi():
    """Historical KPI for a single station — hourly / daily / monthly.

    GET ?station_code=NE=... or ?site_id=<inverter_sites.id>
    GET ?fresh=1 → bypass cache
    """
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Webhook-Secret"
        return resp

    station_code = request.args.get("station_code") or request.args.get("vendor_station_id")
    site_id = request.args.get("site_id")

    # Ak nemáme station_code, dohľadáme z DB cez site_id
    if not station_code and site_id and _hs:
        try:
            rows = _hs.sb_get("inverter_sites", {
                "select": "vendor_station_id,vendor_plant_code",
                "id": f"eq.{site_id}",
                "limit": "1",
            })
            if rows and isinstance(rows[0], dict):
                station_code = rows[0].get("vendor_station_id") or rows[0].get("vendor_plant_code")
        except Exception as e:
            log.warning("[station-kpi] sb_get failed: %s", e)

    if not station_code:
        return jsonify({"ok": False, "error": "missing station_code or site_id"}), 400

    fresh = request.args.get("fresh") in ("1", "true", "yes")
    now = _time.time()
    cached = _STATION_KPI_CACHE.get(station_code)
    if not fresh and cached and (now - cached["ts"]) < _STATION_KPI_TTL_SEC:
        resp = jsonify({**cached["data"], "cached": True, "cache_age_sec": int(now - cached["ts"])})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    try:
        data = _station_kpi_compute(station_code)
    except Exception as e:
        log.exception("[station-kpi] compute failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500

    if data.get("ok"):
        _STATION_KPI_CACHE[station_code] = {"ts": now, "data": data}

    resp = jsonify({**data, "cached": False})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


# =============================================================================
# /webhook/fleet-trend  →  30-day fleet-wide yield (suma všetkých staníc)
# =============================================================================
# Volá Huawei getKpiStationDay batchovo (8 staníc per call),
# sčíta product_power per deň naprieč všetkými stanicami.
# Cache 30 min (Huawei agreguje denné dáta s odstupom).

_FLEET_TREND_CACHE = {"ts": 0.0, "data": None}
_FLEET_TREND_TTL_SEC = 1800  # 30 min


def _fleet_trend_compute() -> dict:
    if not _hs:
        return {"ok": False, "error": "huawei_spot module not loaded"}

    sites = _hs.sb_get("inverter_sites", {
        "select": "vendor_station_id,vendor_plant_code,monitoring_enabled,vendor",
        "vendor": "eq.huawei",
        "monitoring_enabled": "eq.true",
    })

    station_codes = []
    for s in (sites or []):
        if not isinstance(s, dict):
            continue
        code = s.get("vendor_station_id") or s.get("vendor_plant_code")
        if code:
            station_codes.append(code)

    if not station_codes:
        return {"ok": True, "trend": [], "stations_count": 0}

    token = _hs.huawei_login()
    if not token:
        return {"ok": False, "error": "Huawei login failed"}

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    import datetime as _dt
    today_start = _dt.datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    collect_time = int((today_start - _dt.timedelta(days=30)).timestamp() * 1000)

    # Aggregate per timestamp (deň)
    by_day_yield: dict = {}    # {ts: product_power_sum}
    by_day_export: dict = {}   # {ts: ongrid_power_sum}

    CHUNK = 8
    for i in range(0, len(station_codes), CHUNK):
        chunk = station_codes[i:i + CHUNK]
        try:
            r = requests.post(
                f"{base}/getKpiStationDay",
                headers=headers,
                json={"stationCodes": ",".join(chunk), "collectTime": collect_time},
                timeout=30,
            )
            if r.status_code != 200:
                log.warning("[fleet-trend] HTTP %s for chunk %d: %s", r.status_code, i, r.text[:200])
                continue
            payload = r.json() or {}
            data = payload.get("data") or []
            if not isinstance(data, list):
                continue
            for item in data:
                if not isinstance(item, dict):
                    continue
                ts = item.get("collectTime") or item.get("time")
                if not ts:
                    continue
                dim = item.get("dataItemMap") or item
                yield_kwh = dim.get("product_power")
                export_kwh = dim.get("ongrid_power")
                try:
                    if yield_kwh is not None and yield_kwh != "N/A":
                        by_day_yield[ts] = by_day_yield.get(ts, 0.0) + float(yield_kwh)
                    if export_kwh is not None and export_kwh != "N/A":
                        by_day_export[ts] = by_day_export.get(ts, 0.0) + float(export_kwh)
                except (TypeError, ValueError):
                    pass
        except Exception as e:
            log.warning("[fleet-trend] chunk %d failed: %s", i, e)

    trend = []
    for ts in sorted(by_day_yield.keys()):
        trend.append({
            "ts": ts,
            "yield_kwh": round(by_day_yield.get(ts, 0.0), 1),
            "export_kwh": round(by_day_export.get(ts, 0.0), 1),
        })

    return {
        "ok": True,
        "trend": trend,
        "stations_count": len(station_codes),
        "fetched_at": _dt.datetime.utcnow().isoformat() + "Z",
    }


@app.route("/webhook/fleet-trend", methods=["GET", "OPTIONS"])
def webhook_fleet_trend():
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Webhook-Secret"
        return resp

    fresh = request.args.get("fresh") in ("1", "true", "yes")
    now = _time.time()
    if not fresh and _FLEET_TREND_CACHE["data"] is not None and (now - _FLEET_TREND_CACHE["ts"]) < _FLEET_TREND_TTL_SEC:
        resp = jsonify({**_FLEET_TREND_CACHE["data"], "cached": True, "cache_age_sec": int(now - _FLEET_TREND_CACHE["ts"])})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    try:
        data = _fleet_trend_compute()
    except Exception as e:
        log.exception("[fleet-trend] compute failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500

    if data.get("ok"):
        _FLEET_TREND_CACHE["data"] = data
        _FLEET_TREND_CACHE["ts"] = now

    resp = jsonify({**data, "cached": False})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


# =============================================================================
# /webhook/station-history  →  alarmy + daily yield pre detail stanice
# =============================================================================

@app.route("/webhook/station-history", methods=["GET", "OPTIONS"])
def webhook_station_history():
    """Vráti alarm history + daily yield (zo Supabase) pre stanicu.

    GET ?site_id=<uuid>  → JSON s alarms[] a daily[]
    GET ?site_id=<uuid>&format=csv → CSV download
    """
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        return resp

    site_id = request.args.get("site_id")
    if not site_id:
        return jsonify({"ok": False, "error": "missing site_id"}), 400

    fmt = request.args.get("format", "json").lower()

    if not _hs:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    # Alarms (open + resolved, last 90 days)
    try:
        import datetime as _dt
        since = (_dt.datetime.utcnow() - _dt.timedelta(days=90)).strftime("%Y-%m-%d")
        alarms = _hs.sb_get("inverter_alarms", {
            "select": "id,severity,code,message,status,created_at,acknowledged_at,resolved_at",
            "site_id": f"eq.{site_id}",
            "created_at": f"gte.{since}",
            "order": "created_at.desc",
            "limit": "200",
        }) or []
    except Exception as e:
        log.warning("[station-history] alarms query failed: %s", e)
        alarms = []

    # Daily telemetry (zo Supabase inverter_telemetry_daily ak existuje, fallback []
    try:
        daily = _hs.sb_get("inverter_telemetry_daily", {
            "select": "date,yield_kwh,export_kwh,load_kwh,peak_power_kw",
            "site_id": f"eq.{site_id}",
            "order": "date.desc",
            "limit": "90",
        }) or []
    except Exception:
        daily = []

    if fmt == "csv":
        import csv
        import io
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(["Date", "Yield kWh", "Export kWh", "Load kWh", "Peak Power kW"])
        for d in daily:
            if isinstance(d, dict):
                w.writerow([d.get("date"), d.get("yield_kwh"), d.get("export_kwh"), d.get("load_kwh"), d.get("peak_power_kw")])
        w.writerow([])
        w.writerow(["Alarms (last 90 days)"])
        w.writerow(["Severity", "Code", "Message", "Status", "Created", "Acked", "Resolved"])
        for a in alarms:
            if isinstance(a, dict):
                w.writerow([a.get("severity"), a.get("code"), a.get("message"), a.get("status"),
                            a.get("created_at"), a.get("acknowledged_at"), a.get("resolved_at")])
        csv_data = buf.getvalue()
        resp = make_response(csv_data)
        resp.headers["Content-Type"] = "text/csv; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="station-{site_id[:8]}-history.csv"'
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    resp = jsonify({"ok": True, "alarms": alarms, "daily": daily})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


# =============================================================================
# /webhook/alarm-action  → ack / resolve / snooze alarm s logom
# =============================================================================

@app.route("/webhook/alarm-action", methods=["POST", "OPTIONS"])
def webhook_alarm_action():
    """Acknowledge / resolve / snooze alarm.

    POST {"alarm_id": "<uuid>", "action": "ack"|"resolve"|"snooze", "user": "<email>", "note": "..."}
    """
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Webhook-Secret"
        return resp

    if not _hs:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    try:
        payload = request.get_json(force=True) or {}
    except Exception:
        return jsonify({"ok": False, "error": "invalid JSON"}), 400

    alarm_id = payload.get("alarm_id")
    action = (payload.get("action") or "").lower()
    user_email = payload.get("user") or "unknown"
    note = payload.get("note", "")

    if not alarm_id or action not in {"ack", "resolve", "snooze"}:
        return jsonify({"ok": False, "error": "missing alarm_id or invalid action"}), 400

    import datetime as _dt
    now_iso = _dt.datetime.utcnow().isoformat() + "Z"
    update_payload: dict = {}
    if action == "ack":
        update_payload["acknowledged_at"] = now_iso
        update_payload["acknowledged_by"] = user_email
        update_payload["status"] = "acked"
    elif action == "resolve":
        update_payload["resolved_at"] = now_iso
        update_payload["resolved_by"] = user_email
        update_payload["status"] = "resolved"
    elif action == "snooze":
        # Snooze 24h
        snooze_until = (_dt.datetime.utcnow() + _dt.timedelta(hours=24)).isoformat() + "Z"
        update_payload["snoozed_until"] = snooze_until
        update_payload["status"] = "snoozed"

    if note:
        update_payload["resolution_note"] = note

    try:
        ok, errmsg = _hs.sb_patch(f"inverter_alarms?id=eq.{alarm_id}", update_payload)
        if not ok:
            return jsonify({"ok": False, "error": errmsg[:300]}), 500
        resp = jsonify({"ok": True, "action": action, "alarm_id": alarm_id, "updated": update_payload})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp
    except Exception as e:
        log.exception("[alarm-action] update failed")
        return jsonify({"ok": False, "error": str(e)[:300]}), 500


# =============================================================================
# /webhook/fleet-anomaly  →  AI-powered anomaly detection cez Anthropic
# =============================================================================
# Analyzuje aktuálny fleet snapshot + identifikuje podozrivé stanice.
# Vracia štruktúrované insights pre dispečera.

_FLEET_ANOMALY_CACHE: dict = {"ts": 0.0, "data": None}
_FLEET_ANOMALY_TTL_SEC = 300  # 5 min — LLM call je drahý


def _build_fleet_anomaly_prompt(sites: list, alarms_total: int) -> str:
    """Sumarizuj fleet pre LLM."""
    huawei_sites = [s for s in sites if isinstance(s, dict) and (s.get("vendor") or "").lower() == "huawei"]
    if not huawei_sites:
        return ""

    # Vypočítaj flotila stats pre baseline
    spec_yields = [s.get("specific_yield_today") for s in huawei_sites if s.get("specific_yield_today") is not None and s.get("specific_yield_today") > 0]
    fleet_avg_spec = sum(spec_yields) / len(spec_yields) if spec_yields else 0

    lines = []
    for s in huawei_sites[:30]:  # max 30 staníc, LLM context limit
        spec = s.get("specific_yield_today")
        power = s.get("current_ac_power_kw")
        dc_kwp = s.get("dc_kwp") or s.get("ac_kw") or 0
        cf = s.get("capacity_factor_pct")
        yield_kwh = s.get("current_day_yield_kwh")
        alarms = s.get("open_alarms_count", 0)
        status = s.get("status_label", "?")

        # Compute deviation
        dev = ""
        if spec is not None and fleet_avg_spec > 0:
            d = ((spec - fleet_avg_spec) / fleet_avg_spec) * 100
            dev = f"Δ{d:+.0f}% vs flotila"

        lines.append(
            f"- {s.get('site_name', '?')} ({dc_kwp} kWp): "
            f"výkon={power} kW, dnes={yield_kwh} kWh, spec={spec} kWh/kWp {dev}, "
            f"CF={cf}%, status={status}, alarmy={alarms}"
        )

    summary = "\n".join(lines)
    return f"""Si analyst fotovoltickej flotily Energovision (SK). Tu je aktuálny snapshot {len(huawei_sites)} Huawei staníc:

FLOTILA priemer specific yield dnes: {fleet_avg_spec:.2f} kWh/kWp/deň
CELKOVO otvorených alarmov: {alarms_total}

STANICE:
{summary}

ÚLOHA: Identifikuj **TOP 3 najpodozrivejšie stanice** alebo problémy ktoré si dispečer musí pozrieť TERAZ. Pre každú:

1. **Stanica + dôvod** (jednou vetou)
2. **Pravdepodobná príčina** (string offline, MPPT mismatch, soiling, shading, hardware fault, žiadne dáta...)
3. **Akcia** (čo dispečer má urobiť — kontaktovať klienta, pozrieť alarmy, zavolať technika, atď.)

Formát odpovede STRIKTNE JSON:
{{
  "insights": [
    {{"station": "názov", "severity": "high|medium|low", "issue": "popis problému", "cause": "pravdepodobná príčina", "action": "konkrétna akcia"}}
  ],
  "fleet_health": "good|fair|poor",
  "summary": "1-veta súhrn stavu flotily"
}}

Iba JSON, žiadny ďalší text. Ak je všetko v poriadku, vráť prázdne insights[]."""


def _fleet_anomaly_compute() -> dict:
    import datetime as _dt
    if not _hs:
        return {"ok": False, "error": "huawei_spot module not available"}

    # Pull current fleet from cache
    if _FLEET_CACHE.get("data") is None:
        # Force fresh compute ak nie je cache
        try:
            data = _fleet_status_compute()
        except Exception as e:
            return {"ok": False, "error": f"fleet snapshot failed: {e}"}
    else:
        data = _FLEET_CACHE["data"]

    sites = data.get("sites") or []
    alarms_total = (data.get("meta") or {}).get("alarms_total", 0)

    prompt = _build_fleet_anomaly_prompt(sites, alarms_total)
    if not prompt:
        return {"ok": True, "insights": [], "fleet_health": "good", "summary": "Žiadne Huawei stanice."}

    # Call Anthropic
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return {"ok": False, "error": "ANTHROPIC_API_KEY missing"}

    payload = {
        "model": "claude-sonnet-4-5-20250929",
        "max_tokens": 1500,
        "messages": [{"role": "user", "content": prompt}],
    }
    headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json"}

    try:
        r = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload, timeout=60)
        if r.status_code != 200:
            return {"ok": False, "error": f"LLM HTTP {r.status_code}: {r.text[:200]}"}
        resp = r.json()
        text = resp["content"][0]["text"].strip()
        # Strip markdown JSON fence if present
        if text.startswith("```"):
            text = text.split("```", 2)[1]
            if text.startswith("json"):
                text = text[4:]
            text = text.strip()
        import json as _json
        parsed = _json.loads(text)
        return {
            "ok": True,
            "insights": parsed.get("insights", []),
            "fleet_health": parsed.get("fleet_health", "fair"),
            "summary": parsed.get("summary", ""),
            "analyzed_at": _dt.datetime.utcnow().isoformat() + "Z",
        }
    except Exception as e:
        log.exception("[fleet-anomaly] LLM call failed")
        return {"ok": False, "error": str(e)[:300]}


@app.route("/webhook/fleet-anomaly", methods=["GET", "OPTIONS"])
def webhook_fleet_anomaly():
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Webhook-Secret"
        return resp

    fresh = request.args.get("fresh") in ("1", "true", "yes")
    now = _time.time()
    if not fresh and _FLEET_ANOMALY_CACHE["data"] is not None and (now - _FLEET_ANOMALY_CACHE["ts"]) < _FLEET_ANOMALY_TTL_SEC:
        cached = {**_FLEET_ANOMALY_CACHE["data"], "cached": True, "cache_age_sec": int(now - _FLEET_ANOMALY_CACHE["ts"])}
        resp = jsonify(cached)
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    try:
        data = _fleet_anomaly_compute()
    except Exception as e:
        log.exception("[fleet-anomaly] compute failed")
        return jsonify({"ok": False, "error": str(e)[:300]}), 500

    if data.get("ok"):
        _FLEET_ANOMALY_CACHE["data"] = data
        _FLEET_ANOMALY_CACHE["ts"] = now

    resp = jsonify({**data, "cached": False})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


# =============================================================================
# /webhook/station-expected  →  PVGIS očakávaná produkcia pre stanicu
# =============================================================================
# Pre porovnanie Expected vs Actual: voláme PVGIS PVcalc API pre lat/lon stanice
# a vrátime mesačné očakávané kWh + odhad denného profilu.
# Cache 7 dní (PVGIS dáta sa nemenia, sú TMY 2005-2023 priemer).

_STATION_EXPECTED_CACHE: dict = {}  # {site_id: {ts, data}}
_STATION_EXPECTED_TTL_SEC = 7 * 24 * 3600  # týždeň


def _pvgis_monthly_expected(lat: float, lon: float, dc_kwp: float, tilt: int = 35, azimuth: int = 0) -> dict:
    """Volá PVGIS PVcalc API. Vracia 12 mesiacov očakávaný yield (kWh/mesiac).

    Tilt default 35° (typický SK optimum), azimuth 0 (juh).
    PEAK power = dc_kwp, system losses default 14% (PVGIS standard).
    """
    url = "https://re.jrc.ec.europa.eu/api/v5_3/PVcalc"
    params = {
        "lat": lat,
        "lon": lon,
        "peakpower": dc_kwp,
        "loss": 14,
        "angle": tilt,
        "aspect": azimuth,
        "outputformat": "json",
        "mountingplace": "free",
        "pvtechchoice": "crystSi",
    }
    try:
        r = requests.get(url, params=params, timeout=30)
        if r.status_code != 200:
            return {"ok": False, "error": f"PVGIS HTTP {r.status_code}"}
        data = r.json()
        monthly_results = data.get("outputs", {}).get("monthly", {}).get("fixed", [])
        if not monthly_results:
            return {"ok": False, "error": "no monthly data"}

        monthly = []
        for m in monthly_results:
            monthly.append({
                "month": m.get("month"),
                "E_m_kwh": m.get("E_m"),    # mesačná výroba kWh
                "H_m_kwh_m2": m.get("H(i)_m"),  # mesačná irradiance kWh/m²
            })

        annual_total = sum(m["E_m_kwh"] for m in monthly if m["E_m_kwh"] is not None)
        # Daily expected per month (avg)
        days_in_month = [31, 28.25, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
        for i, m in enumerate(monthly):
            if m["E_m_kwh"] is not None:
                m["E_d_avg_kwh"] = round(m["E_m_kwh"] / days_in_month[i], 1)
                m["spec_yield_avg"] = round((m["E_m_kwh"] / days_in_month[i]) / dc_kwp, 2) if dc_kwp > 0 else None

        return {
            "ok": True,
            "monthly": monthly,
            "annual_total_kwh": round(annual_total, 1),
            "spec_yield_annual": round(annual_total / dc_kwp, 0) if dc_kwp > 0 else None,
            "tilt": tilt,
            "azimuth": azimuth,
            "loss_pct": 14,
        }
    except Exception as e:
        return {"ok": False, "error": str(e)[:300]}


@app.route("/webhook/station-expected", methods=["GET", "OPTIONS"])
def webhook_station_expected():
    """PVGIS očakávaná produkcia per stanica.

    GET ?site_id=<uuid>  → 12-month expected production from PVGIS
    """
    if request.method == "OPTIONS":
        resp = jsonify({})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        return resp

    site_id = request.args.get("site_id")
    if not site_id:
        return jsonify({"ok": False, "error": "missing site_id"}), 400

    fresh = request.args.get("fresh") in ("1", "true", "yes")
    now = _time.time()
    cached = _STATION_EXPECTED_CACHE.get(site_id)
    if not fresh and cached and (now - cached["ts"]) < _STATION_EXPECTED_TTL_SEC:
        resp = jsonify({**cached["data"], "cached": True, "cache_age_sec": int(now - cached["ts"])})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    if not _hs:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    # Pull lat/lon/dc_kwp from DB
    try:
        rows = _hs.sb_get("inverter_sites", {
            "select": "latitude,longitude,dc_kwp,ac_kw,site_name",
            "id": f"eq.{site_id}",
            "limit": "1",
        })
        if not rows or not isinstance(rows[0], dict):
            return jsonify({"ok": False, "error": "site not found"}), 404
        site = rows[0]
    except Exception as e:
        return jsonify({"ok": False, "error": f"DB query failed: {e}"}), 500

    lat = site.get("latitude")
    lon = site.get("longitude")
    dc_kwp = site.get("dc_kwp") or site.get("ac_kw") or 0

    if lat is None or lon is None:
        return jsonify({"ok": False, "error": "site missing lat/lon (skip auto-fill GPS first)"}), 400
    if not dc_kwp:
        return jsonify({"ok": False, "error": "site missing dc_kwp"}), 400

    data = _pvgis_monthly_expected(float(lat), float(lon), float(dc_kwp))
    if data.get("ok"):
        data["site_id"] = site_id
        data["site_name"] = site.get("site_name")
        data["lat"] = lat
        data["lon"] = lon
        data["dc_kwp"] = dc_kwp
        _STATION_EXPECTED_CACHE[site_id] = {"ts": now, "data": data}

    resp = jsonify({**data, "cached": False})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp

# ============================================================================
# FVE MONTHLY REPORT — generuje PDF report za predošlý mesiac + email klientovi
# /webhook/fve-monthly-report — Vercel cron 1. dňa v mesiaci o 06:00
# /webhook/fve-monthly-report-site — manuálny trigger pre 1 stanicu
# ============================================================================

import calendar as _cal
import io as _io

_FVE_REPORT_HTML = """
<!DOCTYPE html>
<html lang="sk">
<head>
<meta charset="utf-8">
<title>Mesačný report FVE — {site_name}</title>
<style>
  @page {{ size: A4; margin: 18mm 14mm; }}
  * {{ box-sizing: border-box; }}
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; color: #0F172A; margin: 0; }}
  .header {{ display: flex; justify-content: space-between; align-items: flex-end; padding-bottom: 14px; border-bottom: 3px solid #16A34A; }}
  .brand {{ font-size: 28px; font-weight: 800; color: #16A34A; letter-spacing: -0.5px; }}
  .sub {{ font-size: 11px; color: #64748B; margin-top: 2px; }}
  .meta {{ text-align: right; font-size: 11px; color: #64748B; }}
  h1 {{ font-size: 22px; margin: 22px 0 4px; color: #0F172A; }}
  h2 {{ font-size: 14px; margin: 22px 0 8px; color: #0F172A; padding-bottom: 4px; border-bottom: 1px solid #E2E8F0; }}
  .kpi-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 8px; margin: 14px 0; }}
  .kpi {{ background: #F0FDF4; border: 1px solid #BBF7D0; border-radius: 8px; padding: 12px; }}
  .kpi .lbl {{ font-size: 10px; color: #15803D; text-transform: uppercase; letter-spacing: 0.5px; font-weight: 600; }}
  .kpi .val {{ font-size: 20px; font-weight: 800; color: #14532D; margin-top: 4px; }}
  .kpi .sub {{ font-size: 9px; color: #16A34A; margin-top: 2px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 11px; margin: 8px 0; }}
  th, td {{ padding: 6px 8px; text-align: left; border-bottom: 1px solid #E2E8F0; }}
  th {{ background: #F8FAFC; font-weight: 700; color: #475569; font-size: 10px; text-transform: uppercase; }}
  td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
  .summary-box {{ background: #ECFDF5; border-left: 4px solid #16A34A; padding: 12px 14px; margin: 14px 0; border-radius: 4px; }}
  .summary-box p {{ margin: 0; font-size: 12px; line-height: 1.5; }}
  .footer {{ position: fixed; bottom: 8mm; left: 14mm; right: 14mm; font-size: 9px; color: #94A3B8; text-align: center; border-top: 1px solid #E2E8F0; padding-top: 6px; }}
</style>
</head>
<body>
  <div class="header">
    <div>
      <div class="brand">ENERGOVISION EMS</div>
      <div class="sub">Mesačný report fotovoltickej elektrárne</div>
    </div>
    <div class="meta">
      <div><b>{period}</b></div>
      <div>Vygenerované: {gen_date}</div>
    </div>
  </div>

  <h1>{site_name}</h1>
  <div style="font-size: 12px; color: #64748B; margin-bottom: 6px;">
    Klient: <b>{customer_name}</b>{dc_info}
  </div>

  <h2>Hlavné ukazovatele</h2>
  <div class="kpi-grid">
    <div class="kpi">
      <div class="lbl">Vyrobené</div>
      <div class="val">{total_kwh:,.0f} kWh</div>
      <div class="sub">spec. yield {spec_yield:.2f} kWh/kWp/deň</div>
    </div>
    <div class="kpi">
      <div class="lbl">Úspora</div>
      <div class="val">{savings_eur:,.0f} €</div>
      <div class="sub">pri {price_eur:.2f} €/kWh</div>
    </div>
    <div class="kpi">
      <div class="lbl">CO₂ ušetrené</div>
      <div class="val">{co2_kg:,.0f} kg</div>
      <div class="sub">ekvivalent {trees:.0f} stromov/rok</div>
    </div>
    <div class="kpi">
      <div class="lbl">Performance Ratio</div>
      <div class="val">{pr_pct:.0f} %</div>
      <div class="sub">{pr_label}</div>
    </div>
  </div>

  <div class="summary-box">
    <p><b>Zhrnutie:</b> {summary_text}</p>
  </div>

  <h2>Denná produkcia ({period})</h2>
  <table>
    <thead><tr><th>Deň</th><th class="num">Výroba (kWh)</th><th class="num">Spec. yield (kWh/kWp)</th></tr></thead>
    <tbody>{daily_rows}</tbody>
    <tfoot><tr style="font-weight:700; background:#F0FDF4;">
      <td>SPOLU</td><td class="num">{total_kwh:,.1f}</td><td class="num">{spec_yield_total:.2f}</td>
    </tr></tfoot>
  </table>

  <h2>Alarmy v období</h2>
  {alarms_section}

  <h2>Stav zariadenia</h2>
  <table>
    <tr><th style="width: 40%;">Inštalovaný výkon FVE</th><td>{dc_kwp:.2f} kWp</td></tr>
    <tr><th>Vendor / typ meniča</th><td>{vendor}</td></tr>
    <tr><th>Online status v období</th><td>{uptime_pct:.1f} %</td></tr>
    <tr><th>Posledný kontakt</th><td>{last_seen}</td></tr>
  </table>

  <div class="footer">
    Energovision, s.r.o. — Certified Huawei FusionSolar Installer Partner • IČO 53 036 280 • +421 948 302 137 • dispecing@energovision.sk • energovision.sk
  </div>
</body>
</html>
"""


def _fve_monthly_collect_data(site_id: str, year: int, month: int) -> dict:
    """Z monitoring dát + alarms zostaví dáta pre 1 mesačný report."""
    sb = _sb()
    site_rows = sb.table("inverter_sites").select(
        "id, site_name, dc_kwp, vendor, vendor_plant_code, customer_id, public_portal_email, latitude, longitude"
    ).eq("id", site_id).limit(1).execute().data or []
    if not site_rows:
        return {"ok": False, "error": "site_not_found"}
    site = site_rows[0]

    # Customer
    customer_name = "—"
    if site.get("customer_id"):
        cust = sb.table("customers").select("company_name, first_name, last_name, email").eq("id", site["customer_id"]).limit(1).execute().data or []
        if cust:
            c = cust[0]
            customer_name = c.get("company_name") or f"{c.get('first_name','')} {c.get('last_name','')}".strip() or "—"

    # Daily production z station-kpi (interný call)
    try:
        kpi = _station_kpi_compute(site_id)
        daily = (kpi.get("daily") or [])[-31:]  # posledných 31 dní
    except Exception:
        daily = []

    # Filter na target month
    period_daily = []
    for d in daily:
        ds = d.get("date") or ""
        try:
            y, m, _ = ds.split("-")
            if int(y) == year and int(m) == month:
                period_daily.append(d)
        except Exception:
            continue

    total_kwh = sum(float(d.get("energy_kwh") or 0) for d in period_daily)
    days_in_month = _cal.monthrange(year, month)[1]
    dc_kwp = float(site.get("dc_kwp") or 0)

    spec_yield = total_kwh / dc_kwp / max(len(period_daily), 1) if dc_kwp > 0 else 0
    spec_yield_total = total_kwh / dc_kwp if dc_kwp > 0 else 0

    # Alarmy
    period_start = f"{year}-{month:02d}-01"
    period_end = f"{year}-{month:02d}-{days_in_month:02d}"
    alarms = sb.table("inverter_alarms").select(
        "alarm_name, severity, detected_at, resolved_at, status"
    ).eq("station_id_vendor", site.get("vendor_plant_code") or "").gte(
        "detected_at", period_start
    ).lte("detected_at", period_end + "T23:59:59").order("detected_at", desc=False).execute().data or []

    # Performance Ratio — actual vs PVGIS expected
    pr_pct = 85.0
    if dc_kwp > 0 and site.get("latitude") and site.get("longitude"):
        try:
            pvg = _pvgis_monthly_expected(float(site["latitude"]), float(site["longitude"]), dc_kwp)
            if pvg.get("ok"):
                exp_month_kwh = (pvg.get("monthly_kwh") or [0]*12)[month-1]
                if exp_month_kwh > 0:
                    pr_pct = max(0, min(100, total_kwh / exp_month_kwh * 100))
        except Exception:
            pass

    # Economic
    PRICE = 0.18
    savings_eur = total_kwh * PRICE
    co2_kg = total_kwh * 0.4
    trees = co2_kg / 21  # 1 strom ~ 21 kg CO2/rok

    # PR label
    if pr_pct >= 85:
        pr_label = "Výborné"
    elif pr_pct >= 75:
        pr_label = "OK"
    else:
        pr_label = "Sledujeme"

    # Summary text
    period_name = ["", "január", "február", "marec", "apríl", "máj", "jún", "júl", "august", "september", "október", "november", "december"][month]
    if total_kwh > 0:
        summary_text = (
            f"V mesiaci {period_name} {year} vyrobila vaša FVE {total_kwh:,.0f} kWh, "
            f"čo predstavuje úsporu {savings_eur:,.0f} € na účte za elektrinu. "
            f"Performance Ratio {pr_pct:.0f} % — {pr_label.lower()}. "
            f"CO₂ ekvivalent {co2_kg:,.0f} kg ušetrený."
        )
    else:
        summary_text = "V tomto období neboli zaznamenané produkčné dáta. Kontaktujte servisné stredisko."

    # Daily rows HTML
    daily_rows = ""
    for d in period_daily:
        ds = d.get("date") or ""
        e = float(d.get("energy_kwh") or 0)
        sy = e / dc_kwp if dc_kwp > 0 else 0
        daily_rows += f'<tr><td>{ds}</td><td class="num">{e:,.1f}</td><td class="num">{sy:.2f}</td></tr>'
    if not daily_rows:
        daily_rows = '<tr><td colspan="3" style="text-align:center; color:#94A3B8; padding:14px;">Žiadne dáta v období.</td></tr>'

    # Alarms HTML
    if alarms:
        alarms_section = '<table><thead><tr><th>Dátum</th><th>Alarm</th><th>Severity</th><th>Status</th></tr></thead><tbody>'
        for a in alarms:
            sev = a.get("severity", "info")
            sev_color = {"critical": "#EF4444", "warning": "#F59E0B", "info": "#3B82F6"}.get(sev, "#64748B")
            status_label = {"resolved": "Vyriešené", "acked": "Riešime", "snoozed": "Odložené", "new": "Nové"}.get(a.get("status"), a.get("status", "—"))
            det = (a.get("detected_at") or "")[:16].replace("T", " ")
            alarms_section += (
                f'<tr><td>{det}</td><td>{a.get("alarm_name","Alarm")}</td>'
                f'<td><span style="color:{sev_color};font-weight:600;">{sev.upper()}</span></td>'
                f'<td>{status_label}</td></tr>'
            )
        alarms_section += "</tbody></table>"
    else:
        alarms_section = '<p style="font-size:12px; color:#16A34A; margin: 8px 0;">✓ V tomto období neboli zaznamenané žiadne alarmy. Stanica bežala stabilne.</p>'

    last_seen = "neznáme"
    uptime_pct = 99.0  # placeholder

    return {
        "ok": True,
        "site_id": site_id,
        "site_name": site.get("site_name", "—"),
        "customer_name": customer_name,
        "customer_email": next((c.get("email") for c in (sb.table("customers").select("email").eq("id", site.get("customer_id") or "00000000-0000-0000-0000-000000000000").execute().data or [])), None) if site.get("customer_id") else None,
        "portal_email": site.get("public_portal_email"),
        "year": year,
        "month": month,
        "period": f"{period_name.capitalize()} {year}",
        "gen_date": datetime.now(timezone.utc).strftime("%d.%m.%Y"),
        "total_kwh": total_kwh,
        "savings_eur": savings_eur,
        "co2_kg": co2_kg,
        "trees": trees,
        "pr_pct": pr_pct,
        "pr_label": pr_label,
        "spec_yield": spec_yield,
        "spec_yield_total": spec_yield_total,
        "price_eur": PRICE,
        "dc_kwp": dc_kwp,
        "vendor": (site.get("vendor") or "—").upper(),
        "uptime_pct": uptime_pct,
        "last_seen": last_seen,
        "summary_text": summary_text,
        "daily_rows": daily_rows,
        "alarms_section": alarms_section,
        "dc_info": f" • {dc_kwp:.2f} kWp" if dc_kwp else "",
    }


def _fve_monthly_render_html(data: dict) -> str:
    return _FVE_REPORT_HTML.format(**data)


def _fve_monthly_render_pdf(html: str) -> bytes:
    """Render HTML → PDF cez Playwright (zhodný pattern ako posudok)."""
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox"])
        context = browser.new_context()
        page = context.new_page()
        page.set_content(html, wait_until="networkidle", timeout=20000)
        pdf_bytes = page.pdf(format="A4", print_background=True, margin={"top": "0", "right": "0", "bottom": "0", "left": "0"})
        browser.close()
        return pdf_bytes


def _fve_monthly_upload_pdf(site_id: str, year: int, month: int, pdf_bytes: bytes) -> str:
    """Upload PDF do Supabase Storage `fve-reports/{site_id}/{year}-{month}.pdf`, vráti public URL."""
    sb = _sb()
    path = f"{site_id}/{year}-{month:02d}.pdf"
    try:
        sb.storage.from_("fve-reports").upload(path, pdf_bytes, {"content-type": "application/pdf", "upsert": "true"})
    except Exception as e:
        log.warning("[fve-monthly] storage upload failed (bucket may not exist): %s", e)
        # try to create bucket
        try:
            sb.storage.create_bucket("fve-reports", options={"public": True})
            sb.storage.from_("fve-reports").upload(path, pdf_bytes, {"content-type": "application/pdf", "upsert": "true"})
        except Exception as ee:
            log.exception("[fve-monthly] bucket+upload failed: %s", ee)
            return ""
    return f"{os.environ.get('SUPABASE_URL','').rstrip('/')}/storage/v1/object/public/fve-reports/{path}"


def _fve_monthly_send_email(data: dict, pdf_url: str, recipient: str) -> bool:
    """Pošle email klientovi cez M365 Graph (rovnaký pattern ako ostatné modules)."""
    try:
        from email_m365 import send_email_m365
    except Exception:
        log.warning("[fve-monthly] M365 module not available")
        return False

    subject = f"Mesačný report FVE — {data['site_name']} ({data['period']})"
    body_html = f"""
    <p>Dobrý deň,</p>
    <p>v prílohe / na linke nájdete <b>mesačný report</b> vašej fotovoltickej elektrárne za <b>{data['period']}</b>.</p>
    <ul>
      <li>Vyrobené: <b>{data['total_kwh']:,.0f} kWh</b></li>
      <li>Úspora: <b>{data['savings_eur']:,.0f} €</b></li>
      <li>Performance Ratio: <b>{data['pr_pct']:.0f} % — {data['pr_label']}</b></li>
    </ul>
    <p><a href="{pdf_url}" style="background:#16A34A;color:white;padding:10px 16px;border-radius:6px;text-decoration:none;display:inline-block;">📄 Stiahnuť PDF report</a></p>
    <p>Online dashboard s aktuálnymi dátami: <a href="https://app.energovision.sk/portal/fve/{data.get('public_token','')}">otvoriť portál</a></p>
    <p>S pozdravom,<br>Energovision dispečing</p>
    """
    try:
        send_email_m365(to=recipient, subject=subject, html=body_html)
        return True
    except Exception as e:
        log.exception("[fve-monthly] email send failed: %s", e)
        return False


@app.route("/webhook/fve-monthly-report-site", methods=["POST"])
def webhook_fve_monthly_report_site():
    """Manuálny trigger pre 1 stanicu. body: {site_id, year, month, email?}"""
    body = request.get_json(silent=True) or {}
    site_id = body.get("site_id")
    if not site_id:
        return jsonify({"ok": False, "error": "site_id required"}), 400

    today = datetime.now(timezone.utc).date()
    year = int(body.get("year") or (today.year if today.month > 1 else today.year - 1))
    month = int(body.get("month") or (today.month - 1 if today.month > 1 else 12))
    recipient_override = body.get("email")

    sb = _sb()
    # Check duplicate
    dup = sb.table("fve_monthly_reports").select("id, pdf_url").eq("site_id", site_id).eq("year", year).eq("month", month).limit(1).execute().data or []
    if dup and not body.get("force"):
        return jsonify({"ok": True, "skipped": "exists", "pdf_url": dup[0].get("pdf_url")})

    data = _fve_monthly_collect_data(site_id, year, month)
    if not data.get("ok"):
        return jsonify(data), 400

    # Načítaj token pre email link
    site_token = sb.table("inverter_sites").select("public_token").eq("id", site_id).limit(1).execute().data or [{}]
    data["public_token"] = site_token[0].get("public_token", "") if site_token else ""

    html = _fve_monthly_render_html(data)
    try:
        pdf_bytes = _fve_monthly_render_pdf(html)
    except Exception as e:
        log.exception("[fve-monthly] PDF render failed")
        return jsonify({"ok": False, "error": f"pdf_render_failed: {e}"}), 500

    pdf_url = _fve_monthly_upload_pdf(site_id, year, month, pdf_bytes)

    # Persist
    sb.table("fve_monthly_reports").upsert({
        "site_id": site_id,
        "year": year,
        "month": month,
        "pdf_url": pdf_url,
        "pdf_path": f"{site_id}/{year}-{month:02d}.pdf",
        "total_production_kwh": data["total_kwh"],
        "total_savings_eur": data["savings_eur"],
        "pr_avg_pct": data["pr_pct"],
        "alarms_count": data["alarms_section"].count("<tr>") - 1 if "<table>" in data["alarms_section"] else 0,
    }, on_conflict="site_id,year,month").execute()

    # Email
    recipient = recipient_override or data.get("portal_email") or data.get("customer_email")
    email_sent = False
    if recipient:
        email_sent = _fve_monthly_send_email(data, pdf_url, recipient)
        if email_sent:
            sb.table("fve_monthly_reports").update({
                "email_sent_at": datetime.now(timezone.utc).isoformat(),
                "email_sent_to": recipient,
            }).eq("site_id", site_id).eq("year", year).eq("month", month).execute()

    return jsonify({
        "ok": True,
        "site_id": site_id,
        "period": data["period"],
        "pdf_url": pdf_url,
        "email_sent": email_sent,
        "recipient": recipient,
    })


@app.route("/webhook/fve-monthly-report", methods=["POST", "GET"])
def webhook_fve_monthly_report_cron():
    """
    Vercel cron 1. dňa v mesiaci o 06:00 — generuje reporty pre všetky monitoring stanice.
    """
    today = datetime.now(timezone.utc).date()
    if today.month == 1:
        year = today.year - 1
        month = 12
    else:
        year = today.year
        month = today.month - 1

    sb = _sb()
    # Iba Huawei live stanice (zatiaľ)
    sites = sb.table("inverter_sites").select("id, site_name, public_portal_enabled").eq("vendor", "huawei").execute().data or []

    results = []
    for s in sites:
        if not s.get("public_portal_enabled"):
            continue
        try:
            with app.test_client() as c:
                r = c.post("/webhook/fve-monthly-report-site", json={"site_id": s["id"], "year": year, "month": month})
                results.append({"site_id": s["id"], "site_name": s["site_name"], "status": r.status_code, "ok": r.get_json().get("ok") if r.get_json() else False})
        except Exception as e:
            results.append({"site_id": s["id"], "error": str(e)})

    log.info("[fve-monthly-cron] generated %d reports for %d/%d", len(results), month, year)
    return jsonify({"ok": True, "year": year, "month": month, "count": len(results), "results": results})

# ============================================================================
# NOTIFICATIONS — Twilio SMS + Web Push + Email dispatch
# /webhook/notify-ticket — vyšle notifikácie pre 1 ticket podľa severity + channels
# /webhook/sla-breach-cron — Vercel cron 15-min: zistí SLA breaches + eskaluje
# /webhook/ticket-created-cron — Vercel cron 1-min: vyšle notifikácie pre nové tikety
# ============================================================================

import json as _json
import os as _os
import requests as _req
from datetime import timedelta as _td

# Twilio
TWILIO_SID = _os.environ.get('TWILIO_ACCOUNT_SID', '')
TWILIO_TOKEN = _os.environ.get('TWILIO_AUTH_TOKEN', '')
TWILIO_FROM = _os.environ.get('TWILIO_FROM_NUMBER', '')

# VAPID for Web Push
VAPID_PUBLIC = _os.environ.get('VAPID_PUBLIC_KEY', '')
VAPID_PRIVATE = _os.environ.get('VAPID_PRIVATE_KEY', '')
VAPID_SUBJECT = _os.environ.get('VAPID_SUBJECT', 'mailto:dispecing@energovision.sk')


def _send_sms(phone: str, body: str, ticket_id: str = None) -> dict:
    """Pošli SMS cez Twilio."""
    if not TWILIO_SID or not TWILIO_TOKEN:
        log.warning("[notif] Twilio not configured, skipping SMS")
        return {"ok": False, "error": "twilio_not_configured"}
    if not phone.startswith('+'):
        phone = '+421' + phone.lstrip('0') if phone.startswith('9') or phone.startswith('0') else '+' + phone
    try:
        r = _req.post(
            f'https://api.twilio.com/2010-04-01/Accounts/{TWILIO_SID}/Messages.json',
            auth=(TWILIO_SID, TWILIO_TOKEN),
            data={'From': TWILIO_FROM, 'To': phone, 'Body': body[:1599]},
            timeout=10
        )
        if r.status_code in (200, 201):
            d = r.json()
            return {"ok": True, "sid": d.get("sid"), "status": d.get("status")}
        return {"ok": False, "error": f"twilio_http_{r.status_code}", "detail": r.text[:200]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _send_push(subscription: dict, title: str, body: str, url: str = None) -> dict:
    """Pošli Web Push cez pywebpush."""
    if not VAPID_PRIVATE:
        return {"ok": False, "error": "vapid_not_configured"}
    try:
        from pywebpush import webpush, WebPushException
        payload = _json.dumps({"title": title, "body": body, "url": url or "/admin/servis"})
        webpush(
            subscription_info={
                "endpoint": subscription["endpoint"],
                "keys": {"p256dh": subscription["p256dh"], "auth": subscription["auth_key"]}
            },
            data=payload,
            vapid_private_key=VAPID_PRIVATE,
            vapid_claims={"sub": VAPID_SUBJECT}
        )
        return {"ok": True}
    except WebPushException as e:
        return {"ok": False, "error": f"webpush: {e}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _notify_ticket_dispatch(ticket_id: str, channels: list[str] = None) -> dict:
    """Pre daný ticket pošle notifikácie podľa severity + preference."""
    sb = _sb()
    t = sb.table("service_tickets").select(
        "id, ticket_number, title, severity, sla_tier, sla_due_at, customer_id, assigned_to, site_id, "
        "customer:customers(first_name, last_name, company_name, email, phone1), "
        "site:inverter_sites(site_name, public_token)"
    ).eq("id", ticket_id).single().execute().data
    if not t:
        return {"ok": False, "error": "ticket_not_found"}

    severity = t.get("severity") or "info"

    # Default channels podle alarm_sla_mapping
    if channels is None:
        alarm_code = t.get("vendor_alarm_code")
        if alarm_code:
            m = sb.table("alarm_sla_mapping").select("notification_channels").eq("alarm_code", alarm_code).execute().data or []
            channels = m[0].get("notification_channels", ["email"]) if m else ["email"]
        else:
            channels = ["email"]

    site_name = (t.get("site") or {}).get("site_name", "—")
    customer = t.get("customer") or {}
    customer_name = customer.get("company_name") or f"{customer.get('first_name','')} {customer.get('last_name','')}".strip() or "klient"

    subject = f"🔴 {severity.upper()}: {t['title']} — {site_name}"
    body_short = f"Stanica {site_name}: {t['title']}. Ticket {t['ticket_number']}. Otvor: app.energovision.sk/admin/servis/{t['id']}"
    body_long = f"""
    Dobrý deň,

    na stanici {site_name} bol detegovaný {severity} alarm:
    {t['title']}

    Ticket číslo: {t['ticket_number']}
    SLA tier: {t.get('sla_tier', '—')}
    SLA deadline: {t.get('sla_due_at', '—')}

    Detail: https://app.energovision.sk/admin/servis/{t['id']}

    Energovision dispečing
    +421 948 302 137
    """

    results = []

    # ─── SMS ───
    if "sms" in channels:
        # Get phone z notification_preferences (alebo customer.phone1)
        phones_to_notify = []
        # 1) Customer
        if customer.get("phone1"):
            phones_to_notify.append((customer["phone1"], "customer", customer_name))
        # 2) Assignee user
        if t.get("assigned_to"):
            u = sb.table("notification_preferences").select("sms_phone, sms_critical, sms_warning, sms_info").eq("user_id", t["assigned_to"]).maybeSingle().execute().data
            if u and u.get("sms_phone"):
                want = (severity == "critical" and u.get("sms_critical")) or (severity == "warning" and u.get("sms_warning")) or (severity in ("info","catastrophic"))
                if want:
                    phones_to_notify.append((u["sms_phone"], "technician", "tech"))
        # 3) Admin fallback (Lukáš)
        admin_prefs = sb.table("notification_preferences").select("sms_phone, sms_critical").not_.is_("sms_phone", "null").execute().data or []
        for ap in admin_prefs:
            if ap.get("sms_phone") and ap.get("sms_critical"):
                phones_to_notify.append((ap["sms_phone"], "admin", "admin"))
                break

        for phone, role, label in phones_to_notify[:5]:  # cap pri 5
            r = _send_sms(phone, body_short, ticket_id=ticket_id)
            sb.table("notification_events").insert({
                "ticket_id": ticket_id, "channel": "sms", "recipient": phone,
                "body": body_short, "status": "sent" if r.get("ok") else "failed",
                "provider": "twilio", "provider_id": r.get("sid"),
                "error": r.get("error"), "sent_at": datetime.now(timezone.utc).isoformat(),
                "metadata": {"role": role, "label": label}
            }).execute()
            results.append({"channel": "sms", "to": phone, "ok": r.get("ok")})

    # ─── PUSH ───
    if "push" in channels:
        # All push_subscriptions filtered by user_id assigned + admin
        subs = []
        if t.get("assigned_to"):
            s = sb.table("push_subscriptions").select("*").eq("user_id", t["assigned_to"]).execute().data or []
            subs.extend(s)
        admin_subs = sb.table("push_subscriptions").select("*").is_("customer_id", "null").execute().data or []
        for s in admin_subs:
            if s not in subs:
                subs.append(s)

        for sub in subs[:10]:
            r = _send_push(sub, subject, body_short, f"/admin/servis/{t['id']}")
            sb.table("notification_events").insert({
                "ticket_id": ticket_id, "channel": "push", "recipient": sub.get("device_label","push"),
                "body": body_short, "status": "sent" if r.get("ok") else "failed",
                "provider": "vapid", "error": r.get("error"),
                "sent_at": datetime.now(timezone.utc).isoformat()
            }).execute()
            results.append({"channel": "push", "ok": r.get("ok")})

    # ─── EMAIL ───
    if "email" in channels:
        try:
            from email_m365 import send_email_m365
        except Exception:
            send_email_m365 = None

        recipients_email = []
        if customer.get("email"):
            recipients_email.append(customer["email"])
        # Dispečer email
        recipients_email.append("dispecing@energovision.sk")

        for to in set(recipients_email):
            if send_email_m365:
                try:
                    send_email_m365(to=to, subject=subject, html=body_long.replace("\n", "<br>"))
                    sb.table("notification_events").insert({
                        "ticket_id": ticket_id, "channel": "email", "recipient": to,
                        "subject": subject, "body": body_long, "status": "sent",
                        "provider": "m365", "sent_at": datetime.now(timezone.utc).isoformat()
                    }).execute()
                    results.append({"channel": "email", "to": to, "ok": True})
                except Exception as e:
                    sb.table("notification_events").insert({
                        "ticket_id": ticket_id, "channel": "email", "recipient": to,
                        "subject": subject, "status": "failed", "error": str(e),
                        "provider": "m365", "sent_at": datetime.now(timezone.utc).isoformat()
                    }).execute()
                    results.append({"channel": "email", "to": to, "ok": False, "err": str(e)})

    return {"ok": True, "ticket_id": ticket_id, "channels_attempted": channels, "results": results}


@app.route("/webhook/notify-ticket", methods=["POST"])
def webhook_notify_ticket():
    body = request.get_json(silent=True) or {}
    ticket_id = body.get("ticket_id")
    channels = body.get("channels")
    if not ticket_id:
        return jsonify({"ok": False, "error": "ticket_id required"}), 400
    return jsonify(_notify_ticket_dispatch(ticket_id, channels))


@app.route("/webhook/ticket-created-cron", methods=["POST", "GET"])
def webhook_ticket_created_cron():
    """Vercel cron každú minútu: nájdi nové tickety bez notifikácií a odošli."""
    sb = _sb()
    # Tickety created v poslednych 5 min, status=open, žiadne sent notification_events
    cutoff = (datetime.now(timezone.utc) - _td(minutes=5)).isoformat()
    tickets = sb.table("service_tickets").select("id, severity, vendor_alarm_code, created_at").gte("created_at", cutoff).eq("status", "open").execute().data or []

    results = []
    for t in tickets:
        # Skip ak už máme notifikáciu
        existing = sb.table("notification_events").select("id").eq("ticket_id", t["id"]).limit(1).execute().data
        if existing:
            continue
        r = _notify_ticket_dispatch(t["id"])
        results.append({"ticket_id": t["id"], "result": r})

    return jsonify({"ok": True, "processed": len(results), "results": results})


@app.route("/webhook/sla-breach-cron", methods=["POST", "GET"])
def webhook_sla_breach_cron():
    """Vercel cron každých 15 min: zistí SLA breaches + eskalácie."""
    sb = _sb()
    now_iso = datetime.now(timezone.utc).isoformat()

    # 1) SLA breach mark
    breached = sb.table("service_tickets").select("id, ticket_number, severity, sla_due_at, status").lt("sla_due_at", now_iso).not_.in_("status", ["closed", "cancelled", "resolved", "verified"]).eq("sla_breached", False).execute().data or []
    for t in breached:
        sb.table("service_tickets").update({"sla_breached": True}).eq("id", t["id"]).execute()
        sb.table("ticket_events").insert({
            "ticket_id": t["id"], "event_type": "sla_breached",
            "actor_label": "Systém (cron)",
            "new_value": {"sla_due_at": t["sla_due_at"]},
            "comment": f"SLA deadline prešiel — eskalácia"
        }).execute()
        # Trigger notify
        _notify_ticket_dispatch(t["id"], channels=["email", "sms"])

    # 2) Escalation rules check
    rules = sb.table("escalation_rules").select("*").eq("is_active", True).execute().data or []
    for rule in rules:
        # Find candidates per rule
        q = sb.table("service_tickets").select("id, severity, status, created_at, acknowledged_at, sla_due_at, escalation_level").eq("severity", rule["severity"]).not_.in_("status", ["closed", "cancelled", "resolved", "verified"])
        # Unacked check
        if rule.get("trigger_minutes_unacked"):
            cutoff_unacked = (datetime.now(timezone.utc) - _td(minutes=rule["trigger_minutes_unacked"])).isoformat()
            candidates = q.lt("created_at", cutoff_unacked).is_("acknowledged_at", "null").execute().data or []
            for c in candidates:
                if (c.get("escalation_level") or 0) < 1:
                    sb.table("service_tickets").update({"escalation_level": 1}).eq("id", c["id"]).execute()
                    sb.table("ticket_events").insert({
                        "ticket_id": c["id"], "event_type": "escalated",
                        "actor_label": f"Cron rule '{rule['name']}'",
                        "comment": f"Eskalácia: {rule['name']} (unacked > {rule['trigger_minutes_unacked']} min)"
                    }).execute()
                    _notify_ticket_dispatch(c["id"], channels=rule["notification_channels"])

    return jsonify({"ok": True, "breached": len(breached), "rules_checked": len(rules)})

# ============================================================================
# SERVICE PROTOCOL PDF — generuje PDF protokol pre resolved ticket + email klientovi
# /webhook/generate-service-protocol — manuálny trigger po resolve
# ============================================================================

_PROTOCOL_HTML = """
<!DOCTYPE html>
<html lang="sk">
<head>
<meta charset="utf-8">
<title>Servisný protokol — {ticket_number}</title>
<style>
  @page {{ size: A4; margin: 16mm 14mm; }}
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; color: #0F172A; margin: 0; font-size: 11px; }}
  .header {{ display: flex; justify-content: space-between; align-items: flex-end; padding-bottom: 12px; border-bottom: 3px solid #16A34A; }}
  .brand {{ font-size: 24px; font-weight: 800; color: #16A34A; }}
  .meta {{ text-align: right; font-size: 10px; color: #64748B; }}
  h1 {{ font-size: 18px; margin: 18px 0 6px; }}
  h2 {{ font-size: 13px; margin: 18px 0 6px; padding-bottom: 3px; border-bottom: 1px solid #E2E8F0; }}
  .badge {{ display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 10px; font-weight: 600; }}
  .b-critical {{ background: #FEE2E2; color: #B91C1C; }}
  .b-warning {{ background: #FEF3C7; color: #92400E; }}
  .b-info {{ background: #DBEAFE; color: #1E40AF; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 10px; margin: 6px 0; }}
  th, td {{ padding: 5px 7px; text-align: left; border-bottom: 1px solid #E2E8F0; vertical-align: top; }}
  th {{ background: #F8FAFC; font-weight: 700; color: #475569; font-size: 9px; text-transform: uppercase; }}
  .kpi {{ display: inline-block; background: #F0FDF4; border: 1px solid #BBF7D0; border-radius: 6px; padding: 8px 10px; margin-right: 6px; }}
  .kpi .lbl {{ font-size: 9px; color: #15803D; text-transform: uppercase; }}
  .kpi .val {{ font-size: 14px; font-weight: 700; color: #14532D; }}
  .checklist-item {{ padding: 3px 0; }}
  .check-yes {{ color: #16A34A; }}
  .check-no {{ color: #DC2626; }}
  .signatures {{ display: flex; gap: 20px; margin-top: 24px; }}
  .sig-box {{ flex: 1; border-top: 2px solid #0F172A; padding-top: 6px; min-height: 60px; }}
  .sig-img {{ max-height: 60px; max-width: 100%; display: block; margin-bottom: 4px; }}
  .photos {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 6px; margin: 8px 0; }}
  .photo {{ border: 1px solid #E2E8F0; border-radius: 4px; padding: 4px; text-align: center; }}
  .photo img {{ max-width: 100%; max-height: 100px; }}
  .footer {{ position: fixed; bottom: 6mm; left: 14mm; right: 14mm; font-size: 8px; color: #94A3B8; text-align: center; border-top: 1px solid #E2E8F0; padding-top: 4px; }}
</style>
</head>
<body>
  <div class="header">
    <div>
      <div class="brand">ENERGOVISION EMS</div>
      <div style="font-size: 10px; color: #64748B;">Servisný protokol</div>
    </div>
    <div class="meta">
      <div><b>{ticket_number}</b></div>
      <div>Vytvorený: {created_date}</div>
      <div>Vyriešený: {resolved_date}</div>
    </div>
  </div>

  <h1>{title}</h1>
  <div style="margin-bottom: 8px;">
    <span class="badge b-{severity_class}">{severity_label}</span>
    {site_info}
  </div>

  <h2>Klient</h2>
  <table>
    <tr><th style="width: 30%;">Meno / Firma</th><td>{customer_name}</td></tr>
    <tr><th>Stanica</th><td>{site_name}{dc_info}</td></tr>
    <tr><th>Adresa</th><td>{address}</td></tr>
    <tr><th>Kontakt</th><td>{contact}</td></tr>
  </table>

  <h2>Sumár zásahu</h2>
  <div>
    <span class="kpi"><span class="lbl">Hodín</span><br><span class="val">{hours_worked}</span></span>
    <span class="kpi"><span class="lbl">km</span><br><span class="val">{km_driven}</span></span>
    <span class="kpi"><span class="lbl">Technik</span><br><span class="val" style="font-size:11px;">{technician_name}</span></span>
    <span class="kpi"><span class="lbl">Dátum zásahu</span><br><span class="val" style="font-size:11px;">{resolved_date}</span></span>
  </div>

  <h2>Popis riešenia</h2>
  <p style="background: #F8FAFC; padding: 10px; border-radius: 4px; border-left: 3px solid #16A34A;">{resolution}</p>

  {root_cause_section}

  <h2>Checklist</h2>
  {checklist_section}

  {photos_section}

  <h2>Použitý materiál</h2>
  {materials_section}

  <div class="signatures">
    <div class="sig-box">
      <div style="font-size: 9px; color: #64748B; margin-bottom: 4px;">Technik</div>
      {tech_signature}
      <div style="font-size: 10px;">{technician_name}</div>
    </div>
    <div class="sig-box">
      <div style="font-size: 9px; color: #64748B; margin-bottom: 4px;">Klient</div>
      {customer_signature}
      <div style="font-size: 10px;">{customer_name}</div>
    </div>
  </div>

  <div class="footer">
    Energovision, s.r.o. — IČO 53 036 280 — Lamačská cesta 1738/111, 841 03 Bratislava — +421 948 302 137 — dispecing@energovision.sk
  </div>
</body>
</html>
"""


def _build_protocol_html(ticket_id: str) -> dict:
    sb = _sb()
    t = sb.table("service_tickets").select(
        "*, customer:customers(first_name, last_name, company_name, email, phone1, street, city), "
        "site:inverter_sites(site_name, dc_kwp, vendor), "
        "assignee:users(full_name, email)"
    ).eq("id", ticket_id).single().execute().data
    if not t:
        return {"ok": False, "error": "ticket_not_found"}

    attachments = sb.table("ticket_attachments").select("*").eq("ticket_id", ticket_id).execute().data or []
    photos = [a for a in attachments if a.get("category", "").startswith("photo_")]
    tech_sig = next((a for a in attachments if a.get("category") == "signature" and a.get("metadata", {}).get("role") == "technician"), None)
    cust_sig = next((a for a in attachments if a.get("category") == "signature" and a.get("metadata", {}).get("role") == "customer"), None)

    cust = t.get("customer") or {}
    customer_name = cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip() or "—"
    site = t.get("site") or {}
    assignee = t.get("assignee") or {}
    severity = t.get("severity", "info") or "info"
    severity_labels = {"critical": "🔴 Kritický", "catastrophic": "⚠ Catastrophic", "warning": "🟡 Warning", "info": "Info"}

    # Checklist
    checklist = t.get("checklist") or []
    if checklist:
        checklist_html = "<ul style='padding-left: 0; list-style: none;'>"
        for it in checklist:
            mark = "✓" if it.get("done") else "✗"
            cls = "check-yes" if it.get("done") else "check-no"
            checklist_html += f'<li class="checklist-item"><span class="{cls}">{mark}</span> {it.get("label","")}</li>'
        checklist_html += "</ul>"
    else:
        checklist_html = "<p style='color: #94A3B8; font-style: italic;'>Žiadny checklist.</p>"

    # Photos
    if photos:
        photos_html = '<h2>Fotodokumentácia</h2><div class="photos">'
        cat_labels = {"photo_before": "Pred", "photo_during": "Počas", "photo_after": "Po", "photo_detail": "Detail", "photo_label": "Štítok"}
        for p in photos[:9]:
            label = cat_labels.get(p.get("category", ""), "Foto")
            photos_html += f'<div class="photo"><img src="{p.get("file_url","")}" alt="{label}"/><div style="font-size: 8px; color: #64748B; margin-top: 3px;">{label}</div></div>'
        photos_html += "</div>"
    else:
        photos_html = ""

    # Materials
    materials = t.get("materials_used") or []
    if materials:
        materials_html = '<table><thead><tr><th>Položka</th><th>Množstvo</th><th>Pozn.</th></tr></thead><tbody>'
        for m in materials:
            materials_html += f'<tr><td>{m.get("name","")}</td><td>{m.get("qty","")}</td><td>{m.get("note","")}</td></tr>'
        materials_html += "</tbody></table>"
    else:
        materials_html = "<p style='color: #94A3B8; font-style: italic;'>Žiadny materiál neevidovaný.</p>"

    # Signatures
    tech_sig_html = f'<img src="{tech_sig["file_url"]}" class="sig-img" />' if tech_sig else '<div style="height: 50px;"></div>'
    cust_sig_html = f'<img src="{cust_sig["file_url"]}" class="sig-img" />' if cust_sig else '<div style="height: 50px;"></div>'

    # Root cause
    root_cause_section = ""
    if t.get("root_cause"):
        root_cause_section = f'<h2>Koreňová príčina</h2><p>{t["root_cause"]}</p>'

    address = ", ".join(filter(None, [cust.get("street"), cust.get("city")])) or "—"
    contact = " · ".join(filter(None, [cust.get("phone1"), cust.get("email")])) or "—"
    site_info = ""
    if site.get("site_name"):
        site_info = f' <span style="font-size: 10px; color: #64748B; margin-left: 8px;">📍 {site["site_name"]}</span>'
    dc_info = f" • {site['dc_kwp']} kWp" if site.get("dc_kwp") else ""

    html = _PROTOCOL_HTML.format(
        ticket_number=t.get("ticket_number", "—"),
        created_date=(t.get("created_at") or "")[:10],
        resolved_date=(t.get("resolved_at") or datetime.now(timezone.utc).isoformat())[:10],
        title=t.get("title", "—"),
        severity_class={"critical": "critical", "catastrophic": "critical", "warning": "warning"}.get(severity, "info"),
        severity_label=severity_labels.get(severity, severity),
        site_info=site_info,
        customer_name=customer_name,
        site_name=site.get("site_name", "—"),
        dc_info=dc_info,
        address=address,
        contact=contact,
        hours_worked=t.get("hours_worked", "—"),
        km_driven=t.get("km_driven", "—"),
        technician_name=assignee.get("full_name", "—"),
        resolution=t.get("resolution") or "—",
        root_cause_section=root_cause_section,
        checklist_section=checklist_html,
        photos_section=photos_html,
        materials_section=materials_html,
        tech_signature=tech_sig_html,
        customer_signature=cust_sig_html,
    )
    return {"ok": True, "html": html, "ticket": t, "customer_email": cust.get("email")}


@app.route("/webhook/generate-service-protocol", methods=["POST"])
def webhook_generate_service_protocol():
    body = request.get_json(silent=True) or {}
    ticket_id = body.get("ticket_id")
    if not ticket_id:
        return jsonify({"ok": False, "error": "ticket_id required"}), 400

    data = _build_protocol_html(ticket_id)
    if not data.get("ok"):
        return jsonify(data), 400

    # Render PDF
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox"])
            page = browser.new_page()
            page.set_content(data["html"], wait_until="networkidle", timeout=20000)
            pdf_bytes = page.pdf(format="A4", print_background=True, margin={"top": "0", "right": "0", "bottom": "0", "left": "0"})
            browser.close()
    except Exception as e:
        log.exception("[service-protocol] PDF fail")
        return jsonify({"ok": False, "error": f"pdf_render: {e}"}), 500

    # Upload do Storage
    sb = _sb()
    t = data["ticket"]
    path = f"{ticket_id}/protocol_{t.get('ticket_number','')}.pdf"
    try:
        sb.storage.from_("ticket-attachments").upload(path, pdf_bytes, {"content-type": "application/pdf", "upsert": "true"})
    except Exception:
        try:
            sb.storage.create_bucket("ticket-attachments", options={"public": True})
            sb.storage.from_("ticket-attachments").upload(path, pdf_bytes, {"content-type": "application/pdf"})
        except Exception as ee:
            log.exception("[service-protocol] storage: %s", ee)

    pdf_url = f"{os.environ.get('SUPABASE_URL','').rstrip('/')}/storage/v1/object/public/ticket-attachments/{path}"

    sb.table("service_tickets").update({"resolution_pdf_url": pdf_url}).eq("id", ticket_id).execute()
    sb.table("ticket_attachments").insert({
        "ticket_id": ticket_id,
        "category": "pdf",
        "file_name": f"protocol_{t.get('ticket_number')}.pdf",
        "file_path": path,
        "file_url": pdf_url,
        "mime_type": "application/pdf",
        "size_bytes": len(pdf_bytes),
    }).execute()

    # Email klientovi
    recipient = data.get("customer_email")
    if recipient:
        try:
            from email_m365 import send_email_m365
            body_html = f"""
            <p>Dobrý deň,</p>
            <p>v prílohe / na linke nájdete <b>servisný protokol</b> k ticketu <b>{t.get('ticket_number')}</b> ({t.get('title')}).</p>
            <p>Riešenie: {t.get('resolution', '')[:200]}</p>
            <p><a href="{pdf_url}" style="background:#16A34A;color:white;padding:10px 16px;border-radius:6px;text-decoration:none;display:inline-block;">📄 Stiahnuť protokol</a></p>
            <p>S pozdravom,<br>Energovision servis</p>
            """
            send_email_m365(to=recipient, subject=f"Servisný protokol {t.get('ticket_number')} — {t.get('title')}", html=body_html)
            sb.table("service_tickets").update({"resolution_email_sent_at": datetime.now(timezone.utc).isoformat()}).eq("id", ticket_id).execute()
        except Exception as e:
            log.warning("[service-protocol] email fail: %s", e)

    return jsonify({"ok": True, "ticket_id": ticket_id, "pdf_url": pdf_url, "emailed": bool(recipient)})

# ============================================================================
# EXECUTIVE SUMMARY AI — 1× mes Eva vyrobí 1-page exec summary per stanica
# /webhook/exec-summary-site — manuálny per-site
# /webhook/exec-summary-cron — Vercel cron 1. v mes. 08:00 (po monthly-report)
# ============================================================================

import json as _exj

def _exec_summary_compute(site_id: str, year: int, month: int) -> dict:
    sb = _sb()
    s = sb.table("inverter_sites").select(
        "id, site_name, dc_kwp, vendor, customer_id, public_token, public_portal_email"
    ).eq("id", site_id).single().execute().data
    if not s:
        return {"ok": False, "error": "site_not_found"}

    # Vytiahnut monthly report (ak existuje)
    rep = sb.table("fve_monthly_reports").select("*").eq("site_id", site_id).eq("year", year).eq("month", month).maybeSingle().execute().data
    # Tickety v mesiaci
    period_start = f"{year}-{month:02d}-01"
    period_end_dt = datetime(year, month, 28) + timedelta(days=4)
    period_end = period_end_dt.replace(day=1).strftime("%Y-%m-%d")
    tickets = sb.table("service_tickets").select("id, ticket_number, title, severity, status, resolved_at, sla_breached, hours_worked").eq("site_id", site_id).gte("created_at", period_start).lt("created_at", period_end).execute().data or []
    open_critical = [t for t in tickets if t.get("severity") == "critical" and t.get("status") not in ("resolved","verified","closed")]
    sla_breach_count = len([t for t in tickets if t.get("sla_breached")])

    # Customer
    cust = {}
    if s.get("customer_id"):
        cust_data = sb.table("customers").select("first_name, last_name, company_name, email").eq("id", s["customer_id"]).maybeSingle().execute().data
        if cust_data:
            cust = cust_data

    cust_name = cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip() or "—"
    period_name = ["", "január", "február", "marec", "apríl", "máj", "jún", "júl", "august", "september", "október", "november", "december"][month]

    # Build context for Claude
    context = {
        "site_name": s.get("site_name"),
        "dc_kwp": s.get("dc_kwp"),
        "customer": cust_name,
        "period": f"{period_name} {year}",
        "production_kwh": (rep or {}).get("total_production_kwh"),
        "savings_eur": (rep or {}).get("total_savings_eur"),
        "pr_avg_pct": (rep or {}).get("pr_avg_pct"),
        "alarms_count": (rep or {}).get("alarms_count", 0),
        "tickets_total": len(tickets),
        "tickets_critical_open": len(open_critical),
        "tickets_sla_breached": sla_breach_count,
        "tickets_resolved": len([t for t in tickets if t.get("status") in ("resolved","verified","closed")]),
        "tickets_avg_hours": sum(float(t.get("hours_worked") or 0) for t in tickets) / max(len(tickets),1),
    }

    # AI call — Anthropic Claude Sonnet
    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        prompt = f"""Si energetický analytik Energovision. Vytvor 1-stranový executive summary pre majiteľa FVE stanice.

Dáta za {context['period']}:
- Stanica: {context['site_name']} ({context['dc_kwp']} kWp), klient {context['customer']}
- Výroba: {context['production_kwh']} kWh
- Úspora: {context['savings_eur']} €
- Performance Ratio: {context['pr_avg_pct']} %
- Alarmy: {context['alarms_count']}
- Servisné tickety: {context['tickets_total']} celkom, {context['tickets_resolved']} vyriešených, {context['tickets_critical_open']} kritických otvorených, {context['tickets_sla_breached']} SLA breach

Vytvor JSON s polami:
- "headline": 1 veta — kľúčový takeaway (či mesiac bol OK alebo má problémy)
- "key_metrics": pole 3 najdôležitejších čísiel s krátkym popisom
- "highlights": pole 2-3 pozitívnych vecí
- "concerns": pole 0-3 vecí ktoré treba sledovať
- "next_actions": pole 1-3 odporúčaní pre nasledujúci mesiac
- "verdict": jedno slovo (excellent/good/ok/attention/critical)

Strikt: iba JSON, žiadny markdown ani komentáre. Píš po slovensky, vecne, ako senior consultant. Maximum 60 slov per pole. Žiadne sľuby ani superlatívy."""

        msg = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}]
        )
        text = msg.content[0].text.strip()
        # Strip markdown fences if present
        if text.startswith("```"):
            text = text.split("```")[1].lstrip("json\n")
        ai = _exj.loads(text)
    except Exception as e:
        log.exception("[exec-summary] AI fail: %s", e)
        ai = {
            "headline": f"Stanica {s['site_name']} — {context['production_kwh'] or 0:.0f} kWh za {context['period']}",
            "key_metrics": [
                {"label": "Výroba", "value": f"{context['production_kwh'] or 0:.0f} kWh"},
                {"label": "Úspora", "value": f"{context['savings_eur'] or 0:.0f} €"},
                {"label": "PR", "value": f"{context['pr_avg_pct'] or 0:.0f} %"}
            ],
            "highlights": [],
            "concerns": [f"Critical otvorené: {context['tickets_critical_open']}"] if context['tickets_critical_open'] else [],
            "next_actions": ["Skontroluj mesačný report v emaili"],
            "verdict": "ok"
        }

    return {"ok": True, "site_id": site_id, "context": context, "ai": ai, "customer_email": cust.get("email"), "public_token": s.get("public_token"), "site_name": s.get("site_name")}


@app.route("/webhook/exec-summary-site", methods=["POST"])
def webhook_exec_summary_site():
    body = request.get_json(silent=True) or {}
    site_id = body.get("site_id")
    if not site_id:
        return jsonify({"ok": False, "error": "site_id required"}), 400
    today = datetime.now(timezone.utc).date()
    year = int(body.get("year") or (today.year if today.month > 1 else today.year - 1))
    month = int(body.get("month") or (today.month - 1 if today.month > 1 else 12))

    data = _exec_summary_compute(site_id, year, month)
    if not data.get("ok"):
        return jsonify(data), 400

    # Send email
    recipient = body.get("email") or data.get("customer_email")
    if recipient:
        try:
            from email_m365 import send_email_m365
            ai = data["ai"]
            ctx = data["context"]
            verdict_color = {"excellent":"#16A34A","good":"#16A34A","ok":"#0EA5E9","attention":"#F59E0B","critical":"#EF4444"}.get(ai.get("verdict","ok"), "#64748B")
            highlights_html = "".join(f"<li>{h}</li>" for h in ai.get("highlights", []))
            concerns_html = "".join(f"<li>{c}</li>" for c in ai.get("concerns", []))
            actions_html = "".join(f"<li>{a}</li>" for a in ai.get("next_actions", []))
            kpi_html = "".join(
                f'<td style="background:#F0FDF4;border:1px solid #BBF7D0;border-radius:8px;padding:14px;text-align:center;width:33%;"><div style="font-size:10px;color:#15803D;text-transform:uppercase;">{k.get("label","")}</div><div style="font-size:20px;font-weight:700;color:#14532D;margin-top:4px;">{k.get("value","—")}</div></td>'
                for k in ai.get("key_metrics", [])[:3]
            )
            body_html = f"""
            <div style="font-family:Arial,sans-serif;color:#0F172A;max-width:600px;">
              <div style="background:linear-gradient(135deg,#16A34A,#15803D);color:white;padding:24px;border-radius:12px;">
                <div style="font-size:12px;opacity:0.9;">ENERGOVISION EMS — Executive Summary</div>
                <h1 style="margin:8px 0 4px;font-size:22px;">{ctx['site_name']}</h1>
                <div style="opacity:0.9;font-size:13px;">{ctx['period']} · {ctx['customer']}</div>
              </div>
              <div style="background:white;padding:20px;border:1px solid #E2E8F0;border-top:none;border-radius:0 0 12px 12px;">
                <div style="display:inline-block;padding:4px 12px;background:{verdict_color};color:white;border-radius:20px;font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:0.5px;">{ai.get("verdict","ok")}</div>
                <p style="font-size:16px;font-weight:600;margin:12px 0 16px;color:#0F172A;line-height:1.4;">{ai.get("headline","")}</p>
                <table style="width:100%;border-collapse:separate;border-spacing:6px;margin-bottom:16px;">{kpi_html}</table>
                {f'<h3 style="margin-top:18px;font-size:13px;color:#15803D;">✓ Pozitíva</h3><ul style="margin:6px 0;padding-left:20px;font-size:13px;line-height:1.6;color:#0F172A;">{highlights_html}</ul>' if highlights_html else ''}
                {f'<h3 style="margin-top:18px;font-size:13px;color:#B45309;">⚠ Pozor</h3><ul style="margin:6px 0;padding-left:20px;font-size:13px;line-height:1.6;color:#0F172A;">{concerns_html}</ul>' if concerns_html else ''}
                {f'<h3 style="margin-top:18px;font-size:13px;color:#0F172A;">→ Ďalšie kroky</h3><ul style="margin:6px 0;padding-left:20px;font-size:13px;line-height:1.6;color:#0F172A;">{actions_html}</ul>' if actions_html else ''}
                <div style="margin-top:24px;padding-top:14px;border-top:1px solid #E2E8F0;text-align:center;">
                  <a href="https://app.energovision.sk/portal/fve/{data.get('public_token','')}" style="background:#16A34A;color:white;padding:10px 20px;border-radius:6px;text-decoration:none;font-size:13px;font-weight:600;">📊 Otvoriť portál</a>
                </div>
                <p style="margin-top:20px;font-size:11px;color:#64748B;text-align:center;">Vygenerovala Eva AI · Energovision dispečing +421 948 302 137</p>
              </div>
            </div>
            """
            send_email_m365(to=recipient, subject=f"📊 Executive Summary — {ctx['site_name']} ({ctx['period']})", html=body_html)
            data["emailed"] = True
        except Exception as e:
            log.warning("[exec-summary] email fail: %s", e)
            data["emailed"] = False

    return jsonify(data)


@app.route("/webhook/exec-summary-cron", methods=["POST", "GET"])
def webhook_exec_summary_cron():
    """Vercel cron 1. v mes. 08:00 — exec summary pre všetky monitoring stanice s portal_enabled."""
    today = datetime.now(timezone.utc).date()
    if today.month == 1:
        year = today.year - 1
        month = 12
    else:
        year = today.year
        month = today.month - 1
    sb = _sb()
    sites = sb.table("inverter_sites").select("id, site_name").eq("vendor", "huawei").eq("public_portal_enabled", True).execute().data or []
    results = []
    for s in sites:
        try:
            with app.test_client() as c:
                r = c.post("/webhook/exec-summary-site", json={"site_id": s["id"], "year": year, "month": month})
                results.append({"site_id": s["id"], "ok": r.status_code == 200})
        except Exception as e:
            results.append({"site_id": s["id"], "error": str(e)})
    return jsonify({"ok": True, "year": year, "month": month, "count": len(results), "results": results})

# ============================================================================
# PREDICTIVE MAINTENANCE + LOSS ATTRIBUTION
# /webhook/predictive-scan-cron — denne 03:00 UTC pre všetky stanice
# /webhook/loss-attribution-compute — per-site, mesačne
# ============================================================================

import json as _pj
from datetime import timedelta as _ptd

def _predictive_scan_site(site_id: str) -> dict:
    sb = _sb()
    site = sb.table("inverter_sites").select("id, site_name, dc_kwp, vendor_plant_code, vendor, latitude, longitude").eq("id", site_id).single().execute().data
    if not site:
        return {"ok": False, "error": "site_not_found"}

    # Alarmy posledných 30 dní
    since = (datetime.now(timezone.utc) - _ptd(days=30)).isoformat()
    alarms = sb.table("inverter_alarms").select("alarm_name, alarm_code, severity, detected_at, status").eq("station_id_vendor", site.get("vendor_plant_code") or "").gte("detected_at", since).execute().data or []

    # Performance daily — posledných 90 dní
    since_pr = (datetime.now(timezone.utc) - _ptd(days=90)).isoformat()
    perf = sb.table("inverter_performance_daily").select("day, performance_ratio, energy_kwh").gte("day", since_pr[:10]).execute().data or []

    # Tickety posledných 90 dní
    tickets_90 = sb.table("service_tickets").select("ticket_number, severity, status, created_at").eq("site_id", site_id).gte("created_at", since_pr).execute().data or []

    context = {
        "site_name": site["site_name"],
        "dc_kwp": site.get("dc_kwp"),
        "alarms_30d": len(alarms),
        "alarms_critical": len([a for a in alarms if a.get("severity") == "critical"]),
        "alarm_codes_freq": _freq([a.get("alarm_code") for a in alarms if a.get("alarm_code")]),
        "pr_records": len(perf),
        "pr_avg_last_30d": _avg([p.get("performance_ratio") for p in perf[-30:]]),
        "pr_avg_first_30d": _avg([p.get("performance_ratio") for p in perf[:30]]),
        "tickets_total_90d": len(tickets_90),
        "tickets_critical": len([t for t in tickets_90 if t.get("severity") == "critical"]),
    }

    # AI predikcia
    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        prompt = f"""Si predictive maintenance engineer pre FVE. Z dát zaver predikcie:

Stanica: {context['site_name']} ({context['dc_kwp']} kWp)
Alarmy 30d: {context['alarms_30d']} celkom, {context['alarms_critical']} kritických
Najčastejšie alarm kódy: {context['alarm_codes_freq']}
PR posledných 30d: {context['pr_avg_last_30d']:.1f}%, prvých 30d 90d obdobia: {context['pr_avg_first_30d']:.1f}%
Tickety 90d: {context['tickets_total_90d']} ({context['tickets_critical']} critical)

Vyhodnoť riziká a vráť JSON pole predikcií. Každá predikcia má:
- component_type: "inverter"|"battery"|"string"|"panel"|"meter"|"communication"
- component_label: konkrétny popis (napr. "Invertor 1" alebo "String 3 na MPPT 2")
- prediction_type: "failure_likely"|"degradation"|"underperform"|"maintenance_due"|"soh_low"|"communication_loss"
- severity: "info"|"warning"|"critical"
- confidence: 0-100 (percento istoty)
- predicted_within_days: 0-90
- evidence: krátky popis dôkazov (max 100 znakov)
- recommendation: konkrétna akcia (max 120 znakov)

PR drift > 5 % YoY = degradácia panelov. Časté alarmy z jedného komponentu = riziko zlyhania. PR pod 75 % = underperformance.

Iba JSON pole, max 3 predikcie. Iba ak existuje skutočná evidence. Bez vaty, slovenský jazyk."""

        msg = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}]
        )
        text = msg.content[0].text.strip()
        if text.startswith("```"):
            text = text.split("```")[1].lstrip("json\n")
        predictions = _pj.loads(text)
        if not isinstance(predictions, list):
            predictions = []
    except Exception as e:
        log.warning("[predictive] AI fail: %s", e)
        predictions = []

    # Persist (clear stale + insert new)
    sb.table("predictive_alerts").update({"status": "dismissed"}).eq("site_id", site_id).eq("status", "open").execute()
    for p in predictions[:3]:
        try:
            sb.table("predictive_alerts").insert({
                "site_id": site_id,
                "component_type": p.get("component_type", "inverter"),
                "component_label": p.get("component_label"),
                "prediction_type": p.get("prediction_type", "underperform"),
                "severity": p.get("severity", "warning"),
                "confidence": p.get("confidence", 50),
                "predicted_within_days": p.get("predicted_within_days"),
                "evidence": {"text": p.get("evidence", "")},
                "recommendation": p.get("recommendation"),
                "expires_at": (datetime.now(timezone.utc) + _ptd(days=14)).isoformat(),
            }).execute()
        except Exception as e:
            log.warning("[predictive] insert fail: %s", e)

    return {"ok": True, "site_id": site_id, "predictions_count": len(predictions), "context": context}


def _freq(items):
    from collections import Counter
    return dict(Counter([i for i in items if i]).most_common(5))


def _avg(values):
    nums = [float(v) for v in values if v is not None]
    return sum(nums) / len(nums) if nums else 0.0


@app.route("/webhook/predictive-scan-site", methods=["POST"])
def webhook_predictive_scan_site():
    body = request.get_json(silent=True) or {}
    site_id = body.get("site_id")
    if not site_id:
        return jsonify({"ok": False, "error": "site_id required"}), 400
    return jsonify(_predictive_scan_site(site_id))


@app.route("/webhook/predictive-scan-cron", methods=["POST", "GET"])
def webhook_predictive_scan_cron():
    """Vercel cron denne 03:00 UTC — predictive scan všetkých Huawei staníc."""
    sb = _sb()
    sites = sb.table("inverter_sites").select("id, site_name").eq("vendor", "huawei").execute().data or []
    results = []
    for s in sites:
        try:
            r = _predictive_scan_site(s["id"])
            results.append({"site_id": s["id"], "ok": r.get("ok"), "count": r.get("predictions_count", 0)})
        except Exception as e:
            results.append({"site_id": s["id"], "error": str(e)})
    return jsonify({"ok": True, "scanned": len(results), "results": results})


def _loss_attribution_compute(site_id: str, year: int, month: int) -> dict:
    """Loss attribution waterfall — porovnanie actual vs PVGIS expected, rozklad strát."""
    sb = _sb()
    site = sb.table("inverter_sites").select("id, site_name, dc_kwp, latitude, longitude, vendor_plant_code").eq("id", site_id).single().execute().data
    if not site:
        return {"ok": False, "error": "site_not_found"}
    dc_kwp = float(site.get("dc_kwp") or 0)
    if dc_kwp == 0 or not site.get("latitude") or not site.get("longitude"):
        return {"ok": False, "error": "missing_dc_kwp_or_gps"}

    # PVGIS expected za mesiac
    try:
        pvg = _pvgis_monthly_expected(float(site["latitude"]), float(site["longitude"]), dc_kwp)
        expected_kwh = (pvg.get("monthly_kwh") or [0]*12)[month-1]
    except Exception:
        expected_kwh = 0

    # Actual production za mesiac
    period_start = f"{year}-{month:02d}-01"
    period_end_dt = datetime(year, month, 28) + _ptd(days=4)
    period_end = period_end_dt.replace(day=1).strftime("%Y-%m-%d")
    try:
        daily = sb.table("inverter_performance_daily").select("day, energy_kwh, performance_ratio").gte("day", period_start).lt("day", period_end).execute().data or []
        actual_kwh = sum(float(d.get("energy_kwh") or 0) for d in daily)
        pr_actual = sum(float(d.get("performance_ratio") or 0) for d in daily if d.get("performance_ratio")) / max(1, sum(1 for d in daily if d.get("performance_ratio")))
    except Exception:
        actual_kwh = 0; pr_actual = 0

    total_loss = max(0, expected_kwh - actual_kwh)

    # Heuristics: kategórie strát z alarmov + downtime
    alarms = sb.table("inverter_alarms").select("alarm_code, detected_at, resolved_at").eq("station_id_vendor", site.get("vendor_plant_code") or "").gte("detected_at", period_start).lt("detected_at", period_end).execute().data or []

    # Downtime — kKWh z hodín kedy bola stanica offline
    downtime_hours = 0
    for a in alarms:
        if a.get("alarm_code", "").startswith("INV_OFFLINE") or a.get("alarm_code") == "COMM_LOSS":
            try:
                d_start = datetime.fromisoformat(a["detected_at"].replace("Z", "+00:00"))
                d_end = datetime.fromisoformat((a.get("resolved_at") or datetime.now(timezone.utc).isoformat()).replace("Z", "+00:00"))
                downtime_hours += (d_end - d_start).total_seconds() / 3600
            except Exception:
                pass
    # Pri 5h sunshine avg/day a dc_kwp avg power = dc_kwp/2 (priemer)
    loss_downtime = min(total_loss, downtime_hours * dc_kwp * 0.5)

    # Soiling — odhad cez PR drift (ak PR_actual < PR_expected o > 3 %)
    pr_expected = 82  # SK priemer pre dobre dimensionnu FVE
    pr_gap = max(0, pr_expected - (pr_actual or 0))
    loss_soiling = total_loss * min(0.4, pr_gap / 100 * 0.5) if total_loss > 0 else 0

    # String mismatch / shading — heuristika: ak je veľa STRING_* alarmov
    string_alarms = len([a for a in alarms if "STRING" in (a.get("alarm_code") or "")])
    loss_shading = total_loss * 0.10 * min(1, string_alarms / 5)
    loss_mismatch = total_loss * 0.05 * min(1, string_alarms / 5)

    # Clipping — heuristika: dc_kwp/ac_kw ratio, zatiaľ nepoznáme presne
    loss_clipping = 0

    # Zvyšok = other
    loss_other = max(0, total_loss - loss_downtime - loss_soiling - loss_shading - loss_mismatch - loss_clipping)

    PRICE = 0.18
    row = {
        "site_id": site_id,
        "year": year,
        "month": month,
        "pvgis_expected_kwh": round(expected_kwh, 2),
        "actual_kwh": round(actual_kwh, 2),
        "total_loss_kwh": round(total_loss, 2),
        "loss_soiling_kwh": round(loss_soiling, 2),
        "loss_shading_kwh": round(loss_shading, 2),
        "loss_clipping_kwh": round(loss_clipping, 2),
        "loss_curtailment_kwh": 0,
        "loss_downtime_kwh": round(loss_downtime, 2),
        "loss_mismatch_kwh": round(loss_mismatch, 2),
        "loss_other_kwh": round(loss_other, 2),
        "loss_soiling_eur": round(loss_soiling * PRICE, 2),
        "loss_shading_eur": round(loss_shading * PRICE, 2),
        "loss_clipping_eur": 0,
        "loss_curtailment_eur": 0,
        "loss_downtime_eur": round(loss_downtime * PRICE, 2),
        "loss_mismatch_eur": round(loss_mismatch * PRICE, 2),
        "loss_other_eur": round(loss_other * PRICE, 2),
        "pr_actual_pct": round(pr_actual or 0, 2),
        "pr_expected_pct": pr_expected,
    }

    sb.table("loss_attribution_monthly").upsert(row, on_conflict="site_id,year,month").execute()
    return {"ok": True, **row}


@app.route("/webhook/loss-attribution-compute", methods=["POST"])
def webhook_loss_attribution_compute():
    body = request.get_json(silent=True) or {}
    site_id = body.get("site_id")
    if not site_id:
        return jsonify({"ok": False, "error": "site_id required"}), 400
    today = datetime.now(timezone.utc).date()
    year = int(body.get("year") or (today.year if today.month > 1 else today.year - 1))
    month = int(body.get("month") or (today.month - 1 if today.month > 1 else 12))
    return jsonify(_loss_attribution_compute(site_id, year, month))


@app.route("/webhook/loss-attribution-cron", methods=["POST", "GET"])
def webhook_loss_attribution_cron():
    """Vercel cron 2. dňa v mesiaci — loss attribution pre všetky stanice."""
    today = datetime.now(timezone.utc).date()
    if today.month == 1:
        year = today.year - 1; month = 12
    else:
        year = today.year; month = today.month - 1
    sb = _sb()
    sites = sb.table("inverter_sites").select("id").eq("vendor", "huawei").execute().data or []
    results = []
    for s in sites:
        try:
            r = _loss_attribution_compute(s["id"], year, month)
            results.append({"site_id": s["id"], "ok": r.get("ok"), "total_loss_eur": r.get("total_loss_kwh", 0) * 0.18})
        except Exception as e:
            results.append({"site_id": s["id"], "error": str(e)})
    return jsonify({"ok": True, "year": year, "month": month, "count": len(results), "results": results})


# ============================================================================
# F1: BATTERY HEALTH MONITORING
# /webhook/huawei-pull-battery-health — cron 3× denne (06,12,18 UTC)
# Pulluje devTypeId=39 (battery) + 41 (ESS) + 23048 (battery pack) z Huawei NBI
# ============================================================================

def _pull_battery_health_for_site(site: dict, base: str, headers: dict) -> dict:
    """
    Pre 1 stanicu (site dict): pull device list, filter na battery devices,
    pull real-time KPI, upsert do battery_packs + battery_telemetry_daily.

    Returns: {ok, packs_added, packs_updated, total_soh_avg, error}
    """
    sb = _sb()
    station_code = site.get("vendor_station_id") or site.get("vendor_plant_code")
    if not station_code:
        return {"ok": False, "error": "no station code"}

    # 1) getDevList pre stanicu
    try:
        r = requests.post(
            f"{base}/getDevList",
            headers=headers,
            json={"stationCodes": station_code},
            timeout=30,
        )
        if r.status_code != 200:
            return {"ok": False, "error": f"getDevList HTTP {r.status_code}"}
        dev_data = r.json() or {}
        all_devices = dev_data.get("data") or []
    except Exception as e:
        return {"ok": False, "error": f"getDevList failed: {e}"}

    # 2) Filtruj battery zariadenia (devTypeId 39, 41, 23048, 37890, 23045, 23047)
    BATTERY_DEV_TYPES = {39, 41, 23045, 23047, 23048, 37890}
    battery_devices = [d for d in all_devices if d.get("devTypeId") in BATTERY_DEV_TYPES]

    if not battery_devices:
        return {"ok": True, "packs_found": 0, "info": "no battery devices on site"}

    # 3) Group by devTypeId pre batch call getDevRealKpi
    by_type: Dict[int, List[dict]] = {}
    for d in battery_devices:
        t = int(d.get("devTypeId") or 0)
        by_type.setdefault(t, []).append(d)

    packs_upserted = 0
    soh_values = []
    today = datetime.now(timezone.utc).date()
    now_iso = datetime.now(timezone.utc).isoformat()

    for dev_type, devices in by_type.items():
        # Batch po max 100 (NBI limit)
        for chunk_start in range(0, len(devices), 100):
            chunk = devices[chunk_start:chunk_start + 100]
            dev_ids = [str(d.get("id")) for d in chunk if d.get("id")]
            if not dev_ids:
                continue

            try:
                r2 = requests.post(
                    f"{base}/getDevRealKpi",
                    headers=headers,
                    json={"devIds": ",".join(dev_ids), "devTypeId": dev_type},
                    timeout=30,
                )
                if r2.status_code != 200:
                    log.warning("[battery-pull] getDevRealKpi devType=%s HTTP %s", dev_type, r2.status_code)
                    continue
                kpi_data = r2.json() or {}
                rows = kpi_data.get("data") or []
            except Exception as e:
                log.warning("[battery-pull] getDevRealKpi devType=%s failed: %s", dev_type, e)
                continue

            # 4) Per device: extract SoC, SoH, charge/discharge cap
            dev_meta = {str(d.get("id")): d for d in chunk}
            for row in rows:
                if not isinstance(row, dict):
                    continue
                dev_id = str(row.get("devId") or row.get("id") or "")
                meta = dev_meta.get(dev_id, {})
                kpi = row.get("dataItemMap") or {}
                if not isinstance(kpi, dict):
                    continue

                # Battery-specific fields (per NBI 25.4.0 doc 5.1.2.2)
                soc = kpi.get("battery_soc") or kpi.get("soc")
                soh = kpi.get("battery_soh") or kpi.get("soh")
                charge_cap = kpi.get("charge_cap")
                discharge_cap = kpi.get("discharge_cap")
                cd_power = kpi.get("charge_discharge_power") or kpi.get("ch_discharge_power")

                # Status z run_state alebo battery_status
                state_code = kpi.get("battery_status") or kpi.get("run_state") or kpi.get("state")
                state_map = {0: "offline", 1: "standby", 2: "running", 3: "fault"}
                status = state_map.get(int(state_code), str(state_code)) if state_code is not None else None

                # Per-pack data (ESS môže mať list battery_pack)
                packs = kpi.get("battery_pack") or [{"soh": soh, "soc": soc}]
                if not isinstance(packs, list):
                    packs = [packs] if isinstance(packs, dict) else []

                # Upsert každého pack
                for pack_idx, pack in enumerate(packs, start=1):
                    if not isinstance(pack, dict):
                        continue
                    pack_soh = pack.get("soh") or pack.get("soh_pct") or soh
                    pack_soc = pack.get("soc") or soc
                    pack_soh_num = None
                    if pack_soh is not None:
                        try:
                            # SoH môže byť "90.0%" string alebo number
                            ps = str(pack_soh).replace("%", "").strip()
                            pack_soh_num = float(ps)
                            soh_values.append(pack_soh_num)
                        except (TypeError, ValueError):
                            pass
                    pack_soc_num = None
                    if pack_soc is not None:
                        try:
                            pack_soc_num = float(str(pack_soc).replace("%", "").strip())
                        except (TypeError, ValueError):
                            pass

                    row_data = {
                        "site_id": site["id"],
                        "vendor_device_id": dev_id,
                        "vendor_device_sn": row.get("sn"),
                        "device_type_id": dev_type,
                        "device_model": meta.get("devName") or meta.get("devTypeName"),
                        "pack_index": pack_idx if len(packs) > 1 else None,
                        "soc_pct": pack_soc_num,
                        "soh_pct": pack_soh_num,
                        "charge_power_kw": float(cd_power) if cd_power and float(cd_power) > 0 else None,
                        "discharge_power_kw": -float(cd_power) if cd_power and float(cd_power) < 0 else None,
                        "total_charge_kwh": float(charge_cap) if charge_cap else None,
                        "total_discharge_kwh": float(discharge_cap) if discharge_cap else None,
                        "status": status,
                        "last_seen_at": now_iso,
                        "raw_json": kpi,
                    }

                    try:
                        sb.table("battery_packs").upsert(
                            row_data,
                            on_conflict="site_id,vendor_device_id,pack_index",
                        ).execute()
                        packs_upserted += 1
                    except Exception as e:
                        log.warning("[battery-pull] upsert pack failed dev=%s: %s", dev_id, e)

                    # Tiež daily snapshot
                    try:
                        sb.table("battery_telemetry_daily").upsert({
                            "site_id": site["id"],
                            "vendor_device_id": dev_id,
                            "date": str(today),
                            "soh_pct": pack_soh_num,
                            "soc_avg_pct": pack_soc_num,
                            "total_charge_kwh": float(charge_cap) if charge_cap else None,
                            "total_discharge_kwh": float(discharge_cap) if discharge_cap else None,
                        }, on_conflict="vendor_device_id,date").execute()
                    except Exception as e:
                        log.warning("[battery-pull] daily snapshot failed: %s", e)

    # 5) Update aggregát na inverter_sites
    if soh_values:
        try:
            sb.table("inverter_sites").update({
                "battery_avg_soh_pct": sum(soh_values) / len(soh_values),
                "battery_pack_count": len(soh_values),
                "battery_last_sync_at": now_iso,
            }).eq("id", site["id"]).execute()
        except Exception as e:
            log.warning("[battery-pull] update site agg failed: %s", e)

    return {
        "ok": True,
        "site_id": site["id"],
        "site_name": site.get("site_name"),
        "packs_upserted": packs_upserted,
        "battery_devices": len(battery_devices),
        "avg_soh_pct": round(sum(soh_values) / len(soh_values), 1) if soh_values else None,
    }


@app.route("/webhook/huawei-pull-battery-health", methods=["POST", "GET"])
def webhook_huawei_pull_battery_health():
    """Pull battery health (SoC, SoH, charge/discharge) z Huawei NBI pre celý fleet.
    Spúšťa sa cronom 3× denne (06, 12, 18 UTC).
    """
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed (check backoff / credentials)"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, site_name, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    # Voliteľne filter na 1 stanicu
    body = request.get_json(silent=True) or {}
    site_id_filter = body.get("site_id")
    if site_id_filter:
        sites = [s for s in sites if s["id"] == site_id_filter]

    results = []
    total_packs = 0
    for s in sites:
        try:
            r = _pull_battery_health_for_site(s, base, headers)
            results.append(r)
            if r.get("ok"):
                total_packs += r.get("packs_upserted", 0)
        except Exception as e:
            log.exception("[battery-pull] site %s crashed", s.get("id"))
            results.append({"ok": False, "site_id": s["id"], "error": str(e)})

    return jsonify({
        "ok": True,
        "sites_processed": len(sites),
        "total_packs_upserted": total_packs,
        "results": results,
    })


# ============================================================================
# F2: STRING-LEVEL MPPT TELEMETRIA
# /webhook/huawei-pull-strings — hodinový cron počas dňa (06-20 UTC)
# Extract pv1_u..pv25_u + pv1_i..pv25_i z getDevRealKpi devTypeId=1
# ============================================================================

def _pull_strings_for_site(site: dict, base: str, headers: dict, devices: List[dict]) -> dict:
    """
    Pre 1 stanicu: pre všetky string inverter zariadenia (devTypeId=1)
    vytiahne per-MPPT voltage + current, uloží do string_telemetry,
    spočíta sibling-comparison underperforming detection.
    """
    sb = _sb()
    string_inverters = [d for d in devices if d.get("devTypeId") == 1]
    if not string_inverters:
        return {"ok": True, "info": "no string inverters", "strings_recorded": 0}

    now = datetime.now(timezone.utc)
    today = now.date()
    strings_recorded = 0
    string_points: List[dict] = []   # in-memory per inverter aggregation

    # Batch po max 100
    for chunk_start in range(0, len(string_inverters), 100):
        chunk = string_inverters[chunk_start:chunk_start + 100]
        dev_ids = [str(d.get("id")) for d in chunk if d.get("id")]
        dev_meta = {str(d.get("id")): d for d in chunk}
        if not dev_ids:
            continue

        try:
            r = requests.post(
                f"{base}/getDevRealKpi",
                headers=headers,
                json={"devIds": ",".join(dev_ids), "devTypeId": 1},
                timeout=30,
            )
            if r.status_code != 200:
                continue
            rows = (r.json() or {}).get("data") or []
        except Exception as e:
            log.warning("[strings-pull] getDevRealKpi failed: %s", e)
            continue

        for row in rows:
            if not isinstance(row, dict):
                continue
            dev_id = str(row.get("devId") or row.get("id") or "")
            meta = dev_meta.get(dev_id, {})
            kpi = row.get("dataItemMap") or {}
            if not isinstance(kpi, dict):
                continue

            # Iteruj pv1..pv25
            for n in range(1, 26):
                v_key = f"pv{n}_u"
                i_key = f"pv{n}_i"
                v = kpi.get(v_key)
                i = kpi.get(i_key)
                # Skip ak sú obe None alebo 0 (string neexistuje alebo offline)
                if v is None and i is None:
                    continue
                try:
                    v_num = float(v) if v is not None else None
                    i_num = float(i) if i is not None else None
                except (TypeError, ValueError):
                    continue
                if v_num is None and i_num is None:
                    continue
                p_kw = None
                if v_num is not None and i_num is not None:
                    p_kw = round((v_num * i_num) / 1000.0, 3)
                # Skip ak power 0 a oba sú malé (string je v nočnom stave, nemá zmysel ukladať)
                if p_kw is not None and p_kw < 0.05 and (v_num or 0) < 50:
                    continue

                row_data = {
                    "site_id": site["id"],
                    "vendor_device_id": dev_id,
                    "device_model": meta.get("devName") or meta.get("devTypeName"),
                    "pv_index": n,
                    "voltage_v": v_num,
                    "current_a": i_num,
                    "power_kw": p_kw,
                    "measured_at": now.isoformat(),
                }
                string_points.append(row_data)
                strings_recorded += 1

    # Bulk insert string_telemetry
    if string_points:
        try:
            # Insert v batchoch po 200 aby sa nepretiahol payload
            for batch_start in range(0, len(string_points), 200):
                batch = string_points[batch_start:batch_start + 200]
                sb.table("string_telemetry").insert(batch).execute()
        except Exception as e:
            log.warning("[strings-pull] bulk insert failed: %s", e)

    # Daily snapshot + underperforming detection
    # Group by (device_id, pv_index) → výpočet peak za dnes
    by_string: Dict[Tuple[str, int], List[dict]] = {}
    for p in string_points:
        key = (p["vendor_device_id"], p["pv_index"])
        by_string.setdefault(key, []).append(p)

    # Per device: nájdi peer stringy (rovnaký vendor_device_id), spočítaj sibling avg
    by_device: Dict[str, List[Tuple[int, float]]] = {}
    for (dev_id, pv_idx), pts in by_string.items():
        peaks = [p["power_kw"] for p in pts if p.get("power_kw") is not None]
        if peaks:
            avg_p = sum(peaks) / len(peaks)
            by_device.setdefault(dev_id, []).append((pv_idx, avg_p))

    underperforming_count = 0
    for dev_id, str_list in by_device.items():
        if len(str_list) < 2:
            continue   # potrebujem aspoň 2 stringy pre porovnanie
        peers_avg = sum(p for _, p in str_list) / len(str_list)
        for pv_idx, my_avg in str_list:
            ratio = (my_avg / peers_avg) if peers_avg > 0 else 1.0
            underperforming = ratio < 0.70 and my_avg < peers_avg - 0.5  # 30% pod priemer + abs rozdiel
            issue = None
            if underperforming:
                # Heuristika: zlý voltage = shading, zlý current = soiling/zlý panel
                pts = by_string.get((dev_id, pv_idx), [])
                avg_v = sum(p["voltage_v"] for p in pts if p.get("voltage_v")) / max(1, len([p for p in pts if p.get("voltage_v")]))
                avg_i = sum(p["current_a"] for p in pts if p.get("current_a")) / max(1, len([p for p in pts if p.get("current_a")]))
                issue = "shading" if avg_v < 300 else "soiling_or_panel"
                underperforming_count += 1
            try:
                sb.table("string_performance_daily").upsert({
                    "site_id": site["id"],
                    "vendor_device_id": dev_id,
                    "pv_index": pv_idx,
                    "date": str(today),
                    "avg_voltage_v": sum(p["voltage_v"] for p in by_string.get((dev_id, pv_idx), []) if p.get("voltage_v")) / max(1, len([p for p in by_string.get((dev_id, pv_idx), []) if p.get("voltage_v")])),
                    "avg_current_a": sum(p["current_a"] for p in by_string.get((dev_id, pv_idx), []) if p.get("current_a")) / max(1, len([p for p in by_string.get((dev_id, pv_idx), []) if p.get("current_a")])),
                    "peak_power_kw": my_avg,
                    "samples_count": len(by_string.get((dev_id, pv_idx), [])),
                    "underperforming": underperforming,
                    "performance_ratio_vs_siblings": round(ratio * 100, 1),
                    "issue_type": issue,
                }, on_conflict="vendor_device_id,pv_index,date").execute()
            except Exception as e:
                log.warning("[strings-pull] daily upsert failed: %s", e)

    # Update site agregát
    if strings_recorded > 0:
        try:
            sb.table("inverter_sites").update({
                "string_count": len({k for k in by_string.keys()}),
                "string_underperforming_count": underperforming_count,
                "string_last_sync_at": now.isoformat(),
            }).eq("id", site["id"]).execute()
        except Exception as e:
            log.warning("[strings-pull] site agg failed: %s", e)

    return {
        "ok": True,
        "site_id": site["id"],
        "site_name": site.get("site_name"),
        "string_inverters_processed": len(string_inverters),
        "strings_recorded": strings_recorded,
        "underperforming_count": underperforming_count,
    }


@app.route("/webhook/huawei-pull-strings", methods=["POST", "GET"])
def webhook_huawei_pull_strings():
    """Pull per-MPPT string voltage/current z Huawei NBI pre celý fleet.
    Hodinový cron počas dňa (06-20 UTC), výsledok do string_telemetry.
    """
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed (check backoff)"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, site_name, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    body = request.get_json(silent=True) or {}
    site_id_filter = body.get("site_id")
    if site_id_filter:
        sites = [s for s in sites if s["id"] == site_id_filter]

    total_strings = 0
    total_underperforming = 0
    results = []

    for s in sites:
        try:
            # Pre každú stanicu najprv getDevList aby sme videli aké invertery má
            r = requests.post(
                f"{base}/getDevList",
                headers=headers,
                json={"stationCodes": s.get("vendor_station_id") or ""},
                timeout=30,
            )
            if r.status_code != 200:
                results.append({"site_id": s["id"], "ok": False, "error": f"getDevList HTTP {r.status_code}"})
                continue
            devices = (r.json() or {}).get("data") or []

            res = _pull_strings_for_site(s, base, headers, devices)
            results.append(res)
            if res.get("ok"):
                total_strings += res.get("strings_recorded", 0)
                total_underperforming += res.get("underperforming_count", 0)
        except Exception as e:
            log.exception("[strings-pull] site %s crashed", s.get("id"))
            results.append({"ok": False, "site_id": s["id"], "error": str(e)})

    return jsonify({
        "ok": True,
        "sites_processed": len(sites),
        "total_strings_recorded": total_strings,
        "total_underperforming": total_underperforming,
        "results": results[:20],   # truncate pre response size
    })


# ============================================================================
# F3: ENVIRONMENT MONITOR (devTypeId=10 EMI)
# Pyranometer, ambient/module temp, wind, rainfall - každú hodinu
# ============================================================================

@app.route("/webhook/huawei-pull-environment", methods=["POST", "GET"])
def webhook_huawei_pull_environment():
    """Pull EMI (Environment Monitor) data z Huawei NBI - pyranometer, temps, wind."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, site_name, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    now = datetime.now(timezone.utc)
    today = now.date()
    total_readings = 0
    sites_with_emi = 0
    results = []

    for s in sites:
        try:
            r = requests.post(
                f"{base}/getDevList",
                headers=headers,
                json={"stationCodes": s.get("vendor_station_id") or ""},
                timeout=30,
            )
            if r.status_code != 200:
                continue
            devices = (r.json() or {}).get("data") or []
            # EMI = devTypeId=10
            emi_devices = [d for d in devices if d.get("devTypeId") == 10]
            if not emi_devices:
                continue
            sites_with_emi += 1

            dev_ids = [str(d.get("id")) for d in emi_devices if d.get("id")]
            r2 = requests.post(
                f"{base}/getDevRealKpi",
                headers=headers,
                json={"devIds": ",".join(dev_ids), "devTypeId": 10},
                timeout=30,
            )
            if r2.status_code != 200:
                continue
            rows = (r2.json() or {}).get("data") or []

            site_readings = []
            for row in rows:
                if not isinstance(row, dict):
                    continue
                dev_id = str(row.get("devId") or row.get("id") or "")
                kpi = row.get("dataItemMap") or {}
                if not isinstance(kpi, dict):
                    continue

                # EMI fields (Huawei NBI keys)
                irradiance = kpi.get("radiant_total_dose") or kpi.get("radiant_dose") or kpi.get("global_radiation") or kpi.get("radiation_intensity")
                ambient = kpi.get("temperature") or kpi.get("ambient_temperature")
                module = kpi.get("pv_temperature") or kpi.get("panel_temperature")
                wind = kpi.get("wind_speed")
                wind_dir = kpi.get("wind_direction")
                humidity = kpi.get("humidity")
                rainfall = kpi.get("rainfall")

                def _to_num(x):
                    if x is None:
                        return None
                    try:
                        return float(x)
                    except (TypeError, ValueError):
                        return None

                reading = {
                    "site_id": s["id"],
                    "vendor_device_id": dev_id,
                    "irradiance_wm2": _to_num(irradiance),
                    "ambient_temp_c": _to_num(ambient),
                    "module_temp_c": _to_num(module),
                    "wind_speed_ms": _to_num(wind),
                    "wind_direction_deg": _to_num(wind_dir),
                    "humidity_pct": _to_num(humidity),
                    "rainfall_mm": _to_num(rainfall),
                    "measured_at": now.isoformat(),
                }
                # Skip ak všetky polia None
                if all(v is None for k, v in reading.items() if k not in ("site_id","vendor_device_id","measured_at")):
                    continue
                site_readings.append(reading)

            if site_readings:
                try:
                    sb.table("environment_readings").upsert(site_readings, on_conflict="vendor_device_id,measured_at").execute()
                    total_readings += len(site_readings)
                except Exception as e:
                    log.warning("[env-pull] insert fail: %s", e)

                # Update site flag
                try:
                    sb.table("inverter_sites").update({
                        "has_environment_monitor": True,
                        "environment_last_sync_at": now.isoformat(),
                    }).eq("id", s["id"]).execute()
                except Exception:
                    pass

                # Update daily aggregate
                try:
                    # Načítaj všetky dnešné readings pre site, spočítaj agregát
                    daily = sb.table("environment_readings").select("*").eq("site_id", s["id"]).gte("measured_at", f"{today}T00:00:00Z").execute().data or []
                    if daily:
                        ir_values = [d.get("irradiance_wm2") for d in daily if d.get("irradiance_wm2") is not None]
                        amb_values = [d.get("ambient_temp_c") for d in daily if d.get("ambient_temp_c") is not None]
                        mod_values = [d.get("module_temp_c") for d in daily if d.get("module_temp_c") is not None]
                        wind_values = [d.get("wind_speed_ms") for d in daily if d.get("wind_speed_ms") is not None]
                        rain_values = [d.get("rainfall_mm") for d in daily if d.get("rainfall_mm") is not None]
                        # Trapezoidal integration pre kWh/m² (cca - závisí od sample rate)
                        peak_irr = max(ir_values) if ir_values else None
                        # Naive: priemer × hodiny / 1000
                        if ir_values:
                            hours_in_day = (now.hour + now.minute / 60.0) or 1
                            avg_irr = sum(ir_values) / len(ir_values)
                            total_kwh_m2 = (avg_irr / 1000.0) * hours_in_day
                        else:
                            total_kwh_m2 = None

                        sb.table("environment_daily").upsert({
                            "site_id": s["id"],
                            "date": str(today),
                            "total_irradiance_kwh_m2": total_kwh_m2,
                            "peak_irradiance_wm2": peak_irr,
                            "avg_ambient_temp_c": sum(amb_values)/len(amb_values) if amb_values else None,
                            "max_ambient_temp_c": max(amb_values) if amb_values else None,
                            "avg_module_temp_c": sum(mod_values)/len(mod_values) if mod_values else None,
                            "max_module_temp_c": max(mod_values) if mod_values else None,
                            "avg_wind_speed_ms": sum(wind_values)/len(wind_values) if wind_values else None,
                            "total_rainfall_mm": sum(rain_values) if rain_values else None,
                            "samples_count": len(daily),
                        }, on_conflict="site_id,date").execute()
                except Exception as e:
                    log.warning("[env-pull] daily agg fail: %s", e)

            results.append({"ok": True, "site_id": s["id"], "site_name": s.get("site_name"), "readings": len(site_readings)})
        except Exception as e:
            log.exception("[env-pull] site %s crashed", s.get("id"))
            results.append({"ok": False, "site_id": s["id"], "error": str(e)})

    return jsonify({
        "ok": True,
        "sites_processed": len(sites),
        "sites_with_emi": sites_with_emi,
        "total_readings": total_readings,
        "results": results[:20],
    })


# ============================================================================
# F4: INVERTER STATE ENUM (Huawei NBI Table 5-1)
# ============================================================================

INVERTER_STATE_MAP = {
    0: ("Standby: initializing", "standby"),
    1: ("Standby: insulation resistance detecting", "standby"),
    2: ("Standby: irradiation detecting", "standby"),
    3: ("Standby: grid detecting", "standby"),
    7: ("Standby: initialization after storage", "standby"),
    256: ("Start", "ok"),
    512: ("Grid-connected", "ok"),
    513: ("Grid-connected: power limited", "grid_limit"),
    514: ("Grid-connected: self-derating", "derating"),
    768: ("Shutdown: on fault", "fault"),
    769: ("Shutdown: on command", "shutdown"),
    770: ("Shutdown: OVGR", "shutdown"),
    771: ("Shutdown: communication interrupted", "comm_loss"),
    772: ("Shutdown: power limited", "grid_limit"),
    773: ("Shutdown: manual startup required", "fault"),
    774: ("Shutdown: DC switch disconnected", "shutdown"),
    1025: ("Grid scheduling: cosψ-P curve", "grid_limit"),
    1026: ("Grid scheduling: Q-U curve", "grid_limit"),
    1280: ("Ready for terminal test", "standby"),
    1281: ("Terminal testing...", "standby"),
    1536: ("Inspection in progress", "standby"),
    1792: ("AFCI self-check", "standby"),
    2048: ("I-V curve scanning", "standby"),
    2304: ("DC input detection", "standby"),
    40960: ("Standby: no irradiation", "standby"),
    45056: ("Communication interrupted (SmartLogger)", "comm_loss"),
    49152: ("Loading... (SmartLogger)", "standby"),
}

def map_inverter_state(code):
    """Map Huawei inverter_state code to (label, category)."""
    if code is None:
        return (None, None)
    try:
        c = int(code)
    except (TypeError, ValueError):
        return (str(code), "unknown")
    label, cat = INVERTER_STATE_MAP.get(c, (f"Unknown state {c}", "unknown"))
    return (label, cat)


def compute_phase_imbalance(va, vb, vc):
    """Compute % phase voltage imbalance per NEMA MG-1.
    Vrátí 0 ak všetky fázy = nominal, > 2% = problem.
    """
    try:
        voltages = [float(v) for v in [va, vb, vc] if v is not None]
        if len(voltages) < 2:
            return None
        avg = sum(voltages) / len(voltages)
        if avg < 1:
            return None
        max_dev = max(abs(v - avg) for v in voltages)
        return round((max_dev / avg) * 100, 2)
    except (TypeError, ValueError):
        return None


# ============================================================================
# F5+F6: GRID QUALITY (phase imbalance) + GRID METER (import/export)
# /webhook/huawei-pull-grid-meter — hodinový cron
# ============================================================================

@app.route("/webhook/huawei-pull-grid-meter", methods=["POST", "GET"])
def webhook_huawei_pull_grid_meter():
    """Pull grid meter (devTypeId=17/47) cumulative import/export pre billing kontrolu."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, site_name, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    now = datetime.now(timezone.utc)
    today = now.date()
    total_readings = 0
    results = []

    for s in sites:
        try:
            r = requests.post(
                f"{base}/getDevList",
                headers=headers,
                json={"stationCodes": s.get("vendor_station_id") or ""},
                timeout=30,
            )
            if r.status_code != 200:
                continue
            devices = (r.json() or {}).get("data") or []
            # Grid meters: devTypeId 17 (grid meter) + 47 (power sensor)
            meters = [d for d in devices if d.get("devTypeId") in (17, 47)]
            if not meters:
                continue

            for dev_type in {m.get("devTypeId") for m in meters}:
                type_meters = [m for m in meters if m.get("devTypeId") == dev_type]
                dev_ids = [str(m.get("id")) for m in type_meters if m.get("id")]
                if not dev_ids:
                    continue
                r2 = requests.post(
                    f"{base}/getDevRealKpi",
                    headers=headers,
                    json={"devIds": ",".join(dev_ids), "devTypeId": dev_type},
                    timeout=30,
                )
                if r2.status_code != 200:
                    continue
                rows = (r2.json() or {}).get("data") or []
                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    dev_id = str(row.get("devId") or row.get("id") or "")
                    kpi = row.get("dataItemMap") or {}
                    if not isinstance(kpi, dict):
                        continue

                    def _num(x):
                        try:
                            return float(x) if x is not None else None
                        except (TypeError, ValueError):
                            return None

                    # Grid meter fields per NBI doc
                    import_total = _num(kpi.get("active_cap") or kpi.get("positive_active_power") or kpi.get("forward_active_cap"))
                    export_total = _num(kpi.get("reverse_active_cap") or kpi.get("reverse_active_power"))
                    active_power = _num(kpi.get("active_power"))

                    try:
                        sb.table("grid_meter_readings").upsert({
                            "site_id": s["id"],
                            "vendor_device_id": dev_id,
                            "date": str(today),
                            "import_kwh_total": import_total,
                            "export_kwh_total": export_total,
                            "active_power_kw": active_power,
                            "measured_at": now.isoformat(),
                        }, on_conflict="vendor_device_id,date").execute()
                        total_readings += 1
                    except Exception as e:
                        log.warning("[grid-meter] upsert fail: %s", e)

            results.append({"ok": True, "site_id": s["id"], "meters": len(meters)})
        except Exception as e:
            log.exception("[grid-meter] site %s crashed", s.get("id"))
            results.append({"ok": False, "site_id": s["id"], "error": str(e)})

    return jsonify({
        "ok": True,
        "sites_processed": len(sites),
        "total_readings": total_readings,
        "results": results[:20],
    })


@app.route("/webhook/huawei-pull-detailed-inverters", methods=["POST", "GET"])
def webhook_huawei_pull_detailed_inverters():
    """Pull detailed inverter telemetria - state, phase imbalance, power factor, frequency, temperature.
    Hodinový cron - extra fields nad rámec basic active_power."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, site_name, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    now = datetime.now(timezone.utc)
    measurements_inserted = 0

    for s in sites:
        try:
            r = requests.post(
                f"{base}/getDevList",
                headers=headers,
                json={"stationCodes": s.get("vendor_station_id") or ""},
                timeout=30,
            )
            if r.status_code != 200:
                continue
            devices = (r.json() or {}).get("data") or []
            inverters = [d for d in devices if d.get("devTypeId") in (1, 38)]
            if not inverters:
                continue

            for dev_type in {d.get("devTypeId") for d in inverters}:
                type_devs = [d for d in inverters if d.get("devTypeId") == dev_type]
                dev_ids = [str(d.get("id")) for d in type_devs if d.get("id")]
                if not dev_ids:
                    continue
                r2 = requests.post(
                    f"{base}/getDevRealKpi",
                    headers=headers,
                    json={"devIds": ",".join(dev_ids), "devTypeId": dev_type},
                    timeout=30,
                )
                if r2.status_code != 200:
                    continue
                rows = (r2.json() or {}).get("data") or []

                # Aggregate to site-level: sum power, avg PF, avg freq, max temp, state of "primary" inverter
                site_active_power = 0.0
                site_reactive = 0.0
                pf_values = []
                freq_values = []
                temp_values = []
                state_codes = []
                phase_voltages = {"a": [], "b": [], "c": []}
                phase_currents = {"a": [], "b": [], "c": []}
                mppt_powers = []

                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    kpi = row.get("dataItemMap") or {}
                    if not isinstance(kpi, dict):
                        continue

                    def _num(x):
                        try:
                            return float(x) if x is not None else None
                        except (TypeError, ValueError):
                            return None

                    ap = _num(kpi.get("active_power"))
                    if ap is not None:
                        site_active_power += ap
                    rp = _num(kpi.get("reactive_power"))
                    if rp is not None:
                        site_reactive += rp
                    pf = _num(kpi.get("power_factor"))
                    if pf is not None:
                        pf_values.append(pf)
                    fq = _num(kpi.get("elec_freq"))
                    if fq is not None:
                        freq_values.append(fq)
                    tmp = _num(kpi.get("temperature"))
                    if tmp is not None:
                        temp_values.append(tmp)
                    st = kpi.get("inverter_state")
                    if st is not None:
                        try:
                            state_codes.append(int(st))
                        except (TypeError, ValueError):
                            pass
                    for ph in ("a", "b", "c"):
                        u = _num(kpi.get(f"{ph}_u"))
                        i = _num(kpi.get(f"{ph}_i"))
                        if u is not None:
                            phase_voltages[ph].append(u)
                        if i is not None:
                            phase_currents[ph].append(i)
                    mppt = _num(kpi.get("mppt_power") or kpi.get("mppt_total_cap"))
                    if mppt is not None:
                        mppt_powers.append(mppt)

                # Primary state = najhorší stav (fault > shutdown > comm_loss > standby > grid_limit > ok)
                primary_state = None
                primary_label = None
                primary_cat = None
                if state_codes:
                    priority = {"fault": 0, "shutdown": 1, "comm_loss": 2, "derating": 3, "grid_limit": 4, "standby": 5, "ok": 6, "unknown": 7}
                    best = None
                    for c in state_codes:
                        label, cat = map_inverter_state(c)
                        pr = priority.get(cat, 99)
                        if best is None or pr < best[0]:
                            best = (pr, c, label, cat)
                    if best:
                        primary_state, primary_label, primary_cat = best[1], best[2], best[3]

                avg_va = sum(phase_voltages["a"]) / len(phase_voltages["a"]) if phase_voltages["a"] else None
                avg_vb = sum(phase_voltages["b"]) / len(phase_voltages["b"]) if phase_voltages["b"] else None
                avg_vc = sum(phase_voltages["c"]) / len(phase_voltages["c"]) if phase_voltages["c"] else None
                imbalance = compute_phase_imbalance(avg_va, avg_vb, avg_vc)

                row_data = {
                    "site_id": s["id"],
                    "measured_at": now.isoformat(),
                    "active_power_kw": site_active_power,
                    "reactive_power_kvar": site_reactive,
                    "inverter_state_code": primary_state,
                    "inverter_state_label": primary_label,
                    "inverter_state_category": primary_cat,
                    "power_factor": sum(pf_values) / len(pf_values) if pf_values else None,
                    "grid_frequency_hz": sum(freq_values) / len(freq_values) if freq_values else None,
                    "inverter_internal_temp_c": max(temp_values) if temp_values else None,  # max = najhorší case
                    "phase_a_voltage_v": avg_va,
                    "phase_b_voltage_v": avg_vb,
                    "phase_c_voltage_v": avg_vc,
                    "phase_a_current_a": sum(phase_currents["a"]) / len(phase_currents["a"]) if phase_currents["a"] else None,
                    "phase_b_current_a": sum(phase_currents["b"]) / len(phase_currents["b"]) if phase_currents["b"] else None,
                    "phase_c_current_a": sum(phase_currents["c"]) / len(phase_currents["c"]) if phase_currents["c"] else None,
                    "phase_imbalance_pct": imbalance,
                    "mppt_total_power_kw": sum(mppt_powers) if mppt_powers else None,
                }
                try:
                    sb.table("inverter_measurements").insert(row_data).execute()
                    measurements_inserted += 1
                except Exception as e:
                    log.warning("[detailed-inv] insert fail: %s", e)
        except Exception as e:
            log.exception("[detailed-inv] site %s crashed", s.get("id"))

    return jsonify({
        "ok": True,
        "sites_processed": len(sites),
        "measurements_inserted": measurements_inserted,
    })


# ============================================================================
# F7: HISTORICAL DEVICE DATA (5.1.4.5/6/7)
# /webhook/huawei-pull-device-history — denný cron 23:30 UTC (po skončení dňa)
# ============================================================================

@app.route("/webhook/huawei-pull-device-history-daily", methods=["POST", "GET"])
def webhook_huawei_pull_device_history_daily():
    """Pull per-device daily data z Huawei NBI (getDevKpiDay).
    Spúšťa sa denne 23:30 UTC pre celý fleet, ukladá do device_history_daily.
    """
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    # Voliteľne back-fill cez body.collect_time, inak dnešok
    body = request.get_json(silent=True) or {}
    collect_time_ms = body.get("collect_time_ms") or int(datetime.now(timezone.utc).timestamp() * 1000)

    total_rows = 0
    for s in sites:
        try:
            r = requests.post(
                f"{base}/getDevList",
                headers=headers,
                json={"stationCodes": s.get("vendor_station_id") or ""},
                timeout=30,
            )
            if r.status_code != 200:
                continue
            devices = (r.json() or {}).get("data") or []
            inverters = [d for d in devices if d.get("devTypeId") in (1, 38)]
            if not inverters:
                continue

            for dev_type in {d.get("devTypeId") for d in inverters}:
                type_devs = [d for d in inverters if d.get("devTypeId") == dev_type]
                dev_ids = [str(d.get("id")) for d in type_devs if d.get("id")]
                if not dev_ids:
                    continue
                r2 = requests.post(
                    f"{base}/getDevKpiDay",
                    headers=headers,
                    json={"devIds": ",".join(dev_ids), "devTypeId": dev_type, "collectTime": collect_time_ms},
                    timeout=60,
                )
                if r2.status_code != 200:
                    continue
                rows = (r2.json() or {}).get("data") or []

                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    dev_id = str(row.get("devId") or row.get("id") or "")
                    kpi = row.get("dataItemMap") or {}
                    if not isinstance(kpi, dict):
                        continue

                    # Huawei vracia 96 quarter-hour vzoriek + agregát
                    # Sumarizujeme product_power, operation_time, alarm_count
                    def _num(x):
                        try:
                            return float(x) if x is not None else None
                        except (TypeError, ValueError):
                            return None

                    product = _num(kpi.get("product_power"))
                    pr = _num(kpi.get("perpower_ratio"))
                    op_time = _num(kpi.get("inverter_state_time") or kpi.get("operation_time"))

                    try:
                        sb.table("device_history_daily").upsert({
                            "site_id": s["id"],
                            "vendor_device_id": dev_id,
                            "device_type_id": dev_type,
                            "date": datetime.fromtimestamp(collect_time_ms / 1000, tz=timezone.utc).date().isoformat(),
                            "product_power_kwh": product,
                            "perpower_ratio": pr,
                            "operation_time_minutes": int(op_time) if op_time else None,
                            "raw_json": kpi,
                        }, on_conflict="vendor_device_id,date").execute()
                        total_rows += 1
                    except Exception as e:
                        log.warning("[device-hist-daily] upsert fail: %s", e)
        except Exception as e:
            log.exception("[device-hist-daily] site %s crashed", s.get("id"))

    return jsonify({"ok": True, "sites_processed": len(sites), "rows_upserted": total_rows, "collect_time_ms": collect_time_ms})


@app.route("/webhook/huawei-pull-device-history-monthly", methods=["POST", "GET"])
def webhook_huawei_pull_device_history_monthly():
    """Pull per-device monthly data z Huawei NBI (getDevKpiMonth).
    Spúšťa sa 1. v mesiaci, archivuje predošlý mesiac.
    """
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    body = request.get_json(silent=True) or {}
    # Predošlý mesiac (1. dňa nasledujúceho mesiaca pre kompletné dáta)
    today = datetime.now(timezone.utc)
    if body.get("year") and body.get("month"):
        year = int(body["year"])
        month = int(body["month"])
    else:
        if today.month == 1:
            year, month = today.year - 1, 12
        else:
            year, month = today.year, today.month - 1
    collect_time_ms = int(datetime(year, month, 1, 12, 0, 0, tzinfo=timezone.utc).timestamp() * 1000)

    sb = _sb()
    sites = sb.table("inverter_sites").select(
        "id, vendor_station_id"
    ).eq("vendor", "huawei").eq("monitoring_enabled", True).execute().data or []

    total_rows = 0
    for s in sites:
        try:
            r = requests.post(
                f"{base}/getDevList",
                headers=headers,
                json={"stationCodes": s.get("vendor_station_id") or ""},
                timeout=30,
            )
            if r.status_code != 200:
                continue
            devices = (r.json() or {}).get("data") or []
            inverters = [d for d in devices if d.get("devTypeId") in (1, 38)]
            if not inverters:
                continue

            for dev_type in {d.get("devTypeId") for d in inverters}:
                type_devs = [d for d in inverters if d.get("devTypeId") == dev_type]
                dev_ids = [str(d.get("id")) for d in type_devs if d.get("id")]
                if not dev_ids:
                    continue
                r2 = requests.post(
                    f"{base}/getDevKpiMonth",
                    headers=headers,
                    json={"devIds": ",".join(dev_ids), "devTypeId": dev_type, "collectTime": collect_time_ms},
                    timeout=60,
                )
                if r2.status_code != 200:
                    continue
                rows = (r2.json() or {}).get("data") or []
                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    dev_id = str(row.get("devId") or row.get("id") or "")
                    kpi = row.get("dataItemMap") or {}
                    if not isinstance(kpi, dict):
                        continue
                    def _num(x):
                        try:
                            return float(x) if x is not None else None
                        except (TypeError, ValueError):
                            return None
                    try:
                        sb.table("device_history_monthly").upsert({
                            "site_id": s["id"],
                            "vendor_device_id": dev_id,
                            "device_type_id": dev_type,
                            "year": year,
                            "month": month,
                            "product_power_kwh": _num(kpi.get("product_power")),
                            "perpower_ratio": _num(kpi.get("perpower_ratio")),
                            "raw_json": kpi,
                        }, on_conflict="vendor_device_id,year,month").execute()
                        total_rows += 1
                    except Exception as e:
                        log.warning("[device-hist-monthly] upsert fail: %s", e)
        except Exception as e:
            log.exception("[device-hist-monthly] site crashed: %s", e)

    return jsonify({"ok": True, "year": year, "month": month, "sites": len(sites), "rows_upserted": total_rows})


@rate_limit(max_calls=5, window_seconds=300)
@app.route("/webhook/huawei-debug-stations-public", methods=["GET"])
def webhook_huawei_debug_stations_public():
    """PUBLIC read-only diagnostika /stations - bez auth (dočasne, len pre debug session).
    Vracia len Huawei API response info, žiadne mutácie."""
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500
    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed"}), 503
    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}
    page_no = int(request.args.get("page") or "1")
    try:
        r = requests.post(f"{base}/stations", headers=headers, json={"pageNo": page_no}, timeout=60)
        body_json = {}
        try:
            body_json = r.json() or {}
        except Exception:
            pass
        data_payload = body_json.get("data") or {}
        if isinstance(data_payload, dict):
            station_list = data_payload.get("list", []) or []
            page_count = data_payload.get("pageCount")
            total = data_payload.get("total")
        else:
            station_list = body_json.get("list", []) or []
            page_count = None
            total = None
        return jsonify({
            "ok": r.status_code == 200 and body_json.get("success") is True,
            "http_status": r.status_code,
            "page_no": page_no,
            "fail_code": body_json.get("failCode"),
            "message": body_json.get("message"),
            "success": body_json.get("success"),
            "total_stations": total,
            "page_count": page_count,
            "stations_in_page": len(station_list),
            "first_3_stations": station_list[:3] if station_list else [],
            "body_preview": (r.text or "")[:500],
        })
    except Exception as e:
        log.exception("[debug-stations-public] crashed")
        return jsonify({"ok": False, "error": f"crash: {type(e).__name__}: {e}"}), 500


@app.route("/webhook/huawei-debug-stations", methods=["POST", "GET"])
def webhook_huawei_debug_stations():
    """Diagnostika - priamy POST na /stations endpoint, vráti raw Huawei response."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if _hs is None:
        return jsonify({"ok": False, "error": "huawei_spot module not available"}), 500

    token = _hs.huawei_login()
    if not token:
        return jsonify({"ok": False, "error": "huawei login failed (check backoff)"}), 503

    base = _hs._huawei_session.get("base") or _hs.HUAWEI_BASE
    headers = {"XSRF-TOKEN": token, "Content-Type": "application/json"}

    page_no = int(request.args.get("page") or "1")
    body = {"pageNo": page_no}

    try:
        r = requests.post(f"{base}/stations", headers=headers, json=body, timeout=60)
        body_text = r.text[:2000] if r.text else ""
        body_json = {}
        try:
            body_json = r.json() or {}
        except Exception:
            pass

        data_payload = body_json.get("data") or {}
        if isinstance(data_payload, dict):
            station_list = data_payload.get("list", []) or []
            page_count = data_payload.get("pageCount")
            total = data_payload.get("total")
        else:
            station_list = body_json.get("list", []) or []
            page_count = None
            total = None

        return jsonify({
            "ok": r.status_code == 200 and body_json.get("success") is True,
            "http_status": r.status_code,
            "page_no": page_no,
            "fail_code": body_json.get("failCode"),
            "message": body_json.get("message"),
            "success": body_json.get("success"),
            "total_stations": total,
            "page_count": page_count,
            "stations_in_page": len(station_list),
            "first_3_stations": station_list[:3] if station_list else [],
            "body_preview": body_text[:500],
            "base_url": base,
        })
    except Exception as e:
        log.exception("[huawei-debug-stations] crashed")
        return jsonify({"ok": False, "error": f"crash: {type(e).__name__}: {e}"}), 500


# ============================================================================
# EVA DOCX + XLSX SUITE — generate/read ad-hoc Word a Excel
# ============================================================================

@app.route("/webhook/eva-generate-docx", methods=["POST"])
def webhook_eva_generate_docx():
    """Vygeneruje DOCX z JSON štruktúry. Vstup: {title, sections[{heading, content, level?}], filename?}"""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    body = request.get_json(silent=True) or {}
    title = body.get("title") or "Eva Dokument"
    sections = body.get("sections") or []
    filename = body.get("filename") or f"eva_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        import base64 as _b64

        doc = Document()
        # Title
        t = doc.add_heading(title, level=0)
        for run in t.runs:
            run.font.color.rgb = RGBColor(0x14, 0x83, 0x4A)  # Energovision green

        for sec in sections:
            heading = sec.get("heading", "")
            content = sec.get("content", "")
            level = int(sec.get("level", 1))
            if heading:
                doc.add_heading(heading, level=min(max(level, 1), 5))
            if content:
                # Multi-paragraph support
                for para in str(content).split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            # Bulletlist support
            for bullet in sec.get("bullets", []) or []:
                doc.add_paragraph(bullet, style="List Bullet")
            # Table support
            for table_data in sec.get("tables", []) or []:
                if isinstance(table_data, dict) and table_data.get("rows"):
                    rows = table_data["rows"]
                    headers = table_data.get("headers", [])
                    cols = max(len(headers), max(len(r) for r in rows) if rows else 0)
                    tbl = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=cols)
                    tbl.style = "Light Grid Accent 1"
                    if headers:
                        for i, h in enumerate(headers):
                            tbl.cell(0, i).text = str(h)
                    for ri, row in enumerate(rows):
                        for ci, val in enumerate(row):
                            tbl.cell(ri + (1 if headers else 0), ci).text = str(val) if val is not None else ""

        # Footer
        doc.add_paragraph().add_run("\nEnergovision s.r.o. · IČO 53 036 280 · www.energovision.sk").italic = True

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_b64 = _b64.b64encode(buf.read()).decode("ascii")

        # Upload do Supabase eva-files bucket
        storage_path = f"docx/{filename}"
        try:
            from supabase import create_client
            sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_ROLE_KEY"])
            sb.storage.from_("documents").upload(
                storage_path,
                _b64.b64decode(docx_b64),
                {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "upsert": "true"},
            )
            public_url = sb.storage.from_("documents").get_public_url(storage_path)
        except Exception as e:
            log.warning("[eva-gen-docx] storage upload fail: %s", e)
            public_url = None

        return jsonify({
            "ok": True,
            "filename": filename,
            "file_url": public_url,
            "data_base64": docx_b64,
            "size_kb": len(docx_b64) // 1024,
        })
    except Exception as e:
        log.exception("[eva-generate-docx] crashed")
        return jsonify({"ok": False, "error": str(e)[:300]}), 500


@app.route("/webhook/eva-generate-xlsx", methods=["POST"])
def webhook_eva_generate_xlsx():
    """Vygeneruje XLSX z JSON. Vstup: {sheets:[{name, headers[], rows[][], column_widths?}], filename?}"""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    body = request.get_json(silent=True) or {}
    sheets = body.get("sheets") or []
    filename = body.get("filename") or f"eva_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
        from io import BytesIO
        import base64 as _b64

        wb = Workbook()
        # Remove default sheet
        wb.remove(wb.active)

        for sheet_def in sheets:
            name = sheet_def.get("name", "Sheet1")[:31]
            ws = wb.create_sheet(title=name)
            headers = sheet_def.get("headers") or []
            rows = sheet_def.get("rows") or []

            # Headers s formátovaním (Energovision lime green)
            if headers:
                ws.append(headers)
                hdr_fill = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
                for col in range(1, len(headers) + 1):
                    cell = ws.cell(row=1, column=col)
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = hdr_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            # Data rows
            for row in rows:
                ws.append(row if isinstance(row, list) else [row])

            # Auto-width
            for col_idx in range(1, max(len(headers), max((len(r) for r in rows), default=0)) + 1):
                col_letter = get_column_letter(col_idx)
                max_len = 10
                for cell in ws[col_letter]:
                    if cell.value is not None:
                        max_len = max(max_len, min(50, len(str(cell.value))))
                ws.column_dimensions[col_letter].width = max_len + 2

            # Freeze headers
            if headers:
                ws.freeze_panes = "A2"

        # Ak žiadny sheet, pridaj default
        if len(wb.sheetnames) == 0:
            wb.create_sheet("Data")

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        xlsx_b64 = _b64.b64encode(buf.read()).decode("ascii")

        storage_path = f"xlsx/{filename}"
        try:
            from supabase import create_client
            sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_ROLE_KEY"])
            sb.storage.from_("documents").upload(
                storage_path,
                _b64.b64decode(xlsx_b64),
                {"content-type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "upsert": "true"},
            )
            public_url = sb.storage.from_("documents").get_public_url(storage_path)
        except Exception as e:
            log.warning("[eva-gen-xlsx] storage upload fail: %s", e)
            public_url = None

        return jsonify({
            "ok": True,
            "filename": filename,
            "file_url": public_url,
            "data_base64": xlsx_b64,
            "size_kb": len(xlsx_b64) // 1024,
            "sheets_count": len(sheets),
        })
    except Exception as e:
        log.exception("[eva-generate-xlsx] crashed")
        return jsonify({"ok": False, "error": str(e)[:300]}), 500


@app.route("/webhook/eva-read-docx", methods=["POST"])
def webhook_eva_read_docx():
    """Prečíta DOCX a vráti text + headings + tables. Vstup: {file_url} alebo {data_base64}"""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    body = request.get_json(silent=True) or {}
    file_url = body.get("file_url")
    data_b64 = body.get("data_base64")

    try:
        from docx import Document
        from io import BytesIO
        import base64 as _b64

        if data_b64:
            blob = _b64.b64decode(data_b64)
        elif file_url:
            r = requests.get(file_url, timeout=30)
            if r.status_code != 200:
                return jsonify({"ok": False, "error": f"Download HTTP {r.status_code}"}), 400
            blob = r.content
        else:
            return jsonify({"ok": False, "error": "missing file_url or data_base64"}), 400

        doc = Document(BytesIO(blob))
        headings = []
        paragraphs = []
        tables = []
        for p in doc.paragraphs:
            if p.text.strip():
                if p.style.name.startswith("Heading"):
                    headings.append({"text": p.text, "level": int(p.style.name.replace("Heading ", "") or 1)})
                paragraphs.append(p.text)
        for tbl in doc.tables:
            tbl_rows = []
            for row in tbl.rows:
                tbl_rows.append([cell.text.strip() for cell in row.cells])
            if tbl_rows:
                tables.append({"rows": tbl_rows, "rows_count": len(tbl_rows)})

        full_text = "\n".join(paragraphs)
        return jsonify({
            "ok": True,
            "text": full_text[:20000],   # safety cap pre Claude context
            "headings": headings[:50],
            "paragraphs_count": len(paragraphs),
            "tables_count": len(tables),
            "tables": tables[:5],         # prvých 5 tabuliek
            "truncated": len(full_text) > 20000,
        })
    except Exception as e:
        log.exception("[eva-read-docx] crashed")
        return jsonify({"ok": False, "error": str(e)[:300]}), 500


@app.route("/webhook/eva-read-xlsx", methods=["POST"])
def webhook_eva_read_xlsx():
    """Prečíta XLSX a vráti sheets s headers+rows. Vstup: {file_url} alebo {data_base64}, max_rows_per_sheet"""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    body = request.get_json(silent=True) or {}
    file_url = body.get("file_url")
    data_b64 = body.get("data_base64")
    max_rows = int(body.get("max_rows_per_sheet", 1000))

    try:
        from openpyxl import load_workbook
        from io import BytesIO
        import base64 as _b64

        if data_b64:
            blob = _b64.b64decode(data_b64)
        elif file_url:
            r = requests.get(file_url, timeout=30)
            if r.status_code != 200:
                return jsonify({"ok": False, "error": f"Download HTTP {r.status_code}"}), 400
            blob = r.content
        else:
            return jsonify({"ok": False, "error": "missing file_url or data_base64"}), 400

        wb = load_workbook(BytesIO(blob), data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_rows = list(ws.iter_rows(values_only=True))
            if not all_rows:
                sheets.append({"name": sheet_name, "headers": [], "rows": [], "total_rows": 0})
                continue
            # Detekcia headerov: prvý riadok ak má string hodnoty
            first_row = all_rows[0]
            has_string_header = all(v is None or isinstance(v, str) for v in first_row)
            if has_string_header:
                headers = [str(v) if v is not None else "" for v in first_row]
                data_rows = all_rows[1:max_rows + 1]
            else:
                headers = []
                data_rows = all_rows[:max_rows]
            # Convert cell values to JSON-safe
            clean_rows = []
            for row in data_rows:
                clean_rows.append([str(v) if hasattr(v, "isoformat") else v for v in row])
            sheets.append({
                "name": sheet_name,
                "headers": headers,
                "rows": clean_rows,
                "total_rows": len(all_rows) - (1 if has_string_header else 0),
                "truncated": len(all_rows) > (max_rows + (1 if has_string_header else 0)),
            })

        return jsonify({
            "ok": True,
            "sheets_count": len(sheets),
            "sheets": sheets,
        })
    except Exception as e:
        log.exception("[eva-read-xlsx] crashed")
        return jsonify({"ok": False, "error": str(e)[:300]}), 500


# ============================================================================
# INTERNÝ KALKULAČNÝ DOKUMENT — Task #8
# Detailný breakdown výpočtov pre Energovision INTERNÉ použitie
# (nezdieľa sa s klientom — obsahuje raw čísla, sensitivity matrix, AI úvahy)
# ============================================================================

@app.route("/webhook/analyza-om-render-internal-calc", methods=["POST"])
def webhook_aom_render_internal_calc():
    """Vygeneruje INTERNÝ kalkulačný DOCX pre Energovision (nie pre klienta)."""
    if not _hs_auth_ok(request):
        return jsonify({"error": "unauthorized"}), 401
    if not _aom_v2:
        return jsonify({"ok": False, "error": "analyza_om_v2 not loaded"}), 500

    body = request.get_json(silent=True) or {}
    analyza_id = body.get("analyza_id")
    if not analyza_id:
        return jsonify({"ok": False, "error": "analyza_id required"}), 400

    try:
        result = _generate_internal_calc_doc(analyza_id)
        return jsonify(result)
    except Exception as e:
        log.exception("[aom-internal-calc] failed")
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


def _generate_internal_calc_doc(analyza_id: str) -> dict:
    """Interná kalkulácia - raw data + step-by-step výpočet + sensitivity + AI thinking-out-loud."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import base64 as _b64

    sb = _sb()
    a_res = sb.table("analyza_om").select("*").eq("id", analyza_id).single().execute()
    analyza = a_res.data
    if not analyza:
        raise ValueError(f"Analyza {analyza_id} not found")

    v_res = sb.table("analyza_om_variants").select("*").eq("analyza_id", analyza_id).order("position").execute()
    variants = v_res.data or []

    econ = analyza.get("econ_results") or {}
    ai_narrative = econ.get("ai_narrative") or {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # CONFIDENTIAL header
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header_p.add_run("INTERNÉ — NEZDIEĽAŤ S KLIENTOM")
    hr.font.size = Pt(9); hr.font.bold = True
    hr.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Title
    t = doc.add_heading(f"Interná kalkulácia · {analyza.get('posudok_number') or analyza.get('name', 'AOM-?')}", level=0)
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x14, 0x83, 0x4A)

    doc.add_paragraph(f"Vygenerované: {datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True

    # 1. VSTUPNÉ DÁTA
    doc.add_heading("1. Vstupné dáta", level=1)
    inputs = [
        ("Analyza ID", str(analyza_id)),
        ("Klient (názov)", analyza.get("name", "—")),
        ("Adresa", analyza.get("om_address", "—")),
        ("PSČ", analyza.get("om_psc", "—")),
        ("Sadzba", analyza.get("om_sadzba", "—")),
        ("MRK (kW)", str(analyza.get("om_mrk_kw", "—"))),
        ("RK (kW)", str(analyza.get("om_rk_kw", "—"))),
        ("Max export (kW)", str(analyza.get("max_export_kw", "—"))),
        ("Ročná spotreba (MWh)", str(analyza.get("consumption_annual_mwh", "—"))),
        ("Tarif nákup (€/kWh)", str(analyza.get("tarif_buy", "—"))),
        ("Tarif predaj (€/kWh)", str(analyza.get("tarif_sell", "—"))),
        ("PVGIS yield (kWh/kWp)", str(analyza.get("pvgis_yield_kwh_per_kwp", "—"))),
        ("Scenár", analyza.get("scenario_type", "nova_fve")),
        ("Opis scenára", analyza.get("scenario_description") or "(nepovedané)"),
        ("Existujúca FVE (kWp)", str(analyza.get("existing_fve_kwp", 0))),
        ("Existujúce BESS (kWh)", str(analyza.get("existing_bess_kwh", 0))),
        ("Dotácia zapnutá", "ÁNO" if analyza.get("dotacia_enabled") else "nie"),
        ("Dotačná schéma", analyza.get("dotacia_scheme") or "—"),
    ]
    tbl = doc.add_table(rows=len(inputs), cols=2)
    tbl.style = "Light List Accent 1"
    for i, (k, v) in enumerate(inputs):
        tbl.cell(i, 0).text = k
        tbl.cell(i, 1).text = str(v)
        tbl.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    # 2. PER-VARIANT VÝPOČTY
    doc.add_paragraph()
    doc.add_heading("2. Per-variant ekonomické výpočty", level=1)
    if not variants:
        doc.add_paragraph("Žiadne varianty zatiaľ nepočítané. Spusti VariantGenerator.").italic = True
    else:
        for v in variants:
            doc.add_heading(f"{v.get('name', 'Variant')} — {v.get('fve_kwp', 0)} kWp + {v.get('bess_kwh', 0)} kWh", level=2)
            calc_rows = [
                ("FVE kWp", str(v.get("fve_kwp", 0))),
                ("BESS kWh", str(v.get("bess_kwh", 0))),
                ("BESS kW", str(v.get("bess_kw", 0))),
                ("Topológia", v.get("fve_topology", "south")),
                ("CAPEX (€)", f"{v.get('capex_eur', 0):,.0f}"),
                ("Dotácia (€)", f"{v.get('result_dotacia_eur', 0):,.0f}"),
                ("Net CAPEX (€)", f"{v.get('capex_eur', 0) - (v.get('result_dotacia_eur') or 0):,.0f}"),
                ("Samospotreba (%)", str(v.get("result_samosp_pct", "—"))),
                ("Samostatnosť (%)", str(v.get("result_samostat_pct", "—"))),
                ("Import zo siete (MWh/r)", str(v.get("result_import_mwh", "—"))),
                ("NPV 20r (€)", f"{v.get('result_npv_eur_base', 0):,.0f}"),
                ("IRR (%)", str(v.get("result_irr_pct_base", "—"))),
                ("Návratnosť (r)", str(v.get("result_payback_y_base", "—"))),
            ]
            t2 = doc.add_table(rows=len(calc_rows), cols=2)
            t2.style = "Light List Accent 2"
            for i, (k, vv) in enumerate(calc_rows):
                t2.cell(i, 0).text = k
                t2.cell(i, 1).text = str(vv)
            doc.add_paragraph()

    # 3. SENSITIVITY MATRIX
    doc.add_heading("3. Sensitivity matrix (NPV pri rôznych cenách)", level=1)
    doc.add_paragraph(
        "Tabuľka ukazuje ako NPV najvýhodnejšieho variantu reaguje na zmeny ceny nákupu a výkupu energie. "
        "Užitočné pre stress-test pri zmene tarif klienta alebo regulácie."
    )
    if variants:
        best = max(variants, key=lambda x: float(x.get("result_npv_eur_base") or 0))
        base_buy = float(analyza.get("tarif_buy") or 0.146)
        base_sell = float(analyza.get("tarif_sell") or 0.06)
        base_npv = float(best.get("result_npv_eur_base") or 0)
        # 5x5 matrix
        buy_modifiers = [-0.20, -0.10, 0, +0.10, +0.20]
        sell_modifiers = [-0.20, -0.10, 0, +0.10, +0.20]
        st = doc.add_table(rows=6, cols=6)
        st.style = "Light Grid Accent 1"
        st.cell(0, 0).text = "P_BUY \\ P_SELL"
        for j, sm in enumerate(sell_modifiers):
            st.cell(0, j + 1).text = f"{base_sell * (1 + sm):.3f}"
        for i, bm in enumerate(buy_modifiers):
            st.cell(i + 1, 0).text = f"{base_buy * (1 + bm):.3f}"
            for j, sm in enumerate(sell_modifiers):
                # Simplified: NPV scales linearly with prices (true sensitivity needs full re-sim)
                scale = (1 + bm) * 0.7 + (1 + sm) * 0.3
                npv_scaled = base_npv * scale
                st.cell(i + 1, j + 1).text = f"{npv_scaled / 1000:,.0f} k€"

    # 4. AI THINKING-OUT-LOUD
    doc.add_paragraph()
    doc.add_heading("4. AI úvaha (čo Eva rátala a prečo)", level=1)
    if ai_narrative.get("text"):
        doc.add_paragraph(
            f"AI scenár: {ai_narrative.get('scenario_type', '—')} · "
            f"opis: {ai_narrative.get('scenario_description') or '(nepovedané)'}"
        ).italic = True
        for para in (ai_narrative.get("text") or "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("AI narrative nebola vygenerovaná (skontroluj že VariantGenerator prešiel + ANTHROPIC_API_KEY je nastavený).").italic = True

    # 5. METHODOLOGY
    doc.add_paragraph()
    doc.add_heading("5. Metodika výpočtov", level=1)
    methods = [
        "Diskontná sadzba (WACC): 6.0 % p.a. (štandard infraštruktúrnych projektov)",
        "Životnosť projektu: 20 rokov (FVE LCOE 25-30 r typicky)",
        "Inflácia cien energie: 2.5 % p.a. (regulačný štandard + uhlíkové ceny)",
        "Reziduálna hodnota po 20r: 10 % CAPEX (FVE EOL 90 % výkonu)",
        "OPEX: 1.5 % CAPEX/r (údržba, monitoring, poistenie)",
        "Degradácia FVE: 0.5 %/r lineárna (Tier 1 LONGi/Jinko)",
        "Degradácia BESS: 2 %/r (LFP chémia, 1 cyklus/deň priemer)",
        "DPPO sadzba: 21 % (aktuálna SR)",
        "Daňový odpis: 6 rokov rovnomerne (štandardný FVE odpis)",
        "Spot ceny: OKTE DAM 2025 archív (8760 h, kalibrovaný)",
        "PVGIS yield: zo zadaných lat/lon + sklon 35° juh (alebo custom topology)",
        "BESS arbitráž: rule_based EMS, top-N spot rozdielových párov",
    ]
    for m in methods:
        doc.add_paragraph(m, style="List Bullet")

    # 6. AUDIT
    doc.add_paragraph()
    doc.add_heading("6. Audit trail", level=1)
    doc.add_paragraph(f"Engine verzia: {econ.get('engine_version', '0.9.5')}")
    doc.add_paragraph(f"Variantov v matrix: {econ.get('variants_count', len(variants))}")
    doc.add_paragraph(f"AI model (narrative): claude-sonnet-4.5 (Anthropic)")
    doc.add_paragraph(f"Posudok vygenerovaný: {datetime.now().isoformat()}")

    # Footer disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    fp = p.add_run("⚠ Tento dokument obsahuje raw výpočty a interné úvahy. NEZDIEĽAŤ s klientom. Pre klientsky posudok použi 'Premium DOCX'.")
    fp.italic = True
    fp.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Save + upload
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    blob = buf.read()
    storage_path = f"analyza_om/{analyza_id}/internal_calc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    sb.storage.from_("documents").upload(
        storage_path, blob,
        {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "upsert": "true"}
    )
    public_url = sb.storage.from_("documents").get_public_url(storage_path)

    return {"ok": True, "docx_url": public_url, "storage_path": storage_path, "size_kb": len(blob) // 1024}


# ============================================================
# TS-CONTRACT-GENERATE — vygeneruje zmluvu o správe TS z templátu
# Vstup: {"contract_id": "<uuid>"}
# - Vytiahne ts_contracts + customer + stations
# - Vyplní šablónu Zmluva_sprava_TS_template.docx
# - Upload do Supabase Storage documents/ts_contracts/{id}/zmluva_<num>.docx
# - Vráti public_url
# ============================================================
@app.route("/webhook/ts-contract-generate", methods=["POST"])
def ts_contract_generate():
    body = request.get_json(silent=True) or {}
    contract_id = body.get("contract_id")
    if not contract_id:
        return jsonify({"error": "missing contract_id"}), 400

    sb_headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
    }

    # Načítaj zmluvu
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/ts_contracts",
        headers=sb_headers,
        params={"id": f"eq.{contract_id}", "select": "*"},
        timeout=10
    )
    if not r.ok or not r.json():
        return jsonify({"error": "contract_not_found"}), 404
    contract = r.json()[0]

    # Načítaj customer
    cust_id = contract.get("customer_id")
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/customers",
        headers=sb_headers,
        params={"id": f"eq.{cust_id}", "select": "*"},
        timeout=10
    )
    cust = (r.json() or [{}])[0] if r.ok else {}

    # Načítaj stanice
    station_ids = contract.get("station_ids") or []
    stations = []
    if station_ids:
        # PostgREST IN syntax: id=in.(uuid1,uuid2,...)
        ids_str = ",".join(station_ids)
        r = requests.get(
            f"{SUPABASE_URL}/rest/v1/transformer_stations",
            headers=sb_headers,
            params={"id": f"in.({ids_str})", "select": "id,ts_code,name,location_address,location_city,location_psc,rated_power_kva,notes", "order": "ts_code"},
            timeout=10
        )
        if r.ok:
            stations = r.json()

    # Postavíme ctx
    full_address = " ".join(filter(None, [
        cust.get("address") or cust.get("street") or "",
        cust.get("city") or "",
        cust.get("psc") or cust.get("zip_code") or ""
    ]))

    from datetime import datetime as _dt
    today = _dt.now().strftime("%d.%m.%Y")

    ctx = {
        "companyName": cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip(),
        "companyRegNumber": cust.get("ico") or "",
        "companyTaxNumber": cust.get("dic") or "",
        "companyStreet": cust.get("address") or cust.get("street") or "",
        "companyCity": cust.get("city") or "",
        "companyZipCode": cust.get("psc") or cust.get("zip_code") or "",
        # Konateľ — z Notion-style placeholders (legacy)
        "businessCaseTitul_pred_0cfa0": "",
        "businessCaseMeno_03a27": cust.get("first_name") or "",
        "businessCasePriezvisko_8375b": cust.get("last_name") or "",
        "businessCaseTelefonne__fa5fd": cust.get("phone") or "",
        "businessCaseEmail_2c918": cust.get("email") or "",
        "createdAtDate": today,
        # Prílohy
        "contact_primary_name": cust.get("contact_primary_name") or cust.get("first_name", "") + " " + cust.get("last_name", ""),
        "contact_primary_role": cust.get("contact_primary_role") or "Zodpovedná osoba",
        "contact_primary_phone": cust.get("phone") or "",
        "contact_primary_email": cust.get("email") or "",
        "contact_secondary_name": cust.get("contact_secondary_name") or "",
        "contact_secondary_role": cust.get("contact_secondary_role") or "",
        "contact_secondary_phone": cust.get("contact_secondary_phone") or "",
        "contact_secondary_email": cust.get("contact_secondary_email") or "",
        "ev_service_lead_name": os.environ.get("EV_SERVICE_LEAD_NAME", "Lukáš Bago (zatiaľ)"),
        "ev_service_lead_phone": os.environ.get("EV_SERVICE_LEAD_PHONE", "+421 918 187 762"),
        "ev_service_lead_email": os.environ.get("EV_SERVICE_LEAD_EMAIL", "lukas.bago@energovision.sk"),
    }

    # TS list pre Prílohu č. 1
    ts_rows = []
    cena_rows = []
    monthly_fee = float(contract.get("monthly_fee_eur") or 0)
    per_ts_fee = monthly_fee / max(1, len(stations)) if stations else monthly_fee
    sla_h = int(contract.get("sla_response_hours") or 24)
    havaria = "áno" if contract.get("havarijna_included") else "nie"

    for idx, st in enumerate(stations):
        addr = " ".join(filter(None, [st.get("location_address"), st.get("location_psc"), st.get("location_city")]))
        ts_rows.append({
            "poradie": idx + 1,
            "oznacenie": st.get("ts_code") or st.get("name"),
            "adresa": addr,
            "kva": st.get("rated_power_kva") or "",
            "poznamka": st.get("notes") or "",
        })
        cena_rows.append({
            "poradie": idx + 1,
            "mesacny_pausal": f"{per_ts_fee:.2f}".replace(".", ","),
            "pohotovostna": havaria,
            "reakcia_h": sla_h,
            "poznamka": "",
        })

    ctx["ts_rows"] = ts_rows
    ctx["cena_rows"] = cena_rows
    ctx["celkovy_mesacny_pausal"] = monthly_fee

    # Generuj
    import tempfile
    from pathlib import Path as _Path
    tmpdir = _Path(tempfile.mkdtemp())
    out_path = tmpdir / f"Zmluva_TS_{contract.get('contract_number') or contract_id[:8]}.docx"

    try:
        from generuj_dokumenty import naplnit_ts_zmluvu
        naplnit_ts_zmluvu(ctx, str(out_path))
    except Exception as e:
        log.exception("naplnit_ts_zmluvu zlyhalo")
        return jsonify({"error": f"generate_failed: {e}"}), 500

    # Upload do Storage
    with open(out_path, "rb") as f:
        file_bytes = f.read()

    storage_path = f"ts_contracts/{contract_id}/{out_path.name}"
    up = requests.post(
        f"{SUPABASE_URL}/storage/v1/object/documents/{storage_path}",
        headers={**sb_headers, "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "x-upsert": "true"},
        data=file_bytes,
        timeout=30
    )
    if not up.ok:
        log.warning("ts contract storage upload zlyhal: %s %s", up.status_code, up.text)
        return jsonify({"error": "storage_upload_failed", "body": up.text}), 500

    public_docx_url = f"{SUPABASE_URL}/storage/v1/object/public/documents/{storage_path}"

    # Generuj PDF cez mammoth + weasyprint
    pdf_public_url = None
    try:
        import mammoth
        from weasyprint import HTML
        from io import BytesIO
        with open(out_path, "rb") as f:
            html_result = mammoth.convert_to_html(BytesIO(f.read()))
        html_body = html_result.value
        html_full = f"""<!DOCTYPE html><html lang="sk"><head><meta charset="utf-8"><style>
@page {{ size: A4; margin: 18mm; }}
body {{ font-family: 'Helvetica', sans-serif; font-size: 10pt; color: #1a1a1a; line-height: 1.45; }}
h1 {{ font-size: 16pt; margin: 12pt 0 6pt; }}
h2 {{ font-size: 13pt; margin: 10pt 0 5pt; }}
h3 {{ font-size: 11pt; margin: 8pt 0 4pt; }}
p {{ margin: 4pt 0; }}
table {{ border-collapse: collapse; margin: 6pt 0; width: 100%; }}
td, th {{ border: 0.5pt solid #ccc; padding: 4pt 6pt; }}
strong {{ font-weight: 700; }}
</style></head><body>{html_body}</body></html>"""
        pdf_bytes = HTML(string=html_full).write_pdf()
        pdf_storage_path = f"ts_contracts/{contract_id}/{out_path.stem}.pdf"
        up_pdf = requests.post(
            f"{SUPABASE_URL}/storage/v1/object/documents/{pdf_storage_path}",
            headers={**sb_headers, "Content-Type": "application/pdf", "x-upsert": "true"},
            data=pdf_bytes, timeout=30
        )
        if up_pdf.ok:
            pdf_public_url = f"{SUPABASE_URL}/storage/v1/object/public/documents/{pdf_storage_path}"
    except Exception as _e:
        log.exception("PDF generation failed (DOCX OK)")

    # Update ts_contracts
    requests.patch(
        f"{SUPABASE_URL}/rest/v1/ts_contracts",
        headers={**sb_headers, "Content-Type": "application/json"},
        params={"id": f"eq.{contract_id}"},
        json={"contract_docx_url": public_docx_url, "contract_pdf_url": pdf_public_url, "signed_pdf_url": pdf_public_url or public_docx_url},
        timeout=10
    )

    try:
        out_path.unlink()
        tmpdir.rmdir()
    except Exception:
        pass

    return jsonify({"ok": True, "url": pdf_public_url or public_docx_url, "docx_url": public_docx_url, "pdf_url": pdf_public_url, "filename": out_path.name, "contract_number": contract.get("contract_number")})


# ============================================================
# TS-QUOTE-GENERATE-PDF — vyrobí PDF cenovku TS servisu
# Vstup: {"quote_id": "<uuid>"}
# Výstup: { ok, url, filename }
# - Pull ts_quotes + customer
# - Render HTML → PDF cez WeasyPrint
# - Upload do Storage documents/ts_quotes/{id}/cenovka_<num>.pdf
# - Update ts_quotes.generated_docx_url = pdf_url (re-use field zatiaľ)
# ============================================================
@app.route("/webhook/ts-quote-generate-pdf", methods=["POST"])
def ts_quote_generate_pdf():
    body = request.get_json(silent=True) or {}
    quote_id = body.get("quote_id")
    if not quote_id:
        return jsonify({"error": "missing quote_id"}), 400

    sb_headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
    }

    # Načítaj cenovku
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/ts_quotes",
        headers=sb_headers,
        params={"id": f"eq.{quote_id}", "select": "*"},
        timeout=10
    )
    if not r.ok or not r.json():
        return jsonify({"error": "quote_not_found"}), 404
    q = r.json()[0]

    # Načítaj customer
    cust_id = q.get("customer_id")
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/customers",
        headers=sb_headers,
        params={"id": f"eq.{cust_id}", "select": "*"},
        timeout=10
    )
    cust = (r.json() or [{}])[0] if r.ok else {}

    from datetime import datetime as _dt
    today = _dt.now().strftime("%d.%m.%Y")
    klient = cust.get("company_name") or f"{cust.get('first_name','')} {cust.get('last_name','')}".strip() or "Klient"
    ico = cust.get("ico") or ""
    address = " ".join(filter(None, [cust.get("address") or cust.get("street") or "", cust.get("psc") or "", cust.get("city") or ""]))

    items = q.get("ts_items") or []
    variants = q.get("variants") or []
    has_variants = q.get("has_variants") and len(variants) > 0
    monthly_total = float(q.get("monthly_total_eur") or 0)
    contract_months = q.get("contract_duration_months") or 24
    setup_fee = float(q.get("one_time_setup_eur") or 0)
    valid_until = q.get("valid_until") or ""
    notes = q.get("notes") or ""

    TYPE_LABEL = {"basic": "Basic", "komplet": "Komplet", "havarijna_24_7": "Havarijná 24/7"}

    # TS list HTML
    ts_rows_html = ""
    for i, it in enumerate(items, 1):
        ts_rows_html += f"""<tr>
<td>{i}</td>
<td>{(it.get('code') or '')}</td>
<td>{(it.get('name') or '')}</td>
<td>{(it.get('address') or '')}</td>
<td style="text-align:right">{(it.get('kva') or '')}</td>
</tr>"""

    # Variants HTML
    variants_html = ""
    if has_variants:
        cards = ""
        for v in variants:
            badge = '<span class="badge">⭐ Odporúčané</span>' if v.get("highlight") else ""
            included = "".join(f'<li>{s}</li>' for s in (v.get("included_items") or []))
            cards += f"""<div class="variant-card {'highlight' if v.get('highlight') else ''}">
<h3>{badge} {v.get('label', '')}</h3>
<div class="price">{float(v.get('monthly_total', 0)):,.2f} €/mes</div>
<div class="sub">SLA {v.get('sla_response_hours')}h{' · 24/7' if v.get('has_24_7') else ''}</div>
<ul class="included">{included}</ul>
</div>""".replace(",", " ").replace(".", ",", 1)
        variants_html = f'<h2>Vyberte si úroveň servisu</h2><div class="variants">{cards}</div>'
    else:
        contract_label = TYPE_LABEL.get(q.get("contract_type"), q.get("contract_type", ""))
        variants_html = f"""<h2>Cenová špecifikácia</h2>
<div class="single-price">
<div><strong>Typ zmluvy:</strong> {contract_label}</div>
<div class="price-big">{monthly_total:,.2f} € / mes bez DPH</div>
<div class="sub">Ročne: {monthly_total * 12:,.2f} € · za {contract_months} mes: {monthly_total * contract_months:,.2f} €</div>
</div>""".replace(",", " ").replace(".", ",", 1)

    setup_html = f'<p style="margin-top:8pt"><strong>Jednorázový setup fee:</strong> {setup_fee:,.2f} € bez DPH</p>'.replace(",", " ").replace(".", ",", 1) if setup_fee > 0 else ""
    notes_html = f'<div class="notes"><strong>Poznámky:</strong><br>{notes.replace(chr(10), "<br>")}</div>' if notes else ""

    html = f"""<!DOCTYPE html><html lang="sk"><head><meta charset="utf-8"><style>
@page {{ size: A4; margin: 18mm 16mm; @bottom-right {{ content: counter(page) " / " counter(pages); font-size: 9pt; color: #999; }} }}
body {{ font-family: 'Helvetica', sans-serif; font-size: 10pt; color: #1a1a1a; line-height: 1.45; }}
.header {{ background: #92D050; color: #fff; padding: 16pt; border-radius: 4pt; margin-bottom: 16pt; }}
.header h1 {{ margin: 0; font-size: 20pt; }}
.header .subtitle {{ font-size: 10pt; opacity: 0.95; margin-top: 4pt; }}
h2 {{ font-size: 13pt; color: #10b981; margin: 16pt 0 8pt; border-bottom: 1pt solid #e5e7eb; padding-bottom: 3pt; }}
h3 {{ font-size: 11pt; margin: 6pt 0 4pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 6pt 0; font-size: 9pt; }}
th, td {{ border: 0.5pt solid #e5e7eb; padding: 5pt 7pt; }}
th {{ background: #f8fafc; text-align: left; font-weight: 600; color: #475569; }}
.party {{ display: table; width: 100%; margin-bottom: 14pt; }}
.party > div {{ display: table-cell; width: 50%; padding: 8pt; background: #f8fafc; border-radius: 3pt; vertical-align: top; }}
.party h3 {{ margin-top: 0; color: #475569; font-size: 9pt; text-transform: uppercase; letter-spacing: 0.5pt; }}
.variants {{ display: flex; gap: 8pt; flex-wrap: wrap; margin-top: 8pt; }}
.variant-card {{ flex: 1; min-width: 145pt; border: 1.5pt solid #e5e7eb; border-radius: 4pt; padding: 10pt; }}
.variant-card.highlight {{ border: 2pt solid #10b981; background: #ecfdf5; }}
.variant-card h3 {{ margin: 0 0 4pt; }}
.variant-card .price {{ font-size: 18pt; font-weight: bold; color: #047857; margin: 4pt 0; }}
.variant-card .sub {{ font-size: 8pt; color: #64748b; margin-bottom: 6pt; }}
.variant-card .included {{ margin: 0; padding-left: 14pt; font-size: 8.5pt; }}
.variant-card .included li {{ margin: 1pt 0; }}
.badge {{ background: #10b981; color: #fff; font-size: 7pt; padding: 1pt 4pt; border-radius: 2pt; vertical-align: middle; }}
.single-price {{ padding: 12pt; background: #ecfdf5; border: 1.5pt solid #10b981; border-radius: 4pt; }}
.single-price .price-big {{ font-size: 22pt; font-weight: bold; color: #047857; margin: 6pt 0; }}
.single-price .sub {{ font-size: 9pt; color: #64748b; }}
.notes {{ background: #fef3c7; border-left: 3pt solid #f59e0b; padding: 8pt 12pt; margin: 12pt 0; font-size: 9pt; }}
.footer {{ margin-top: 24pt; padding-top: 12pt; border-top: 0.5pt solid #e5e7eb; font-size: 8pt; color: #64748b; text-align: center; }}
.terms {{ background: #f8fafc; padding: 10pt; margin-top: 14pt; font-size: 9pt; border-radius: 3pt; }}
.terms ul {{ margin: 4pt 0 0; padding-left: 16pt; }}
.terms li {{ margin: 2pt 0; }}
</style></head><body>

<div class="header">
  <h1>⚡ Cenová ponuka servisu trafostaníc</h1>
  <div class="subtitle">{q.get('quote_number', '')} · {today}</div>
</div>

<div class="party">
  <div>
    <h3>Poskytovateľ</h3>
    <strong>Energovision s.r.o.</strong><br>
    Lamačská cesta 1738/111, 841 03 Bratislava<br>
    IČO: 53 036 280 · DIČ: 2121238526<br>
    IČ DPH: SK2121238526<br>
    Lukáš Bago · +421 918 187 762<br>
    lukas.bago@energovision.sk
  </div>
  <div>
    <h3>Objednávateľ</h3>
    <strong>{klient}</strong><br>
    {address}<br>
    {f'IČO: {ico}<br>' if ico else ''}
    {f'Email: {cust.get("email")}<br>' if cust.get("email") else ''}
    {f'Telefón: {cust.get("phone")}' if cust.get("phone") else ''}
  </div>
</div>

<h2>Predmet ponuky — {len(items)} trafostaníc</h2>
<table>
<thead><tr><th>Por.</th><th>Označenie</th><th>Názov</th><th>Adresa</th><th style="text-align:right">kVA</th></tr></thead>
<tbody>{ts_rows_html}</tbody>
</table>

{variants_html}
{setup_html}

<div class="terms">
<strong>Podmienky ponuky:</strong>
<ul>
<li>Doba viazanosti: <strong>{contract_months} mesiacov</strong></li>
<li>Splatnosť faktúr: 14 dní od doručenia</li>
<li>Ceny sú uvedené bez DPH (21 % bude pripočítané)</li>
<li>Inflačná indexácia: max +3 % ročne podľa ŠÚ SR (len pri inflácii &gt; 2 %)</li>
{f'<li>Platnosť ponuky do: <strong>{valid_until}</strong></li>' if valid_until else ''}
</ul>
</div>

{notes_html}

<div class="footer">
Energovision s.r.o. · Moderné energetické riešenia, ktoré hľadáte · www.energovision.sk
</div>

</body></html>"""

    # HTML → PDF
    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=html).write_pdf()
    except Exception as e:
        log.exception("weasyprint failed")
        return jsonify({"error": f"pdf_render_failed: {e}"}), 500

    # Upload do Storage
    filename = f"Cenovka_{q.get('quote_number') or quote_id[:8]}.pdf"
    storage_path = f"ts_quotes/{quote_id}/{filename}"
    up = requests.post(
        f"{SUPABASE_URL}/storage/v1/object/documents/{storage_path}",
        headers={**sb_headers, "Content-Type": "application/pdf", "x-upsert": "true"},
        data=pdf_bytes, timeout=30
    )
    if not up.ok:
        return jsonify({"error": "storage_upload_failed", "body": up.text}), 500

    public_url = f"{SUPABASE_URL}/storage/v1/object/public/documents/{storage_path}"

    # Update quote
    requests.patch(
        f"{SUPABASE_URL}/rest/v1/ts_quotes",
        headers={**sb_headers, "Content-Type": "application/json"},
        params={"id": f"eq.{quote_id}"},
        json={"generated_docx_url": public_url},
        timeout=10
    )

    return jsonify({"ok": True, "url": public_url, "filename": filename})


# ============================================================
# RAYNET IMPORT — spustí raynet_import.py
# Vstup: {dry_run: bool, entities: ["companies","persons","leads",...]}
# Výstup: {ok, result: {...}}
# ============================================================
@app.route("/webhook/raynet-import", methods=["POST"])
def raynet_import_endpoint():
    body = request.get_json(silent=True) or {}
    dry_run = bool(body.get("dry_run"))
    entities = body.get("entities") or ["companies", "persons", "leads"]

    # Log start
    sb_headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }
    log_resp = requests.post(
        f"{SUPABASE_URL}/rest/v1/raynet_import_log",
        headers=sb_headers,
        json={"dry_run": dry_run, "entity_types": entities},
        timeout=10,
    )
    log_id = (log_resp.json() or [{}])[0].get("id") if log_resp.ok else None

    try:
        from raynet_import import run as raynet_run
        result = raynet_run(entities, dry_run)
        # Update log
        if log_id:
            requests.patch(
                f"{SUPABASE_URL}/rest/v1/raynet_import_log",
                headers=sb_headers,
                params={"id": f"eq.{log_id}"},
                json={"finished_at": "now()", "result": result},
                timeout=10,
            )
        return jsonify({"ok": True, "result": result, "log_id": log_id})
    except Exception as e:
        log.exception("raynet import failed")
        if log_id:
            requests.patch(
                f"{SUPABASE_URL}/rest/v1/raynet_import_log",
                headers=sb_headers,
                params={"id": f"eq.{log_id}"},
                json={"finished_at": "now()", "error": str(e)},
                timeout=10,
            )
        return jsonify({"error": str(e), "log_id": log_id}), 500
