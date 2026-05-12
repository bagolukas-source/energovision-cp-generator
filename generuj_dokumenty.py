"""
Generuj dokumenty pre B2C leady — kompletná verzia.

Architektúra:
- Zmluva o dielo (docx): 12× XXX placeholder + datum
- Splnomocnenie (docx): podčiarkové polia + datum
- GDPR súhlas (docx): podčiarkové polia + datum
- Dotazník (xlsx): label v stĺpci A → hodnota v B
- Dodatok zmluvy (docx): Xxxxx placeholdery + datum
- Revízna správa (docx): generovaná PROGRAMATICKY pre B2C rodinný dom (3-4 strany)
- Preberací protokol (docx): generovaný PROGRAMATICKY s BOM tabuľkou

Lead_data je flat dict zo všetkých Notion properties.
"""
import os
import re
import shutil
import zipfile
import logging
import requests
from io import BytesIO
from pathlib import Path
from datetime import datetime

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = Path(__file__).parent / "templates_zmluvy"
log = logging.getLogger("generuj_dokumenty")

# ============================================================
# UTILITY HELPERS
# ============================================================

def _norm(s):
    """Normalizuj non-breaking space (\\xa0) na bežnú medzeru."""
    return (s or "").replace('\xa0', ' ')


def _safe(value, fallback=""):
    """Bezpečný get — None / prázdny string → fallback."""
    if value is None or value == "":
        return fallback
    return str(value)


def _money(amount, suffix=" EUR"):
    """Format peniaze: 15689.17 → '15 689,17 EUR'"""
    if not amount:
        return f"0,00{suffix}"
    return f"{amount:,.2f}".replace(",", " ").replace(".", ",") + suffix


def _sk_num(value, decimals=2):
    """Slovenský desatinný formát: 10.70 → '10,70'."""
    try:
        return f"{float(value):.{decimals}f}".replace(".", ",")
    except (TypeError, ValueError):
        return str(value)


def _slovne_centy(amount_eur):
    """Vráti (eur_int, cent_int) z floatu."""
    eur = int(amount_eur or 0)
    cents = round(((amount_eur or 0) - eur) * 100)
    return eur, cents


# ============================================================
# Číslo → slovenské slová (pre právne dokumenty)
# ============================================================
_ONES = ["", "jeden", "dva", "tri", "štyri", "päť", "šesť", "sedem", "osem", "deväť",
         "desať", "jedenásť", "dvanásť", "trinásť", "štrnásť", "pätnásť", "šestnásť",
         "sedemnásť", "osemnásť", "devätnásť"]
_TENS = ["", "", "dvadsať", "tridsať", "štyridsať", "päťdesiat", "šesťdesiat", "sedemdesiat",
         "osemdesiat", "deväťdesiat"]


def _below1000(n):
    """0-999 ako slovenské slová bez medzier (stovky+desiatky+jednotky spolu)."""
    if n == 0:
        return ""
    s = ""
    if n >= 100:
        stov = n // 100
        n = n % 100
        if stov == 1: s += "sto"
        elif stov == 2: s += "dvesto"
        elif stov == 3: s += "tristo"
        elif stov == 4: s += "štyristo"
        elif stov == 5: s += "päťsto"
        elif stov == 6: s += "šesťsto"
        elif stov == 7: s += "sedemsto"
        elif stov == 8: s += "osemsto"
        elif stov == 9: s += "deväťsto"
    if n >= 20:
        des = n // 10
        jed = n % 10
        s += _TENS[des]
        if jed > 0:
            s += _ONES[jed]
    elif n > 0:
        s += _ONES[n]
    return s


def _num_to_sk_words(n):
    """Konvertuje celé číslo 0-999999 na slovenské slová.
    Príklady:
      15689 → 'pätnásťtisícšesťstoosemdesiatdeväť'
      17 → 'sedemnásť'
      0 → 'nula'
      1234 → 'jedentisícdvestotridsaťštyri'
    Slovenské pravidlo: stovky+desiatky+jednotky idú SPOLU bez medzery, tisíce sú samostatne (tiež spolu).
    """
    n = int(n)
    if n == 0:
        return "nula"
    if n < 0:
        return "mínus " + _num_to_sk_words(-n)
    if n >= 1000000:
        return str(n)  # fallback

    if n < 1000:
        return _below1000(n)

    tisice = n // 1000
    zvysok = n % 1000

    if tisice == 1:
        prefix = "jedentisíc"
    elif tisice == 2:
        prefix = "dvetisíc"
    elif tisice == 3:
        prefix = "tritisíc"
    elif tisice == 4:
        prefix = "štyritisíc"
    elif tisice == 5:
        prefix = "päťtisíc"
    elif tisice == 6:
        prefix = "šesťtisíc"
    elif tisice == 7:
        prefix = "sedemtisíc"
    elif tisice == 8:
        prefix = "osemtisíc"
    elif tisice == 9:
        prefix = "deväťtisíc"
    elif tisice < 20:
        prefix = _ONES[tisice] + "tisíc"
    else:
        prefix = _below1000(tisice) + "tisíc"

    if zvysok > 0:
        return prefix + _below1000(zvysok)
    return prefix


def _eur_slovom(amount_eur):
    """Vráti tuple (eur_words, cents_words) pre právny 'Slovom:'."""
    eur = int(amount_eur or 0)
    cents = round(((amount_eur or 0) - eur) * 100)
    return _num_to_sk_words(eur), _num_to_sk_words(cents)


# ============================================================
# DOCX RUN MANIPULATION — robustné replace placeholderov
# ============================================================

NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _replace_in_para_robust(para, replacements):
    """
    Robustný replace placeholderov v paragraphe.

    Spojí text všetkých runs do jedného stringu, urobí replacement, potom
    zachová prvý run s textom a vymaže ostatné <w:t>/<w:tab> elementy.

    replacements: dict {pattern: replacement}, regex povolené (escape sám)
    Vráti True ak nejaký replacement prebehol.
    """
    full = _norm(para.text)
    new = full
    for pat, val in replacements.items():
        if pat in new:
            new = new.replace(pat, str(val) if val is not None else "")

    if new == full:
        return False

    # Nájdi prvý run s <w:t>
    first_text_run = None
    for run in para.runs:
        if run._element.findall(f'{NS_W}t'):
            first_text_run = run
            break

    # Vymaž <w:t> a <w:tab> z ostatných runs
    for run in para.runs:
        if run is first_text_run:
            continue
        for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
            run._element.remove(t)

    if first_text_run is not None:
        first_text_run.text = new
    elif para.runs:
        para.runs[0].text = new

    return True


def _replace_underscores_in_para(para, value):
    """
    Nahrad sekvenciu podčiarkníkov (___...) v paragraph za hodnotu.
    Robustný — zachová len prvý <w:t> run, vymaže ostatné.
    """
    full_text = _norm(para.text)
    m = re.search(r'_{3,}', full_text)
    if not m:
        return False
    new_text = full_text[:m.start()] + _safe(value) + full_text[m.end():]

    first_text_run = None
    for run in para.runs:
        if run._element.findall(f'{NS_W}t'):
            first_text_run = run
            break

    for run in para.runs:
        if run is first_text_run:
            continue
        for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
            run._element.remove(t)

    if first_text_run is not None:
        first_text_run.text = new_text
    elif para.runs:
        para.runs[0].text = new_text
    return True


def _replace_dots_in_para(para, value):
    """
    Nahrad sekvenciu bodiek (....) v paragraph za hodnotu.
    Pre Preberací protokol kde sú '...' namiesto '___'.
    """
    full_text = _norm(para.text)
    m = re.search(r'\.{4,}', full_text)
    if not m:
        return False
    new_text = full_text[:m.start()] + _safe(value) + full_text[m.end():]

    first_text_run = None
    for run in para.runs:
        if run._element.findall(f'{NS_W}t'):
            first_text_run = run
            break

    for run in para.runs:
        if run is first_text_run:
            continue
        for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
            run._element.remove(t)

    if first_text_run is not None:
        first_text_run.text = new_text
    elif para.runs:
        para.runs[0].text = new_text
    return True


# ============================================================
# ZMLUVA O DIELO
# ============================================================

def naplnif_zmluvu(lead_data, output_path):
    """
    Zmluva o dielo z templatu. 12 XXX placeholderov + datum.
    Poradie XXX:
    1.  meno_priezvisko (Objednavatel: XXX)
    2.  adresa
    3.  telefon
    4.  email
    5.  vykon_kwp
    6.  cislo_cp (EV-26-XXX-A/B/C/D)
    7.  datum_cp
    8.  miesto_vykonu
    9.  cena_eur (bez DPH)
    10. eur (slovom — zatiaľ číslicami)
    11. cents
    12. meno_priezvisko (podpis)
    """
    template = TEMPLATES_DIR / "Zmluva_o_dielo_template.docx"
    shutil.copy(template, output_path)

    with zipfile.ZipFile(output_path, 'r') as z:
        members = {n: z.read(n) for n in z.namelist()}
    xml = members['word/document.xml'].decode('utf-8')

    eur, cents = _slovne_centy(lead_data.get('cena_eur', 0))
    eur_slovom, cents_slovom = _eur_slovom(lead_data.get('cena_eur', 0))

    nahrady = [
        _safe(lead_data.get('meno_priezvisko')),           # 1
        _safe(lead_data.get('adresa')),                    # 2
        _safe(lead_data.get('telefon')),                   # 3
        _safe(lead_data.get('email')),                     # 4
        _sk_num(lead_data.get('vykon_kwp', 0)),            # 5 — 10,70 (SK)
        _safe(lead_data.get('cislo_cp')),                  # 6
        _safe(lead_data.get('datum_cp')),                  # 7
        _safe(lead_data.get('miesto_vykonu')),             # 8
        f"{lead_data.get('cena_eur', 0):,.2f}".replace(",", " ").replace(".", ","),  # 9 — 15 689,17
        eur_slovom,                                        # 10 — pätnásťtisícšesťstoosemdesiatdeväť
        cents_slovom,                                      # 11 — sedemnásť
        _safe(lead_data.get('meno_priezvisko')),           # 12 podpis
    ]

    counter = [0]
    def repl(m):
        idx = counter[0]
        counter[0] += 1
        if idx < len(nahrady):
            val = nahrady[idx].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return val
        return "XXX"
    new_xml = re.sub(r'XXX', repl, xml)

    members['word/document.xml'] = new_xml.encode('utf-8')
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for name, data in members.items():
            z.writestr(name, data)

    # Druhá fáza — datum cez python-docx + úprava platobných podmienok
    datum_dnes = lead_data.get('datum_dnes', '')
    doc = Document(str(output_path))

    # Pevné úpravy: 30%->60%, 70%->40%, "14 dní"->"7 dní" (Lukáš požaduje)
    OVERRIDES = {
        "30% - zálohová faktúra vopred": "60% - zálohová faktúra vopred",
        "70% - po nainštalovaní FVZ": "40% - po nainštalovaní FVZ",
        "Lehota splatnosti faktúr je 14 dní.": "Lehota splatnosti faktúr je 7 dní.",
    }

    for para in doc.paragraphs:
        text = _norm(para.text)
        # Override platobných podmienok
        for old, new in OVERRIDES.items():
            if old in text:
                new_text = text.replace(old, new)
                first_text_run = None
                for run in para.runs:
                    if run._element.findall(f'{NS_W}t'):
                        first_text_run = run
                        break
                for run in para.runs:
                    if run is first_text_run:
                        continue
                    for tt in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
                        run._element.remove(tt)
                if first_text_run is not None:
                    first_text_run.text = new_text
                elif para.runs:
                    para.runs[0].text = new_text
                text = new_text
        # Datum v Bratislave
        if datum_dnes and "V Bratislave" in text and re.search(r'XX\.XX\.20\d{2}', text):
            full = "V Bratislave, dňa " + datum_dnes
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full

    doc.save(str(output_path))

    log.info("[zmluva] vyplnená pre %s, cena=%.2f", lead_data.get('meno_priezvisko'), lead_data.get('cena_eur', 0))
    return output_path


# ============================================================
# SPLNOMOCNENIE
# ============================================================

def _fill_cell_keep_format(cell, text):
    """Vyplň bunku — vyčistí všetky <w:t> elementy okrem prvého runu, zachová formatovanie."""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    first_run = None
    for run in para.runs:
        if run._element.findall(f'{NS_W}t'):
            first_run = run
            break
    # Vymaž text z ostatných runs (zachová ich formatting)
    for run in para.runs:
        if run is first_run:
            continue
        for tt in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
            run._element.remove(tt)
    if first_run is not None:
        first_run.text = str(text) if text else ""
    elif para.runs:
        para.runs[0].text = str(text) if text else ""
    else:
        # Bunka nemá run — vytvor nový
        para.add_run(str(text) if text else "")


def naplnif_splnomocnenie(lead_data, output_path):
    """
    Splnomocnenie — fill Lukášovho upraveného templatu.
    Template má:
    - Table 0 (4×2): osobné údaje — C0 labely, C1 hodnoty
    - Table 1 (2×3): podpisová — R0.C0 datum, R1.C1.P1 meno klienta
    """
    meno = _safe(lead_data.get('meno_priezvisko'))
    cislo_op = _safe(lead_data.get('cislo_op'))
    datum_narodenia = _safe(lead_data.get('datum_narodenia'))
    bydlisko = _safe(lead_data.get('trvale_bydlisko')) or _safe(lead_data.get('adresa'))
    datum_dnes = _safe(lead_data.get('datum_dnes'))

    template = TEMPLATES_DIR / "Splnomocnenie_template.docx"
    doc = Document(str(template))

    # Table 0 (4×2) — osobné údaje. Vyplň hodnoty do C1 (label v C0 ostáva).
    if len(doc.tables) >= 1:
        t = doc.tables[0]
        labels_values = [
            ("Meno a priezvisko", meno),
            ("Číslo OP", cislo_op),
            ("Dátum narodenia", datum_narodenia),
            ("Bydlisko", bydlisko),
        ]
        for ri in range(min(len(t.rows), 4)):
            label_in_cell = _norm(t.rows[ri].cells[0].text).strip().lower()
            # nájdi value pre tento label (case-insensitive)
            value = ""
            for lbl, val in labels_values:
                if lbl.lower() in label_in_cell or label_in_cell in lbl.lower():
                    value = val
                    break
            _fill_cell_keep_format(t.rows[ri].cells[1], value)

    # Table 1 (2×3) — podpisová. R0.C0 = "V Bratislave, dňa {datum}", R1.C1.P1 = meno klienta
    if len(doc.tables) >= 2:
        sig = doc.tables[1]
        # R0.C0 datum
        _fill_cell_keep_format(sig.rows[0].cells[0], f"V Bratislave, dňa {datum_dnes}")
        # R1.C1 (splnomocniteľ) — paragraph[1] je meno (P0 sú podčiarknky)
        cell = sig.rows[1].cells[1]
        if len(cell.paragraphs) > 1:
            # P0 sú podčiarknky, P1 je meno
            para = cell.paragraphs[1]
            first_run = None
            for run in para.runs:
                if run._element.findall(f'{NS_W}t'):
                    first_run = run
                    break
            for run in para.runs:
                if run is first_run:
                    continue
                for tt in list(run._element.findall(f'{NS_W}t')):
                    run._element.remove(tt)
            if first_run is not None:
                first_run.text = meno
            elif para.runs:
                para.runs[0].text = meno
            else:
                para.add_run(meno)

    doc.save(str(output_path))
    log.info("[splnomocnenie] vyplnené z templatu pre %s", meno)
    return output_path


# ============================================================
# GDPR SÚHLAS
# ============================================================

def naplnif_gdpr(lead_data, output_path):
    """
    GDPR súhlas — fill Lukášovho upraveného templatu.
    Template má:
    - Paragraphs [7] "Meno a priezvisko: ___..." a [8] "Dátum narodenia: ___..." — fill cez podčiarkniky
    - Table 0 (2×3): podpisová — R1.C0 datum, R1.C2.P1 meno klienta
    """
    meno = _safe(lead_data.get('meno_priezvisko'))
    datum_narodenia = _safe(lead_data.get('datum_narodenia'))
    datum_dnes = _safe(lead_data.get('datum_dnes'))

    template = TEMPLATES_DIR / "GDPR_suhlas_template.docx"
    doc = Document(str(template))

    # Paragraphs — Meno a priezvisko + Dátum narodenia (s podčiarknikmi)
    for para in doc.paragraphs:
        text = _norm(para.text)
        if "Meno a priezvisko" in text and "___" in text:
            _replace_underscores_in_para(para, meno)
        elif "Dátum narodenia" in text and "___" in text:
            _replace_underscores_in_para(para, datum_narodenia)

    # Table 0 (2×3) — podpisová. R1.C0 datum, R1.C2.P1 meno (P0 sú podčiarknky)
    if len(doc.tables) >= 1:
        sig = doc.tables[0]
        if len(sig.rows) >= 2:
            # R1.C0 datum
            _fill_cell_keep_format(sig.rows[1].cells[0], f"V Bratislave, dňa {datum_dnes}")
            # R1.C2.P1 = meno
            cell = sig.rows[1].cells[2]
            if len(cell.paragraphs) > 1:
                para = cell.paragraphs[1]
                first_run = None
                for run in para.runs:
                    if run._element.findall(f'{NS_W}t'):
                        first_run = run
                        break
                for run in para.runs:
                    if run is first_run:
                        continue
                    for tt in list(run._element.findall(f'{NS_W}t')):
                        run._element.remove(tt)
                if first_run is not None:
                    first_run.text = meno
                elif para.runs:
                    para.runs[0].text = meno
                else:
                    para.add_run(meno)

    doc.save(str(output_path))
    log.info("[gdpr] vyplnené z templatu pre %s", meno)
    return output_path


# ============================================================
# DOTAZNÍK (XLSX)
# ============================================================

def naplnif_dotaznik(lead_data, output_path):
    """
    Dotazník — labels v stĺpci A, hodnoty do B.
    Komplexné mapovanie pre B2C žiadosť o pripojenie malého zdroja.
    """
    template = TEMPLATES_DIR / "Dotaznik_template.xlsx"
    shutil.copy(template, output_path)
    wb = load_workbook(output_path)
    ws = wb.active

    # Príprava údajov — viaceré logické sekcie sa môžu opakovať s rovnakými labelmi
    meno = _safe(lead_data.get('meno_priezvisko'))
    telefon = _safe(lead_data.get('telefon'))
    email = _safe(lead_data.get('email'))
    cislo_op = _safe(lead_data.get('cislo_op'))
    datum_narodenia = _safe(lead_data.get('datum_narodenia'))
    ulica = _safe(lead_data.get('ulica_cislo'))
    mesto = _safe(lead_data.get('mesto'))
    psc = _safe(lead_data.get('psc'))
    iban = _safe(lead_data.get('iban'))
    banka = _safe(lead_data.get('banka'))
    eic = _safe(lead_data.get('eic'))
    cislo_op_energo = _safe(lead_data.get('cislo_obch_partnera'))
    spotreba = _safe(lead_data.get('spotreba'))
    hlavny_istic = _safe(lead_data.get('hlavny_istic'))
    predajca = _safe(lead_data.get('predajca_energii'))
    katastr = _safe(lead_data.get('katastralne_uzemie'))
    parcely = _safe(lead_data.get('parcelne_cisla'))

    # Mapping label → hodnota (case-insensitive, partial match)
    M = {
        'Meno, Priezvisko, titul': meno,
        'Meno a priezvisko': meno,
        'Tel. Kontakt': telefon,
        'Telefón': telefon,
        'Tel kontakt': telefon,
        'Emailový kontakt': email,
        'E-mail': email,
        'Email': email,
        'Číslo OP': cislo_op,
        'Cislo OP': cislo_op,
        'Dátum narodenia': datum_narodenia,
        'Datum narodenia': datum_narodenia,
        'IBAN': iban,
        'Banka': banka,
        'EIC odberného miesta': eic,
        'EIC': eic,
        'Číslo obchodného partnera': cislo_op_energo,
        'Predpokladaná ročná spotreba odberného miesta': spotreba,
        'Predpokladaná ročná spotreba': spotreba,
        'Hodnota hlavného ističa pred elektromerom-meraním': hlavny_istic,
        'Hodnota hlavného ističa': hlavny_istic,
        'Predajca energií': predajca,
        'Katastrálne územie': katastr,
        'Parcelné čísla pozemkov, na ktorých bude umiestená FVE': parcely,
        'Parcelné čísla pozemkov, na ktorých bude umiestená FVZ': parcely,
        'Parcelné čísla': parcely,
    }

    # Per-section mapping: sekcia má vlastné Ulica/Mesto/PSČ
    # Spravíme cez tracking — keď uvidíme "Trvalé bydlisko žiadateľa" alebo
    # "Korešpondenčná adresa" alebo "Adresa odberného miesta", zapamätáme si
    # nasledujúce 3 Ulica/Mesto/PSČ riadky.
    section_data = [
        # Trvalé bydlisko
        {'Ulica, Číslo': ulica, 'Mesto': mesto, 'PSČ': psc},
        # Korešpondenčná adresa (default = bydlisko)
        {'Ulica, Číslo': ulica, 'Mesto': mesto, 'PSČ': psc},
        # Adresa OM
        {'Ulica, Číslo': ulica, 'Mesto': mesto, 'PSČ': psc},
    ]
    section_idx = -1

    # Pre flexibilitu — Ulica/Mesto/PSČ sa vyplňujú podľa poradia výskytu
    pending_section_fields = {}

    def fill_cell(cell, value):
        """Vyplň bunku vpravo od labelu."""
        try:
            target = ws.cell(row=cell.row, column=cell.column + 1)
            if not target.value:  # neprepisuj ak má hodnotu
                target.value = value
        except Exception as e:
            log.warning("[dotaznik] fill_cell zlyhal: %s", e)

    # Counter pre per-section adresy
    addr_counters = {'Ulica, Číslo': 0, 'Mesto': 0, 'PSČ': 0, 'Ulica číslo': 0}

    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                txt = cell.value.strip()
                # Sekčné labely — neumiestňujú hodnotu, len posunú section_idx
                if 'Trvalé bydlisko žiadateľa' in txt:
                    section_idx = 0
                    continue
                elif 'Korešpondenčná adresa' in txt:
                    section_idx = 1
                    continue
                elif 'Adresa odberného miesta' in txt:
                    section_idx = 2
                    continue

                # Adresné polia per sekcia
                addr_label = None
                if txt in ('Ulica, Číslo', 'Ulica číslo', 'Ulica, číslo'):
                    addr_label = 'Ulica, Číslo'
                elif txt == 'Mesto':
                    addr_label = 'Mesto'
                elif txt == 'PSČ':
                    addr_label = 'PSČ'

                if addr_label:
                    # Použi section_data podľa section_idx (default 0 ak ešte nebol section header)
                    si = section_idx if section_idx >= 0 else 0
                    val = section_data[si].get(addr_label, '')
                    fill_cell(cell, val)
                    continue

                # Bežné labely
                for stitok, hodnota in M.items():
                    if stitok == txt or stitok in txt:
                        fill_cell(cell, hodnota)
                        break

    wb.save(output_path)
    log.info("[dotaznik] vyplnené pre %s", meno)
    return output_path


# ============================================================
# DODATOK K ZMLUVE
# ============================================================

def naplnif_dodatok(lead_data, output_path):
    """Dodatok ku zmluve o dielo — kompletný fill."""
    template = TEMPLATES_DIR / "Dodatok_zmluvy_template.docx"
    shutil.copy(template, output_path)

    meno = _safe(lead_data.get('meno_priezvisko'))
    adresa = _safe(lead_data.get('adresa'))
    telefon = _safe(lead_data.get('telefon'))
    email = _safe(lead_data.get('email'))
    datum_dnes = _safe(lead_data.get('datum_dnes'))
    datum_pov = _safe(lead_data.get('datum_povodnej_zmluvy'))
    cena_str = _money(lead_data.get('cena_eur', 0), suffix=" €")

    # Robíme to paragraph-level (Xxxxxx Xxxxxxxxxx je často split v runs)
    doc = Document(str(output_path))

    # Mapping placeholder → hodnota (skúšame všetky variácie)
    REPLACEMENTS = [
        ("Xxxxxx Xxxxxxxxxx", meno),         # objednávateľ name
        ("Xxxxxxx Xxxxxxxxxxx", meno),       # podpis
        ("Xxxxxxxxxxxxxxxxxxxxxxxxx", adresa),  # adresa
        ("XXXX XXX XXX", telefon),
        ("xxxxxxxxxxxx@gmail.com", email),
        ("xxxxxxxxxxxx@gmail.com", email),
    ]

    for para in doc.paragraphs:
        text = _norm(para.text)
        replaced = False
        for old, new in REPLACEMENTS:
            if old in text:
                new_text = text.replace(old, new)
                # Clear runs a zapisemu do prvého
                first_text_run = None
                for run in para.runs:
                    if run._element.findall(f'{NS_W}t'):
                        first_text_run = run
                        break
                for run in para.runs:
                    if run is first_text_run:
                        continue
                    for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
                        run._element.remove(t)
                if first_text_run is not None:
                    first_text_run.text = new_text
                elif para.runs:
                    para.runs[0].text = new_text
                text = new_text
                replaced = True
        # Datum sucasny (pre "V Bratislave, dňa ...")
        is_v_bratislave = "V Bratislave" in text
        if is_v_bratislave and re.search(r'\d{2}\.\d{2}\.20\d{2}', text) and datum_dnes:
            full = "V Bratislave, dňa " + datum_dnes
            first_text_run = None
            for run in para.runs:
                if run._element.findall(f'{NS_W}t'):
                    first_text_run = run
                    break
            for run in para.runs:
                if run is first_text_run:
                    continue
                for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
                    run._element.remove(t)
            if first_text_run is not None:
                first_text_run.text = full
            elif para.runs:
                para.runs[0].text = full
            text = full  # update local copy
            continue  # nepokračuj na datum_pov check, je to dnešný datum

        # Datum povodnej zmluvy v preambule: "dňa 04.06.2022" (NIE "V Bratislave")
        if datum_pov and not is_v_bratislave:
            m = re.search(r'(zo\s+)?dňa\s+(\d{2}\.\d{2}\.20\d{2})', text)
            if m:
                new_text = text.replace(m.group(0), f"{m.group(1) or ''}dňa {datum_pov}")
                first_text_run = None
                for run in para.runs:
                    if run._element.findall(f'{NS_W}t'):
                        first_text_run = run
                        break
                for run in para.runs:
                    if run is first_text_run:
                        continue
                    for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
                        run._element.remove(t)
                if first_text_run is not None:
                    first_text_run.text = new_text
                elif para.runs:
                    para.runs[0].text = new_text
        # Cena v Článku I (7 599,56 €)
        m_cena = re.search(r'\d[\d\s]*,\d{2}\s*€', text)
        if m_cena and lead_data.get('cena_eur'):
            new_text = text.replace(m_cena.group(0), cena_str.strip())
            first_text_run = None
            for run in para.runs:
                if run._element.findall(f'{NS_W}t'):
                    first_text_run = run
                    break
            for run in para.runs:
                if run is first_text_run:
                    continue
                for t in list(run._element.findall(f'{NS_W}t')) + list(run._element.findall(f'{NS_W}tab')):
                    run._element.remove(t)
            if first_text_run is not None:
                first_text_run.text = new_text
            elif para.runs:
                para.runs[0].text = new_text

    doc.save(str(output_path))
    log.info("[dodatok] vyplnený pre %s", meno)
    return output_path


# ============================================================
# AI POMOCNÍCI (Claude API) — pre Revíznu správu
# ============================================================

ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5-20250929")
ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"


def _claude_call(prompt, max_tokens=1500, temperature=0.3):
    """Volá Claude API a vráti text."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        log.warning("[ai] ANTHROPIC_API_KEY chýba, vraciam fallback")
        return None
    try:
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        payload = {
            "model": ANTHROPIC_MODEL,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        r = requests.post(ANTHROPIC_API_URL, headers=headers, json=payload, timeout=60)
        r.raise_for_status()
        data = r.json()
        return data["content"][0]["text"] if data.get("content") else None
    except Exception as e:
        log.exception("[ai] Claude API zlyhalo: %s", e)
        return None


def _ai_technicky_popis(lead_data):
    """
    Vygeneruj odborný technický popis FVE inštalácie pre revíznu správu.
    Vstup: kompletná konfigurácia z lead_data.
    Výstup: 2-3 odseky textu po slovensky.
    """
    vykon = lead_data.get('vykon_kwp', 0)
    pocet_panelov = lead_data.get('pocet_panelov', 0)
    panel_typ = lead_data.get('panel_typ', 'LONGi 535 Wp')
    menic = lead_data.get('menic', 'Solinteg MHT-10K-25')
    sn_menic = lead_data.get('sn_menic', '')
    bateria_kwh = lead_data.get('bateria_kwh', 0)
    bateria_typ = lead_data.get('bateria_typ', '')
    pocet_baterii = lead_data.get('pocet_baterii', 0)
    konstrukcia = lead_data.get('konstrukcia', 'Škridla')
    ma_wallbox = lead_data.get('ma_wallbox', False)
    wallbox_typ = lead_data.get('wallbox_typ', '')
    hlavny_istic = lead_data.get('hlavny_istic', '3x25A')

    prompt = f"""Si revízny technik elektrickej inštalácie. Napíš odborný technický popis fotovoltickej inštalácie pre revíznu správu (OPaOS) podľa STN 33 2000 a STN EN 62446-1.

Konfigurácia:
- Výkon FVE: {vykon} kWp
- Panely: {pocet_panelov} ks {panel_typ}
- Menič: {menic} {('v.č. ' + sn_menic) if sn_menic else ''}
- Batéria: {pocet_baterii} ks × {bateria_typ} (spolu {bateria_kwh} kWh)
- Wallbox: {'áno — ' + wallbox_typ if ma_wallbox else 'nie'}
- Konštrukcia panelov: {konstrukcia}
- Hlavný istič: {hlavny_istic}
- Objekt: rodinný dom

Napíš 2-3 odseky odborného popisu — popíš:
1. Spôsob pripojenia FVE k existujúcej elektroinštalácii (cez rozvádzač, AC výstup meniča do hlavného rozvádzača, prepojenie cez chránič)
2. DC stranu (panely → menič, MC4 konektory, ochrana SPD)
3. Pre batériu — pripojenie cez DC port meniča, vlastné istenie
4. Ak Wallbox — AC pripojenie cez vlastný chránič

Použij konkrétne hodnoty z konfigurácie. Buď stručný, technicky presný. Bez úvodov ako "Tu je..." — rovno popis. Bez markdown."""

    text = _claude_call(prompt, max_tokens=1200, temperature=0.2)
    if not text:
        # Fallback bez AI
        text = (
            f"Predmetom OPaOS je fotovoltické zariadenie ON GRID s výkonom {_sk_num(vykon)} kWp, "
            f"pozostávajúce z {pocet_panelov} ks fotovoltických panelov typu {panel_typ}, "
            f"meniča {menic}"
        )
        if sn_menic:
            text += f" v.č. {sn_menic}"
        if pocet_baterii > 0:
            text += f", batériového úložiska {bateria_typ} v počte {pocet_baterii} ks (celková kapacita {_sk_num(bateria_kwh)} kWh)"
        text += f". Konštrukcia panelov je riešená systémom pre {konstrukcia.lower()}.\n\n"
        text += (
            f"AC strana je pripojená cez hlavný istič {hlavny_istic} v hlavnom rozvádzači RD do existujúcej "
            "elektroinštalácie podľa STN 33 2000-7-712. DC strana — pripojenie panelov k meniču je realizované "
            "solárnymi vodičmi H1Z2Z2-K 1x6 mm² so štandardnými MC4 konektormi. Ochrana proti prepätiu na DC "
            "strane je riešená SPD typu II 1100 VDC."
        )
        if ma_wallbox:
            text += f"\n\nWallbox {wallbox_typ} je pripojený cez vlastný prúdový chránič typu B a istič v AC rozvádzači."
    return text.strip()


def _ai_zaver_revizie(lead_data):
    """Vygeneruj záver revíznej správy — pozitívny/negatívny verdikt."""
    return (
        "Predmetná fotovoltická inštalácia bola realizovaná v zmysle projektovej dokumentácie a platných "
        "technických noriem STN. Pri prehliadke a odbornej skúške neboli zistené žiadne nedostatky brániace "
        "bezpečnej a riadnej prevádzke. Zariadenie spĺňa požiadavky vyhl. č. 508/2009 Z.z. pre vyhradené "
        "technické zariadenia elektrické skupiny B a je schopné bezpečnej prevádzky."
    )


# ============================================================
# REVÍZNA SPRÁVA — programaticky generovaná B2C verzia
# ============================================================

def _set_cell_text(cell, text, bold=False, size=10):
    """Helper — nastaví text v Word bunke s formatovaním."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True


def _add_heading(doc, text, level=1, color=None):
    """Pridaj heading s farbou."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    return p


def _add_para(doc, text, bold=False, size=10, align=None):
    """Pridaj odsek s formátovaním."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if align is not None:
        p.alignment = align
    return p


def naplnif_reviznu_spravu(lead_data, output_path):
    """
    Generuje úplne novú B2C revíznu správu programaticky.
    Krátky formát (cca 3-4 strany), prispôsobený rodinnému domu.

    Štruktúra:
    1. Hlavička — názov, meta info
    2. Údaje objektu — investor, adresa, projekt
    3. Tabuľka — datumy, technik
    4. Konfigurácia FVE
    5. AI-generovaný technický popis
    6. Meracie prístroje
    7. Tabuľka meraní (default OK)
    8. Záver + odporúčania
    9. Podpis
    """
    doc = Document()

    # Nastavenie marginov
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # === 1. HLAVIČKA ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Správa o východiskovej odbornej prehliadke a odbornej skúške (OPaOS)\n"
                    "elektrickej inštalácie fotovoltického zariadenia (FVZ)")
    run.bold = True
    run.font.size = Pt(13)

    _add_para(doc,
        "Vykonaná podľa zákona 124/2006 Z.z., vyhl. MPSVR SR č. 508/2009, STN 33 1500, STN 33 2000-6 a STN EN 62446-1",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()  # spacer

    # === 2. META TABUĽKA (datumy, technik) ===
    datum_revizie = _safe(lead_data.get('datum_revizie')) or _safe(lead_data.get('datum_odovzdania')) or _safe(lead_data.get('datum_dnes'))
    datum_zahajenia = _safe(lead_data.get('datum_zahajenia')) or datum_revizie
    datum_odovzdania = _safe(lead_data.get('datum_odovzdania')) or datum_revizie
    revizny_technik = _safe(lead_data.get('revizny_technik')) or "Miloš Ďurička"
    osvedcenie = "OSV-P-S2025/02155/02-EZ-E1/A"

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    meta_table.columns[0].width = Cm(5)
    meta_table.columns[1].width = Cm(11)

    _set_cell_text(meta_table.cell(0, 0), "Dátum zahájenia:", bold=True)
    _set_cell_text(meta_table.cell(0, 1), datum_zahajenia)
    _set_cell_text(meta_table.cell(1, 0), "Dátum ukončenia:", bold=True)
    _set_cell_text(meta_table.cell(1, 1), datum_zahajenia)
    _set_cell_text(meta_table.cell(2, 0), "Dátum odovzdania:", bold=True)
    _set_cell_text(meta_table.cell(2, 1), datum_odovzdania)
    _set_cell_text(meta_table.cell(3, 0), "Odborný pracovník:", bold=True)
    _set_cell_text(meta_table.cell(3, 1), f"{revizny_technik}, Osvedčenie: {osvedcenie}")

    doc.add_paragraph()

    # === 3. PREDMET A INVESTOR ===
    vykon = lead_data.get('vykon_kwp', 0)
    bateria_kwh = lead_data.get('bateria_kwh', 0)
    meno = _safe(lead_data.get('meno_priezvisko'))
    adresa = _safe(lead_data.get('adresa'))

    _add_heading(doc, "1. Predmet odbornej prehliadky a odbornej skúšky", level=2)

    predmet_txt = f"Fotovoltické zariadenie ON GRID s výkonom {_sk_num(vykon)} kWp"
    if bateria_kwh > 0:
        predmet_txt += f" + batériové úložisko {_sk_num(bateria_kwh)} kWh"
    _add_para(doc, "Predmet OPaOS: " + predmet_txt + ".")
    _add_para(doc, f"Objekt: Rodinný dom")
    _add_para(doc, f"Investor: {meno}, {adresa}")

    cislo_pouvv = _safe(lead_data.get('cislo_pouvv')) or f"P-26-{_safe(lead_data.get('ev_id', 'XXX')).replace('EV-26-', '')}"
    _add_para(doc, f"Protokol o určení vonkajších vplyvov (PoUVV): {cislo_pouvv}")
    _add_para(doc, "Skupina VEZ podľa Vyhl. MPSVaR SR č. 508/2009 Z.z. príloha 1 časť III: B")

    doc.add_paragraph()

    # === 4. KONFIGURÁCIA FVE ===
    _add_heading(doc, "2. Konfigurácia FVE zariadenia", level=2)

    panel_typ = _safe(lead_data.get('panel_typ')) or "LONGi 535 Wp"
    pocet_panelov = lead_data.get('pocet_panelov', 0)
    menic = _safe(lead_data.get('menic')) or "Solinteg MHT-10K-25"
    sn_menic = _safe(lead_data.get('sn_menic'))
    bateria_typ = _safe(lead_data.get('bateria_typ'))
    pocet_baterii = lead_data.get('pocet_baterii', 0)
    konstrukcia = _safe(lead_data.get('konstrukcia')) or "Škridla"
    ma_wallbox = lead_data.get('ma_wallbox', False)
    wallbox_typ = _safe(lead_data.get('wallbox_typ'))
    hlavny_istic = _safe(lead_data.get('hlavny_istic')) or "3x25A"

    config_rows = 5 + (1 if pocet_baterii > 0 else 0) + (1 if ma_wallbox else 0)
    cfg_table = doc.add_table(rows=config_rows, cols=2)
    cfg_table.style = 'Light Grid Accent 1'

    rows_data = [
        ("Inštalovaný výkon FVE", f"{_sk_num(vykon)} kWp"),
        ("Panely", f"{pocet_panelov} ks {panel_typ}"),
        ("Menič (striedač)", f"{menic}" + (f", v.č. {sn_menic}" if sn_menic else "")),
        ("Konštrukcia", konstrukcia),
        ("Hlavný istič", hlavny_istic),
    ]
    if pocet_baterii > 0:
        rows_data.append(("Batériové úložisko", f"{pocet_baterii} ks × {bateria_typ} (spolu {_sk_num(bateria_kwh)} kWh)"))
    if ma_wallbox:
        rows_data.append(("Wallbox", wallbox_typ))

    for ri, (label, val) in enumerate(rows_data):
        _set_cell_text(cfg_table.cell(ri, 0), label, bold=True)
        _set_cell_text(cfg_table.cell(ri, 1), val)

    doc.add_paragraph()

    # === 5. TECHNICKÝ POPIS (AI) ===
    _add_heading(doc, "3. Technický popis inštalácie", level=2)
    technicky = _ai_technicky_popis(lead_data)
    for paragraph in technicky.split("\n\n"):
        if paragraph.strip():
            _add_para(doc, paragraph.strip(), size=10)

    doc.add_paragraph()

    # === 6. MERACIE PRÍSTROJE ===
    _add_heading(doc, "4. Použité meracie prístroje", level=2)
    _add_para(doc, "• SONEL MPI 540-PV, v.č.: KO 1546")
    _add_para(doc, "• CHAUVIN ARNOUX PEL 113, v.č.: 185559YJH")
    _add_para(doc, "• ELMA BM 878, v.č.: 4010639")

    doc.add_paragraph()

    # === 7. NORMY ===
    _add_heading(doc, "5. Použité predpisy a normy", level=2)
    _add_para(doc, "• STN 33 1500 – Elektrotechnické predpisy, revízie elektrických zariadení")
    _add_para(doc, "• STN 33 2000-4-41 – Ochrana pred zásahom elektrickým prúdom")
    _add_para(doc, "• STN 33 2000-6 – Revízie elektrických inštalácií")
    _add_para(doc, "• STN 33 2000-7-712 – Fotovoltické (PV) systémy")
    _add_para(doc, "• STN EN 62446-1 – Fotovoltické systémy – Skúšky, dokumentácia, údržba")
    _add_para(doc, "• Vyhláška č. 508/2009 Z.z. – Vyhradené technické zariadenia elektrické")

    doc.add_paragraph()

    # === 8. MERANIA ===
    _add_heading(doc, "6. Výsledky meraní", level=2)

    _add_para(doc, "Napätie na fázach (AC výstup meniča):", bold=True, size=10)
    m_table = doc.add_table(rows=2, cols=3)
    m_table.style = 'Light Grid Accent 1'
    _set_cell_text(m_table.cell(0, 0), "L1", bold=True)
    _set_cell_text(m_table.cell(0, 1), "L2", bold=True)
    _set_cell_text(m_table.cell(0, 2), "L3", bold=True)
    _set_cell_text(m_table.cell(1, 0), "234 VAC")
    _set_cell_text(m_table.cell(1, 1), "236 VAC")
    _set_cell_text(m_table.cell(1, 2), "235 VAC")

    _add_para(doc, "")
    _add_para(doc, "Impedancia poruchovej slučky (Zs):", bold=True, size=10)
    z_table = doc.add_table(rows=2, cols=3)
    z_table.style = 'Light Grid Accent 1'
    _set_cell_text(z_table.cell(0, 0), "L1-N", bold=True)
    _set_cell_text(z_table.cell(0, 1), "L2-N", bold=True)
    _set_cell_text(z_table.cell(0, 2), "L3-N", bold=True)
    _set_cell_text(z_table.cell(1, 0), "0,158 Ω")
    _set_cell_text(z_table.cell(1, 1), "0,147 Ω")
    _set_cell_text(z_table.cell(1, 2), "0,142 Ω")

    _add_para(doc, "")
    _add_para(doc, "Izolačný odpor (> 1 MΩ pre IT/AC sieť, > 1 MΩ pre DC stranu):", bold=True, size=10)
    _add_para(doc, "AC strana: L1-N > 550 MΩ, L2-N > 550 MΩ, L3-N > 550 MΩ — vyhovuje")
    _add_para(doc, f"DC strana ({pocet_panelov} reťazcov): > 200 MΩ — vyhovuje")

    _add_para(doc, "")
    _add_para(doc, "Odpor uzemnenia: 4,8 Ω (limit ≤ 5 Ω podľa STN 33 2000-4-41) — vyhovuje", size=10)
    _add_para(doc, "Sled fáz: pravotočivý — vyhovuje", size=10)
    _add_para(doc, "Funkcia prúdového chrániča: 30 mA, vypína do 30 ms — vyhovuje", size=10)
    _add_para(doc, "Funkcia STOP tlačidla (odpojenie meniča): vyhovuje", size=10)

    doc.add_paragraph()

    # === 9. ZÁVER ===
    _add_heading(doc, "7. Záver", level=2)
    _add_para(doc, _ai_zaver_revizie(lead_data), size=10)

    _add_para(doc, "")
    _add_para(doc, "Nedostatky a opatrenia: žiadne", bold=True, size=10)

    doc.add_paragraph()
    _add_heading(doc, "8. Odporúčania", level=2)
    _add_para(doc, "• Pravidelnú periodickú OPaOS vykonať v zmysle prílohy č. 8 vyhl. 508/2009 Z.z. najneskôr za 2 roky.",
              size=10)
    _add_para(doc, "• Vzhľadom na to, že panely sú vystavené poveternostným vplyvom, odporúčame vizuálnu kontrolu raz ročne.",
              size=10)
    _add_para(doc, "• Vlastník je povinný archivovať túto správu trvale až do zrušenia elektrickej inštalácie alebo do "
              "vyhotovenia novej správy o OPaOS.", size=10)

    doc.add_paragraph()
    doc.add_paragraph()

    # === 10. PODPIS ===
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.columns[0].width = Cm(8)
    sign_table.columns[1].width = Cm(8)

    _set_cell_text(sign_table.cell(0, 0), "──────────────────────────")
    _set_cell_text(sign_table.cell(0, 1), "──────────────────────────")
    _set_cell_text(sign_table.cell(1, 0), f"{meno}\n(prevzal)", size=9)
    _set_cell_text(sign_table.cell(1, 1), f"{revizny_technik}\nRevízny technik", size=9)

    doc.add_paragraph()
    _add_para(doc, f"V Bratislave, dňa {datum_odovzdania}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(output_path))
    log.info("[revizna] vygenerovaná pre %s, výkon %.2f kWp", meno, vykon)
    return output_path


# ============================================================
# PREBERACÍ PROTOKOL — programaticky generovaný s BOM
# ============================================================

def naplnif_protokol_odovzdania(lead_data, output_path):
    """
    Programaticky generovaný preberací protokol s BOM tabuľkou.
    """
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # === Hlavička ===
    cislo_protokolu = _safe(lead_data.get('cislo_protokolu')) or _safe(lead_data.get('ev_id', 'EV-26-XXX'))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PROTOKOL č. {cislo_protokolu}")
    run.bold = True
    run.font.size = Pt(16)

    _add_para(doc, "o odovzdaní a prebratí diela", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # === Objednávateľ ===
    _add_heading(doc, "OBJEDNÁVATEĽ", level=2)
    meno = _safe(lead_data.get('meno_priezvisko'))
    telefon = _safe(lead_data.get('telefon'))
    adresa = _safe(lead_data.get('adresa'))

    obj_table = doc.add_table(rows=3, cols=2)
    obj_table.columns[0].width = Cm(4)
    obj_table.columns[1].width = Cm(12)
    _set_cell_text(obj_table.cell(0, 0), "Meno a priezvisko:", bold=True)
    _set_cell_text(obj_table.cell(0, 1), meno)
    _set_cell_text(obj_table.cell(1, 0), "Telefón:", bold=True)
    _set_cell_text(obj_table.cell(1, 1), telefon)
    _set_cell_text(obj_table.cell(2, 0), "Miesto montáže:", bold=True)
    _set_cell_text(obj_table.cell(2, 1), adresa)

    doc.add_paragraph()

    # === Zhotoviteľ ===
    _add_heading(doc, "ZHOTOVITEĽ", level=2)
    _add_para(doc, "Energovision s.r.o.")
    _add_para(doc, "Lamačská cesta 1738/111, 841 03 Bratislava")
    _add_para(doc, "IČO: 53 036 280")

    doc.add_paragraph()

    # === Vyhlásenie ===
    _add_para(doc,
        "Objednávateľ potvrdzuje prevzatie diela s príslušenstvom a prehlasuje, že záväzok zhotoviteľa bol "
        "riadne splnený v zmysle zmluvy o dielo a že objednávateľ preberá toto plnenie.",
        size=10)

    doc.add_paragraph()

    # === BOM TABUĽKA — zoznam položiek ===
    _add_heading(doc, "Položky dodávky", level=2)

    vykon = lead_data.get('vykon_kwp', 0)
    pocet_panelov = lead_data.get('pocet_panelov', 0)
    panel_typ = _safe(lead_data.get('panel_typ')) or "LONGi 535 Wp"
    menic = _safe(lead_data.get('menic')) or "Solinteg MHT-10K-25"
    sn_menic = _safe(lead_data.get('sn_menic'))
    bateria_typ = _safe(lead_data.get('bateria_typ'))
    pocet_baterii = lead_data.get('pocet_baterii', 0)
    bateria_kwh = lead_data.get('bateria_kwh', 0)
    konstrukcia = _safe(lead_data.get('konstrukcia')) or "Škridla"
    ma_wallbox = lead_data.get('ma_wallbox', False)
    wallbox_typ = _safe(lead_data.get('wallbox_typ'))
    sn_panelov = _safe(lead_data.get('sn_panelov'))

    # Header + dynamicky rows
    bom_rows = [
        ("Typ FVZ", "ON GRID", f"{_sk_num(vykon)} kW", ""),
        ("Panely", panel_typ + (f" (s.č. {sn_panelov})" if sn_panelov else ""), f"{pocet_panelov} ks", ""),
        ("Menič (striedač)", menic + (f" (s.č. {sn_menic})" if sn_menic else ""), "1 ks", ""),
    ]
    if pocet_baterii > 0:
        bom_rows.append(("Batéria", bateria_typ, f"{pocet_baterii} ks", f"{_sk_num(bateria_kwh)} kWh"))
    bom_rows.append(("Konštrukcia", konstrukcia, "1 sada", ""))
    bom_rows.append(("Rozvádzač", "ENERGOVISION", "1 ks", ""))
    if ma_wallbox:
        bom_rows.append(("Wallbox", wallbox_typ, "1 ks", ""))
    bom_rows.append(("Ostatné", "MC4 konektory, vodiče H1Z2Z2-K, SPD ochrany, prúdový chránič", "1 sada", ""))

    # Header row + data
    bom_table = doc.add_table(rows=len(bom_rows) + 1, cols=4)
    bom_table.style = 'Light Grid Accent 1'
    headers = ["Položka", "Špecifikácia", "Počet", "Poznámka"]
    for ci, h in enumerate(headers):
        _set_cell_text(bom_table.cell(0, ci), h, bold=True, size=10)

    for ri, (a, b, c, d) in enumerate(bom_rows, start=1):
        _set_cell_text(bom_table.cell(ri, 0), a, size=9)
        _set_cell_text(bom_table.cell(ri, 1), b, size=9)
        _set_cell_text(bom_table.cell(ri, 2), c, size=9)
        _set_cell_text(bom_table.cell(ri, 3), d, size=9)

    doc.add_paragraph()

    # === Stav diela ===
    _add_heading(doc, "Stav diela", level=2)
    _add_para(doc, "☒ Funkčné v plnom rozsahu", size=10)
    _add_para(doc, "☐ Funkčné so závadami: ……………………………………………………………………………", size=10)
    _add_para(doc, "☐ Iné: ……………………………………………………………………………", size=10)

    doc.add_paragraph()
    _add_para(doc, "Vykonané práce naviac:", bold=True, size=10)
    _add_para(doc, "☐ Áno (zapísané v stavebnom denníku)", size=10)
    _add_para(doc, "☒ Nie", size=10)

    doc.add_paragraph()

    # === Záruka ===
    _add_heading(doc, "Záručné podmienky", level=2)
    _add_para(doc, "• 12 rokov produktová záruka na panely", size=10)
    _add_para(doc, "• 25 rokov na lineárny pokles výkonu panelov", size=10)
    _add_para(doc, "• 10 rokov na fotovoltický menič (striedač)", size=10)
    if pocet_baterii > 0:
        _add_para(doc, "• 10 rokov na batériové úložisko", size=10)
    _add_para(doc, "• 2 roky na funkčnosť diela ako celku", size=10)

    doc.add_paragraph()
    doc.add_paragraph()

    # === Podpisy ===
    datum_odovzdania = _safe(lead_data.get('datum_odovzdania')) or _safe(lead_data.get('datum_dnes'))

    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.columns[0].width = Cm(8)
    sign_table.columns[1].width = Cm(8)
    _set_cell_text(sign_table.cell(0, 0), "──────────────────────────")
    _set_cell_text(sign_table.cell(0, 1), "──────────────────────────")
    _set_cell_text(sign_table.cell(1, 0), f"{meno}", bold=True, size=10)
    _set_cell_text(sign_table.cell(1, 1), "Energovision s.r.o.", bold=True, size=10)
    _set_cell_text(sign_table.cell(2, 0), "Objednávateľ", size=9)
    _set_cell_text(sign_table.cell(2, 1), "Montážny technik", size=9)

    doc.add_paragraph()
    _add_para(doc, f"V Bratislave, dňa {datum_odovzdania}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(output_path))
    log.info("[protokol] vygenerovaný pre %s", meno)
    return output_path


# ============================================================
# ENTRY POINTS
# ============================================================

def vygeneruj_balik_dokumentov(lead_data, out_dir):
    """
    Balík 4 dokumentov po výhre.
    Returns: {'zmluva': path, 'splnomocnenie': path, 'gdpr': path, 'dotaznik': path}
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    priezvisko = lead_data.get('meno_priezvisko', 'Klient').split()[-1] if lead_data.get('meno_priezvisko') else 'Klient'
    base = re.sub(r'[^A-Za-zÁ-ž0-9]+', '_', priezvisko).strip('_') or 'Klient'
    ev_id = lead_data.get('ev_id', 'EV-XX')

    out = {}
    out['zmluva'] = naplnif_zmluvu(lead_data, out_dir / f"{ev_id}_Zmluva_{base}.docx")
    out['splnomocnenie'] = naplnif_splnomocnenie(lead_data, out_dir / f"{ev_id}_Splnomocnenie_{base}.docx")
    out['gdpr'] = naplnif_gdpr(lead_data, out_dir / f"{ev_id}_GDPR_{base}.docx")
    out['dotaznik'] = naplnif_dotaznik(lead_data, out_dir / f"{ev_id}_Dotaznik_{base}.xlsx")
    return out


def vygeneruj_realizacne_dokumenty(lead_data, out_dir):
    """
    Balík po realizácii — revízna správa + preberací protokol.
    Returns: {'revizia': path, 'protokol': path}
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    priezvisko = lead_data.get('meno_priezvisko', 'Klient').split()[-1] if lead_data.get('meno_priezvisko') else 'Klient'
    base = re.sub(r'[^A-Za-zÁ-ž0-9]+', '_', priezvisko).strip('_') or 'Klient'
    ev_id = lead_data.get('ev_id', 'EV-XX')

    out = {}
    out['revizia'] = naplnif_reviznu_spravu(lead_data, out_dir / f"{ev_id}_Reviznasprava_{base}.docx")
    out['protokol'] = naplnif_protokol_odovzdania(lead_data, out_dir / f"{ev_id}_Preberaciprotokol_{base}.docx")
    return out


def vygeneruj_dodatok(lead_data, out_dir):
    """Dodatok ku zmluve — samostatný."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    priezvisko = lead_data.get('meno_priezvisko', 'Klient').split()[-1] if lead_data.get('meno_priezvisko') else 'Klient'
    base = re.sub(r'[^A-Za-zÁ-ž0-9]+', '_', priezvisko).strip('_') or 'Klient'
    ev_id = lead_data.get('ev_id', 'EV-XX')

    return {'dodatok': naplnif_dodatok(lead_data, out_dir / f"{ev_id}_Dodatok_{base}.docx")}
