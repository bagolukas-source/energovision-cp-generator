# -*- coding: utf-8 -*-
"""Generátor dokumentov pre Správu trafostaníc (DOCX).
Vstup: dict `ts` (transformer_stations row vrátane tech_details) + voliteľne contract/inspection.
Fáza 1: Preberací protokol + Zmluva. Ďalej: MPP, Revízna správa."""
from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ENERGO = {
    "nazov": "Energovision s.r.o.",
    "sidlo": "Lamačská cesta 1738/111, 841 03 Bratislava",
    "office": "Tomášikova 19, 821 02 Bratislava",
    "ico": "53 036 280", "dic": "2121238526", "icdph": "SK2121238526",
    "orsr": "OR OS Bratislava I, odd: Sro, vložka č. 158744/B",
    "tel": "+421 948 302 137", "email": "info@energovision.sk",
    "banka": "Tatra banka, a.s.", "iban": "SK48 1100 0000 0029 4708 4971", "swift": "TATRSKBX",
}
GREEN = RGBColor(0x1B, 0x5E, 0x3F)


def _h(doc, text, size=14, color=GREEN, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, after=6):
    p = doc.add_paragraph(); p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def _kv_table(doc, rows, widths=(60, 40)):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = str(k); c[1].text = "" if v is None else str(v)
        for run in c[0].paragraphs[0].runs: run.bold = True
    return t


def _two_col(doc, left_title, left_lines, right_title, right_lines):
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    for cell, title, lines in ((t.rows[0].cells[0], left_title, left_lines),
                               (t.rows[0].cells[1], right_title, right_lines)):
        p = cell.paragraphs[0]; r = p.add_run(title); r.bold = True; r.font.color.rgb = GREEN
        for ln in lines:
            cp = cell.add_paragraph(ln); cp.paragraph_format.space_after = Pt(0)
    return t


# ============================================================
# 1) PREBERACÍ PROTOKOL
# ============================================================
def generate_preberaci_protokol(ts: dict) -> bytes:
    td = ts.get("tech_details") or {}
    prev = td.get("prevadzkovatel") or {}
    doc = Document()
    _h(doc, "PREBERACÍ PROTOKOL", size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph("Týmto preberacím protokolom preberáme uvedenú transformátorovú stanicu "
                          "do našej správy na základe servisnej zmluvy.")
    p.paragraph_format.space_after = Pt(12)

    _two_col(doc,
        "Dodávateľ – zhotoviteľ:",
        [ENERGO["nazov"], ENERGO["sidlo"], f"IČO: {ENERGO['ico']}", f"DIČ: {ENERGO['dic']}",
         f"IČ DPH: {ENERGO['icdph']}", ENERGO["orsr"]],
        "Doručovacia adresa / kontakt:",
        [ENERGO["nazov"], ENERGO["office"], f"tel.: {ENERGO['tel']}", ENERGO["email"],
         f"{ENERGO['banka']}  SWIFT: {ENERGO['swift']}", f"IBAN: {ENERGO['iban']}"])
    doc.add_paragraph()

    _h(doc, "Objednávateľ:", size=12)
    _kv_table(doc, [
        ("Názov spoločnosti", prev.get("nazov") or ts.get("name")),
        ("IČO / DIČ", f"{prev.get('ico','—')} / {prev.get('dic','—')}"),
        ("Adresa inštalácie zariadenia", ts.get("location_address") or prev.get("sidlo")),
        ("Mesto / PSČ", f"{ts.get('location_city','')} {ts.get('location_psc','')}".strip()),
        ("Kontaktná osoba", prev.get("kontakt")),
        ("Telefón", prev.get("tel")),
    ])
    doc.add_paragraph()

    _h(doc, "Technická špecifikácia:", size=12)
    tr = td.get("transformator") or {}
    _kv_table(doc, [
        ("Označenie", ts.get("ts_code")),
        ("Názov / typ TS", ts.get("ts_type")),
        ("Umiestnenie", f"{ts.get('location_address','')}, {ts.get('location_psc','')} {ts.get('location_city','')}".strip(", ")),
        ("Výkon transformátora", f"{ts.get('rated_power_kva','—')} kVA"),
        ("Napätie VN/NN", f"{ts.get('vn_voltage_kv','—')} kV / {ts.get('nn_voltage_v','—')} V"),
        ("Počet kusov", "1"),
    ])
    doc.add_paragraph(); doc.add_paragraph()

    _two_col(doc, "Za objednávateľa:", ["", "Meno: ............................", "Dátum: ............................",
                                        "Podpis: ............................"],
             "Za Energovision s.r.o.:", ["", "Meno: Lukáš Bago", "Dátum: ............................",
                                         "Podpis: ............................"])
    b = BytesIO(); doc.save(b); return b.getvalue()


# ============================================================
# 2) ZMLUVA o prevádzkovaní a servise TS
# ============================================================
def generate_zmluva(ts: dict, contract: dict = None) -> bytes:
    td = ts.get("tech_details") or {}; prev = td.get("prevadzkovatel") or {}
    c = contract or {}
    pausal = c.get("monthly_fee_eur"); sla = c.get("sla_response_hours")
    doc = Document()
    _h(doc, "ZMLUVA O PREVÁDZKOVANÍ A SERVISE TRANSFORMÁTOROVEJ STANICE", size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
    p=doc.add_paragraph("uzatvorená podľa § 269 ods. 2 zákona č. 513/1991 Zb. Obchodný zákonník v znení neskorších predpisov ako zmluva o poskytovaní služieb")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True; p.runs[0].font.size=Pt(9)
    doc.add_paragraph()
    _h(doc,"Zmluvné strany",size=12)
    _two_col(doc,"Poskytovateľ:",
        [ENERGO["nazov"], f"IČO: {ENERGO['ico']}", f"Sídlo: {ENERGO['sidlo']}", "Koná: Lukáš Bago",
         f"Tel.: +421 918 187 762", "E-mail: lukas.bago@energovision.sk", "(ďalej „Poskytovateľ“)"],
        "Objednávateľ:",
        [prev.get("nazov") or ts.get("name") or "—", f"IČO: {prev.get('ico','—')}", f"DIČ: {prev.get('dic','—')}",
         f"Sídlo: {prev.get('sidlo','—')}", f"Koná: {prev.get('kontakt','—')}", f"Tel.: {prev.get('tel','—')}", "(ďalej „Objednávateľ“)"])
    doc.add_paragraph()

    def art(n,t): _h(doc,f"Článok {n} — {t}",size=12,after=2)
    art("I","Predmet zmluvy")
    doc.add_paragraph("Poskytovateľ sa zaväzuje zabezpečovať pre Objednávateľa odbornú prevádzku, dohľad a servis transformátorovej stanice špecifikovanej v Prílohe č. 1. Rozsah služieb je uvedený v tejto zmluve a jej prílohách.")
    art("II","Rozsah poskytovaných služieb")
    for b in ["pravidelné kontroly stavu TS podľa platných technických predpisov","odborné prehliadky a skúšky podľa STN a súvisiacich predpisov",
              "periodické revízie prostredníctvom odborne spôsobilej osoby","vedenie a archivácia prevádzkovej a technickej dokumentácie TS",
              "návrhy na odstránenie zistených závad","súčinnosť pri komunikácii s prevádzkovateľom DS a orgánmi dozoru",
              "odborné pokyny pre obsluhu TS","účasť na plánovaných odstávkach a spúšťaní TS podľa dohody"]:
        doc.add_paragraph(b, style="List Bullet")
    art("III","Cena a platobné podmienky")
    doc.add_paragraph(f"Odmena je stanovená ako paušálna odmena vo výške {('%.2f € / mesiac' % float(pausal)) if pausal else '… € / mesiac (Príloha č. 1)'} a fakturuje sa mesačne.")
    doc.add_paragraph("Jednorazové služby a zásahy mimo paušálu sa fakturujú podľa platného cenníka. Materiál sa fakturuje podľa skutočnej nákupnej ceny zvýšenej o zmluvnú prirážku 15 %.")
    doc.add_paragraph("Splatnosť faktúr je 14 kalendárnych dní odo dňa doručenia. Pri omeškaní má Poskytovateľ právo na úrok z omeškania 0,05 % z dlžnej sumy za každý deň.")
    if sla: doc.add_paragraph(f"Garantovaná reakčná doba (SLA): do {int(sla)} hodín od nahlásenia.")
    art("IV","Práva a povinnosti zmluvných strán")
    doc.add_paragraph("Poskytovateľ vykonáva činnosti s odbornou starostlivosťou a v súlade s platnými predpismi. Objednávateľ poskytuje súčinnosť a prístup k TS.")
    art("V","Trvanie a ukončenie zmluvy")
    doc.add_paragraph("Zmluva sa uzatvára na dobu neurčitú. Vypovedať ju možno písomne s výpovednou lehotou podľa dohody zmluvných strán.")
    art("VI","Záverečné ustanovenia")
    doc.add_paragraph("Zmluva nadobúda platnosť a účinnosť dňom podpisu oboch zmluvných strán. Mení sa písomnými dodatkami. Vyhotovuje sa v dvoch rovnopisoch.")
    doc.add_paragraph()
    _two_col(doc,"Za Poskytovateľa:",["","Dátum: ............","Podpis: ............","Lukáš Bago, Energovision s.r.o."],
             "Za Objednávateľa:",["","Dátum: ............","Podpis: ............", prev.get("kontakt","")])
    doc.add_paragraph()
    _h(doc,"Príloha č. 1 — Zoznam a špecifikácia TS + cena",size=11)
    _kv_table(doc,[("Označenie TS",ts.get("ts_code")),("Umiestnenie",f"{ts.get('location_address','')}, {ts.get('location_psc','')} {ts.get('location_city','')}".strip(", ")),
                   ("Výkon / napätie",f"{ts.get('rated_power_kva','—')} kVA · {ts.get('vn_voltage_kv','—')}/{(ts.get('nn_voltage_v') or 0)/1000 if ts.get('nn_voltage_v') else '—'} kV"),
                   ("Mesačný paušál",f"{('%.2f €' % float(pausal)) if pausal else '… €'}")])
    b=BytesIO(); doc.save(b); return b.getvalue()


# ============================================================
# 3) MPP — Miestny prevádzkový predpis
# ============================================================
def generate_mpp(ts: dict) -> bytes:
    td = ts.get("tech_details") or {}; prev = td.get("prevadzkovatel") or {}
    vn = td.get("vn_rozvadzac") or {}; tr = td.get("transformator") or {}; nn = td.get("nn_rozvadzac") or {}
    su = td.get("sustavy") or {}; pr = td.get("pripojenie") or {}; rt = td.get("revizny_technik") or {}
    doc = Document()
    _h(doc,"MIESTNY PREVÁDZKOVÝ PREDPIS",size=16,align=WD_ALIGN_PARAGRAPH.CENTER)
    _h(doc,f"pre transformačnú stanicu VN/NN — {ts.get('ts_code','TS')}",size=12,align=WD_ALIGN_PARAGRAPH.CENTER,color=None,bold=False)
    p=doc.add_paragraph(prev.get("nazov") or ""); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    _kv_table(doc,[("Vypracoval","......................................"),
                   ("Za prevádzkovateľa TS schválil", rt.get("meno","......................................")),
                   ("Za Západoslovenská distribučná, a.s. schválil","......................................"),
                   ("Dátum schválenia","......................................")])
    def sec(n,t): _h(doc,f"{n}. {t}",size=12,after=2)
    sec(1,"Všeobecné ustanovenia")
    doc.add_paragraph("MPP je platný od dátumu schválenia až do odvolania. Majiteľ TS je povinný zabezpečiť obsluhu a údržbu TS ustanovením prevádzkovateľa TS. Pracovníci poverení prevádzkovaním TS musia byť preukázateľne oboznámení s týmto MPP.")
    sec(2,"Všeobecný popis transformačnej stanice")
    doc.add_paragraph(f"TS {ts.get('vn_voltage_kv','22')}/{(ts.get('nn_voltage_v') or 400)/1000 if ts.get('nn_voltage_v') else '0,4'} kV ({ts.get('ts_code','')}) je {ts.get('ts_type','kioskového prevedenia')}. Slúži na napojenie prevádzky {prev.get('nazov','')}.")
    if pr: doc.add_paragraph(f"Pripojenie: {pr.get('ds','DS')}, VN vedenie {pr.get('vn_vedenie','—')}, prívod z {pr.get('privod_z_ts','—')}. {pr.get('majetkova_hranica','')}")
    sec(3,"Popis častí transformačnej stanice")
    _h(doc,"3.1 VN rozvádzač",size=11,color=None)
    doc.add_paragraph(f"Výrobca/typ: {vn.get('vyrobca','')} {vn.get('typ','')}; počet polí: {vn.get('pocet_poli','—')}. {vn.get('polia','')}. VN káble: {vn.get('vn_kable','—')}.")
    _h(doc,"3.2 Transformátor",size=11,color=None)
    doc.add_paragraph(f"Typ {tr.get('typ','')}, výrobca {tr.get('vyrobca','')}, výkon {tr.get('vykon_kva',ts.get('rated_power_kva','—'))} kVA, napätie {tr.get('napatie','')}, výr. č. {tr.get('vyrobne_cislo','')}, chladenie {tr.get('chladenie','')}, ochrana {tr.get('ochrana','')}.")
    _h(doc,"3.3 NN rozvádzač",size=11,color=None)
    doc.add_paragraph(f"Označenie {nn.get('oznacenie','')}, hlavný istič {nn.get('hlavny_istic','')}, menovitý prúd {nn.get('in_a','—')} A, počet vývodov {nn.get('vyvody','—')}.")
    if su: doc.add_paragraph(f"Sústavy: VN {su.get('vn','')}; NN {su.get('nn','')}.")
    sec(4,"Predpisy pre manipuláciu")
    doc.add_paragraph("Práce na elektrickej inštalácii môže vykonávať iba odborne spôsobilá osoba (Vyhl. MPSVaR č. 508/2009). Pred prácou dodržať pravidlá: vypni — zaisti proti zapnutiu — odskúšaj — uzemni/skratuj — oddeľ živé a neživé časti.")
    sec(5,"Zabezpečenie pracoviska")
    doc.add_paragraph("Na práce na VN zariadeniach musí byť vystavený „B — príkaz“. Pracovisko musí byť vypnuté a uzemnené zo všetkých strán, vyznačené výstražnými tabuľkami a zabezpečené proti neoprávnenej manipulácii.")
    sec(6,"Ochranné a pracovné pomôcky")
    doc.add_paragraph("Vybavenie TS musí zodpovedať STN 38 1981. Pomôcky odskúšať v predpísaných lehotách; raz ročne vizuálna prehliadka mechanického stavu.")
    sec(7,"Údržba transformačnej stanice")
    doc.add_paragraph("Odborná prehliadka a skúška TS sa vykonáva raz za 5 rokov (príloha č. 8 k Vyhl. 508/2009). Pravidelne kontrolovať prúdové zaťaženie, napätie, čistotu izolátorov, kontakty a svorky, tesnosť veka, vetranie.")
    sec(8,"Plán opatrení")
    doc.add_paragraph(f"Pri poruche/nezvyčajnom jave (sršanie, únik oleja) ihneď nahlásiť prevádzkovateľovi TS: {rt.get('tel','—')} a urobiť záznam do prevádzkového denníka. Pri požiari postupovať podľa požiarneho poriadku, pri veľkom rozsahu zabezpečiť odpojenie TS na prívode v koordinácii so ZSD.")
    sec(9,"Prvá pomoc pri zásahu elektrickým prúdom")
    doc.add_paragraph("Vyslobodiť postihnutého z dosahu prúdu (vypnúť hlavný vypínač), pri zástave dychu zahájiť umelé dýchanie, pri nehmatateľnom pulze kombinovať s masážou srdca, privolať RLP (155/112).")
    sec(10,"Zoznam dôležitých kontaktov")
    _kv_table(doc,[("Tiesňové volanie","112"),("Polícia","158"),("Hasiči","150"),("RLP","155"),
                   ("Prevádzkovateľ TS",f"{rt.get('meno','')} {rt.get('tel','')}"),("ZSD poruchová (nonstop)","0800 111 567")])
    sec(11,"Zoznam použitých skratiek")
    doc.add_paragraph("MPP — miestny prevádzkový predpis · NN — nízke napätie · VN — vysoké napätie · TS — transformačná stanica · DS — distribučná sústava · ZSD — Západoslovenská distribučná, a.s.")
    b=BytesIO(); doc.save(b); return b.getvalue()


# ============================================================
# 4) REVÍZNA SPRÁVA (OPaOS) — skelet, technik doplní merania
# ============================================================
def generate_revizna(ts: dict, inspection: dict = None) -> bytes:
    td = ts.get("tech_details") or {}; prev = td.get("prevadzkovatel") or {}
    su = td.get("sustavy") or {}; rt = td.get("revizny_technik") or {}; ins = inspection or {}
    doc = Document()
    _h(doc,"SPRÁVA O ODBORNEJ PREHLIADKE A ODBORNEJ SKÚŠKE ELEKTRICKEJ INŠTALÁCIE",size=13,align=WD_ALIGN_PARAGRAPH.CENTER)
    p=doc.add_paragraph("vykonaná v zmysle § 13 vyhlášky MPSVaR SR č. 508/2009 Z.z. a podľa STN 33 1500, STN 33 2000-6")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(9)
    doc.add_paragraph()
    _kv_table(doc,[
        ("Druh", ins.get("inspection_type","Opakovaná OPaOS")),
        ("Dátum vykonania", ins.get("performed_date","..............")),
        ("Revízny technik", rt.get("meno","..............")),
        ("Číslo osvedčenia", rt.get("cislo_osvedcenia","..............")),
        ("Číslo revíznej správy", ins.get("cislo_spravy","RS-...........")),
        ("Prevádzkovateľ / adresa", f"{prev.get('nazov','')} — {ts.get('location_address','')}, {ts.get('location_city','')}"),
        ("Názov zariadenia", ts.get("ts_code")),
        ("Sústava VN", su.get("vn","..............")),
        ("Sústava NN", su.get("nn","..............")),
    ])
    _h(doc,"Celkový posudok:",size=12)
    doc.add_paragraph("Revidované zariadenie v rozsahu uvedenom v tejto správe je / nie je schopné bezpečnej prevádzky ku dňu vykonania. (doplní revízny technik)")
    def kap(n,t): _h(doc,f"Kapitola {n} — {t}",size=12,after=2)
    kap(1,"Úvod")
    doc.add_paragraph(f"Podľa STN 33 1500 a STN 33 2000-6 bola vykonaná odborná prehliadka a odborná skúška TS {ts.get('ts_code','')}. {ts.get('ts_type','')}. (popis doplní technik)")
    kap(2,"Normatívne odkazy")
    doc.add_paragraph("STN 33 0360, STN 33 1500, STN 33 2000-1, STN 33 2000-4-41, STN 33 2000-5-51, STN 33 2000-6, STN 33 3201 a súvisiace.")
    kap(3,"Požiadavky na dokumentáciu")
    for b in ["Vyhlásenia o zhode rozvádzača","Protokol o kusovej skúške rozvádzača","Osvedčenie použitých káblov","Protokol o určení prostredia a vonkajších vplyvov","Projektová dokumentácia","Protokoly o skúške transformátora a VN rozvádzača"]:
        doc.add_paragraph(b, style="List Bullet")
    kap(4,"Overovanie — prehliadka a skúška (STN 33 2000-6)")
    for b in ["spôsob ochrany pred zásahom elektrickým prúdom","ochrana pred účinkami tepla / šírením požiaru","prúdová zaťažiteľnosť a úbytok napätia vodičov","nastavenie ochranných a monitorovacích prístrojov","označenie neutrálnych a ochranných vodičov","schémy a výstražné nápisy","pripojenie a pospájanie vodičov","prístupnosť pre údržbu"]:
        doc.add_paragraph(b+" — ......", style="List Bullet")
    kap(5,"Merania")
    _h(doc,"Izolačné odpory",size=11,color=None)
    t=doc.add_table(rows=1,cols=5); t.style="Table Grid"
    for i,h in enumerate(["Obvod","Druh vedenia","Izolačný odpor","Istenie [A]","Posudok"]):
        t.rows[0].cells[i].text=h; t.rows[0].cells[i].paragraphs[0].runs[0].bold=True
    for _ in range(4):
        r=t.add_row().cells
        for ci in range(5): r[ci].text="......"
    _h(doc,"Impedancia slučky / uzemnenie",size=11,color=None)
    doc.add_paragraph("Namerané hodnoty doplní revízny technik (prístroj, metóda, výsledok).")
    kap(6,"Záver")
    doc.add_paragraph("Najbližšiu pravidelnú OPaOS vykonať v termíne podľa Vyhl. 508/2009, príloha č. 8 (interval 5 rokov pre TS).")
    doc.add_paragraph()
    _h(doc,"VYHLÁSENIE O ZODPOVEDNOSTI OSOBY ZODPOVEDNEJ ZA ODBORNÚ PREHLIADKU",size=11)
    doc.add_paragraph(f"Ja, {rt.get('meno','..............')}, zodpovedný za revíziu elektrickej inštalácie, vyhlasujem, že zodpovednosť je obmedzená na vyššie opísanú činnosť.")
    doc.add_paragraph("Podpis: ............    Dátum: ............")
    b=BytesIO(); doc.save(b); return b.getvalue()
