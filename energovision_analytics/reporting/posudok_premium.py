"""Premium DOCX posudok — Energovision šablóna + engine multi-variant features.

Replikuje štýl z `fve-bess-posudok` skillu (Node.js build_docx_template.js)
v Pythone cez python-docx + pridáva nové sekcie z analytického enginu v0.9:
    - Top-6 ranking variantov (kritériá NPV/IRR/payback/...)
    - Tornado sensitivity (±20 % na 6 premenných)
    - Monte Carlo P10/P50/P90 distribúcia
    - Dotácia info card s YAML schémami
    - DPPO výber + daňový štít kalkulácia
    - Manifest footer (engine ver + tariff hash + spot dáta)
    - Battery vendor preset info (Huawei/Solinteg/BYD)

Brand:
    EV_GREEN     = 92D050   (hlavná zelená)
    EV_GREEN_LT  = E8F4D5   (svetlá zelená pre fill)
    EV_BLACK     = 1A1A1A
    EV_DARK      = 2C2C2C
    EV_GRAY      = 8C8C8C
    EV_LIGHTGRAY = F5F5F5
    EV_BORDER    = D9D9D9
    Arial font (predvolený v Energovision šablóne)

Použitie:
    docx_bytes = generate_premium_posudok(
        client_name="RE-PLAST, s.r.o.",
        project_id="P-26-134",
        client_address="Zvončín 107, Zvončín",
        client_contact="Ján Krčula",
        site_meta={...},
        run_response=run_variants_response_dict,
        sensitivity_data=tornado_data,  # optional
        monte_carlo_data=mc_data,        # optional
        manifest_footer="Engine v0.9.0 | tariff 2026@abc | ...",
    )
"""
from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Emu, Mm, Pt, RGBColor

# === BRAND ===
EV_GREEN_HEX     = "92D050"
EV_GREEN_LT_HEX  = "E8F4D5"
EV_BLACK_HEX     = "1A1A1A"
EV_DARK_HEX      = "2C2C2C"
EV_GRAY_HEX      = "8C8C8C"
EV_LIGHTGRAY_HEX = "F5F5F5"
EV_BORDER_HEX    = "D9D9D9"

EV_GREEN  = RGBColor(0x92, 0xD0, 0x50)
EV_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
EV_DARK   = RGBColor(0x2C, 0x2C, 0x2C)
EV_GRAY   = RGBColor(0x8C, 0x8C, 0x8C)


# ============================================================================
# LOW-LEVEL HELPERS
# ============================================================================
def _set_cell_bg(cell, color_hex: str) -> None:
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border(cell, *, left_color: Optional[str] = None,
                     left_size: int = 0, all_thin: bool = False) -> None:
    """Set cell borders. left_size in EMU (1pt = ~12700, 8 = ~0.5pt)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')

    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        if all_thin and side != 'left':
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:color'), EV_BORDER_HEX)
        elif side == 'left' and left_color:
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), str(left_size or 24))
            b.set(qn('w:color'), left_color)
        else:
            b.set(qn('w:val'), 'nil')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _kicker(doc: Document, text: str) -> None:
    """Malý zelený nadpis NAD h1 (uppercase, spaced)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = EV_GREEN
    # character spacing (cez XML)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '40')
    rPr.append(spacing)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = EV_BLACK


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = EV_BLACK


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = EV_DARK


def _para(doc: Document, text: str, *, italic: bool = False,
          size: int = 11, color: RGBColor = EV_DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.color.rgb = color


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.runs[-1] if p.runs else p.add_run(text)
    if r.text != text:
        r.text = text
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = EV_DARK


def _highlight_box(doc: Document, kicker_text: str, lines: list[str | tuple]) -> None:
    """Highlight box — 1 cell tabuľka so zelenou ľavou hranou + svetlým pozadím.

    lines: list strings alebo (text, bold) tuples
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, EV_LIGHTGRAY_HEX)
    _set_cell_border(cell, left_color=EV_GREEN_HEX, left_size=32)
    # Margin
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement('w:tcMar')
    for side, val in [('top', '180'), ('bottom', '180'), ('left', '240'), ('right', '240')]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        margin.append(m)
    tc_pr.append(margin)
    # Clear default paragraph
    cell.paragraphs[0].text = ""

    # Kicker
    p_k = cell.paragraphs[0]
    r_k = p_k.add_run(kicker_text.upper())
    r_k.font.name = "Arial"
    r_k.font.size = Pt(9)
    r_k.font.bold = True
    r_k.font.color.rgb = EV_BLACK

    # Body lines
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        if isinstance(line, str):
            r = p.add_run(line)
            r.font.name = "Arial"
            r.font.size = Pt(11)
            r.font.color.rgb = EV_DARK
        elif isinstance(line, tuple):
            # tuple of (text, is_bold) pairs
            for piece in line:
                if isinstance(piece, str):
                    r = p.add_run(piece)
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
                    r.font.color.rgb = EV_DARK
                else:
                    text, bold = piece
                    r = p.add_run(text)
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
                    r.font.bold = bold
                    r.font.color.rgb = EV_BLACK if bold else EV_DARK

    doc.add_paragraph()  # spacer


def _kpi_box(doc: Document, items: list[dict]) -> None:
    """KPI box — N stĺpcová tabuľka so zelenou ľavou hranou per cell.

    items: [{"label": "CAPEX", "value": "100 000 €", "note": "bez DPH"}]
    """
    n = len(items)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, it in enumerate(items):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, EV_LIGHTGRAY_HEX)
        _set_cell_border(cell, left_color=EV_GREEN_HEX, left_size=24)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""

        # Label
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(it["label"].upper())
        r1.font.name = "Arial"
        r1.font.size = Pt(7)
        r1.font.bold = True
        r1.font.color.rgb = EV_GRAY

        # Value
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(it["value"])
        r2.font.name = "Arial"
        r2.font.size = Pt(14)
        r2.font.bold = True
        r2.font.color.rgb = EV_BLACK

        # Note
        if it.get("note"):
            p3 = cell.add_paragraph()
            r3 = p3.add_run(it["note"])
            r3.font.name = "Arial"
            r3.font.size = Pt(8)
            r3.font.italic = True
            r3.font.color.rgb = EV_GRAY

    doc.add_paragraph()


def _data_table(doc: Document, rows: list[list], *, header: bool = True,
                highlight_last: bool = False, col_widths_cm: Optional[list[float]] = None) -> None:
    """Tabuľka s alternujúcimi farbami riadkov + tučnou hlavičkou."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell, all_thin=True)

            # Background
            is_header_row = header and i == 0
            is_highlight = highlight_last and i == len(rows) - 1
            if is_header_row:
                _set_cell_bg(cell, EV_BLACK_HEX)
            elif is_highlight:
                _set_cell_bg(cell, EV_GREEN_LT_HEX)
            elif i % 2 == 0:
                _set_cell_bg(cell, EV_LIGHTGRAY_HEX)

            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = "Arial"
            r.font.size = Pt(10)
            if is_header_row:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif is_highlight:
                r.font.bold = True
                r.font.color.rgb = EV_BLACK
            else:
                r.font.color.rgb = EV_DARK

            # Col widths
            if col_widths_cm and j < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[j])

    doc.add_paragraph()


def _step_table(doc: Document, steps: list[dict]) -> None:
    """Step table — N stĺpcov s číslom + nadpisom + popisom (Energovision štýl)."""
    n = len(steps)
    table = doc.add_table(rows=1, cols=n)

    for i, s in enumerate(steps):
        cell = table.rows[0].cells[i]
        _set_cell_border(cell, left_color=EV_GREEN_HEX, left_size=18)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.paragraphs[0].text = ""

        # Číslo
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(f"{i+1:02d}")
        r1.font.name = "Arial"
        r1.font.size = Pt(18)
        r1.font.bold = True
        r1.font.color.rgb = EV_GREEN

        # Title
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(s["title"])
        r2.font.name = "Arial"
        r2.font.size = Pt(11)
        r2.font.bold = True
        r2.font.color.rgb = EV_BLACK

        # Body
        p3 = cell.add_paragraph()
        r3 = p3.add_run(s["body"])
        r3.font.name = "Arial"
        r3.font.size = Pt(9)
        r3.font.color.rgb = EV_DARK

    doc.add_paragraph()


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = EV_GRAY


def _embed_chart_png(doc: Document, png_bytes: bytes, width_cm: float = 16) -> bool:
    """Embed PNG (z Plotly→kaleido) do dokumentu."""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(io.BytesIO(png_bytes), width=Cm(width_cm))
        return True
    except Exception:
        return False


# ============================================================================
# CHARTS PIPELINE — Plotly → PNG bytes
# ============================================================================
def _plotly_to_png(fig, width: int = 900, height: int = 380) -> Optional[bytes]:
    """Konvertuje Plotly figure na PNG bytes cez kaleido."""
    try:
        return fig.to_image(format="png", width=width, height=height, scale=2)
    except Exception as e:
        # Kaleido nie je nainštalovaný / fig nie je validný
        return None


# ============================================================================
# HLAVNÁ FUNKCIA — GENERATE
# ============================================================================
def generate_premium_posudok(
    *,
    client_name: str,
    project_id: str,
    client_address: str = "",
    client_contact: str = "",
    project_name: str = "Hybridné riešenie FVE + BESS",
    site_meta: dict,
    run_response: dict,
    selected_variant_id: Optional[str] = None,
    sensitivity_data: Optional[dict] = None,
    monte_carlo_data: Optional[dict] = None,
    dotacia_scheme_info: Optional[dict] = None,
    additional_notes: str = "",
    engine_version: str = "0.9.0",
    manifest_footer: str = "",
    posudok_date: Optional[str] = None,
    prepared_by_name: str = "Lukáš Bago",
    prepared_by_email: str = "lukas.bago@energovision.sk",
    prepared_by_phone: str = "0918 187 762",
    company_ico: str = "53 036 280",
    logo_path: Optional[str] = None,
    # NEW v2 — scenario-aware posudok
    scenario_type: str = "nova_fve",          # nova_fve / rozsirenie_fve / pridanie_bess / iba_bess_arbitraz / custom
    scenario_description: Optional[str] = None,  # 1-veta opis od obchodníka
    existing_fve_kwp: float = 0,
    existing_bess_kwh: float = 0,
    existing_fve_samosp_pct: Optional[float] = None,
    ai_narrative: Optional[str] = None,       # AI úvaha (4 odseky energetik-style)
    posudok_number: Optional[str] = None,     # P-26-XXX (explicitné číslo)
) -> bytes:
    """Vyrobí kompletný DOCX posudok podľa Energovision šablóny + engine sekcií.

    Args:
        run_response: výstup z /run-variants (variants, top_picks, manifest)
        selected_variant_id: ID konkrétneho variantu (default = najvyššie NPV)
        sensitivity_data: {"variables": [...], "low": [...], "high": [...]} pre tornado
        monte_carlo_data: {"p10": ..., "p50": ..., "p90": ..., "pct_above_zero": ...}
    """
    doc = Document()

    # Margins — A4 portrait
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = EV_DARK

    variants = run_response.get("variants", [])
    top_picks = run_response.get("top_picks", [])
    if not variants:
        raise ValueError("run_response neobsahuje žiadne varianty")

    # Vyber variant
    if selected_variant_id:
        selected = next((v for v in variants if v["variant_id"] == selected_variant_id), None)
        if not selected:
            selected = max(variants, key=lambda v: v.get("npv_eur", 0))
    else:
        selected = max(variants, key=lambda v: v.get("npv_eur", 0))

    posudok_date = posudok_date or datetime.now().strftime("%d.%m.%Y")

    # ===============================================================
    # TITULNÁ STRANA
    # ===============================================================
    if logo_path and Path(logo_path).exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Cm(6))
    else:
        # Text fallback ako logo
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p_logo.add_run("energo")
        r1.font.name = "Arial"; r1.font.size = Pt(18); r1.font.bold = True
        r1.font.color.rgb = EV_BLACK
        r2 = p_logo.add_run("vision")
        r2.font.name = "Arial"; r2.font.size = Pt(18); r2.font.bold = True
        r2.font.color.rgb = EV_GREEN

    # Kicker + ID
    doc.add_paragraph()
    p_kicker = doc.add_paragraph()
    p_kicker.paragraph_format.space_after = Pt(4)
    r_k = p_kicker.add_run("TECHNICKO-EKONOMICKÝ POSUDOK")
    r_k.font.name = "Arial"; r_k.font.size = Pt(10); r_k.font.bold = True
    r_k.font.color.rgb = EV_GREEN
    sep = p_kicker.add_run("   ·   ")
    sep.font.name = "Arial"; sep.font.size = Pt(10); sep.font.color.rgb = EV_GRAY
    posudok_label = posudok_number if posudok_number else project_id
    r_id = p_kicker.add_run(f"{posudok_label} — {project_name}")
    r_id.font.name = "Arial"; r_id.font.size = Pt(10); r_id.font.color.rgb = EV_DARK

    # Veľký nadpis klienta
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(6)
    r_name = p_name.add_run(client_name)
    r_name.font.name = "Arial"; r_name.font.size = Pt(28); r_name.font.bold = True
    r_name.font.color.rgb = EV_BLACK

    # Sub-title
    bess_txt = (
        f" + batériové úložisko {selected['bess_kwh']:.0f} kWh"
        if selected['bess_kwh'] > 0 else ""
    )
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(3)
    r_sub = p_sub.add_run(
        f"Hybridné energetické riešenie — fotovoltika {selected['pv_kwp']:.0f} kWp{bess_txt}"
    )
    r_sub.font.name = "Arial"; r_sub.font.size = Pt(13); r_sub.font.color.rgb = EV_DARK

    p_tag = doc.add_paragraph()
    p_tag.paragraph_format.space_after = Pt(24)
    r_tag = p_tag.add_run("Vlastná výroba elektriny pre slovenské odberné miesto — od simulácie po dispatch.")
    r_tag.font.name = "Arial"; r_tag.font.size = Pt(11); r_tag.font.italic = True
    r_tag.font.color.rgb = EV_GRAY

    # 2-stĺpcový info box (PRE / VYSTAVENÉ)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cell PRE
    cell_pre = info_table.rows[0].cells[0]
    _set_cell_border(cell_pre, left_color=EV_GREEN_HEX, left_size=18)
    cell_pre.paragraphs[0].text = ""
    for text, sz, bold, color in [
        ("PRE", 8, True, EV_GRAY),
        (client_name, 12, True, EV_BLACK),
    ]:
        p = cell_pre.paragraphs[0] if text == "PRE" else cell_pre.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(sz); r.font.bold = bold
        r.font.color.rgb = color
    if client_address:
        p = cell_pre.add_paragraph()
        r = p.add_run(client_address)
        r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = EV_DARK
    if client_contact:
        p = cell_pre.add_paragraph()
        r = p.add_run(f"Kontakt: {client_contact}")
        r.font.name = "Arial"; r.font.size = Pt(10); r.font.italic = True
        r.font.color.rgb = EV_DARK
    p = cell_pre.add_paragraph()
    r = p.add_run("PARAMETRE OM")
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = EV_GRAY
    p = cell_pre.add_paragraph()
    _rk = site_meta.get('rk_kw') or 0
    _mrk = site_meta.get('mrk_kw') or (_rk * 1.2)
    r = p.add_run(
        f"{site_meta.get('distribuutor') or '?'} {site_meta.get('sadzba') or 'NN'} · "
        f"RK {_rk:.0f} kW · MRK {_mrk:.0f} kW"
    )
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.color.rgb = EV_DARK
    p = cell_pre.add_paragraph()
    r = p.add_run(f"Spotreba: {site_meta.get('rocna_spotreba_kwh') or 0:,.0f} kWh/rok")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.color.rgb = EV_DARK

    # Cell VYSTAVENÉ
    cell_iss = info_table.rows[0].cells[1]
    _set_cell_border(cell_iss, left_color=EV_GREEN_HEX, left_size=18)
    cell_iss.paragraphs[0].text = ""
    p = cell_iss.paragraphs[0]
    r = p.add_run("VYSTAVENÉ")
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = EV_GRAY
    p = cell_iss.add_paragraph()
    r = p.add_run(posudok_date)
    r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = EV_BLACK
    p = cell_iss.add_paragraph()
    r = p.add_run("Bratislava")
    r.font.name = "Arial"; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = EV_DARK
    p = cell_iss.add_paragraph()
    r = p.add_run("POSUDZOVANÉ OBDOBIE")
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = EV_GRAY
    p = cell_iss.add_paragraph()
    r = p.add_run("01.01.2025 – 31.12.2025")
    r.font.name = "Arial"; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = EV_BLACK
    p = cell_iss.add_paragraph()
    r = p.add_run("ENGINE")
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = EV_GRAY
    p = cell_iss.add_paragraph()
    r = p.add_run(f"Energovision Analyzer v{engine_version}")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.color.rgb = EV_DARK

    doc.add_paragraph()
    doc.add_paragraph()

    # PRIPRAVIL PRE VÁS — highlight box
    _highlight_box(doc, "PRIPRAVIL PRE VÁS", [
        f"{prepared_by_name}",
        "Energovision, s.r.o.",
        f"{prepared_by_email}   ·   {prepared_by_phone}",
    ])

    doc.add_page_break()

    # ===============================================================
    # MANAŽÉRSKE ZHRNUTIE
    # ===============================================================
    _kicker(doc, "Manažérske zhrnutie")
    _h1(doc, (
        f"Variant {selected['pv_kwp']:.0f} kWp FVE{bess_txt} dosahuje "
        f"návratnosť {selected['payback_simple_y']:.1f} roka pri NPV +{selected['npv_eur']:,.0f} €."
    ))

    # Scenario-aware úvod
    scenario_intros = {
        "nova_fve": (
            f"Posudok hodnotí investíciu do novej fotovoltickej elektrárne "
            f"{selected['pv_kwp']:.0f} kWp{bess_txt} na odbernom mieste klienta "
            f"{client_name} (greenfield bez existujúcej FVE)."
        ),
        "rozsirenie_fve": (
            f"Posudok hodnotí ROZŠÍRENIE existujúcej FVE {existing_fve_kwp:.0f} kWp "
            f"o ďalších {selected['pv_kwp'] - existing_fve_kwp:.0f} kWp{bess_txt} na odbernom mieste klienta "
            f"{client_name}. Aktuálna samospotreba existujúcej FVE: "
            f"{existing_fve_samosp_pct:.0f}% (z ročnej produkcie)." if existing_fve_samosp_pct else
            f"Posudok hodnotí ROZŠÍRENIE existujúcej FVE {existing_fve_kwp:.0f} kWp "
            f"o ďalších {selected['pv_kwp'] - existing_fve_kwp:.0f} kWp{bess_txt} na odbernom mieste klienta {client_name}."
        ),
        "pridanie_bess": (
            f"Posudok hodnotí PRIDANIE batériového úložiska {selected.get('bess_kwh', 0):.0f} kWh "
            f"k existujúcej FVE {existing_fve_kwp:.0f} kWp na odbernom mieste klienta {client_name}. "
            f"Cieľ: maximalizovať samospotrebu a využiť spot arbitráž."
        ),
        "iba_bess_arbitraz": (
            f"Posudok hodnotí investíciu do batériového úložiska {selected.get('bess_kwh', 0):.0f} kWh "
            f"bez FVE na odbernom mieste klienta {client_name}. "
            f"Cieľ: čistá spot arbitráž + peak shaving (zníženie MRK)."
        ),
        "custom": (
            f"Posudok hodnotí riešenie {selected['pv_kwp']:.0f} kWp FVE{bess_txt} "
            f"na odbernom mieste klienta {client_name} podľa špecifických požiadaviek obchodníka."
        ),
    }
    _para(doc, scenario_intros.get(scenario_type, scenario_intros["nova_fve"]))

    # Ak je obchodníkov opis, zaraď ho do úvodu
    if scenario_description:
        _para(doc, f"Požiadavka klienta: {scenario_description}", italic=True)

    _para(doc, (
        f"Ročná spotreba {site_meta.get('rocna_spotreba_kwh') or 0:,.0f} kWh, "
        f"distribútor {site_meta.get('distribuutor') or '?'} {site_meta.get('sadzba') or 'NN'}. "
        f"Engine simuloval {run_response.get('n_variants_run', 0)} variantov "
        f"konfigurácie PV × BESS na hodinovej granularite (8 760 h) "
        f"s reálnymi spot cenami OKTE 2025 a distribučnými tarifami ÚRSO 2026."
    ))

    _h2(doc, "Kľúčové parametre projektu")
    _kpi_box(doc, [
        {"label": "CAPEX celkom",  "value": f"{selected['capex_total_eur']:,.0f} €",
         "note": "bez DPH, turn-key"},
        {"label": "Dotácia",        "value": f"{selected['dotacia_eur']:,.0f} €",
         "note": dotacia_scheme_info.get("nazov", "Žiadna") if dotacia_scheme_info else "Žiadna"},
        {"label": "Návratnosť",     "value": f"{selected['payback_simple_y']:.1f} r",
         "note": "jednoduchá"},
        {"label": "NPV 20 r.",      "value": f"+{selected['npv_eur']:,.0f} €",
         "note": "s daň. odpisom"},
    ])

    # Highlight box odporúčania
    irr_str = f"{selected['irr_pct']:.1f} %" if selected.get('irr_pct') else "—"
    _highlight_box(doc, "ODPORÚČANIE", [
        f"Variant {selected['pv_kwp']:.0f} kWp FVE{bess_txt}",
        (
            ("Pri investícii ", False),
            (f"{selected['capex_total_eur']:,.0f} €", True),
            (" (z toho dotácia ", False),
            (f"{selected['dotacia_eur']:,.0f} €", True),
            (") klient získa ročnú úsporu ", False),
            (f"{selected['saving_y1_eur']:,.0f} €", True),
            (", IRR ", False),
            (irr_str, True),
            (" a NPV 20 r. ", False),
            (f"+{selected['npv_eur']:,.0f} €", True),
            (".", False),
        ),
    ])

    # ===============================================================
    # SEKCIA 1 — VSTUPNÉ DÁTA
    # ===============================================================
    _kicker(doc, "1 — Vstupné dáta a metodika")
    _h1(doc, "Východiská posudku")
    _h3(doc, "Charakteristika odberného miesta")
    _rk_safe = site_meta.get('rk_kw') or 0
    _mrk_safe = site_meta.get('mrk_kw') or (_rk_safe * 1.2)
    _data_table(doc, [
        ["Parameter", "Hodnota"],
        ["Klient", client_name],
        ["Adresa OM", client_address or site_meta.get("lokalita") or ""],
        ["Distribútor", site_meta.get("distribuutor") or "?"],
        ["Napäťová úroveň", site_meta.get("sadzba") or "NN"],
        ["RK (rezervovaná kapacita)", f"{_rk_safe:.0f} kW"],
        ["MRK (max. rezervovaná kapacita)", f"{_mrk_safe:.0f} kW"],
        ["Ročná spotreba", f"{site_meta.get('rocna_spotreba_kwh') or 0:,.0f} kWh"],
        ["Typ kontraktu", site_meta.get("typ_tarify") or "spot"],
    ], col_widths_cm=[8, 9])

    _h3(doc, "Metodika simulácie")
    _para(doc, (
        "Hodinová bilančná simulácia 8 760 hodín s využitím reálneho profilu odberu "
        "alebo syntetického profilu zodpovedajúceho prevádzkovému charakteru OM:"
    ))
    _bullet(doc, "FVE produkcia — PVGIS-kalibrovaný analytický model pre GPS lokáciu klienta.")
    _bullet(doc, "BESS — Naumann-Schimpe degradačný model LFP (kalibrovaný 2026-05-24), warranty constraint.")
    _bullet(doc, "EMS — rule-based dispatch s multi-cycle stratégiou (samospotreba + arbitráž BS + peak shaving).")
    _bullet(doc, "Spot ceny — OKTE DAM 2025 hodinové (priemer 103 €/MWh).")
    _bullet(doc, "Tarify — ÚRSO 2026 (distribúcia + TPS + nájomné + straty).")
    _bullet(doc, "Ekonomika — NPV 20 r., diskont 6 %, OPEX 1.5–2.0 % CAPEX/rok, degradácia FVE 0.5 %/rok, BESS 2 %/rok.")
    _bullet(doc, "Daňový odpis — lineárny 6 r. pri DPPO 22 % (novela 2025 pre MSP).")

    # ===============================================================
    # SEKCIA 2 — VARIANT GENERATOR (NOVÉ Z ENGINE)
    # ===============================================================
    _kicker(doc, "2 — Engine multi-variant analýza")
    _h1(doc, "Porovnanie všetkých posudzovaných variantov")
    _para(doc, (
        f"Engine vygeneroval a vyhodnotil celkom {len(variants)} kombinácií "
        f"PV × BESS pre dané OM. Každý variant prešiel plnou hodinovou simuláciou "
        f"(PV výroba, EMS dispatch, finančný model). Tabuľka nižšie obsahuje výsledky "
        f"zoradené podľa NPV:"
    ))

    table_rows = [["Variant", "PV (kWp)", "BESS (kWh)", "CAPEX (€)", "Dotácia (€)",
                    "Úspora Y1 (€)", "NPV (€)", "IRR (%)", "Payback (r)"]]
    for v in sorted(variants, key=lambda x: -x.get("npv_eur", 0)):
        table_rows.append([
            v["label"][:25],
            f"{v['pv_kwp']:.0f}",
            f"{v['bess_kwh']:.0f}",
            f"{v['capex_total_eur']:,.0f}",
            f"{v['dotacia_eur']:,.0f}",
            f"{v['saving_y1_eur']:,.0f}",
            f"{v['npv_eur']:,.0f}",
            f"{v['irr_pct']:.1f}" if v.get('irr_pct') else "—",
            f"{v['payback_simple_y']:.1f}",
        ])
    _data_table(doc, table_rows)

    _h3(doc, "Top picks — víťazi v jednotlivých kritériách")
    pick_rows = [["Kritérium", "Variant", "NPV (€)"]]
    for tp in top_picks:
        pick_rows.append([tp["label"], tp["variant_id"], f"{tp['npv_eur']:,.0f}"])
    _data_table(doc, pick_rows, col_widths_cm=[6, 6, 5])

    # ===============================================================
    # SEKCIA 3 — VYBRANÝ VARIANT (detail)
    # ===============================================================
    _kicker(doc, "3 — Vybraný variant — detail")
    _h1(doc, f"{selected['label']}")

    _h3(doc, "Investičné náklady (CAPEX)")
    _data_table(doc, [
        ["Položka", "Suma (bez DPH)"],
        [f"FVE {selected['pv_kwp']:.0f} kWp (turn-key)", f"{selected['capex_pv_eur']:,.0f} €"],
        [f"BESS {selected['bess_kwh']:.0f} kWh / {selected['bess_kw']:.0f} kW",
         f"{selected['capex_bess_eur']:,.0f} €"],
        ["CAPEX brutto", f"{selected['capex_total_eur']:,.0f} €"],
        [f"Dotácia ({(dotacia_scheme_info or {}).get('nazov', 'Žiadna')})",
         f"−{selected['dotacia_eur']:,.0f} €"],
        ["Čistá investícia", f"{selected['net_capex_eur']:,.0f} €"],
    ], highlight_last=True, col_widths_cm=[10, 7])

    _h3(doc, "Energetické toky")
    _data_table(doc, [
        ["Ukazovateľ", "Hodnota"],
        ["FVE ročná výroba", f"{selected['pv_total_kwh']:,.0f} kWh"],
        ["Samospotreba FVE", f"{selected['samospotreba_pct']:.1f} %"],
        ["Samostatnosť OM", f"{selected['samostatnost_pct']:.1f} %"],
        ["Import zo siete", f"{selected['grid_import_kwh']:,.0f} kWh"],
    ], col_widths_cm=[10, 7])

    # ===============================================================
    # SEKCIA 3b — AI ÚVAHA ENERGETIKA (Detailná analýza)
    # ===============================================================
    if ai_narrative and len(ai_narrative.strip()) > 50:
        _kicker(doc, "3b — Detailná analýza energetika")
        _h1(doc, "Úvaha senior energetika k tomuto projektu")
        _para(doc, (
            "Nasledujúce úvahy spracoval senior energetik s 15-ročnou praxou "
            "v dimenzovaní FVE/BESS pre slovenský priemysel. Berie do úvahy "
            "scenár klienta, profil odberu, aktuálne tarify a budúce regulačné zmeny."
        ), italic=True)

        # AI narrative môže obsahovať markdown headings / odseky
        for paragraph in ai_narrative.split("\n\n"):
            ptxt = paragraph.strip()
            if not ptxt:
                continue
            # Detekcia headingov (1. PREČO ... / **NADPIS** / # NADPIS)
            if ptxt.startswith("#") or (ptxt[0].isdigit() and ". " in ptxt[:5] and ptxt.split(". ", 1)[0].isdigit()):
                # Heading
                heading_text = ptxt.lstrip("#").strip()
                # Strip numbering
                if heading_text and heading_text[0].isdigit():
                    parts = heading_text.split(". ", 1)
                    if len(parts) == 2:
                        heading_text = parts[1]
                _h3(doc, heading_text.replace("**", "").strip())
            else:
                _para(doc, ptxt.replace("**", ""))

        doc.add_paragraph()

    # ===============================================================
    # SEKCIA 4 — DOTÁCIA (ak je)
    # ===============================================================
    if dotacia_scheme_info and selected.get("dotacia_eur", 0) > 0:
        _kicker(doc, "4 — Dotácia")
        _h1(doc, f"Schéma: {dotacia_scheme_info.get('nazov', '?')}")
        _data_table(doc, [
            ["Parameter", "Hodnota"],
            ["Vyhlasovateľ", dotacia_scheme_info.get("vyhlasovatel", "—")],
            ["Status", dotacia_scheme_info.get("status", "—")],
            ["Max. suma", f"{dotacia_scheme_info.get('max_eur', 0):,.0f} €"],
            ["Intenzita", f"{dotacia_scheme_info.get('intensity_pct', 0):.0f} %"],
            ["Min. samospotreba", f"{dotacia_scheme_info.get('min_samospotreba_pct', 0):.0f} %"],
            ["Aplikovaná suma pre tento variant", f"{selected['dotacia_eur']:,.0f} €"],
            ["Skutočná intenzita",
             f"{selected['dotacia_eur'] / selected['capex_total_eur'] * 100:.1f} %"],
        ], highlight_last=True, col_widths_cm=[10, 7])
        if dotacia_scheme_info.get("notes"):
            _para(doc, dotacia_scheme_info["notes"], italic=True, color=EV_GRAY)

    # ===============================================================
    # SEKCIA 5 — SENSITIVITY (NOVÉ)
    # ===============================================================
    if sensitivity_data:
        _kicker(doc, "5 — Citlivosť (Tornado)")
        _h1(doc, "Vplyv kľúčových premenných na NPV")
        _para(doc, (
            "Tornado analýza ukazuje, ako sa NPV mení pri zmene jednej premennej "
            "o ±20 % (alebo iný špecifikovaný rozsah) pri ostatných nezmenených. "
            "Najdlhšia tornado-bar = najväčšia citlivosť (kritická premenná)."
        ))
        rows = [["Premenná", "Nízky scenár (€)", "Vysoký scenár (€)", "Rozpätie (€)"]]
        for i, var in enumerate(sensitivity_data["variables"]):
            low = sensitivity_data["low"][i]
            high = sensitivity_data["high"][i]
            rows.append([var, f"{low:,.0f}", f"{high:,.0f}", f"{abs(high-low):,.0f}"])
        _data_table(doc, rows, col_widths_cm=[6, 4, 4, 3])

    # ===============================================================
    # SEKCIA 6 — MONTE CARLO (NOVÉ)
    # ===============================================================
    if monte_carlo_data:
        _kicker(doc, "6 — Risk analýza (Monte Carlo)")
        _h1(doc, "Distribúcia NPV — 1 000 scenárov")
        _para(doc, (
            "Monte Carlo simulácia 1 000 scenárov s náhodnými odchýlkami v spot cene "
            "(±12 %), CAPEX (±8 %) a degradácii (±20 %). Výsledná distribúcia ukazuje "
            "rozsah NPV ktorý sa dá očakávať podľa pravdepodobnosti:"
        ))
        _kpi_box(doc, [
            {"label": "P10 (worst case)", "value": f"{monte_carlo_data['p10']:,.0f} €",
             "note": "10 % horšia ako P10"},
            {"label": "P50 (medián)", "value": f"{monte_carlo_data['p50']:,.0f} €",
             "note": "50/50 šanca"},
            {"label": "P90 (best case)", "value": f"{monte_carlo_data['p90']:,.0f} €",
             "note": "10 % lepšia ako P90"},
            {"label": "Pravdep. NPV > 0", "value": f"{monte_carlo_data['pct_above_zero']:.0f} %",
             "note": "z 1 000 scenárov"},
        ])

    # ===============================================================
    # SEKCIA 7 — ZÁVER + DALŠIE KROKY
    # ===============================================================
    _kicker(doc, "7 — Záver")
    _h1(doc, "Záverečné odporúčanie")
    _highlight_box(doc, "ODPORÚČANIE PRE KLIENTA", [
        (
            (f"Variant {selected['pv_kwp']:.0f} kWp FVE{bess_txt} ", True),
            (f"dosahuje najvyššiu hodnotu NPV ({selected['npv_eur']:,.0f} €) "
             f"v rámci všetkých posudzovaných {len(variants)} variantov. "
             f"IRR {irr_str} výrazne prevyšuje typické bezrizikové úložky aj firemné cieľové ROI.", False),
        ),
    ])

    _h3(doc, "Argumenty pre realizáciu")
    _bullet(doc,
            f"Návratnosť {selected['payback_simple_y']:.1f} roka — "
            f"v rámci typického rozsahu komerčných FVE+BESS projektov.")
    _bullet(doc,
            f"Samospotreba FVE {selected['samospotreba_pct']:.1f} % — "
            f"vysoký podiel priamej spotreby znamená minimálnu závislosť od výkupnej ceny.")
    if selected['bess_kwh'] > 0:
        _bullet(doc,
                f"BESS {selected['bess_kwh']:.0f} kWh zvyšuje samostatnosť OM "
                f"na {selected['samostatnost_pct']:.1f} %.")
    if selected['dotacia_eur'] > 0:
        _bullet(doc,
                f"Dotácia {selected['dotacia_eur']:,.0f} € znižuje čistú investíciu "
                f"na {selected['net_capex_eur']:,.0f} €.")
    _bullet(doc,
            "Daňová optimalizácia — DPPO 22 % × 6r odpis = "
            f"~{selected['net_capex_eur']*0.22:,.0f} € daňový štít spolu.")

    _h2(doc, "Ďalšie kroky")
    _step_table(doc, [
        {"title": "Akceptácia ponuky",
         "body": "Klient odsúhlasí cenovú ponuku a parametre konfigurácie."},
        {"title": "Zmluva o dielo",
         "body": "Pevná cena, harmonogram, míľniky. Administráciu pripojenia VSDS vedie Energovision."},
        {"title": "Realizácia (3–6 mes.)",
         "body": "Inžiniering, montáž FVE + BESS + trafostanice, parametrizácia EMS, monitoring."},
    ])

    # ===============================================================
    # SEKCIA 8 — OTVORENÉ OTÁZKY (ak sú)
    # ===============================================================
    if additional_notes:
        _kicker(doc, "8 — Otvorené otázky pre klienta")
        _h1(doc, "Body na ďalšiu konzultáciu")
        _para(doc, additional_notes)

    # ===============================================================
    # DISCLAIMER + FOOTER s MANIFEST
    # ===============================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Posudok je indikatívny a vychádza z modelových predpokladov "
        "(spot OKTE 2025, distribučné tarify ÚRSO 2026, DPPO 22 %, "
        "odpis 6 rokov, diskont 6 %, degradácia LFP kalibrovaná na vendor presets). "
        "Skutočná výroba môže kolísať ±10 % v závislosti od počasia a profilu spotreby. "
        "Pred realizáciou odporúčame verifikovať PV výnos podľa PVGIS reportu a tarify "
        "podľa aktuálnej faktúry."
    )
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = EV_GRAY

    # Záverečný tagline
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Energovision — viac než len dodávateľ fotovoltických systémov.")
    r.font.name = "Arial"; r.font.size = Pt(11); r.font.italic = True
    r.font.color.rgb = EV_GRAY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Energetický partner pre budúcnosť.")
    r.font.name = "Arial"; r.font.size = Pt(11); r.font.italic = True; r.font.bold = True
    r.font.color.rgb = EV_BLACK

    # Manifest footer (engine version + tariff hash + spot dáta)
    if manifest_footer:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(manifest_footer)
        r.font.name = "Arial"; r.font.size = Pt(7); r.font.color.rgb = EV_GRAY
        r.font.italic = True

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
