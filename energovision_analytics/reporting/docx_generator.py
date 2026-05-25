"""DOCX generator pre Energovision posudky.

Šablóny:
    one_pager — A4, 1-2 strany, len top variant + KPI + 1 graf
    full_posudok — kompletný posudok podľa Energovision šablóny (viď memory/sablona_posudok)

Brand:
    Zelená Energovision: #7AB835
    Tmavá: #1A1A1A
    Font: Calibri (default DOCX) / Arial fallback
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

EV_GREEN = RGBColor(0x7A, 0xB8, 0x35)
EV_DARK = RGBColor(0x1A, 0x1A, 0x1A)
EV_GREY = RGBColor(0x66, 0x66, 0x66)


def _set_cell_bg(cell, color_hex: str) -> None:
    """Set cell background color."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_heading(doc: Document, text: str, level: int = 1, color: RGBColor = EV_DARK) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color


def _kpi_row(table, label: str, value: str, highlight: bool = False) -> None:
    row = table.add_row()
    cell_l, cell_v = row.cells[0], row.cells[1]
    cell_l.text = label
    cell_v.text = value
    if highlight:
        for c in row.cells:
            _set_cell_bg(c, "F4F9EC")  # light green
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True


def generate_one_pager(
    *,
    client_name: str,
    project_name: str,
    variant: dict,
    site_meta: dict,
    dotacia_scheme: Optional[str] = None,
    additional_notes: str = "",
    engine_version: str = "0.8.0",
    manifest_footer: str = "",
) -> bytes:
    """Vyrobí 1-stranový posudok ako DOCX bytes.

    Args:
        variant: dict z RunVariantsResponse.variants[i]
        site_meta: {nazov, psc, distribuutor, lokalita, rocna_spotreba_kwh, rk_kw}
    """
    doc = Document()

    # Margins
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)

    # === HLAVIČKA ===
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = head.add_run("ENERGOVISION")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = EV_GREEN

    sub = doc.add_paragraph()
    sub_r = sub.add_run("Posudok FVE + BESS — odporúčaný variant")
    sub_r.font.size = Pt(11)
    sub_r.font.color.rgb = EV_GREY

    doc.add_paragraph()  # spacer

    # === KLIENT ===
    _add_heading(doc, client_name, level=1)
    p = doc.add_paragraph()
    p.add_run(f"{project_name} · ").bold = True
    p.add_run(
        f"{site_meta.get('lokalita', '')} · "
        f"PSČ {site_meta.get('psc', '')} · "
        f"distribútor {site_meta.get('distribuutor', '')}"
    )
    p2 = doc.add_paragraph()
    p2.add_run(
        f"Ročná spotreba: {site_meta.get('rocna_spotreba_kwh', 0):,.0f} kWh · "
        f"RK: {site_meta.get('rk_kw', 0):.0f} kW"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # === ODPORÚČANÝ VARIANT ===
    bess_txt = (
        f" + BESS {variant['bess_kwh']:.0f} kWh / {variant['bess_kw']:.0f} kW"
        if variant['bess_kwh'] > 0 else " (bez BESS)"
    )
    _add_heading(doc, f"Variant: {variant['pv_kwp']:.0f} kWp FVE{bess_txt}",
                 level=2, color=EV_GREEN)

    # === KPI TABUĽKA ===
    kpi_table = doc.add_table(rows=0, cols=2)
    kpi_table.style = "Light Grid Accent 1"
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _kpi_row(kpi_table, "NPV (20 r, 6 %)",
             f"{variant['npv_eur']:,.0f} €", highlight=True)
    _kpi_row(kpi_table, "IRR",
             f"{variant['irr_pct']:.1f} %" if variant.get('irr_pct') else "—",
             highlight=True)
    _kpi_row(kpi_table, "Návratnosť (jednoduchá)",
             f"{variant['payback_simple_y']:.1f} rokov", highlight=True)
    _kpi_row(kpi_table, "Samospotreba FVE",
             f"{variant['samospotreba_pct']:.1f} %")
    _kpi_row(kpi_table, "Samostatnosť",
             f"{variant['samostatnost_pct']:.1f} %")
    _kpi_row(kpi_table, "Ročná úspora (Y1)",
             f"{variant['saving_y1_eur']:,.0f} €")

    doc.add_paragraph()

    # === INVESTÍCIA ===
    _add_heading(doc, "Investícia", level=3)
    inv_table = doc.add_table(rows=0, cols=2)
    inv_table.style = "Light Grid Accent 1"
    _kpi_row(inv_table, "CAPEX FVE",      f"{variant['capex_pv_eur']:,.0f} €")
    _kpi_row(inv_table, "CAPEX BESS",     f"{variant['capex_bess_eur']:,.0f} €")
    _kpi_row(inv_table, "CAPEX spolu",    f"{variant['capex_total_eur']:,.0f} €")
    if variant.get('dotacia_eur', 0) > 0:
        _kpi_row(inv_table, f"Dotácia ({dotacia_scheme or 'aplikovaná'})",
                 f"−{variant['dotacia_eur']:,.0f} €")
    _kpi_row(inv_table, "Čistá investícia",
             f"{variant['net_capex_eur']:,.0f} €", highlight=True)

    doc.add_paragraph()

    # === ENERGETICKÉ TOKY (textová sumár — full bude mať grafy) ===
    _add_heading(doc, "Energetické toky", level=3)
    energy_p = doc.add_paragraph()
    energy_p.add_run(
        f"FVE vyrobí ročne približne "
    )
    energy_p.add_run(f"{variant['pv_total_kwh']:,.0f} kWh").bold = True
    energy_p.add_run(", z čoho ")
    energy_p.add_run(f"{variant['samospotreba_pct']:.0f} %").bold = True
    energy_p.add_run(
        f" sa využije na samospotrebu. Zo siete sa odoberie ďalších "
    )
    energy_p.add_run(f"{variant['grid_import_kwh']:,.0f} kWh").bold = True
    energy_p.add_run(".")

    # === POZNÁMKY ===
    if additional_notes:
        doc.add_paragraph()
        _add_heading(doc, "Poznámky", level=3)
        doc.add_paragraph(additional_notes)

    # === DISCLAIMER ===
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc_r = disc.add_run(
        "Posudok je indikatívny a vychádza z modelových predpokladov "
        "(spot OKTE 2025, distribučné tarify ÚRSO 2026, DPPO 22 %, "
        "odpis 6 rokov, diskont 6 %). Skutočná výroba môže kolísať ±10 % "
        "v závislosti od počasia a profilu spotreby."
    )
    disc_r.font.size = Pt(8)
    disc_r.font.color.rgb = EV_GREY
    disc_r.italic = True

    # === FOOTER ===
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = foot.add_run(manifest_footer or f"Energovision Analyzer v{engine_version}")
    fr.font.size = Pt(7)
    fr.font.color.rgb = EV_GREY

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_full_posudok(
    *,
    client_name: str,
    project_name: str,
    variant: dict,
    site_meta: dict,
    all_variants: list[dict],
    top_picks: list[dict],
    dotacia_scheme: Optional[str] = None,
    additional_notes: str = "",
    engine_version: str = "0.8.0",
    manifest_footer: str = "",
) -> bytes:
    """Plný posudok podľa Energovision šablóny — TODO.

    Zatiaľ delegujeme na one_pager + extra sekciu pre porovnanie variantov.
    Plnú šablónu (4+ grafy embed, step layout, KPI box) doplníme v Sprint 10.
    """
    doc = Document()

    # Reuse one_pager pre prvé strany
    one_pager_bytes = generate_one_pager(
        client_name=client_name, project_name=project_name,
        variant=variant, site_meta=site_meta,
        dotacia_scheme=dotacia_scheme, additional_notes=additional_notes,
        engine_version=engine_version, manifest_footer=manifest_footer,
    )

    # Reopen + pridať porovnanie variantov
    doc = Document(io.BytesIO(one_pager_bytes))
    doc.add_page_break()

    _add_heading(doc, "Porovnanie všetkých variantov", level=1, color=EV_GREEN)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Variant", "PV (kWp)", "BESS (kWh)", "NPV (€)",
                                 "IRR (%)", "Payback (r)"]):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for v in sorted(all_variants, key=lambda x: -x.get("npv_eur", 0)):
        row = table.add_row().cells
        row[0].text = v["label"]
        row[1].text = f"{v['pv_kwp']:.0f}"
        row[2].text = f"{v['bess_kwh']:.0f}"
        row[3].text = f"{v['npv_eur']:,.0f}"
        row[4].text = f"{v['irr_pct']:.1f}" if v.get('irr_pct') else "—"
        row[5].text = f"{v['payback_simple_y']:.1f}"

    doc.add_paragraph()
    _add_heading(doc, "Top picks (kritériá)", level=2)
    for tp in top_picks:
        p = doc.add_paragraph()
        p.add_run(f"• {tp['label']}: ").bold = True
        p.add_run(f"{tp['variant_id']} (NPV {tp['npv_eur']:,.0f} €)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
